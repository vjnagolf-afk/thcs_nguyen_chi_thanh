# -*- coding: utf-8 -*-
"""
============================================================
XUẤT BẢN WORD - BỘ ĐIỀU PHỐI TRUNG TÂM (WORD EXPORT ENGINE)
FILE: export/word_export_engine.py
============================================================
"""

import io
import re
import logging
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

class WordExportEngine:
    @staticmethod
    def convert_markdown_to_docx_bytes(markdown_text, template_path=None):
        try:
            # 1. LOAD TEMPLATE HOẶC TẠO MỚI AN TOÀN
            if template_path and os.path.exists(template_path):
                doc = Document(template_path)
            else:
                doc = Document()
            
            # 2. THỐNG NHẤT LỀ VĂN BẢN CHUẨN 5512 (Top=1.5, Bot=1.5, Left=2.0, Right=1.5)
            for section in doc.sections:
                section.top_margin = Cm(1.5)
                section.bottom_margin = Cm(1.5)
                section.left_margin = Cm(2.0)
                section.right_margin = Cm(1.5)
            
            # 3. THIẾT LẬP STYLE CƠ BẢN (Times New Roman, 13pt)
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(13)
            
            if not markdown_text:
                markdown_text = "Không có nội dung xuất bản."

            lines = str(markdown_text).split('\n')
            
            for line in lines:
                line_clean = line.strip()
                if not line_clean:
                    continue
                
                # --- Heading Renderer ---
                if line_clean.startswith('#'):
                    level = len(line_clean) - len(line_clean.lstrip('#'))
                    text = line_clean.lstrip('#').strip()
                    p = doc.add_heading(text, level=min(level, 3))
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.bold = True
                        run.font.color.rgb = None 
                        run.font.size = Pt(16) if level == 1 else Pt(14)
                    continue

                # --- Paragraph & List Renderer ---
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Làm sạch dấu ** thừa nếu AI sinh sót lại
                line_clean = line_clean.replace('**', '')

                # Xử lý thụt lề thông minh theo từng loại dòng
                if line_clean.startswith('*'):
                    # Các bước tổ chức thực hiện (*Chuyển giao, *Thực hiện...)
                    p.paragraph_format.left_indent = Cm(0.5)
                    parts = line_clean.lstrip('* ').split(':', 1)
                    if len(parts) == 2:
                        run_bold = p.add_run(parts[0] + ":")
                        run_bold.bold = True
                        p.add_run(parts[1])
                    else:
                        p.add_run(line_clean)
                elif re.match(r'^[a-zA-Z]\)\s+', line_clean) or re.match(r'^\d+\.\s+', line_clean) or line_clean.startswith('-'):
                    # Các mục a), b), c), d) hoặc danh sách gạch đầu dòng
                    p.paragraph_format.left_indent = Cm(0.3)
                    p.add_run(line_clean)
                else:
                    # Đoạn văn bản thông thường (thụt dòng đầu dòng 1cm)
                    p.paragraph_format.first_line_indent = Cm(1.0)
                    
                    # --- Inline Tokenizer (Tách LaTeX Math & Text thường) ---
                    tokens = re.split(r'(\$\$.*?\$\$|\$.*?\$)', line_clean)
                    for token in tokens:
                        if not token:
                            continue
                        if token.startswith('$$') and token.endswith('$$'):
                            run = p.add_run(token[2:-2])
                            run.italic = True
                        elif token.startswith('$') and token.endswith('$'):
                            run = p.add_run(token[1:-1])
                            run.italic = True
                        else:
                            p.add_run(token)

            # 5. LƯU THÀNH BYTES AN TOÀN (Đảm bảo nút tải không bao giờ bị khóa)
            f = io.BytesIO()
            doc.save(f)
            return f.getvalue()
            
        except Exception as e:
            logger.error(f"Lỗi xuất bản file Word: {str(e)}")
            # Tạo một file Word chứa thông báo lỗi để Streamlit không bị treo nút tải
            err_doc = Document()
            err_doc.add_paragraph(f"Đã xảy ra lỗi khi tạo file Word: {str(e)}")
            f = io.BytesIO()
            err_doc.save(f)
            return f.getvalue()
