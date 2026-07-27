# -*- coding: utf-8 -*-
r"""
============================================================
MODULE: modules/ho_tro_gv/xd_cham_viet.py
Nhiệm vụ: Trợ lý Chấm bài Tự luận/Viết (Đa phương thức).
Đọc File (Word, PDF), Đọc ảnh chữ viết tay (JPG, PNG).
Kết nối trực tiếp qua ai_engine_2.py
============================================================
"""

import os
import io
import re
import logging
import streamlit as st
from PIL import Image
from docx import Document

logger = logging.getLogger(__name__)

# ============================================================
# 1. BỘ TRÍCH XUẤT DỮ LIỆU ĐA PHƯƠNG THỨC
# ============================================================
def extract_file_content(uploaded_file):
    """
    Trích xuất chữ và ảnh từ file nộp của học sinh/đáp án của giáo viên.
    Trả về: {"text": "...", "images": [PIL.Image, ...]}
    """
    result = {"text": "", "images": []}
    if not uploaded_file:
        return result

    file_name = uploaded_file.name.lower()
    file_bytes = uploaded_file.getvalue()

    try:
        # XỬ LÝ ẢNH CHỤP (Hỗ trợ AI đọc chữ viết tay)
        if file_name.endswith(('.png', '.jpg', '.jpeg', '.webp')):
            img = Image.open(io.BytesIO(file_bytes))
            result["images"].append(img)
            
        # XỬ LÝ WORD (DOCX)
        elif file_name.endswith('.docx'):
            doc = Document(io.BytesIO(file_bytes))
            texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    texts.append(" | ".join([cell.text.replace("\n", " ").strip() for cell in row.cells]))
            result["text"] = "\n".join(texts)

        # XỬ LÝ PDF (Trích xuất Text và Ảnh bên trong)
        elif file_name.endswith('.pdf'):
            import fitz  # PyMuPDF
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            texts = []
            for i in range(len(doc)):
                page = doc[i]
                texts.append(page.get_text("text"))
                for img_info in page.get_images(full=True):
                    try:
                        base_image = doc.extract_image(img_info[0])
                        img = Image.open(io.BytesIO(base_image["image"]))
                        result["images"].append(img)
                    except: pass
            result["text"] = "\n".join(texts)

        # XỬ LÝ TXT
        elif file_name.endswith('.txt'):
            result["text"] = file_bytes.decode('utf-8')

    except Exception as e:
        logger.error(f"Lỗi trích xuất file {file_name}: {e}")
        raise ValueError(f"Không thể đọc file {file_name}. Lỗi: {e}")

    return result

# ============================================================
# 2. XÂY DỰNG PROMPT CHẤM THI CHUYÊN GIA
# ============================================================
def build_grading_prompt(loai_bai, thang_diem, rubric_text, student_text, student_images):
    prompt_sys = r"""Bạn là một Giáo viên dày dặn kinh nghiệm, tâm huyết, chấm bài công tâm và vô cùng tỉ mỉ.
Nhiệm vụ của bạn là chấm bài làm của học sinh dựa CHÍNH XÁC vào Rubric/Đáp án chuẩn được cung cấp.

[KỶ LUẬT ĐỊNH DẠNG - BẮT BUỘC]
- Tuyệt đối không dùng dấu backtick (`) cho công thức Toán/Lý/Hóa. Mọi công thức phải bọc trong $...$ (VD: $x^2 = 4$, $H_2O$).
- Xuất kết quả bằng định dạng Markdown rõ ràng, dùng Heading để chia các phần.
"""

    if loai_bai == "Trắc nghiệm":
        prompt_task = f"""
# YÊU CẦU CHẤM TRẮC NGHIỆM:
1. Đối chiếu bài làm của học sinh với ĐÁP ÁN CHUẨN.
2. Thống kê rõ: Số câu ĐÚNG, Số câu SAI, Số câu CHƯA LÀM.
3. Liệt kê chi tiết những câu học sinh làm sai và chỉ ra đáp án đúng.
4. Tính điểm chính xác theo thang điểm {thang_diem}. (Điểm = (Số câu đúng / Tổng câu) * {thang_diem}).
5. Đưa ra 1-2 câu nhận xét ngắn gọn, khích lệ.
"""
    elif loai_bai == "Tự luận / Viết":
        prompt_task = f"""
# YÊU CẦU CHẤM TỰ LUẬN / VIẾT:
1. Đối chiếu chặt chẽ từng ý trong bài làm của học sinh với TIÊU CHÍ/RUBRIC.
2. Đánh giá chi tiết từng tiêu chí: Đạt được ý nào? Thiếu ý nào? Phân tích nông/sâu ra sao?
3. Nhặt sạn (nếu có): Chỉ ra các lỗi sai về kiến thức, lỗi chính tả, diễn đạt, ngữ pháp.
4. Chấm điểm từng phần và Tổng hợp điểm số dự kiến trên thang điểm {thang_diem}.
5. Đưa ra lời nhận xét đánh giá tổng quan (Ưu điểm, Hạn chế) và Lời khuyên để học sinh cải thiện.
"""
    else:
        prompt_task = f"""
# YÊU CẦU CHẤM BÀI HỖN HỢP (TRẮC NGHIỆM + TỰ LUẬN):
1. PHẦN TRẮC NGHIỆM: Chỉ ra câu sai, đáp án đúng. Chấm điểm thành phần.
2. PHẦN TỰ LUẬN: Đối chiếu Rubric, đánh giá thiếu/đủ ý, sửa lỗi diễn đạt/kiến thức. Chấm điểm thành phần.
3. TỔNG KẾT: Cộng tổng điểm trên thang {thang_diem}.
4. Lời nhận xét tổng quan và khích lệ học sinh.
"""

    image_note = ""
    if student_images:
        image_note = "\n\n[LƯU Ý QUAN TRỌNG: Học sinh có nộp bài làm dưới dạng HÌNH ẢNH. BẮT BUỘC dùng khả năng thị giác (Vision) của bạn để đọc thật kỹ chữ viết tay hoặc nội dung trên ảnh để chấm điểm]."

    prompt_user = f"""
{prompt_task}

--- ĐÁP ÁN CHUẨN / RUBRIC CHẤM ---
{rubric_text}

--- BÀI LÀM CỦA HỌC SINH (Dạng Text) ---
{student_text if student_text else "(Học sinh nộp bài qua ảnh/file đính kèm)"}
{image_note}
"""
    return prompt_sys, prompt_user

# ============================================================
# 3. KẾT NỐI AI ENGINE 2
# ============================================================
def call_ai_grader(ai_engine, prompt_sys, prompt_user, images=None):
    """
    Giao tiếp với `ai_engine_2.py`.
    Thích ứng thông minh với các hàm (generate_text, generate_content) có trong Engine của dự án.
    """
    full_prompt = f"{prompt_sys}\n\n{prompt_user}"
    
    try:
        # Nếu có ảnh bài làm (Chữ viết tay)
        if images:
            # Ưu tiên tìm hàm truyền list đa phương thức chuẩn SDK mới
            if hasattr(ai_engine, "client") and hasattr(ai_engine.client, "models"):
                contents = [full_prompt] + images
                response = ai_engine.client.models.generate_content(
                    model="gemini-2.5-pro", # Dùng Pro để đọc chữ viết tay tốt nhất
                    contents=contents
                )
                return response.text
            # Nếu ai_engine có xây dựng sẵn hàm cho ảnh
            elif hasattr(ai_engine, "generate_with_images"):
                return ai_engine.generate_with_images(full_prompt, images)
            else:
                st.warning("⚠️ Engine AI hiện tại chưa hỗ trợ hàm đọc ảnh. Chuyển sang chấm bằng văn bản trích xuất được...")
        
        # Fallback về chấm bằng Text thông thường
        if hasattr(ai_engine, "generate_text"):
            return ai_engine.generate_text(full_prompt)
        elif hasattr(ai_engine, "generate"):
            return ai_engine.generate(full_prompt)
        else:
            raise AttributeError("Không tìm thấy hàm generate_text trong ai_engine.")
            
    except Exception as e:
        logger.error(f"Lỗi AI Chấm bài: {e}")
        raise RuntimeError(f"Sự cố khi gọi ai_engine: {e}")

# ============================================================
# 4. GIAO DIỆN HIỂN THỊ (VIEW)
# ============================================================
def render_xd_cham_viet(ai_engine=None):
    if ai_engine is None:
        st.error("❌ Hệ thống chưa kết nối được với AI Engine. Vui lòng kiểm tra lại cấu hình.")
        return

    st.markdown("## 📝 Trợ lý AI Chấm bài Tự luận & Trắc nghiệm")
    st.info("💡 **Tính năng Đa phương thức:** Thầy/Cô có thể dán văn bản, tải lên file Word, PDF hoặc **chụp ảnh bài làm viết tay** của học sinh. AI sẽ tự động đọc chữ, đối chiếu Đáp án/Rubric và chấm điểm chi tiết.")

    # KHU VỰC 1: CẤU HÌNH
    with st.expander("⚙️ Cấu hình Bài chấm", expanded=True):
        col_type, col_score = st.columns(2)
        with col_type:
            loai_bai = st.radio("Loại bài kiểm tra:", ["Trắc nghiệm", "Tự luận / Viết", "Hỗn hợp (Cả 2)"], horizontal=True)
        with col_score:
            thang_diem = st.number_input("Thang điểm tối đa:", min_value=1, max_value=100, value=10)

    # KHU VỰC 2: NHẬP LIỆU
    col_hs, col_gv = st.columns(2)

    with col_gv:
        st.markdown("### 🎯 Tiêu chí / Đáp án chuẩn")
        rubric_file = st.file_uploader("Tải lên Đáp án (Word/PDF):", type=["docx", "pdf", "txt"], key="rubric_file")
        rubric_text = st.text_area("Hoặc nhập Tiêu chí chấm:", height=250, 
                                 placeholder="- Câu 1: A, Câu 2: B...\n- Ý 1: Nêu nguyên nhân (2đ)\n- Ý 2: Hậu quả (3đ)...")

    with col_hs:
        st.markdown("### 📄 Bài làm của Học sinh")
        student_file = st.file_uploader("Tải lên Bài làm (Ảnh chụp, Word, PDF):", 
                                      type=["png", "jpg", "jpeg", "webp", "docx", "pdf", "txt"], 
                                      accept_multiple_files=True, key="student_file")
        student_text = st.text_area("Hoặc dán trực tiếp đoạn văn bài làm:", height=250, 
                                  placeholder="Nhập bài làm của học sinh vào đây...")

    # KHU VỰC 3: XỬ LÝ VÀ CHẤM BÀI
    if st.button("🚀 BẮT ĐẦU CHẤM BÀI", type="primary", use_container_width=True):
        
        # Tiền xử lý dữ liệu Đáp án
        final_rubric = rubric_text
        if rubric_file:
            try:
                extracted_rubric = extract_file_content(rubric_file)["text"]
                final_rubric = f"{extracted_rubric}\n\n{rubric_text}".strip()
            except Exception as e:
                st.error(f"Lỗi đọc file Đáp án: {e}")
                return
                
        if not final_rubric:
            st.warning("⚠️ Vui lòng cung cấp Tiêu chí chấm hoặc Đáp án chuẩn.")
            return

        if not student_file and not student_text.strip():
            st.warning("⚠️ Vui lòng tải lên file bài làm hoặc nhập chữ vào ô Bài làm học sinh.")
            return

        try:
            with st.spinner("⏳ Đang trích xuất và đọc ảnh/chữ viết tay bài làm của học sinh..."):
                final_student_text = student_text
                student_images = []
                
                if student_file:
                    for sf in student_file:
                        extracted = extract_file_content(sf)
                        if extracted["text"]:
                            final_student_text += f"\n\n{extracted['text']}"
                        if extracted["images"]:
                            student_images.extend(extracted["images"])

            with st.spinner("🧠 AI đang chấm bài, nhặt sạn và đối chiếu với Rubric..."):
                # Sinh Prompt
                sys_prompt, user_prompt = build_grading_prompt(
                    loai_bai, thang_diem, final_rubric, final_student_text, student_images
                )
                
                # Gọi Cỗ máy AI (ai_engine_2.py)
                result = call_ai_grader(ai_engine, sys_prompt, user_prompt, student_images)
                
                # Hiển thị
                st.markdown("---")
                st.markdown("### 💯 BẢNG ĐIỂM & NHẬN XÉT CHI TIẾT")
                st.markdown(result, unsafe_allow_html=True)
                
        except Exception as e:
            st.error(f"❌ Quá trình chấm bài thất bại: {e}")
