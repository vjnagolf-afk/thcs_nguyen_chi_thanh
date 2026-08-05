# -*- coding: utf-8 -*-
import io
import os
import time
import logging
import streamlit as st
from PIL import Image
import urllib.parse as urlparse

logger = logging.getLogger(__name__)

# Kết nối bộ xuất Word
try:
    from export.export_word import export_word
except ImportError:
    export_word = None

# Bắt buộc import AIEngine2 để dùng tính năng Multimodal
try:
    from utils.ai_engine_2 import AIEngine2
except ImportError:
    AIEngine2 = None

# Thử import thư viện QR Code
try:
    import qrcode
    HAS_QRCODE = True
except ImportError:
    HAS_QRCODE = False

# ============================================================
# HÀM TRÍCH XUẤT TÀI LIỆU ĐA NGUỒN (PDF, DOCX, ẢNH)
# ============================================================
def extract_content_from_source(uploaded_file):
    if not uploaded_file:
        return "", []
    
    file_name = uploaded_file.name.lower()
    file_bytes = uploaded_file.getvalue()
    extracted_text = ""
    images = []

    try:
        if file_name.endswith(('.png', '.jpg', '.jpeg', '.webp')):
            img = Image.open(io.BytesIO(file_bytes)).convert('RGB')
            images.append(img)
            extracted_text = "[Học sinh/Giáo viên tải lên hình ảnh tài liệu]"
            
        elif file_name.endswith('.docx'):
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    texts.append(" | ".join([cell.text.replace("\n", " ").strip() for cell in row.cells]))
            extracted_text = "\n".join(texts)

        elif file_name.endswith('.pdf'):
            import fitz  # PyMuPDF
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            texts = []
            for i in range(len(doc)):
                page = doc[i]
                texts.append(page.get_text("text"))
                for img_info in page.get_images(full=True):
                    try:
                        base_image = doc.extract_image(img_info[0])
                        img = Image.open(io.BytesIO(base_image["image"])).convert('RGB')
                        images.append(img)
                    except: pass
            extracted_text = "\n".join(texts)

        elif file_name.endswith('.txt'):
            extracted_text = file_bytes.decode('utf-8', errors='ignore')

    except Exception as e:
        logger.error(f"Lỗi trích xuất file {file_name}: {e}")
        
    return extracted_text, images

# HÀM CROSS-ROUTING FALLBACK CHỐNG LỖI 429
def safe_generate_quiz(ai_engine_cu, prompt, extracted_images=[]):
    api_key = None
    for key, val in st.session_state.items():
        if isinstance(val, str) and val.startswith("sk-"):
            api_key = val
            break
            
    if not api_key:
        for k in ["user_api_key", "api_key", "openai_api_key", "sk_key"]:
            if st.session_state.get(k) and str(st.session_state.get(k)).startswith("sk-"):
                api_key = st.session_state.get(k)
                break
                
    if not api_key and "OPENAI_API_KEY" in st.secrets:
        api_key = st.secrets["OPENAI_API_KEY"]

    def run_openai():
        if not api_key:
            raise RuntimeError("Chưa cấu hình API Key OpenAI (sk-) để dự phòng.")
        import openai
        import base64
        client = openai.OpenAI(api_key=str(api_key).strip())
        
        content_array = [{"type": "text", "text": prompt}]
        for img in extracted_images:
            buffered = io.BytesIO()
            img.save(buffered, format="JPEG")
            base64_img = base64.b64encode(buffered.getvalue()).decode('utf-8')
            content_array.append({
                "type": "image_url",
                "image_url": {"url": f"data:image/jpeg;base64,{base64_img}"}
            })
            
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": content_array}],
            temperature=0.7
        )
        return response.choices[0].message.content.strip()

    def run_gemini():
        try:
            from utils.ai_engine_2 import AIEngine2
            engine_v2 = AIEngine2(default_model="gemini-1.5-flash")
            
            if extracted_images:
                contents = [prompt]
                for img in extracted_images:
                    buffered = io.BytesIO()
                    img.save(buffered, format="JPEG")
                    contents.append({
                        "mime_type": "image/jpeg",
                        "data": buffered.getvalue()
                    })
                if hasattr(engine_v2, "generate_multimodal"):
                    res = engine_v2.generate_multimodal(contents)
                else:
                    raise RuntimeError("❌ Thiếu hàm generate_multimodal.")
            else:
                res = engine_v2.generate_text(prompt, temperature=0.7)
                
            if res and not res.startswith("❌") and not res.startswith("⚠️") and "429" not in res and "RESOURCE_EXHAUSTED" not in res:
                return res
            raise RuntimeError("Hạn mức Gemini cạn kiệt.")
        except Exception as e:
            raise RuntimeError(f"Lỗi Gemini: {str(e)}")

    error_msgs = []
    try:
        return run_gemini()
    except Exception as e1:
        error_msgs.append(f"Gemini: {e1}")
        try:
            return run_openai()
        except Exception as e2:
            error_msgs.append(f"OpenAI: {e2}")
            
    raise RuntimeError(f"Hệ thống quá tải hoặc hết hạn mức:\n- {error_msgs[0]}\n- {error_msgs[1]}\n\n👉 Nạp Key OpenAI (sk-) vào hệ thống để chạy ổn định.")


def render_xd_quizizz(ai_engine_cu=None):
    if "quiz_result" not in st.session_state:
        st.session_state["quiz_result"] = None
    if "quiz_topic" not in st.session_state:
        st.session_state["quiz_topic"] = "Bo_Cau_Hoi"

    st.markdown("### ⚡ Trợ lý Tạo Bộ Câu Hỏi & Tổ Chức Quiz Trực Tuyến")
    st.info("💡 **Góc chuyên gia:** Tạo bộ câu hỏi, quét QR code kết nối học sinh, trình chiếu link bất kỳ lên tivi và tự động chấm điểm dựa trên đáp án chuẩn của giáo viên!")

    tab_tao_de, tab_nhung = st.tabs([
        "🛠️ 1. Soạn thảo Bộ câu hỏi (AI Builder)", 
        "🏆 2. Phòng Thi Đấu & Chấm Điểm Tự Động"
    ])

    # ========================================================
    # TAB 1: SOẠN THẢO & TẠO BỘ CÂU HỎI
    # ========================================================
    with tab_tao_de:
        with st.container(border=True):
            st.markdown("#### 🎯 Chọn phương thức nạp dữ liệu (Nguồn AI)")
            
            nguon_nhap = st.radio(
                "Nguồn dữ liệu đầu vào:",
                [
                    "✍️ Nhập chủ đề trực tiếp", 
                    "📁 Trích xuất từ tệp (PDF, Word, Ảnh)", 
                    "📺 Trích xuất từ YouTube (Link video)", 
                    "🌐 Trích xuất từ trang web (URL)"
                ],
                horizontal=True,
                label_visibility="collapsed"
            )

            input_data_content = ""
            uploaded_file = None
            extracted_images = []

            if "Nhập chủ đề" in nguon_nhap:
                input_data_content = st.text_input("Nhập chủ đề bài kiểm tra:", placeholder="VD: Định luật Ôm, Phản ứng hóa học lớp 9...")
            elif "Trích xuất từ tệp" in nguon_nhap:
                uploaded_file = st.file_uploader("Tải lên tài liệu (PDF, Word, TXT, Ảnh):", type=["pdf", "docx", "txt", "png", "jpg", "jpeg"])
                if uploaded_file:
                    with st.spinner("Đang đọc dữ liệu từ tệp..."):
                        input_data_content, extracted_images = extract_content_from_source(uploaded_file)
                        st.success(f"✅ Đã đọc thành công tệp: {uploaded_file.name}")
            elif "YouTube" in nguon_nhap:
                input_data_content = st.text_input("Dán đường dẫn (URL) Video YouTube:", placeholder="https://www.youtube.com/watch?v=...")
            else:
                input_data_content = st.text_input("Dán đường dẫn (URL) Trang web tài liệu:", placeholder="https://vi.wikipedia.org/wiki/...")

            st.markdown("---")
            st.markdown("#### ⚙️ Cấu hình định dạng câu hỏi & Số lượng")
            
            col_cfg1, col_cfg2, col_cfg3 = st.columns(3)
            with col_cfg1:
                so_luong = st.number_input("Tổng số câu hỏi:", min_value=5, max_value=50, value=10, step=5)
            with col_cfg2:
                do_kho = st.selectbox("Mức độ khó:", ["Nhận biết", "Thông hiểu", "Vận dụng", "Hỗn hợp các mức độ"])
            with col_cfg3:
                do_ut_cau_hoi = st.multiselect(
                    "Các dạng câu hỏi bao gồm:",
                    ["Trắc nghiệm (4 lựa chọn)", "Trả lời ngắn", "Đúng / Sai", "Bài luận"],
                    default=["Trắc nghiệm (4 lựa chọn)", "Đúng / Sai"]
                )

            them_chi_tiet = st.text_area("Yêu cầu thêm (Tuỳ chọn):", height=70, placeholder="VD: Tập trung vào phần bài tập tính toán...")
            
            btn_tao_quiz = st.button("🚀 TẠO BỘ CÂU HỎI THÔNG MINH", type="primary", use_container_width=True)

        if btn_tao_quiz:
            if not input_data_content.strip() and not uploaded_file:
                st.warning("⚠️ Vui lòng cung cấp nội dung chủ đề hoặc tải lên tệp tài liệu.")
            else:
                with st.spinner("⏳ AI đang phân tích tài liệu và biên soạn bộ câu hỏi..."):
                    types_str = ", ".join(do_ut_cau_hoi)
                    prompt = f"""BẠN LÀ CHUYÊN GIA BIÊN SOẠN CÂU HỎI TRẮC NGHIỆM VÀ ĐÁNH GIÁ NĂNG LỰC.
- Nguồn: {nguon_nhap}
- Nội dung: {input_data_content[:15000]}
- Số lượng: {so_luong} câu
- Mức độ: {do_kho}
- Dạng câu: {types_str}
- Yêu cầu thêm: {them_chi_tiet if them_chi_tiet else 'Không'}

Biên soạn rõ ràng theo định dạng Markdown chuẩn:
### Câu [Số]: [Nội dung câu hỏi]
- **Dạng câu hỏi:** [...]
- **Các đáp án:** A. ... B. ... C. ... D. ...
- **Đáp án đúng:** [...]
- **Giải thích chi tiết:** [...]
NẾU có công thức Toán/Lý/Hóa, BẮT BUỘC bọc trong dấu `$ ... $`."""
                    
                    try:
                        result = safe_generate_quiz(ai_engine_cu, prompt, extracted_images)
                        st.session_state["quiz_result"] = result
                        st.session_state["quiz_topic"] = "Bo_Cau_Hoi_Quiz"
                    except Exception as e:
                        st.error(f"❌ Lỗi hệ thống: {e}")

        if st.session_state.get("quiz_result"):
            st.markdown("---")
            st.markdown("### 📑 BỘ CÂU HỎI ĐÃ ĐƯỢC BIÊN SOẠN")
            st.markdown(st.session_state["quiz_result"], unsafe_allow_html=True)
            
            st.markdown("### 📥 Lưu trữ & Xuất tệp")
            col_d1, col_d2 = st.columns(2)
            with col_d1:
                st.download_button("📄 Tải bộ câu hỏi (.TXT)", data=st.session_state["quiz_result"], file_name="Bo_Cau_Hoi_Quizizz.txt", mime="text/plain", use_container_width=True)
            with col_d2:
                if export_word:
                    try:
                        export_data = {"ai_generated_content": st.session_state["quiz_result"], "is_dkt": False}
                        word_bytes = export_word(export_data)
                        st.download_button("📘 Tải bộ câu hỏi (.DOCX)", data=word_bytes, file_name="Bo_Cau_Hoi_Quizizz.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, type="primary")
                    except Exception as e:
                        st.error(f"Lỗi xuất Word: {e}")

    # ========================================================
    # TAB 2: PHÒNG THI ĐẤU & CHẤM ĐIỂM TỰ ĐỘNG
    # ========================================================
    with tab_nhung:
        st.markdown("#### 🏆 Phòng Thi Đấu Trực Tuyến & Chấm Điểm Tự Động")
        
        with st.expander("🔑 Cấu hình Máy chủ Pusher (Nhập 1 lần)", expanded=False):
            pusher_app_id = st.text_input("Pusher App ID:", type="password")
            pusher_key = st.text_input("Pusher Key:", type="password")
            pusher_secret = st.text_input("Pusher Secret:", type="password")
            pusher_cluster = st.text_input("Pusher Cluster (VD: ap1):")
        
        query_params = st.query_params
        default_phong = query_params.get("phong", "")
        default_role_idx = 1 if default_phong else 0

        vai_tro = st.radio(
            "Thầy/Cô đang mở máy tính này với vai trò gì?", 
            ["🖥️ Giáo viên (Host - Trình chiếu & Chấm điểm)", "📱 Học sinh (Client - Gửi đáp án)"], 
            index=default_role_idx,
            horizontal=True
        )
        
        st.markdown("---")
        
        if not (pusher_app_id and pusher_key and pusher_secret and pusher_cluster):
            st.warning("⚠️ Cần cấu hình Máy chủ Pusher để kích hoạt tính năng thi đấu chấm điểm.")
            st.markdown("#### 🌐 Khung Trình Chiếu Đơn Thuần (Chưa kết nối)")
            embed_input = st.text_input("Dán Link liên kết hoặc Mã nhúng (Iframe / URL bất kỳ):", placeholder="VD: Link Quizizz, Blooket, Kahoot, hoặc Web tùy ý...")
            target_url = embed_input.strip()
            if "src=" in target_url:
                try:
                    import re
                    match = re.search(r'src=["\'](.*?)["\']', target_url)
                    if match: target_url = match.group(1)
                except: pass
            if target_url:
                st.components.v1.iframe(target_url, height=600, scrolling=True)

        else:
            try:
                import pusher
                pusher_client = pusher.Pusher(
                    app_id=pusher_app_id,
                    key=pusher_key,
                    secret=pusher_secret,
                    cluster=pusher_cluster,
                    ssl=True
                )
                
                # =========================================
                # MÀN HÌNH GIÁO VIÊN
                # =========================================
                if "Giáo viên" in vai_tro:
                    st.markdown("### 👑 BẢNG ĐIỀU KHIỂN & CHẤM ĐIỂM GIÁO VIÊN")
                    
                    col_host1, col_host2 = st.columns([1, 2])
                    
                    with col_host1:
                        st.markdown("#### 1. Mời Học Sinh (QR Code)")
                        ma_phong = st.text_input("Mã Phòng thi:", value="123456")
                        app_url = st.text_input("Link phần mềm này:", placeholder="VD: https://thcsnguyenchithanh-lhd.streamlit.app/xd_quizizz")
                        
                        if app_url and HAS_QRCODE:
                            parsed = urlparse.urlparse(app_url)
                            query_dict = urlparse.parse_qs(parsed.query)
                            query_dict["phong"] = [ma_phong]
                            new_query = urlparse.urlencode(query_dict, doseq=True)
                            deep_link = urlparse.urlunparse((parsed.scheme, parsed.netloc, parsed.path, parsed.params, new_query, parsed.fragment))
                            
                            qr = qrcode.QRCode(version=1, box_size=5, border=2)
                            qr.add_data(deep_link)
                            qr.make(fit=True)
                            img_qr = qr.make_image(fill_color="black", back_color="white")
                            buf = io.BytesIO()
                            img_qr.save(buf, format="PNG")
                            st.image(buf.getvalue(), caption=f"Quét để vào phòng: {ma_phong}", width=200)
                        elif not HAS_QRCODE:
                            st.warning("Cài đặt `pip install qrcode[pil]` để tạo mã QR.")

                        st.markdown("#### 3. Đáp Án Chuẩn (Để chấm điểm)")
                        dap_an_dung = st.selectbox("Chọn đáp án đúng cho câu hiện tại:", ["Chưa chọn", "A", "B", "C", "D"])

                    with col_host2:
                        st.markdown("#### 2. Trình Chiếu Câu Hỏi Lên Tivi")
                        embed_input = st.text_input("Dán Link bất kỳ (Quizizz, Blooket, Web sưu tầm...):", placeholder="VD: Link bài tập được chia sẻ...")
                        target_url = embed_input.strip()
                        if "src=" in target_url:
                            try:
                                import re
                                match = re.search(r'src=["\'](.*?)["\']', target_url)
                                if match: target_url = match.group(1)
                            except: pass
                        if target_url:
                            st.components.v1.iframe(target_url, height=450, scrolling=True)
                        else:
                            st.info("💡 Dán link vào ô trên để hiển thị câu hỏi lên màn hình tivi cho học sinh quan sát.")

                    st.markdown("---")
                    st.markdown("#### 🏆 BẢNG KẾT QUẢ & CHẤM ĐIỂM TRỰC TUYẾN")
                    
                    col_btn1, col_btn2 = st.columns(2)
                    with col_btn1:
                        if st.button("🚀 PHÁT TÍN HIỆU BẮT ĐẦU TRẢ LỜI", type="primary", use_container_width=True):
                            current_time = time.time()
                            pusher_client.trigger(f'phong-{ma_phong}', 'cau_moi', {'thoi_gian': current_time})
                            st.success("✅ Đã phát tín hiệu mở cổng trả lời cho học sinh!")
                    with col_btn2:
                        if st.button("🛑 KHÓA CỔNG TRẢ LỜI", use_container_width=True):
                            pusher_client.trigger(f'phong-{ma_phong}', 'het_gio', {'message': 'stop'})
                            st.error("🛑 Đã khóa cổng trả lời!")

                    # RADAR NHẬN ĐÁP ÁN & HIỂN THỊ DANH SÁCH + TỰ ĐỘNG CHẤM ĐIỂM
                    html_code = f"""
                    <!DOCTYPE html>
                    <html>
                    <head>
                      <script src="https://js.pusher.com/8.0/pusher.min.js"></script>
                      <style>
                        body {{ font-family: sans-serif; font-size: 14px; margin: 0; background: #f9f9f9; color: #333; }}
                        .container {{ display: flex; gap: 15px; }}
                        .box {{ flex: 1; height: 260px; overflow-y: auto; background: #1e1e1e; color: #00ff00; padding: 12px; border-radius: 8px; font-family: monospace; line-height: 1.4; }}
                        .score-board {{ flex: 1; height: 260px; overflow-y: auto; background: #fff; color: #333; padding: 12px; border-radius: 8px; border: 1px solid #ccc; }}
                        h4 {{ margin: 0 0 8px 0; font-size: 14px; color: #555; }}
                        .correct {{ color: #4caf50; font-weight: bold; }}
                        .incorrect {{ color: #f44336; font-weight: bold; }}
                      </style>
                    </head>
                    <body>
                      <div class="container">
                        <div class="box">
                          <h4>📡 Nhật ký kết nối (Live Log):</h4>
                          <div id="log">[Hệ thống] Sẵn sàng nhận đáp án...<br></div>
                        </div>
                        <div class="score-board">
                          <h4>🏆 Danh sách đáp án học sinh gửi về:</h4>
                          <div id="results">Chưa có học sinh gửi đáp án.<br></div>
                        </div>
                      </div>
                      <script>
                        var pusher = new Pusher('{pusher_key}', {{ cluster: '{pusher_cluster}' }});
                        var channel = pusher.subscribe('phong-{ma_phong}');
                        var logDiv = document.getElementById('log');
                        var resDiv = document.getElementById('results');
                        var studentAnswers = {{}};
                        var startTime = 0;
                        var correctAns = "{dap_an_dung}";

                        channel.bind('hs_vao_phong', function(data) {{
                          logDiv.innerHTML += '👋 [' + data.nhom + '] đã vào phòng.<br>';
                          logDiv.scrollTop = logDiv.scrollHeight;
                        }});

                        channel.bind('cau_moi', function(data) {{
                          startTime = data.thoi_gian;
                          studentAnswers = {{}};
                          resDiv.innerHTML = 'Đang chờ đáp án vòng mới...<br>';
                          logDiv.innerHTML += '<br>🚀 BẮT ĐẦU CÂU HỎI MỚI!<br>';
                          logDiv.scrollTop = logDiv.scrollHeight;
                        }});

                        channel.bind('nop_dap_an', function(data) {{
                          var timeTaken = startTime > 0 ? (data.thoi_gian - startTime).toFixed(2) + 's' : '';
                          studentAnswers[data.nhom] = {{ ans: data.dap_an, time: timeTaken }};
                          
                          // Hiển thị lại bảng đáp án
                          var html = '<ul>';
                          for (var student in studentAnswers) {{
                            var item = studentAnswers[student];
                            var status = "";
                            if (correctAns !== "Chưa chọn") {{
                              status = (item.ans === correctAns) ? '<span class="correct"> ✔️ Đúng (+1đ)</span>' : '<span class="incorrect"> ❌ Sai</span>';
                            }}
                            html += '<li><b>' + student + '</b> chọn: <span style="font-size:16px; color:blue;">' + item.ans + '</span> (' + item.time + ')' + status + '</li>';
                          }}
                          html += '</ul>';
                          resDiv.innerHTML = html;
                          logDiv.innerHTML += '🎯 [' + student.nhom + '] vừa chọn ' + data.dap_an + '<br>';
                          logDiv.scrollTop = logDiv.scrollHeight;
                        }});
                      </script>
                    </body>
                    </html>
                    """
                    st.components.v1.html(html_code, height=300, scrolling=True)

                # =========================================
                # MÀN HÌNH HỌC SINH
                # =========================================
                else:
                    st.markdown("### 📱 TAY CẦM ĐIỀU KHIỂN HỌC SINH")
                    
                    if "hs_nhom" not in st.session_state:
                        st.session_state["hs_nhom"] = ""
                    if "hs_phong" not in st.session_state:
                        st.session_state["hs_phong"] = ""

                    if not st.session_state["hs_nhom"]:
                        st.info("Nhập thông tin để kết nối vào phòng thi.")
                        nhap_phong = st.text_input("Mã Phòng (Tự động điền nếu quét QR):", value=default_phong)
                        nhap_ten = st.text_input("Tên Nhóm / Tên Học sinh:")
                        
                        if st.button("🚪 VÀO PHÒNG THI", type="primary", use_container_width=True):
                            if nhap_phong and nhap_ten:
                                st.session_state["hs_phong"] = nhap_phong
                                st.session_state["hs_nhom"] = nhap_ten
                                try:
                                    pusher_client.trigger(
                                        f'phong-{nhap_phong}', 
                                        'hs_vao_phong', 
                                        {'nhom': nhap_ten}
                                    )
                                except: pass
                                st.rerun()
                            else:
                                st.error("Vui lòng nhập đủ Mã phòng và Tên!")
                    
                    else:
                        st.success(f"🟢 Thí sinh: **{st.session_state['hs_nhom']}** | Phòng: **{st.session_state['hs_phong']}**")
                        st.markdown("#### 📺 Quan sát câu hỏi trên tivi và chọn đáp án:")
                        
                        def send_answer(ans):
                            try:
                                pusher_client.trigger(
                                    f'phong-{st.session_state["hs_phong"]}', 
                                    'nop_dap_an', 
                                    {'nhom': st.session_state['hs_nhom'], 'dap_an': ans, 'thoi_gian': time.time()}
                                )
                                st.success(f"✅ Đã gửi đáp án **{ans}** lên hệ thống!")
                            except:
                                st.error("Lỗi kết nối mạng.")

                        st.markdown("""
                        <style>
                        div[data-testid="stButton"] button { height: 90px; font-size: 28px; font-weight: bold; }
                        </style>
                        """, unsafe_allow_html=True)
                        
                        c1, c2 = st.columns(2)
                        with c1:
                            if st.button("🟥 A", use_container_width=True): send_answer("A")
                            if st.button("🟦 C", use_container_width=True): send_answer("C")
                        with c2:
                            if st.button("🟨 B", use_container_width=True): send_answer("B")
                            if st.button("🟩 D", use_container_width=True): send_answer("D")
                            
                        if st.button("Thoát phòng"):
                            st.session_state["hs_nhom"] = ""
                            st.session_state["hs_phong"] = ""
                            st.query_params.clear()
                            st.rerun()

            except ImportError:
                st.error("❌ Thư viện kết nối chưa cài đặt. Vui lòng chạy: `pip install pusher qrcode[pil]`")
            except Exception as e:
                st.error(f"❌ Lỗi cấu hình Pusher: {e}")
