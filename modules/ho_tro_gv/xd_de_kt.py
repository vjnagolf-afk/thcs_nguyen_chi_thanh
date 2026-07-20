import streamlit as st
import sys
import os
from pathlib import Path
from string import Template

def extract_text_from_file(uploaded_file):
    try:
        if uploaded_file.name.endswith('.pdf'):
            from pypdf import PdfReader
            return "\n".join([p.extract_text() for p in PdfReader(uploaded_file).pages if p.extract_text()])
        elif uploaded_file.name.endswith('.docx'):
            import docx
            return "\n".join([para.text for para in docx.Document(uploaded_file).paragraphs])
        elif uploaded_file.name.endswith('.txt'):
            return uploaded_file.read().decode("utf-8")
    except: return ""
    return ""

def load_prompt(filename):
    # Đường dẫn trỏ đúng thư mục prompts/ trong thư mục gốc
    root_path = Path(__file__).resolve().parents[2]
    file_path = root_path / "prompts" / filename
    if not file_path.exists():
        return f"Lỗi: Không tìm thấy file {filename} tại {file_path}"
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def render_xd_de_kt(ai_engine):
    import streamlit as st
    import sys
    from pathlib import Path
    from string import Template

    # ============================================================
    # HÀM ĐỌC NỘI DUNG ĐỀ CƯƠNG
    # ============================================================
    def extract_text_from_file(uploaded_file):

        if uploaded_file is None:
            return ""

        try:
            filename = uploaded_file.name.lower()

            # ----------------------------------------------------
            # ĐỌC PDF
            # ----------------------------------------------------
            if filename.endswith(".pdf"):

                from pypdf import PdfReader

                reader = PdfReader(uploaded_file)
                texts = []

                for page_number, page in enumerate(reader.pages, start=1):

                    text = page.extract_text()

                    if text and text.strip():
                        texts.append(
                            f"\n--- TRANG {page_number} ---\n{text.strip()}"
                        )

                return "\n".join(texts)

            # ----------------------------------------------------
            # ĐỌC DOCX
            # ----------------------------------------------------
            elif filename.endswith(".docx"):

                import docx

                document = docx.Document(uploaded_file)
                texts = []

                # Đọc các đoạn văn
                for para in document.paragraphs:

                    text = para.text.strip()

                    if text:
                        texts.append(text)

                # Đọc toàn bộ bảng
                for table_index, table in enumerate(
                    document.tables,
                    start=1
                ):

                    texts.append(
                        f"\n--- BẢNG {table_index} ---"
                    )

                    for row in table.rows:

                        row_text = " | ".join(
                            cell.text.strip()
                            for cell in row.cells
                        )

                        if row_text.strip():
                            texts.append(row_text)

                return "\n".join(texts)

            # ----------------------------------------------------
            # ĐỌC TXT
            # ----------------------------------------------------
            elif filename.endswith(".txt"):

                return uploaded_file.read().decode(
                    "utf-8",
                    errors="ignore"
                )

        except Exception as e:

            st.error(
                f"❌ Lỗi khi đọc đề cương: {e}"
            )

            return ""

        return ""

    # ============================================================
    # HÀM CHUẨN HÓA ĐỀ CƯƠNG
    # ============================================================
    def normalize_outline(text):

        if not text:
            return ""

        # Xóa các dòng trống liên tiếp
        lines = []

        for line in text.splitlines():

            line = line.strip()

            if line:
                lines.append(line)

        text = "\n".join(lines)

        # Giới hạn độ dài để tránh prompt quá lớn
        # Có thể điều chỉnh nếu cần
        MAX_CHARS = 60000

        if len(text) > MAX_CHARS:

            st.warning(
                f"⚠️ Đề cương có {len(text):,} ký tự. "
                f"Hệ thống sử dụng {MAX_CHARS:,} ký tự đầu tiên."
            )

            text = text[:MAX_CHARS]

        return text

    # ============================================================
    # GIAO DIỆN
    # ============================================================
    st.markdown(
        "### 📝 Soạn thảo Ma trận, Đặc tả & Đề KT "
        "(Chuẩn 5512)"
    )

    # ============================================================
    # 1. THÔNG TIN CHUNG
    # ============================================================
    c1, c2, c3, c4, c5, c6 = st.columns(
        [1, 0.8, 1.2, 1, 2, 0.8]
    )

    mon_hoc = c1.selectbox(
        "Môn",
        [
            "Toán",
            "Ngữ văn",
            "Ngoại ngữ",
            "KHTN",
            "Lịch sử & Địa lý",
            "Tin học",
            "Khác"
        ],
        key="de_kt_mon_hoc"
    )

    lop = c2.selectbox(
        "Lớp",
        [
            "Lớp 6",
            "Lớp 7",
            "Lớp 8",
            "Lớp 9",
            "Lớp 10",
            "Lớp 11"
            "Lớp 12"
        ],
        index=2,
        key="de_kt_lop"
    )

    hinh_thuc = c3.selectbox(
        "Hình thức",
        [
            "Trắc nghiệm & Tự luận",
            "100% Trắc nghiệm",
            "100% Tự luận"
        ],
        key="de_kt_hinh_thuc"
    )

    thoi_gian = c4.selectbox(
        "Thời gian",
        [
            "15 phút",
            "45 phút",
            "90 phút"
        ],
        key="de_kt_thoi_gian"
    )

    ten_de = c5.text_input(
        "Tên bài kiểm tra",
        key="de_kt_ten_de"
    )

    with c6:

        st.markdown(
            "<div style='margin-top: 25px;'></div>",
            unsafe_allow_html=True
        )

        bam_sat = st.checkbox(
            "Bám sát đề cương",
            value=True,
            key="de_kt_bam_sat"
        )

    # ============================================================
    # 2. TẢI ĐỀ CƯƠNG
    # ============================================================
    file_de = st.file_uploader(
        "📚 Tải đề cương / tài liệu làm căn cứ sinh đề",
        type=[
            "pdf",
            "docx",
            "txt"
        ],
        key="de_kt_file_de_cuong"
    )

    # ============================================================
    # 3. CẤU HÌNH TỶ LỆ MỨC ĐỘ NHẬN THỨC
    # ============================================================
    with st.expander(
        "⚙️ Cấu hình Tỷ lệ & Số câu",
        expanded=True
    ):

        r1, r2, r3, r4 = st.columns(4)

        nb = r1.number_input(
            "Nhận biết (%)",
            min_value=0,
            max_value=100,
            value=40,
            key="de_kt_nb"
        )

        th = r2.number_input(
            "Thông hiểu (%)",
            min_value=0,
            max_value=100,
            value=30,
            key="de_kt_th"
        )

        vd = r3.number_input(
            "Vận dụng (%)",
            min_value=0,
            max_value=100,
            value=20,
            key="de_kt_vd"
        )

        vdc = r4.number_input(
            "Vận dụng cao (%)",
            min_value=0,
            max_value=100,
            value=10,
            key="de_kt_vdc"
        )

        tong_ty_le = nb + th + vd + vdc

        if tong_ty_le != 100:

            st.warning(
                f"⚠️ Tổng tỷ lệ mức độ hiện tại là "
                f"{tong_ty_le}%. Phải bằng 100%."
            )

        # --------------------------------------------------------
        # CẤU HÌNH CÁC DẠNG CÂU HỎI
        # --------------------------------------------------------
        st.markdown(
            "#### 📌 Cấu trúc các dạng câu hỏi"
        )

        cols = st.columns(8)

        n_nlc = cols[0].number_input(
            "NLC",
            min_value=0,
            value=10,
            key="de_kt_n_nlc"
        )

        d_nlc = cols[1].number_input(
            "Đ.NLC",
            min_value=0.0,
            value=0.25,
            step=0.25,
            key="de_kt_d_nlc"
        )

        n_ds = cols[2].number_input(
            "Đ/S",
            min_value=0,
            value=2,
            key="de_kt_n_ds"
        )

        d_ds = cols[3].number_input(
            "Đ.Đ/S",
            min_value=0.0,
            value=0.25,
            step=0.25,
            key="de_kt_d_ds"
        )

        n_dk = cols[4].number_input(
            "Điền K",
            min_value=0,
            value=2,
            key="de_kt_n_dk"
        )

        d_dk = cols[5].number_input(
            "Đ.DK",
            min_value=0.0,
            value=0.25,
            step=0.25,
            key="de_kt_d_dk"
        )

        n_ngan = cols[6].number_input(
            "TL Ngắn",
            min_value=0,
            value=2,
            key="de_kt_n_ngan"
        )

        d_ngan = cols[7].number_input(
            "Đ.TLN",
            min_value=0.0,
            value=0.50,
            step=0.25,
            key="de_kt_d_ngan"
        )

        # --------------------------------------------------------
        # TÍNH ĐIỂM CÁC DẠNG CÂU HỎI
        # --------------------------------------------------------
        total_diem_tn = (
            (n_nlc * d_nlc)
            + (n_ds * d_ds)
            + (n_dk * d_dk)
            + (n_ngan * d_ngan)
        )

        tl_cols = st.columns(4)

        num_tl = tl_cols[0].number_input(
            "Số câu Tự luận",
            min_value=1,
            max_value=10,
            value=2,
            key="de_kt_num_tl"
        )

        tl_points = []

        for i in range(num_tl):

            p = tl_cols[1].number_input(
                f"Câu {i + 1} (đ)",
                min_value=0.0,
                value=1.0,
                step=0.25,
                key=f"de_kt_tl_p_{i}"
            )

            tl_points.append(p)

        total_diem_tl = sum(tl_points)

        total_diem = (
            total_diem_tn
            + total_diem_tl
        )

        tl_cols[2].metric(
            "Tổng điểm TN",
            f"{total_diem_tn:.2f}"
        )

        tl_cols[3].metric(
            "Tổng điểm TL",
            f"{total_diem_tl:.2f}"
        )

        st.metric(
            "TỔNG ĐIỂM ĐỀ",
            f"{total_diem:.2f} / 10"
        )

    # ============================================================
    # 4. KIỂM TRA CẤU HÌNH TRƯỚC KHI SINH
    # ============================================================
    if st.button(
        "🚀 TẠO MA TRẬN & ĐỀ THI",
        type="primary",
        use_container_width=True,
        key="de_kt_btn_generate"
    ):

        # --------------------------------------------------------
        # KIỂM TRA TỶ LỆ
        # --------------------------------------------------------
        if tong_ty_le != 100:

            st.error(
                "❌ Tổng tỷ lệ Nhận biết + Thông hiểu "
                "+ Vận dụng + Vận dụng cao phải bằng 100%."
            )

            st.stop()

        # --------------------------------------------------------
        # KIỂM TRA ĐỀ CƯƠNG
        # --------------------------------------------------------
        if bam_sat and file_de is None:

            st.error(
                "❌ Thầy đã chọn 'Bám sát đề cương' "
                "nhưng chưa tải lên đề cương."
            )

            st.stop()

        # --------------------------------------------------------
        # KIỂM TRA TỔNG ĐIỂM
        # --------------------------------------------------------
        if abs(total_diem - 10.0) > 0.01:

            st.error(
                f"❌ Tổng điểm hiện tại là "
                f"{total_diem:.2f}/10. "
                "Vui lòng điều chỉnh lại cấu hình."
            )

            st.stop()

        # ========================================================
        # 5. ĐỌC VÀ XÂY DỰNG PHẠM VI KIẾN THỨC ĐƯỢC PHÉP
        # ========================================================
        with st.spinner(
            "📚 Đang đọc và xây dựng phạm vi kiến thức..."
        ):

            if bam_sat and file_de:

                raw_outline = extract_text_from_file(
                    file_de
                )

                outline_text = normalize_outline(
                    raw_outline
                )

                if not outline_text:

                    st.error(
                        "❌ Không đọc được nội dung đề cương."
                    )

                    st.stop()

                allowed_scope = f"""
============================================================
PHẠM VI KIẾN THỨC ĐƯỢC PHÉP SỬ DỤNG
============================================================

Tài liệu dưới đây là nguồn kiến thức duy nhất
được phép sử dụng để xây dựng đề kiểm tra.

------------------- BẮT ĐẦU ĐỀ CƯƠNG -------------------

{outline_text}

-------------------- KẾT THÚC ĐỀ CƯƠNG ------------------

QUY TẮC PHẠM VI:

1. Chỉ sử dụng kiến thức xuất hiện trong đề cương.

2. Không được tự ý bổ sung kiến thức ngoài đề cương.

3. Không được mở rộng sang bài học, chủ đề hoặc nội dung
   không xuất hiện trong đề cương.

4. Mỗi câu hỏi phải xác định được căn cứ kiến thức
   tương ứng trong đề cương.

5. Nếu một nội dung không thể xác định được căn cứ
   trong đề cương thì KHÔNG ĐƯỢC sử dụng nội dung đó
   để xây dựng câu hỏi.

6. Không được dùng kiến thức nền ngoài phạm vi nếu kiến thức
   đó làm thay đổi bản chất yêu cầu của câu hỏi.

7. Có thể sử dụng các thao tác suy luận, tính toán hoặc
   vận dụng trực tiếp từ kiến thức trong đề cương,
   nhưng không được đưa thêm kiến thức mới.

8. Đề kiểm tra phải phủ đúng các đơn vị kiến thức
   được nêu trong đề cương.

9. Không được tạo câu hỏi chỉ vì kiến thức đó phù hợp
   với môn học nhưng không có trong đề cương.

10. Sau khi tạo đề, phải tự kiểm tra từng câu hỏi:
    - Câu hỏi có căn cứ trong đề cương hay không?
    - Đáp án có dựa trên kiến thức trong đề cương hay không?
    - Có xuất hiện kiến thức ngoài phạm vi hay không?

============================================================
"""

            else:

                allowed_scope = """
============================================================
PHẠM VI KIẾN THỨC
============================================================

Không có đề cương được tải lên.

AI được phép sử dụng kiến thức phù hợp với:

- Chương trình giáo dục hiện hành.
- Môn học được chọn.
- Lớp học được chọn.
- Tên bài kiểm tra.
- Thời lượng kiểm tra.

============================================================
"""

        # ========================================================
        # 6. XÂY DỰNG CẤU HÌNH ĐỀ
        # ========================================================
        exam_configuration = f"""
============================================================
THÔNG TIN CẤU HÌNH ĐỀ KIỂM TRA
============================================================

Môn học: {mon_hoc}

Lớp: {lop}

Tên bài kiểm tra: {ten_de}

Hình thức: {hinh_thuc}

Thời gian: {thoi_gian}

Tổng điểm bắt buộc: 10 điểm.

------------------------------------------------------------
PHÂN BỐ MỨC ĐỘ NHẬN THỨC
------------------------------------------------------------

Nhận biết: {nb}%

Thông hiểu: {th}%

Vận dụng: {vd}%

Vận dụng cao: {vdc}%

------------------------------------------------------------
CẤU TRÚC CÂU HỎI
------------------------------------------------------------

Nhiều lựa chọn:
- Số câu: {n_nlc}
- Điểm mỗi câu: {d_nlc}

Đúng / Sai:
- Số câu: {n_ds}
- Điểm mỗi câu: {d_ds}

Điền khuyết:
- Số câu: {n_dk}
- Điểm mỗi câu: {d_dk}

Trả lời ngắn:
- Số câu: {n_ngan}
- Điểm mỗi câu: {d_ngan}

Tự luận:
- Số câu: {num_tl}
- Tổng điểm: {total_diem_tl:.2f}

Tổng điểm:
- Trắc nghiệm / câu hỏi ngắn: {total_diem_tn:.2f}
- Tự luận: {total_diem_tl:.2f}
- Tổng cộng: {total_diem:.2f}/10

============================================================
"""

        # ========================================================
        # 7. PROMPT KIỂM SOÁT CHẶT
        # ========================================================
        strict_instruction = """
============================================================
YÊU CẦU BẮT BUỘC KHI TẠO ĐỀ
============================================================

Thầy/Cô yêu cầu tạo một bộ hồ sơ kiểm tra gồm:

1. MA TRẬN ĐỀ KIỂM TRA.

2. BẢN ĐẶC TẢ ĐỀ KIỂM TRA.

3. ĐỀ KIỂM TRA.

4. ĐÁP ÁN VÀ HƯỚNG DẪN CHẤM.

------------------------------------------------------------
QUY TRÌNH BẮT BUỘC
------------------------------------------------------------

BƯỚC 1:
Phân tích phạm vi kiến thức được phép sử dụng.

BƯỚC 2:
Liệt kê các chủ đề / bài học / đơn vị kiến thức
có thể sử dụng.

BƯỚC 3:
Phân bố câu hỏi vào ma trận theo:
- Nội dung kiến thức.
- Mức độ nhận thức.
- Dạng câu hỏi.
- Số lượng câu.
- Số điểm.

BƯỚC 4:
Tạo bản đặc tả tương ứng với ma trận.

BƯỚC 5:
Tạo câu hỏi theo đúng ma trận và đặc tả.

BƯỚC 6:
Tạo đáp án và hướng dẫn chấm.

BƯỚC 7:
TỰ KIỂM TRA TOÀN BỘ ĐỀ trước khi trả kết quả.

------------------------------------------------------------
QUY TẮC KIỂM TRA NỘI DUNG
------------------------------------------------------------

Mỗi câu hỏi phải:

- Có căn cứ kiến thức trong phạm vi cho phép.
- Phù hợp với môn học và lớp học.
- Đúng mức độ nhận thức được phân công.
- Đúng dạng câu hỏi.
- Đúng số điểm.
- Không trùng lặp không cần thiết.
- Không chứa kiến thức ngoài phạm vi.

Nếu phát hiện câu hỏi có kiến thức ngoài phạm vi:

1. Loại bỏ câu hỏi đó.
2. Thay bằng câu hỏi khác có căn cứ trong phạm vi.
3. Kiểm tra lại toàn bộ đề.

------------------------------------------------------------
QUY TẮC SỐ LƯỢNG VÀ ĐIỂM
------------------------------------------------------------

Bắt buộc:

- Đúng số lượng câu hỏi.
- Đúng số lượng từng dạng câu hỏi.
- Đúng số điểm từng câu.
- Tổng điểm chính xác là 10 điểm.
- Ma trận, đặc tả, đề và đáp án phải khớp nhau.

Không được tự ý thay đổi:

- Số câu.
- Số điểm.
- Tỷ lệ mức độ.
- Dạng câu hỏi.
- Thời gian kiểm tra.

------------------------------------------------------------
YÊU CẦU ĐẦU RA
------------------------------------------------------------

Trình bày kết quả theo đúng thứ tự:

# PHẦN I. PHẠM VI KIẾN THỨC ĐƯỢC SỬ DỤNG

Liệt kê các chủ đề và đơn vị kiến thức
được sử dụng để xây dựng đề.

# PHẦN II. MA TRẬN ĐỀ KIỂM TRA

Trình bày rõ:
- Nội dung / chủ đề.
- Đơn vị kiến thức.
- Nhận biết.
- Thông hiểu.
- Vận dụng.
- Vận dụng cao.
- Số câu.
- Số điểm.

# PHẦN III. BẢN ĐẶC TẢ

Mỗi yêu cầu cần đạt phải thể hiện:
- Nội dung kiến thức.
- Mức độ nhận thức.
- Yêu cầu cần đạt.
- Dạng câu hỏi.
- Số câu.
- Số điểm.

# PHẦN IV. ĐỀ KIỂM TRA

Trình bày đầy đủ:
- Tên bài kiểm tra.
- Môn học.
- Lớp.
- Thời gian.
- Hướng dẫn làm bài.
- Toàn bộ câu hỏi.
- Các phương án lựa chọn nếu có.

# PHẦN V. ĐÁP ÁN VÀ HƯỚNG DẪN CHẤM

Trình bày:
- Đáp án từng câu.
- Điểm từng câu.
- Hướng dẫn chấm tự luận.
- Tổng điểm.

# PHẦN VI. BẢNG TỰ KIỂM TRA

Tạo bảng kiểm tra:

| STT | Nội dung kiểm tra | Kết quả |
|-----|-------------------|---------|

Bắt buộc kiểm tra:

1. Đề có đúng phạm vi kiến thức không?
2. Có câu hỏi ngoài đề cương không?
3. Đúng số lượng câu không?
4. Đúng tổng điểm 10 không?
5. Ma trận có khớp đặc tả không?
6. Đặc tả có khớp đề không?
7. Đề có khớp đáp án không?
8. Đúng tỷ lệ mức độ không?

Chỉ được kết luận "ĐẠT" khi tất cả các yêu cầu
trên đều được đáp ứng.

============================================================
"""

        # ========================================================
        # 8. TẠO PROMPT CUỐI CÙNG
        # ========================================================
        final_prompt = f"""
{strict_instruction}

{exam_configuration}

{allowed_scope}

============================================================
NHIỆM VỤ
============================================================

Hãy thực hiện đầy đủ quy trình kiểm tra nội bộ trước.

Sau đó tạo:

1. Phạm vi kiến thức được sử dụng.
2. Ma trận.
3. Bản đặc tả.
4. Đề kiểm tra.
5. Đáp án và hướng dẫn chấm.
6. Bảng tự kiểm tra.

ƯU TIÊN TUYỆT ĐỐI:

ĐỘ CHÍNH XÁC VÀ KHẢ NĂNG BÁM SÁT
PHẠM VI KIẾN THỨC ĐƯỢC PHÉP

Không được bỏ qua bước tự kiểm tra.
"""

        # ========================================================
        # 9. GỌI AI ĐÚNG 1 LẦN
        # ========================================================
        with st.spinner(
            "🤖 AI đang xây dựng ma trận, đặc tả và đề kiểm tra..."
        ):

            try:

                result = ai_engine.generate_text(
                    final_prompt
                )

                # ------------------------------------------------
                # KIỂM TRA KẾT QUẢ
                # ------------------------------------------------
                if not result or not result.strip():

                    st.error(
                        "❌ AI trả về kết quả rỗng."
                    )

                    st.stop()

                # ------------------------------------------------
                # LƯU KẾT QUẢ
                # ------------------------------------------------
                st.session_state[
                    "de_kt_content"
                ] = result

                # Lưu thêm thông tin cấu hình
                st.session_state[
                    "de_kt_config"
                ] = {

                    "mon_hoc": mon_hoc,

                    "lop": lop,

                    "ten_de": ten_de,

                    "hinh_thuc": hinh_thuc,

                    "thoi_gian": thoi_gian,

                    "tong_diem": total_diem,

                    "bam_sat": bam_sat,

                    "file_name": (
                        file_de.name
                        if file_de
                        else None
                    )
                }

                st.success(
                    "✅ Đã tạo xong ma trận, đặc tả và đề kiểm tra."
                )

                st.rerun()

            except Exception as e:

                st.error(
                    f"❌ Lỗi khi sinh đề: {e}"
                )

    # ============================================================
    # 10. HIỂN THỊ KẾT QUẢ
    # ============================================================
    if "de_kt_content" in st.session_state:

        st.divider()

        st.markdown(
            "## 📄 KẾT QUẢ ĐỀ KIỂM TRA"
        )

        # --------------------------------------------------------
        # NÚT XÓA
        # --------------------------------------------------------
        if st.button(
            "🗑️ XÓA ĐỀ",
            key="de_kt_delete"
        ):

            del st.session_state[
                "de_kt_content"
            ]

            if "de_kt_config" in st.session_state:

                del st.session_state[
                    "de_kt_config"
                ]

            st.rerun()

        # --------------------------------------------------------
        # HIỂN THỊ NỘI DUNG
        # --------------------------------------------------------
        st.markdown(
            st.session_state[
                "de_kt_content"
            ]
        )

        # --------------------------------------------------------
        # XUẤT WORD
        # --------------------------------------------------------
        try:

            root_path = str(
                Path(__file__).resolve().parents[2]
            )

            if root_path not in sys.path:

                sys.path.insert(
                    0,
                    root_path
                )

            from export.export_word import (
                WordExportEngine
            )

            config = st.session_state.get(
                "de_kt_config",
                {}
            )

            word_bytes = (
                WordExportEngine.export_to_word(
                    {
                        "ai_generated_content":
                            st.session_state[
                                "de_kt_content"
                            ],

                        "is_de_kt": True,

                        "title":
                            config.get(
                                "ten_de",
                                "Đề kiểm tra"
                            )
                    }
                )
            )

            st.download_button(
                "📥 TẢI FILE WORD",
                data=word_bytes,
                file_name="De_Thi.docx",
                use_container_width=True,
                key="de_kt_download_word"
            )

        except Exception as e:

            st.warning(
                f"⚠️ Lỗi xuất Word: {e}"
            )
