# -*- coding: utf-8 -*-
r"""
============================================================
MODULE: modules/quan_ly_to/ke_hoach.py
Nhiệm vụ: Quản lý và Xây dựng Chuyên đề Giáo dục.
Chức năng: Trưng bày 6 mảng chuyên đề trọng tâm, AI hỗ trợ 
khởi tạo khung kế hoạch chi tiết, hỗ trợ xuất file Word (.docx) 
chuẩn hành chính và file .txt.
============================================================
"""

import streamlit as st
import io

# Kiểm tra thư viện hỗ trợ xuất file Word (.docx)
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

def render_ke_hoach():
    st.markdown("### 🎯 Quản lý và Xây dựng Chuyên đề Giáo dục")
    st.caption("Không gian chia sẻ, thiết kế và lưu trữ các chuyên đề chuyên môn sâu rộng, bám sát định hướng đổi mới giáo dục.")

    # 1. TRƯNG BÀY 6 MẢNG CHUYÊN ĐỀ BẰNG TABS
    st.markdown("#### 📚 Hệ thống Chuyên đề Trọng tâm")
    
    tabs_cd = st.tabs([
        "🛠️ PP & Kỹ thuật", 
        "📖 NC Bài học", 
        "📝 Kiểm tra ĐG", 
        "🎯 Phân hóa HS", 
        "💻 CNTT & Số hóa", 
        "🤝 Nghiệp vụ QL"
    ])

    with tabs_cd[0]:
        st.info("""
        **1. Phương pháp và Kỹ thuật dạy học**
        * **Đổi mới phương pháp:** Áp dụng các mô hình như Bàn tay nặn bột, Dạy học theo dự án, Giáo dục STEM/STEAM, phát huy tính tích cực, tự chủ của học sinh.
        * **Dạy học tích hợp và liên môn:** Xây dựng kế hoạch giảng dạy liên kết các kiến thức bộ môn gần gũi với thực tiễn.
        * **Phát triển năng lực:** Tập trung thiết kế bài giảng để phát triển phẩm chất, năng lực cốt lõi theo Chương trình GDPT.
        """)

    with tabs_cd[1]:
        st.success("""
        **2. Nghiên cứu bài học (Lesson Study)**
        * **Xây dựng kế hoạch bài dạy:** Giáo viên cùng nhau soạn, góp ý và hoàn thiện giáo án một chủ đề hoặc bài học khó.
        * **Dự giờ và phân tích bài học:** Tổ chức dạy thực nghiệm, quay phim, quan sát phản ứng, mức độ tiếp thu và khó khăn của học sinh, từ đó rút kinh nghiệm chung.
        """)

    with tabs_cd[2]:
        st.warning("""
        **3. Kiểm tra, đánh giá học sinh**
        * **Đổi mới đánh giá năng lực:** Xây dựng ngân hàng câu hỏi, đề kiểm tra thường xuyên và định kỳ theo hướng phát triển năng lực (tăng cường câu hỏi vận dụng, thực tiễn).
        * **Đánh giá quá trình:** Hướng dẫn cách chấm điểm, nhận xét và hỗ trợ học sinh tiến bộ thông qua các hoạt động trên lớp.
        """)

    with tabs_cd[3]:
        st.error("""
        **4. Hỗ trợ, phân hóa đối tượng**
        * **Phụ đạo học sinh chưa đạt:** Đưa ra các giải pháp cụ thể giúp học sinh lấy lại căn bản, cải thiện kết quả học tập.
        * **Bồi dưỡng học sinh giỏi:** Xây dựng chuyên đề chuyên sâu, các dạng bài tập nâng cao để ôn thi các cấp.
        """)

    with tabs_cd[4]:
        st.info("""
        **5. Ứng dụng CNTT và Chuyển đổi số**
        * **Ứng dụng AI trong giáo dục:** Sử dụng các công cụ Trí tuệ nhân tạo để hỗ trợ soạn giáo án, thiết kế bài tập và hỗ trợ quản lý lớp.
        * **Sử dụng phần mềm dạy học:** Hướng dẫn sử dụng phần mềm tạo bài giảng E-learning, trò chơi tương tác (Quizizz, Kahoot, Padlet), hoặc thí nghiệm ảo.
        """)

    with tabs_cd[5]:
        st.success("""
        **6. Nghiệp vụ và quản lý lớp học**
        * **Công tác chủ nhiệm lớp:** Giải quyết các tình huống sư phạm, giáo dục học sinh cá biệt, hoặc xây dựng tập thể lớp đoàn kết.
        * **Tư vấn tâm lý học đường:** Các phương pháp hỗ trợ học sinh có vấn đề về tâm lý, áp lực học tập.
        """)

    st.markdown("---")

    # 2. TRỢ LÝ AI: LẬP KẾ HOẠCH TRIỂN KHAI CHUYÊN ĐỀ
    st.markdown("#### 🤖 Trợ lý AI: Khởi tạo Khung Kế hoạch Chuyên đề")
    st.caption("Chọn một nhóm chuyên đề và cung cấp ý tưởng, AI sẽ lập ngay một khung kế hoạch chi tiết để tổ chuyên môn xét duyệt.")
    
    # Form yêu cầu AI
    with st.container(border=True):
        col_form1, col_form2 = st.columns([1, 2])
        with col_form1:
            nhom_cd = st.selectbox("Thuộc nhóm chuyên đề:", [
                "Phương pháp và Kỹ thuật dạy học",
                "Nghiên cứu bài học",
                "Kiểm tra, đánh giá học sinh",
                "Hỗ trợ, phân hóa đối tượng",
                "Ứng dụng CNTT và Chuyển đổi số",
                "Nghiệp vụ và quản lý lớp học"
            ])
            nguoi_bao_cao = st.text_input("Người báo cáo/Phụ trách:", placeholder="VD: Cô Huyền Trang")
        with col_form2:
            ten_chuyen_de = st.text_area("Tên chuyên đề & Ý tưởng trọng tâm:", height=110, placeholder="VD: Xây dựng chủ đề STEM 'Thuyền tự hành' để phát triển năng lực giải quyết vấn đề cho học sinh khối 8.")

        if st.button("🚀 Sinh Kế hoạch Chuyên đề", type="primary", use_container_width=True):
            if ten_chuyen_de.strip():
                with st.spinner("AI đang thiết kế khung kế hoạch chuyên đề..."):
                    prompt = f"""
                    HÃY ĐÓNG VAI TỔ TRƯỞNG CHUYÊN MÔN. VIẾT MỘT 'KẾ HOẠCH TRIỂN KHAI CHUYÊN ĐỀ' THẬT CHI TIẾT, KHOA HỌC, BÁM SÁT CÁC TIÊU CHÍ CỦA TRƯỜNG THCS.
                    
                    THÔNG TIN CHUYÊN ĐỀ:
                    - Nhóm chuyên đề: {nhom_cd}
                    - Tên chuyên đề/Ý tưởng: {ten_chuyen_de}
                    - Người phụ trách: {nguoi_bao_cao}
                    
                    CẤU TRÚC KẾ HOẠCH BẮT BUỘC:
                    I. Mục đích, yêu cầu (Phát triển năng lực gì? Giải quyết vấn đề gì?)
                    II. Đối tượng và thời gian thực hiện
                    III. Nội dung chi tiết của chuyên đề
                    IV. Tổ chức thực hiện (Phân công chuẩn bị, tiến trình báo cáo/dự giờ)
                    
                    Trình bày chuyên nghiệp, sử dụng bullet point rõ ràng, văn phong hành chính trang trọng.
                    """
                    
                    khung_ke_hoach = None
                    try:
                        # 1. Thử gọi qua ai_engine
                        ai_engine = st.session_state.get("ai_engine", None)
                        if ai_engine and hasattr(ai_engine, "generate_text"):
                            try:
                                khung_ke_hoach = ai_engine.generate_text(prompt)
                            except Exception:
                                pass 

                        # 2. Dự phòng gọi trực tiếp OpenAI bằng khóa sk-
                        if not khung_ke_hoach:
                            api_key = None
                            for key, val in st.session_state.items():
                                if isinstance(val, str) and val.startswith("sk-"):
                                    api_key = val
                                    break
                            
                            if not api_key:
                                for k in ["user_api_key", "api_key", "openai_api_key", "sk_key"]:
                                    if st.session_state.get(k) and str(st.session_state.get(k)).startswith("sk-"):
                                        api_key = st.session_state.get(k)
                                        break
                            
                            if not api_key and "OPENAI_API_KEY" in st.secrets:
                                api_key = st.secrets["OPENAI_API_KEY"]

                            if api_key:
                                from openai import OpenAI
                                client = OpenAI(api_key=str(api_key).strip())
                                response = client.chat.completions.create(
                                    model="gpt-4o-mini",
                                    messages=[{"role": "user", "content": prompt}]
                                ी]
                                khung_ke_hoach = response.choices[0].message.content
                            else:
                                st.error("❌ Không tìm thấy khóa API `sk-` hợp lệ. Thầy vui lòng kiểm tra lại ô nhập API Key ở menu bên trái.")

                        if khung_ke_hoach:
                            # Làm sạch dấu ** ngay lập tức
                            st.session_state.ket_qua_chuyen_de = khung_ke_hoach.replace("**", "")
                            st.rerun()
                    except Exception as e:
                        st.error(f"Lỗi AI: {e}")
            else:
                st.warning("Thầy vui lòng nhập Tên chuyên đề hoặc Ý tưởng để AI có cơ sở lập kế hoạch nhé!")

    # 3. HIỂN THỊ VÀ TẢI KẾ HOẠCH (HỖ TRỢ CẢ WORD VÀ TEXT)
    if st.session_state.get("ket_qua_chuyen_de"):
        st.markdown("---")
        st.markdown("#### 📄 Khung Kế hoạch đề xuất")
        st.text_area("Chỉnh sửa Kế hoạch (nếu cần):", value=st.session_state.get("ket_qua_chuyen_de", "").replace("**", ""), height=400, key="edit_cd")
        
        dl_col1, dl_col2 = st.columns(2)
        
        with dl_col1:
            if HAS_DOCX:
                try:
                    doc = Document()
                    
                    # Thiết lập lề trang giấy chuẩn hành chính
                    sections = doc.sections
                    for section in sections:
                        section.top_margin = Inches(1.0)
                        section.bottom_margin = Inches(1.0)
                        section.left_margin = Inches(1.2)
                        section.right_margin = Inches(0.8)
                    
                    # Tiêu đề chính
                    p_title = doc.add_paragraph()
                    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_title = p_title.add_run("KẾ HOẠCH TRIỂN KHAI CHUYÊN ĐỀ")
                    run_title.bold = True
                    run_title.font.size = Pt(13)
                    p_title.paragraph_format.space_after = Pt(16)

                    # Xử lý nội dung văn bản (loại bỏ dấu **)
                    raw_text = st.session_state.get("ket_qua_chuyen_de", "").replace("**", "")
                    for line in raw_text.split('\n'):
                        clean_line = line.strip()
                        if not clean_line:
                            continue
                        
                        p = doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(4)
                        p.paragraph_format.line_spacing = 1.15
                        
                        # In đậm các đề mục lớn (I., II., III., IV., V., 1., 2., 3...)
                        if clean_line.startswith(('I.', 'II.', 'III.', 'IV.', 'V.', 'VI.', '1.', '2.', '3.', '4.', '5.')) and len(clean_line) < 90 and not clean_line.startswith('-'):
                            run = p.add_run(clean_line)
                            run.bold = True
                            if clean_line.startswith(('I.', 'II.', 'III.', 'IV.')):
                                p.paragraph_format.space_before = Pt(8)
                                run.font.size = Pt(11)
                            else:
                                run.font.size = Pt(11)
                        else:
                            p.add_run(clean_line)

                    docx_buffer = io.BytesIO()
                    doc.save(docx_buffer)
                    docx_buffer.seek(0)
                    
                    st.download_button(
                        label="⬇️ Tải Kế hoạch Word (.docx)",
                        data=docx_buffer,
                        file_name="Ke_Hoach_Chuyen_De.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary",
                        use_container_width=True
                    )
                except Exception as ex:
                    st.error(f"Lỗi tạo file Word: {ex}")
            else:
                st.warning("⚠️ Cần cài đặt `python-docx` để tải file Word. Chạy lệnh: `pip install python-docx`")

        with dl_col2:
            clean_txt_data = st.session_state.get("ket_qua_chuyen_de", "").replace("**", "")
            st.download_button(
                label="⬇️ Tải Kế hoạch Text (.txt)",
                data=clean_txt_data,
                file_name="Ke_Hoach_Chuyen_De.txt",
                mime="text/plain",
                type="secondary",
                use_container_width=True
            )
