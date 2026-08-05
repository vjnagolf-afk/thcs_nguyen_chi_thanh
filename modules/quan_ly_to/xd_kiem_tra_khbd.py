# -*- coding: utf-8 -*-
r"""
============================================================
MODULE: modules/quan_ly_to/kiem_tra_khbd.py
Nhiệm vụ: Kiểm tra, phê duyệt Kế hoạch bài dạy (Giáo án).
Chức năng: Đọc file PDF/Word, quét lỗi đa chiều theo CV 5512, 
kiểm tra Năng lực số theo Thông tư 02/2025/BGDĐT, phát hiện mâu thuẫn
và hỗ trợ chat thẩm định toàn diện.
============================================================
"""

import streamlit as st
from pypdf import PdfReader
import re
import docx

def render_kiem_tra_khbd(ai_engine=None):
    st.markdown("### 🔎 Kiểm tra, phê duyệt Kế hoạch bài dạy (Giáo án)")
    st.caption("AI đọc trực tiếp file KHBD tải lên, rà soát lỗi đa chiều, kiểm tra Năng lực số (TT 02/2025/BGDĐT) và phát hiện mâu thuẫn logic.")

    # --- KHỞI TẠO BỘ NHỚ TẠM ---
    if "chat_history_khbd" not in st.session_state:
        st.session_state.chat_history_khbd = [{"role": "assistant", "content": "Chào thầy/cô! Hãy nạp file KHBD ở cột bên trái và nhấn **'Quét & Phân tích'** nhé!"}]
    
    if "noidung_khbd" not in st.session_state:
        st.session_state.noidung_khbd = ""
        
    if "ai_analysis_report" not in st.session_state:
        st.session_state.ai_analysis_report = ""

    # --- HỖ TRỢ GỌI AI LINH HOẠT CHO MỌI LOẠI KHÓA ---
    def call_ai(prompt_text):
        # 1. Thử qua ai_engine truyền vào hoặc trong session
        engine = ai_engine or st.session_state.get("ai_engine", None)
        if engine and hasattr(engine, "generate_text"):
            try:
                return engine.generate_text(prompt_text)
            except Exception:
                pass

        # 2. Dự phòng gọi trực tiếp OpenAI bằng khóa sk-
        api_key = None
        for key, val in st.session_state.items():
            if isinstance(val, str) and val.startswith("sk-"):
                api_key = val
                break
        
        if not api_key:
            for k in ["user_api_key", "api_key", "openai_api_key", "sk_key"]:
                if st.session_state.get(k) and str(st.session_state.get(k)).startswith("sk-"):
                    api_key = st.session_state.get(k)
                    break
        
        if not api_key and "OPENAI_API_KEY" in st.secrets:
            api_key = st.secrets["OPENAI_API_KEY"]

        if api_key:
            from openai import OpenAI
            client = OpenAI(api_key=str(api_key).strip())
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt_text}]
            )
            return response.choices[0].message.content
        
        return "❌ Không tìm thấy API Key hoặc AI Engine hợp lệ. Vui lòng kiểm tra lại cấu hình."

    # --- BỐ CỤC GIAO DIỆN CHÍNH ---
    col_left, col_right = st.columns([1, 2.2])

    # ==========================================
    # CỘT TRÁI: CONTROL PANEL 
    # ==========================================
    with col_left:
        with st.container(border=True):
            st.markdown("#### ⚙️ Cấu hình Kiểm tra")
            mon_hoc = st.selectbox("Môn học:", ["Khoa học Tự nhiên", "Toán học", "Ngữ Văn", "Tiếng Anh", "Tin học", "Khác"])
            khoi_lop = st.selectbox("Khối lớp:", ["Lớp 6", "Lớp 7", "Lớp 8", "Lớp 9"])
            
            st.markdown("**Tải file giáo án lên đây:**")
            file_khbd = st.file_uploader("Hỗ trợ định dạng PDF và Word (.docx)", type=["pdf", "docx"], label_visibility="collapsed")
            
            if st.button("🚀 Quét & Phân tích Chuyên Sâu", type="primary", use_container_width=True):
                if file_khbd:
                    valid_mime = ["application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]
                    if file_khbd.type not in valid_mime and not file_khbd.name.endswith(('.pdf', '.docx')):
                        st.error("⚠️ Định dạng file không hợp lệ! Hệ thống chỉ chấp nhận file PDF hoặc DOCX chuẩn.")
                    else:
                        with st.spinner("AI đang đọc file, quét thuật ngữ, rà soát Năng lực số và mâu thuẫn logic..."):
                            try:
                                st.session_state.chat_history_khbd = []
                                st.session_state.ai_analysis_report = ""
                                st.session_state.noidung_khbd = ""
                                
                                extracted_text = ""
                                
                                # XỬ LÝ NẾU LÀ FILE WORD (.docx)
                                if file_khbd.name.endswith(".docx"):
                                    doc = docx.Document(file_khbd)
                                    for para in doc.paragraphs:
                                        if para.text.strip():
                                            extracted_text += para.text + "\n"
                                    for table in doc.tables:
                                        for row in table.rows:
                                            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                                            if row_data:
                                                extracted_text += " | ".join(row_data) + "\n"
                                
                                # XỬ LÝ NẾU LÀ FILE PDF (.pdf)
                                elif file_khbd.name.endswith(".pdf"):
                                    reader = PdfReader(file_khbd)
                                    for page in reader.pages:
                                        text = page.extract_text()
                                        if text:
                                            extracted_text += text + "\n"
                                
                                noidung = extracted_text[:60000]
                                noidung = re.sub(r'\s+', ' ', noidung).strip()
                                st.session_state.noidung_khbd = noidung
                                
                                # PROMPT ĐƯỢC NÂNG CẤP CHUYÊN SÂU ĐỂ KIỂM TRA NĂNG LỰC SỐ & MÂU THUẪN
                                prompt_scan = f"""
                                Bạn là Tổ trưởng chuyên môn dày dặn kinh nghiệm, nắm rất vững Công văn 5512 và Thông tư 02/2025/BGDĐT về chuẩn Năng lực số cho học sinh phổ thông. Hãy phân tích KHBD dưới đây một cách cực kỳ khắt khe và lập Báo cáo chi tiết theo 3 phần rõ ràng:

                                PHẦN 1: ĐÁNH GIÁ TÍCH HỢP NĂNG LỰC SỐ (THÔNG TƯ 02/2025/BGDĐT)
                                - Giáo án có tích hợp thành phần Năng lực số cho học sinh hay không? (Nêu rõ có hoặc không).
                                - Liệt kê chi tiết các công cụ số, phần mềm, học liệu số hoặc phương pháp khai thác thông tin số được (hoặc chưa được) đưa vào giáo án. Nếu thiếu mục này, BẮT BUỘC phải ghi nhận là điểm trừ lớn trong báo cáo.

                                PHẦN 2: PHÁT HIỆN MÂU THUẪN LOGIC TRONG KẾ HOẠCH
                                Kiểm tra kỹ lưỡng sự thống nhất xuyên suốt các phần:
                                - Mâu thuẫn Mục tiêu vs Thiết bị/Học liệu: Phần I (Mục tiêu) có yêu cầu học sinh hình thành/phát triển năng lực số hoặc sử dụng công cụ số, nhưng phần Thiết bị/Học liệu lại hoàn toàn bỏ trống máy tính, máy chiếu, phần mềm, thí nghiệm ảo... => Coi là Mâu thuẫn.
                                - Mâu thuẫn Mục tiêu vs Tổ chức hoạt động: Mục tiêu yêu cầu làm việc nhóm/trực tuyến hoặc sử dụng sản phẩm số, nhưng sang phần Tổ chức hoạt động dạy học (Hoạt động 1, 2, 3, 4) giáo viên lại chỉ cho học sinh làm việc cá nhân thủ công, không hề triển khai... => Coi là Mâu thuẫn.
                                - Các mâu thuẫn hình thức khác (VD: Mục tiêu yêu cầu vẽ sơ đồ nhưng sản phẩm học sinh không có...).
                                *BẮT BUỘC trích dẫn nguyên văn đoạn sai và chỉ rõ mâu thuẫn ở phần nào.*

                                PHẦN 3: RÀ SOÁT LỖI NGÔN NGỮ, THUẬT NGỮ & KIẾN THỨC
                                - 🔤 Lỗi chính tả & Câu từ thực sự (Bỏ qua lỗi dính khoảng trắng tự động, chỉ báo lỗi sai âm vần thực sự).
                                - 📖 Lỗi thuật ngữ chuyên môn khoa học (đặc biệt là môn {mon_hoc}).
                                - 🧠 Lỗi kiến thức khoa học (nếu có).
                                *Nếu không phát hiện lỗi ở mục nào, ghi rõ "Không phát hiện lỗi".*

                                NỘI DUNG KHBD THỰC TẾ CẦN PHÂN TÍCH:
                                '''{st.session_state.noidung_khbd}'''
                                """
                                
                                report = call_ai(prompt_scan)
                                st.session_state.ai_analysis_report = report.replace("**", "")
                                
                                st.session_state.chat_history_khbd.append({"role": "assistant", "content": f"✅ Tôi đã thẩm định xong giáo án **{mon_hoc} {khoi_lop}**. Mời thầy/cô xem **Báo cáo phân tích chuyên sâu** ở tab bên cạnh."})
                                st.rerun()
                            except Exception as e:
                                st.error(f"Lỗi hệ thống hoặc lỗi kết nối AI: {e}")
                else:
                    st.warning("⚠️ Vui lòng tải file KHBD (.pdf hoặc .docx) lên trước.")

    # ==========================================
    # CỘT PHẢI: KHUNG CHAT & DASHBOARD
    # ==========================================
    with col_right:
        tab_chat, tab_report = st.tabs(["💬 Trò chuyện với AI", "📊 Báo cáo Thẩm định Chuyên Sâu"])

        with tab_chat:
            chat_container = st.container(height=450)

            st.markdown("⚡ **Gợi ý kiểm tra nhanh:**")
            q1, q2 = st.columns(2)
            clicked_quick_prompt = None
            
            with q1:
                if st.button("💻 Kiểm tra chi tiết Năng lực số", use_container_width=True): 
                    clicked_quick_prompt = "Hãy rà soát kỹ xem giáo án này đã đáp ứng đầy đủ tiêu chí Năng lực số theo Thông tư 02/2025 chưa? Nếu chưa, hãy chỉ rõ cần bổ sung vào hoạt động nào."
                if st.button("🔄 Gợi ý khắc phục mâu thuẫn", use_container_width=True): 
                    clicked_quick_prompt = "Hãy chỉ ra các điểm mâu thuẫn giữa Mục tiêu, Thiết bị học liệu và Tổ chức hoạt động trong giáo án này, sau đó hướng dẫn tôi cách sửa lại cho logic."
            with q2:
                if st.button("🎯 Kiểm tra Mục tiêu (Bloom)", use_container_width=True): 
                    clicked_quick_prompt = "Hãy rà soát xem phần Mục tiêu đã dùng đúng động từ Bloom chưa? Nếu sai, hãy trích dẫn câu sai và gợi ý cách sửa."
                if st.button("📝 Chỉnh lại văn phong sư phạm", use_container_width=True): 
                    clicked_quick_prompt = "Hãy tìm các đoạn diễn đạt lủng củng hoặc sai thuật ngữ trong giáo án và viết lại chúng sao cho chuẩn văn phong sư phạm."

            st.markdown("---")
            btn_tham_dinh = st.button("🧐 BÁO CÁO THẨM ĐỊNH TOÀN DIỆN (5 TIÊU CHÍ TỔ TRƯỞNG)", use_container_width=True, type="primary")

            user_input = st.chat_input("Hỏi AI về nội dung giáo án đang tải lên...")
            
            prompt = user_input or clicked_quick_prompt
            
            if btn_tham_dinh:
                if not st.session_state.noidung_khbd:
                    st.warning("⚠️ Thầy cần tải file KHBD lên và bấm 'Quét & Phân tích' trước khi thẩm định!")
                else:
                    st.session_state.chat_history_khbd.append({"role": "user", "content": "Hãy thẩm định KHBD này theo 5 tiêu chí của Tổ trưởng chuyên môn."})
            
            elif prompt:
                if not st.session_state.noidung_khbd:
                    st.warning("⚠️ Hãy tải file và Quét KHBD trước khi hỏi AI nhé!")
                else:
                    st.session_state.chat_history_khbd.append({"role": "user", "content": prompt})

            with chat_container:
                for msg in st.session_state.chat_history_khbd:
                    with st.chat_message(msg["role"]): 
                        st.markdown(msg["content"])
                
                if btn_tham_dinh and st.session_state.noidung_khbd:
                    with st.chat_message("assistant"):
                        with st.spinner("🕵️‍♂️ Tổ trưởng AI đang phân tích sâu 5 tiêu chí..."):
                            try:
                                prompt_template = ""
                                try:
                                    with open("prompts/prompt_tham_dinh_khbd.txt", "r", encoding="utf-8") as f:
                                        prompt_template = f.read()
                                except FileNotFoundError:
                                    prompt_template = "Hãy thẩm định chi tiết Kế hoạch bài dạy sau theo 5 tiêu chí giáo dục phổ thông:\n[NOI_DUNG_KHBD_ODAY]"

                                prompt_hoan_thien = prompt_template.replace("[NOI_DUNG_KHBD_ODAY]", st.session_state.noidung_khbd)
                                
                                ket_qua_tham_dinh = call_ai(prompt_hoan_thien).replace("**", "")
                                
                                st.markdown(ket_qua_tham_dinh)
                                st.session_state.chat_history_khbd.append({"role": "assistant", "content": ket_qua_tham_dinh})
                            except Exception as e:
                                st.error(f"Lỗi hệ thống: {e}")

                elif prompt and st.session_state.noidung_khbd:
                    with st.chat_message("assistant"):
                        with st.spinner("AI đang đọc lại giáo án và trả lời..."):
                            try:
                                chat_prompt = f"Dựa vào nội dung KHBD sau:\n\n{st.session_state.noidung_khbd}\n\nThực hiện yêu cầu sau: {prompt}"
                                ai_response = call_ai(chat_prompt).replace("**", "")
                                
                                st.markdown(ai_response)
                                st.session_state.chat_history_khbd.append({"role": "assistant", "content": ai_response})
                            except Exception as e:
                                st.error(f"Lỗi gọi AI: {e}")

        with tab_report:
            if st.session_state.ai_analysis_report:
                st.success("Báo cáo thẩm định chuyên sâu dưới đây được AI tổng hợp TRỰC TIẾP từ dữ liệu file vừa tải lên.")
                st.markdown(st.session_state.ai_analysis_report)
            else:
                st.info("💡 Bảng điều khiển đang trống. Vui lòng tải file giáo án (.docx hoặc .pdf) lên và nhấn 'Quét & Phân tích Chuyên Sâu'.")
