# -*- coding: utf-8 -*-
r"""
============================================================
MODULE: modules/quan_ly_to/bien_ban.py
Nhiệm vụ: Trợ lý Thư ký - Xây dựng Biên bản Sinh hoạt.
Chức năng: AI tự động soạn thảo biên bản họp, loại bỏ dấu ** 
và xuất file Word (.docx) căn chỉnh chuẩn thể thức hành chính.
============================================================
"""

import streamlit as st
from pypdf import PdfReader
import io

# Kiểm tra thư viện hỗ trợ xuất file Word (.docx)
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

def render_bien_ban(ai_engine=None):
    st.markdown("### 📝 Trợ lý Thư ký: Xây dựng Biên bản Sinh hoạt")
    st.caption("AI tự động soạn thảo biên bản họp bám sát cấu trúc dự thảo kế hoạch, hỗ trợ xuất file Word chuẩn hành chính không chứa ký tự thừa.")

    # Lấy danh sách GV từ bộ nhớ (nếu có) để làm menu chọn Chủ tọa/Thư ký
    ds_gv = st.session_state.get("danh_sach_gv", ["Chưa có dữ liệu (Hãy qua thẻ Danh sách)"])
    
    if "ket_qua_bien_ban" not in st.session_state:
        st.session_state.ket_qua_bien_ban = None

    # 1. KHU VỰC THÔNG TIN CUỘC HỌP
    with st.expander("📌 Bước 1: Thông tin cơ bản", expanded=True):
        h1_c1, h1_c2, h1_c3 = st.columns([2, 1.5, 1.5])
        with h1_c1:
            loai_cuoc_hop = st.selectbox("📌 Loại hình sinh hoạt:", [
                "Sinh hoạt chuyên môn định kỳ",
                "Nghiên cứu bài học (Bước 2, 3)",
                "Xây dựng chuyên đề / STEM",
                "Phân tích kết quả kiểm tra",
                "Thống nhất ma trận, đặc tả đề",
                "Hình thức khác..."
            ])
        with h1_c2:
            thoi_gian = st.text_input("⏰ Thời gian:", placeholder="VD: 14h00, 18/07/2026")
        with h1_c3:
            dia_diem = st.text_input("📍 Địa điểm:", placeholder="VD: Văn phòng Trường")

        h2_c1, h2_c2, h2_c3, h2_c4 = st.columns(4)
        with h2_c1:
            chu_toa = st.selectbox("👨‍🏫 Chủ tọa:", ds_gv, index=0)
        with h2_c2:
            index_thu_ky = 1 if len(ds_gv) > 1 else 0
            thu_ky = st.selectbox("✍️ Thư ký:", ds_gv, index=index_thu_ky)
        with h2_c3:
            co_mat = st.text_input("👥 Có mặt:", placeholder="VD: 10/10")
        with h2_c4:
            vang_mat = st.text_input("🚫 Vắng mặt:", placeholder="VD: 0 (hoặc ghi tên)")

    # Lưu thông tin chủ tọa và thư ký vào session để dùng khi xuất file Word
    st.session_state.chu_toa_hien_tai = chu_toa
    st.session_state.thu_ky_hien_tai = thu_ky

    # 2. KHU VỰC NẠP DỰ THẢO
    st.markdown("#### 📄 Bước 2: Nạp Dự thảo kế hoạch / Dàn ý")
    st.info("💡 AI sẽ dò tìm các đề mục lớn (I, II, III...) và mục nhỏ (1, 2, a, b...) trong văn bản này để tạo khung biên bản tương ứng.")
    
    tab_nhap, tab_file = st.tabs(["✍️ Dán văn bản Dự thảo", "📂 Tải file PDF Dự thảo"])
    noidung_du_thao = ""
    
    with tab_nhap:
        text_input = st.text_area("Dán nội dung dự thảo vào đây:", height=150, placeholder="Ví dụ:\nI. Đánh giá công tác tuần qua\n1. Ưu điểm\n2. Tồn tại\nII. Triển khai công tác tuần tới...")
        if text_input:
            noidung_du_thao = text_input
            
    with tab_file:
        uploaded_file = st.file_uploader("Tải lên file dự thảo (PDF)", type=["pdf"])
        if uploaded_file:
            try:
                reader = PdfReader(uploaded_file)
                extracted_text = ""
                for page in reader.pages:
                    extracted_text += page.extract_text() + "\n"
                noidung_du_thao = extracted_text
                st.success("✅ Đã đọc thành công nội dung file PDF!")
            except Exception as e:
                st.error(f"Lỗi đọc file PDF: {e}")

    # 3. NÚT XỬ LÝ AI
    col_btn1, col_btn2 = st.columns([2, 1])
    with col_btn1:
        if st.button("🪄 Viết Biên bản bằng AI", type="primary", use_container_width=True):
            if not noidung_du_thao.strip():
                st.warning("⚠️ Thầy vui lòng cung cấp nội dung hoặc file Dự thảo trước nhé!")
            else:
                with st.spinner("🧠 Thư ký AI đang tổng hợp và soạn thảo biên bản..."):
                    prompt = f"""
                    BẠN LÀ THƯ KÝ TỔ CHUYÊN MÔN TRƯỜNG THCS. HÃY VIẾT MỘT "BIÊN BẢN CUỘC HỌP" CHI TIẾT, MANG VĂN PHONG HÀNH CHÍNH TRANG TRỌNG.
                    
                    THÔNG TIN CHUNG:
                    - Loại hình cuộc họp: {loai_cuoc_hop}
                    - Thời gian: {thoi_gian}
                    - Địa điểm: {dia_diem}
                    - Thành phần: Có mặt {co_mat}, Vắng mặt {vang_mat}
                    - Chủ tọa: {chu_toa}
                    - Thư ký: {thu_ky}
                    
                    NGUYÊN TẮC BẮT BUỘC:
                    Bên dưới là văn bản Dự thảo kế hoạch. Bạn PHẢI bám sát TUYỆT ĐỐI cấu trúc các đề mục của bản dự thảo này (Ví dụ: I, II, III... 1, 2, 3... a, b, c...). Không được tự ý bỏ sót bất kỳ mục nào.
                    
                    YÊU CẦU NỘI DUNG:
                    1. Mở đầu biên bản chuẩn thể thức hành chính, bao gồm đầy đủ Thông tin chung.
                    2. Tại mỗi đề mục, trình bày nội dung của Chủ tọa, sau đó TỰ ĐỘNG THÊM VÀO các ý kiến thảo luận giả định mang tính sư phạm của các thành viên.
                    3. Cuối mỗi mục lớn phải có kết luận chốt lại vấn đề của Chủ tọa.
                    4. Phần cuối biên bản là thời gian kết thúc và chữ ký.
                    
                    DỰ THẢO KẾ HOẠCH:
                    '''{noidung_du_thao}'''
                    """
                    
                    bien_ban = None
                    try:
                        engine_to_use = ai_engine
                        if not engine_to_use and "ai_engine" in st.session_state:
                            engine_to_use = st.session_state.ai_engine

                        if engine_to_use and hasattr(engine_to_use, "generate_text"):
                            try:
                                bien_ban = engine_to_use.generate_text(prompt)
                            except Exception:
                                pass 

                        if not bien_ban:
                            api_key = None
                            for key, val in st.session_state.items():
                                if isinstance(val, str) and val.startswith("sk-"):
                                    api_key = val
                                    break
                            
                            if not api_key:
                                for k in ["user_api_key", "api_key", "openai_api_key", "sk_key"]:
                                    if st.session_state.get(k) and str(st.session_state.get(k)).startswith("sk-"):
                                        api_key = st.session_state.get(k)
                                        break
                            
                            if not api_key and "OPENAI_API_KEY" in st.secrets:
                                api_key = st.secrets["OPENAI_API_KEY"]

                            if api_key:
                                from openai import OpenAI
                                client = OpenAI(api_key=str(api_key).strip())
                                response = client.chat.completions.create(
                                    model="gpt-4o-mini",
                                    messages=[{"role": "user", "content": prompt}]
                                )
                                bien_ban = response.choices[0].message.content
                            else:
                                st.error("❌ Không tìm thấy khóa API `sk-` hợp lệ. Thầy vui lòng kiểm tra lại ô nhập API Key ở menu bên trái.")

                        if bien_ban:
                            # Làm sạch dấu ** ngay từ khi lưu vào session
                            st.session_state.ket_qua_bien_ban = bien_ban.replace("**", "")
                            st.rerun()
                    except Exception as e:
                        st.error(f"❌ Lỗi khi gọi AI: {str(e)}")

    with col_btn2:
        if st.button("🗑️ Xóa / Làm lại", type="secondary", use_container_width=True):
            st.session_state.ket_qua_bien_ban = None
            st.rerun()

    # 4. HIỂN THỊ KẾT QUẢ VÀ TẢI VỀ
    st.markdown("---")
    if st.session_state.ket_qua_bien_ban:
        st.success("🎉 Biên bản đã hoàn thành! Thầy có thể tải về file Word (.docx) được căn chỉnh chuẩn hành chính hoặc file Văn bản (.txt).")
        
        dl_col1, dl_col2 = st.columns(2)
        
        with dl_col1:
            if HAS_DOCX:
                try:
                    doc = Document()
                    
                    # Thiết lập lề trang giấy chuẩn hành chính
                    sections = doc.sections
                    for section in sections:
                        section.top_margin = Inches(1.0)
                        section.bottom_margin = Inches(1.0)
                        section.left_margin = Inches(1.2)
                        section.right_margin = Inches(0.8)
                    
                    # 1. Phần tiêu đề đầu trang (Bảng 2 cột: Tên trường bên trái, Quốc hiệu bên phải)
                    header_table = doc.add_table(rows=1, cols=2)
                    header_table.autofit = False
                    header_table.columns[0].width = Inches(3.0)
                    header_table.columns[1].width = Inches(3.5)
                    
                    cell_left = header_table.cell(0, 0)
                    p_left = cell_left.paragraphs[0]
                    p_left.paragraph_format.space_after = Pt(0)
                    p_left.paragraph_format.line_spacing = 1.15
                    run_l1 = p_left.add_run("TRƯỜNG THCS NGUYỄN CHÍ THANH\n")
                    run_l1.bold = True
                    run_l1.font.size = Pt(10)
                    run_l2 = p_left.add_run("TỔ KHOA HỌC TỰ NHIÊN - GDTC")
                    run_l2.bold = True
                    run_l2.font.size = Pt(10)

                    cell_right = header_table.cell(0, 1)
                    p_right = cell_right.paragraphs[0]
                    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_right.paragraph_format.space_after = Pt(0)
                    p_right.paragraph_format.line_spacing = 1.15
                    run_r1 = p_right.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
                    run_r1.bold = True
                    run_r1.font.size = Pt(10)
                    run_r2 = p_right.add_run("Độc lập - Tự do - Hạnh phúc")
                    run_r2.bold = True
                    run_r2.font.size = Pt(10)

                    # Khoảng cách trước tên biên bản
                    doc.add_paragraph().paragraph_format.space_after = Pt(12)
                    
                    # 2. Tiêu đề chính BIÊN BẢN CUỘC HỌP
                    p_title = doc.add_paragraph()
                    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_title = p_title.add_run("BIÊN BẢN CUỘC HỌP")
                    run_title.bold = True
                    run_title.font.size = Pt(13)
                    p_title.paragraph_format.space_after = Pt(16)

                    # 3. Xử lý và đưa nội dung vào file (Loại bỏ hoàn toàn dấu **)
                    raw_text = st.session_state.ket_qua_bien_ban.replace("**", "")
                    for line in raw_text.split('\n'):
                        clean_line = line.strip()
                        if not clean_line:
                            continue
                        
                        p = doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(4)
                        p.paragraph_format.line_spacing = 1.15
                        
                        # In đậm các đề mục lớn (I., II., III., 1., 2., 3...)
                        if clean_line.startswith(('I.', 'II.', 'III.', '1.', '2.', '3.', '4.', '5.')) and len(clean_line) < 90 and not clean_line.startswith('-'):
                            run = p.add_run(clean_line)
                            run.bold = True
                            if clean_line.startswith(('I.', 'II.', 'III.')):
                                p.paragraph_format.space_before = Pt(8)
                                run.font.size = Pt(11)
                            else:
                                run.font.size = Pt(11)
                        else:
                            p.add_run(clean_line)

                    # 4. Phần chữ ký cuối trang (Bảng 2 cột: Chủ tọa bên trái, Thư ký bên phải)
                    doc.add_paragraph().paragraph_format.space_after = Pt(12)
                    sig_table = doc.add_table(rows=2, cols=2)
                    sig_table.autofit = False
                    sig_table.columns[0].width = Inches(3.25)
                    sig_table.columns[1].width = Inches(3.25)

                    cell_sig1 = sig_table.cell(0, 0)
                    p_s1 = cell_sig1.paragraphs[0]
                    p_s1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r_s1 = p_s1.add_run("CHỦ TỌA")
                    r_s1.bold = True

                    cell_sig2 = sig_table.cell(0, 1)
                    p_s2 = cell_sig2.paragraphs[0]
                    p_s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r_s2 = p_s2.add_run("THƯ KÝ")
                    r_s2.bold = True

                    # Dòng hướng dẫn ký và tên người ký
                    cell_sig1_bot = sig_table.cell(1, 0)
                    p_s1_bot = cell_sig1_bot.paragraphs[0]
                    p_s1_bot.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_s1_bot.add_run("(Ký, ghi rõ họ tên)\n\n\n").italic = True
                    r_name1 = p_s1_bot.add_run(f"{st.session_state.get('chu_toa_hien_tai', '')}")
                    r_name1.bold = True

                    cell_sig2_bot = sig_table.cell(1, 1)
                    p_s2_bot = cell_sig2_bot.paragraphs[0]
                    p_s2_bot.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_s2_bot.add_run("(Ký, ghi rõ họ tên)\n\n\n").italic = True
                    r_name2 = p_s2_bot.add_run(f"{st.session_state.get('thu_ky_hien_tai', '')}")
                    r_name2.bold = True

                    docx_buffer = io.BytesIO()
                    doc.save(docx_buffer)
                    docx_buffer.seek(0)
                    
                    st.download_button(
                        label="⬇️ Tải Biên bản Word (.docx)",
                        data=docx_buffer,
                        file_name="Bien_Ban_SHCM.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary",
                        use_container_width=True
                    )
                except Exception as ex:
                    st.error(f"Lỗi tạo file Word: {ex}")
            else:
                st.warning("⚠️ Cần cài đặt `python-docx` để tải file Word. Chạy lệnh: `pip install python-docx`")

        with dl_col2:
            # Nút tải file Text (đã lọc sạch dấu **)
            clean_txt_data = st.session_state.ket_qua_bien_ban.replace("**", "")
            st.download_button(
                label="⬇️ Tải Biên bản Text (.txt)",
                data=clean_txt_data,
                file_name="Bien_Ban_SHCM.txt",
                mime="text/plain",
                type="secondary",
                use_container_width=True
            )
        
        st.markdown("#### 📜 Nội dung Biên bản")
        st.text_area("Chỉnh sửa biên bản (nếu cần):", value=st.session_state.ket_qua_bien_ban.replace("**", ""), height=600, label_visibility="collapsed")
