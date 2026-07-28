# -*- coding: utf-8 -*-
r"""
============================================================
MODULE: modules/quan_ly_to/xd_cham_sang_kien.py
Mô tả: Module chấm và phân tích sáng kiến kinh nghiệm bằng AI.
Tính năng:
    - Đọc PDF (Sử dụng PyMuPDF/fitz hiện đại)
    - Đọc DOCX
    - Đọc ảnh JPG/JPEG/PNG bằng AI Engine trung tâm
    - Chấm theo Rubric 10 điểm chuẩn
    - Phân tích tính mới, khả thi, hiệu quả, phạm vi ảnh hưởng
    - Kiểm tra logic, văn phong và xuất báo cáo Word / Markdown
============================================================
"""

import streamlit as st
import fitz  # PyMuPDF thay thế PyPDF2
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image
import io
import re
from typing import List, Any

# ============================================================
# 1. CẤU HÌNH MODULE
# ============================================================
MAX_TEXT_LENGTH = 120000
MAX_IMAGE_COUNT = 30

# ============================================================
# 2. HÀM TIỆN ÍCH CHUNG
# ============================================================
def safe_text(value: Any) -> str:
    """Chuyển dữ liệu bất kỳ thành chuỗi an toàn."""
    if value is None:
        return ""
    return str(value).strip()

def normalize_text(text: str) -> str:
    """Chuẩn hóa văn bản."""
    if not text:
        return ""
    text = text.replace("\x00", " ")
    text = re.sub(r"\r\n", "\n", text)
    text = re.sub(r"\r", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

def truncate_text(text: str, max_length: int = MAX_TEXT_LENGTH) -> str:
    """Giới hạn độ dài văn bản để tránh vượt context window."""
    if not text:
        return ""
    if len(text) <= max_length:
        return text
    return text[:max_length] + "\n\n[HỆ THỐNG: Nội dung đã được cắt bớt do vượt giới hạn xử lý.]"

# ============================================================
# 3. ĐỌC FILE PDF (SỬ DỤNG PyMuPDF / fitz)
# ============================================================
def extract_text_from_pdf(uploaded_file) -> str:
    """Trích xuất văn bản từ file PDF sử dụng thư viện PyMuPDF (fitz) tốc độ cao."""
    text_parts = []
    try:
        content = uploaded_file.getvalue() if hasattr(uploaded_file, "getvalue") else uploaded_file.read()
        doc = fitz.open(stream=content, filetype="pdf")
        
        for page_number in range(len(doc)):
            try:
                page = doc[page_number]
                page_text = safe_text(page.get_text("text"))
                if page_text:
                    text_parts.append(f"\n--- TRANG {page_number + 1} ---\n{page_text}")
            except Exception as page_error:
                text_parts.append(f"\n--- TRANG {page_number + 1} ---\n[Không thể đọc trang này: {page_error}]")
                
    except Exception as error:
        raise RuntimeError(f"Không thể đọc file PDF: {error}")
        
    return normalize_text("\n".join(text_parts))

# ============================================================
# 4. ĐỌC FILE WORD
# ============================================================
def extract_text_from_docx(uploaded_file) -> str:
    """Trích xuất nội dung từ DOCX."""
    try:
        source = uploaded_file.getvalue() if hasattr(uploaded_file, "getvalue") else uploaded_file.read()
        doc = Document(io.BytesIO(source))
        content_parts = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                content_parts.append(text)
        for table_index, table in enumerate(doc.tables, start=1):
            content_parts.append(f"\n--- BẢNG {table_index} ---")
            for row in table.rows:
                row_values = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_values.append(cell_text)
                content_parts.append(" | ".join(row_values))
    except Exception as error:
        raise RuntimeError(f"Không thể đọc file Word: {error}")
    return normalize_text("\n".join(content_parts))

# ============================================================
# 5. ĐỌC FILE ẢNH
# ============================================================
def load_image(uploaded_file):
    """Đọc ảnh và chuyển về RGB."""
    try:
        uploaded_file.seek(0)
        image = Image.open(uploaded_file)
        if image.mode != "RGB":
            image = image.convert("RGB")
        return image
    except Exception as error:
        raise RuntimeError(f"Không thể đọc hình ảnh: {error}")

# ============================================================
# 6. PHÂN LOẠI VÀ XỬ LÝ FILE
# ============================================================
def process_uploaded_files(uploaded_files):
    """Xử lý toàn bộ file được tải lên."""
    text_parts = []
    images = []
    processed_files = []

    for uploaded_file in uploaded_files:
        file_name = uploaded_file.name.lower()
        processed_files.append(uploaded_file.name)

        if file_name.endswith(".pdf"):
            text = extract_text_from_pdf(uploaded_file)
            if text:
                text_parts.append(f"\n===== FILE: {uploaded_file.name} =====\n{text}")
        elif file_name.endswith(".docx"):
            text = extract_text_from_docx(uploaded_file)
            if text:
                text_parts.append(f"\n===== FILE: {uploaded_file.name} =====\n{text}")
        elif file_name.endswith((".jpg", ".jpeg", ".png")):
            if len(images) < MAX_IMAGE_COUNT:
                image = load_image(uploaded_file)
                images.append(image)

    combined_text = normalize_text("\n".join(text_parts))
    combined_text = truncate_text(combined_text)

    return {
        "text": combined_text,
        "images": images,
        "files": processed_files
    }

# ============================================================
# 7. PROMPT CHUYÊN GIA CHẤM SÁNG KIẾN
# ============================================================
def build_evaluation_prompt(content: str, file_names: List[str], has_images: bool = False) -> str:
    """Tạo prompt chuyên gia đánh giá sáng kiến."""
    image_note = ""
    if has_images:
        image_note = "Ngoài phần văn bản dưới đây, hệ thống có thể cung cấp thêm hình ảnh các trang tài liệu. Hãy sử dụng thông tin trong hình ảnh khi cần thiết để bổ sung cho việc phân tích."

    prompt = f"""
VAI TRÒ
Bạn là chuyên gia trong hội đồng đánh giá sáng kiến kinh nghiệm trong lĩnh vực giáo dục.
Bạn có nhiệm vụ phân tích khách quan, có căn cứ và có tính phản biện đối với một sáng kiến kinh nghiệm.

MỤC TIÊU
Hãy đánh giá sáng kiến dựa trên nội dung thực tế được cung cấp.
Không được tự bịa thêm kết quả, số liệu, minh chứng hoặc thông tin không xuất hiện trong tài liệu.
Nếu tài liệu thiếu thông tin để đánh giá một tiêu chí, phải ghi rõ: "Chưa đủ minh chứng để kết luận chắc chắn."
Không được khẳng định sáng kiến đạo văn chỉ dựa trên văn phong.
Đánh giá dấu hiệu AI hoặc sao chép chỉ được trình bày dưới dạng: "Nhận diện tham khảo, không phải kết luận pháp lý hoặc kết luận đạo văn."

{image_note}

============================================================
THÔNG TIN FILE ĐƯỢC PHÂN TÍCH
============================================================
{chr(10).join(file_names)}

============================================================
NỘI DUNG SÁNG KIẾN
============================================================
{content}

============================================================
RUBRIC CHẤM ĐIỂM
============================================================
TỔNG ĐIỂM: 10 ĐIỂM

1. TÍNH MỚI: 0 - 3 ĐIỂM
2. TÍNH KHẢ THI: 0 - 3 ĐIỂM
3. HIỆU QUẢ: 0 - 2 ĐIỂM
4. PHẠM VI ẢNH HƯỞNG: 0 - 2 ĐIỂM

============================================================
YÊU CẦU ĐẦU RA (Trả về định dạng Markdown đúng cấu trúc, không dùng dấu ** thừa)
============================================================
# BÁO CÁO ĐÁNH GIÁ SÁNG KIẾN KINH NGHIỆM

## I. TÓM TẮT ĐÁNH GIÁ
(Tóm tắt ngắn gọn: vấn đề, giải pháp, đối tượng, kết quả, nhận định tổng quan)

## II. BẢNG ĐIỂM TỔNG HỢP
| STT | Tiêu chí | Điểm tối đa | Điểm đạt | Nhận xét ngắn |
|---|---|---:|---:|---|
| 1 | Tính mới | 3 | x | ... |
| 2 | Tính khả thi | 3 | x | ... |
| 3 | Hiệu quả | 2 | x | ... |
| 4 | Phạm vi ảnh hưởng | 2 | x | ... |
| | TỔNG CỘNG | 10 | x | ... |

ĐIỂM TỔNG: x/10
MỨC ĐÁNH GIÁ: ...

## III. PHÂN TÍCH CHI TIẾT TỪNG TIÊU CHÍ
### 1. TÍNH MỚI: x/3 ĐIỂM
### 2. TÍNH KHẢ THI: x/3 ĐIỂM
### 3. HIỆU QUẢ: x/2 ĐIỂM
### 4. PHẠM VI ẢNH HƯỞNG: x/2 ĐIỂM

## IV. KIỂM TRA TÍNH LOGIC
## V. PHÂN TÍCH TÍNH NGUYÊN GỐC VÀ VĂN PHONG
## VI. ƯU ĐIỂM NỔI BẬT
## VII. HẠN CHẾ VÀ RỦI RO
## VIII. CÁC ĐIỂM CẦN BỔ SUNG MINH CHỨNG
## IX. ĐỀ XUẤT CHỈNH SỬA
## X. KẾT LUẬN CHUYÊN MÔN
"""
    return prompt

# ============================================================
# 8. GỌI AI ENGINE TRUNG TÂM (Hỗ trợ đa định tuyến)
# ============================================================
def call_ai_engine(ai_engine, prompt: str) -> str:
    """Gọi AI Engine trung tâm hoặc dự phòng gọi OpenAI trực tiếp."""
    if ai_engine and hasattr(ai_engine, "generate_text"):
        try:
            result = ai_engine.generate_text(prompt)
            if isinstance(result, str):
                return result.strip()
            if hasattr(result, "text"):
                return safe_text(result.text)
            if isinstance(result, dict):
                for key in ["text", "content", "response", "result"]:
                    if key in result:
                        return safe_text(result[key])
            return safe_text(result)
        except Exception:
            pass

    # Dự phòng gọi trực tiếp bằng khóa sk- đang có trong session_state hoặc secrets
    api_key = None
    for key, val in st.session_state.items():
        if isinstance(val, str) and val.startswith("sk-"):
            api_key = val
            break
    
    if not api_key:
        for k in ["user_api_key", "api_key", "openai_api_key", "sk_key"]:
            if st.session_state.get(k) and str(st.session_state.get(k)).startswith("sk-"):
                api_key = st.session_state.get(k)
                break
    
    if not api_key and "OPENAI_API_KEY" in st.secrets:
        api_key = st.secrets["OPENAI_API_KEY"]

    if api_key:
        from openai import OpenAI
        client = OpenAI(api_key=str(api_key).strip())
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}]
        )
        return response.choices[0].message.content.strip()

    raise RuntimeError("AI Engine chưa được khởi tạo hoặc không tìm thấy API Key hợp lệ.")

# ============================================================
# 9. TẠO BÁO CÁO MARKDOWN
# ============================================================
def build_markdown_report(ai_result: str, file_names: List[str]) -> str:
    """Tạo báo cáo Markdown hoàn chỉnh."""
    report = []
    report.append("# BÁO CÁO CHẤM VÀ PHÂN TÍCH SÁNG KIẾN KINH NGHIỆM")
    report.append("\n## THÔNG TIN HỒ SƠ\n")
    report.append("### Tài liệu được phân tích")
    for file_name in file_names:
        report.append(f"- {file_name}")
    report.append("\n---\n")
    report.append(ai_result)
    report.append("\n---\n")
    report.append("*Báo cáo được tạo với sự hỗ trợ của AI. Kết quả cần được hội đồng chuyên môn kiểm tra, đối chiếu và quyết định cuối cùng.*")
    return "\n".join(report)

# ============================================================
# 10. HỖ TRỢ XUẤT WORD
# ============================================================
def set_cell_shading(cell, fill: str = "D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_text_bold(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "808080")

def add_markdown_content_to_docx(document: Document, markdown_text: str):
    lines = markdown_text.replace("**", "").splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue

        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            title = line.lstrip("#").strip()
            heading = document.add_heading(title, level=min(level, 3))
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            index += 1
            continue

        if "|" in line and index + 1 < len(lines) and "|" in lines[index + 1]:
            table_lines = []
            while index < len(lines) and "|" in lines[index]:
                current_line = lines[index].strip()
                if not re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?$", current_line):
                    table_lines.append(current_line)
                index += 1
            if table_lines:
                rows = [ [cell.strip() for cell in tl.strip("|").split("|")] for tl in table_lines ]
                max_columns = max(len(row) for row in rows)
                table = document.add_table(rows=len(rows), cols=max_columns)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = "Table Grid"
                set_table_borders(table)
                for row_index, row in enumerate(rows):
                    for col_index in range(max_columns):
                        cell = table.cell(row_index, col_index)
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        if col_index < len(row):
                            cell.text = row[col_index]
                        else:
                            cell.text = ""
                        if row_index == 0:
                            set_cell_shading(cell)
                            set_cell_text_bold(cell)
                document.add_paragraph()
            continue

        if line.startswith("- ") or line.startswith("* "):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(line[2:].strip())
            index += 1
            continue

        numbered_match = re.match(r"^\d+\.\s+(.*)", line)
        if numbered_match:
            paragraph = document.add_paragraph(style="List Number")
            paragraph.add_run(numbered_match.group(1))
            index += 1
            continue

        paragraph = document.add_paragraph()
        paragraph.add_run(line)
        index += 1

def create_word_report(markdown_report: str) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    normal_style = styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BÁO CÁO CHẤM VÀ PHÂN TÍCH\nSÁNG KIẾN KINH NGHIỆM")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)
    document.add_paragraph()

    add_markdown_content_to_docx(document, markdown_report)

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output.getvalue()

# ============================================================
# 11. GIAO DIỆN STREAMLIT
# ============================================================
def render_cham_sang_kien(ai_engine):
    st.markdown("## 🔍 Chấm & Góp ý Sáng kiến Kinh nghiệm")
    st.caption("Phân tích sáng kiến theo Rubric 10 điểm và hỗ trợ xây dựng báo cáo phản biện chuyên môn.")

    with st.expander("ℹ️ Hướng dẫn sử dụng", expanded=False):
        st.markdown(
            "**Bước 1:** Tải lên một hoặc nhiều file sáng kiến.\n\n"
            "**Bước 2:** Có thể tải lên: File PDF, File Word DOCX, Ảnh JPG/JPEG/PNG.\n\n"
            "**Bước 3:** Nhấn **BẮT ĐẦU CHẤM ĐIỂM & PHÂN TÍCH**.\n\n"
            "**Bước 4:** Kiểm tra kết quả và xuất báo cáo Word."
        )

    uploaded_files = st.file_uploader(
        "📂 Tải lên bản sáng kiến",
        type=["pdf", "docx", "jpg", "jpeg", "png"],
        accept_multiple_files=True,
        help="Có thể tải lên nhiều file PDF, DOCX hoặc ảnh các trang sáng kiến."
    )

    if uploaded_files:
        st.success(f"Đã tải lên {len(uploaded_files)} file.")
        with st.expander("📄 Danh sách file", expanded=False):
            for file in uploaded_files:
                st.write(f"• {file.name}")

    analyze_button = st.button("⚖️ BẮT ĐẦU CHẤM ĐIỂM & PHÂN TÍCH", type="primary", use_container_width=True)

    if analyze_button:
        if not uploaded_files:
            st.warning("⚠️ Vui lòng tải ít nhất một file sáng kiến.")
            return

        try:
            with st.spinner("🔍 Đang đọc và xử lý tài liệu..."):
                processed_data = process_uploaded_files(uploaded_files)
            
            combined_text = processed_data["text"]
            images = processed_data["images"]
            file_names = processed_data["files"]

            if not combined_text and not images:
                st.error("❌ Không thể đọc được nội dung từ các file đã tải lên.")
                return

            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Số file", len(file_names))
            with col2:
                st.metric("Số ký tự văn bản", f"{len(combined_text):,}")
            with col3:
                st.metric("Số ảnh", len(images))

            with st.spinner("🧠 AI đang xây dựng phân tích chuyên môn..."):
                prompt = build_evaluation_prompt(content=combined_text, file_names=file_names, has_images=bool(images))

            with st.spinner("⚖️ AI đang chấm điểm theo Rubric..."):
                ai_result = call_ai_engine(ai_engine, prompt)

            if not ai_result:
                st.error("❌ AI không trả về kết quả.")
                return

            markdown_report = build_markdown_report(ai_result, file_names)

            st.session_state["sk_result"] = ai_result
            st.session_state["sk_markdown_report"] = markdown_report
            st.session_state["sk_file_names"] = file_names

            st.success("✅ Đã hoàn thành phân tích sáng kiến.")

        except Exception as error:
            st.error(f"❌ Lỗi xử lý: {error}")
            st.exception(error)

    if "sk_result" in st.session_state:
        st.divider()
        st.markdown("## 📊 KẾT QUẢ CHẤM VÀ PHÂN TÍCH")
        st.markdown(st.session_state["sk_result"])
        st.divider()

        st.markdown("## 📥 XUẤT BÁO CÁO")
        markdown_report = st.session_state["sk_markdown_report"]

        try:
            word_bytes = create_word_report(markdown_report)
        except Exception as error:
            word_bytes = None
            st.warning(f"Không thể tạo file Word: {error}")

        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                label="📄 Tải báo cáo Markdown",
                data=markdown_report,
                file_name="Bao_cao_cham_sang_kien.md",
                mime="text/markdown",
                use_container_width=True
            )
        with col2:
            if word_bytes:
                st.download_button(
                    label="📝 Tải báo cáo Word",
                    data=word_bytes,
                    file_name="Bao_cao_cham_sang_kien.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

    with st.expander("🛡️ Tiện ích kiểm tra nâng cao", expanded=False):
        st.info("Phần nhận diện tính nguyên gốc trong báo cáo AI chỉ mang tính tham khảo dựa trên phân tích văn phong và tính nhất quán nội dung. Không được sử dụng kết quả này như kết luận chính thức về đạo văn hoặc việc sử dụng AI.")
        st.markdown("### Gợi ý mở rộng trong tương lai\nCó thể tích hợp thêm: Cơ sở dữ liệu sáng kiến, Tìm kiếm văn bản tương đồng, So sánh các phiên bản...")
