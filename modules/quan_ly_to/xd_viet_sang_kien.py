# -*- coding: utf-8 -*-

import streamlit as st
import io

try:
    import fitz  # PyMuPDF (Đã thay thế hoàn toàn PyPDF2)
except ImportError:
    pass

from docx import Document

# ============================================================
# ĐỌC FILE PDF (SỬ DỤNG PyMuPDF / fitz)
# ============================================================
def extract_text_from_pdf(uploaded_file):
    text_parts = []
    try:
        content = uploaded_file.getvalue() if hasattr(uploaded_file, "getvalue") else uploaded_file.read()
        doc = fitz.open(stream=content, filetype="pdf")
        for page in doc:
            page_text = page.get_text("text")
            if page_text:
                text_parts.append(page_text.strip())
    except Exception as e:
        return f"[LỖI ĐỌC FILE PDF: {str(e)}]"
    return "\n\n".join(text_parts)

# ============================================================
# ĐỌC FILE WORD
# ============================================================
def extract_text_from_docx(uploaded_file):
    text_parts = []
    try:
        source = uploaded_file.getvalue() if hasattr(uploaded_file, "getvalue") else uploaded_file.read()
        doc = Document(io.BytesIO(source))
        # Đoạn văn
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)
        # Bảng
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))
    except Exception as e:
        return f"[LỖI ĐỌC FILE DOCX: {str(e)}]"
    return "\n".join(text_parts)

# ============================================================
# ĐỌC FILE EXCEL
# ============================================================
def extract_text_from_xlsx(uploaded_file):
    try:
        import openpyxl
        source = uploaded_file.getvalue() if hasattr(uploaded_file, "getvalue") else uploaded_file.read()
        workbook = openpyxl.load_workbook(io.BytesIO(source), data_only=True)
        text_parts = []
        for worksheet in workbook.worksheets:
            text_parts.append(f"\n--- SHEET: {worksheet.title} ---")
            for row in worksheet.iter_rows(values_only=True):
                values = [str(value).strip() for value in row if value is not None and str(value).strip()]
                if values:
                    text_parts.append(" | ".join(values))
        return "\n".join(text_parts)
    except Exception as e:
        return f"[LỖI ĐỌC FILE XLSX: {str(e)}]"

# ============================================================
# ĐỌC FILE TỔNG QUÁT
# ============================================================
def extract_text_from_file(uploaded_file):
    file_name = uploaded_file.name.lower()
    if file_name.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)
    if file_name.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)
    if file_name.endswith(".xlsx"):
        return extract_text_from_xlsx(uploaded_file)
    return f"[ĐỊNH DẠNG FILE CHƯA ĐƯỢC HỖ TRỢ: {uploaded_file.name}]"

# ============================================================
# TẠO PROMPT VIẾT SÁNG KIẾN
# ============================================================
def build_sang_kien_prompt(nam_hoc, doi_tuong, mon_hoc, chu_de, huong_dan_text, du_lieu_text):
    return f"""
Bạn là chuyên gia xây dựng sáng kiến kinh nghiệm trong giáo dục phổ thông Việt Nam.
Bạn đang hỗ trợ giáo viên xây dựng một bản sáng kiến kinh nghiệm có tính thực tiễn, có minh chứng và có thể triển khai trong nhà trường.

============================================================
I. THÔNG TIN BỐI CẢNH
============================================================
Năm học: {nam_hoc}
Môn học: {mon_hoc}
Đối tượng/Lớp: {doi_tuong}
Tên đề tài: {chu_de}

============================================================
II. TÀI LIỆU HƯỚNG DẪN / QUY ĐỊNH
============================================================
{huong_dan_text}

============================================================
III. DỮ LIỆU VÀ MINH CHỨNG THỰC TẾ
============================================================
{du_lieu_text}

============================================================
IV. NHIỆM VỤ
============================================================
Hãy xây dựng bản sáng kiến kinh nghiệm theo hướng:
1. Bám sát bối cảnh giáo dục thực tế.
2. Không tự bịa số liệu.
3. Nếu dữ liệu minh chứng chưa đủ, phải đánh dấu rõ: [CẦN BỔ SUNG MINH CHỨNG]
4. Không sao chép máy móc tài liệu hướng dẫn.
5. Sử dụng văn phong chuyên môn, tự nhiên, phù hợp với giáo viên THCS.

============================================================
V. CẤU TRÚC ĐỀ XUẤT
============================================================
# TÊN ĐỀ TÀI
## I. ĐẶT VẤN ĐỀ
### 1. Lý do chọn đề tài
### 2. Mục đích nghiên cứu
### 3. Đối tượng nghiên cứu
### 4. Phạm vi nghiên cứu
### 5. Phương pháp nghiên cứu

## II. NỘI DUNG SÁNG KIẾN
### 1. Cơ sở lý luận
### 2. Cơ sở thực tiễn
### 3. Thực trạng trước khi áp dụng giải pháp
### 4. Các giải pháp thực hiện
### 5. Tính mới của sáng kiến
### 6. Hiệu quả của sáng kiến

## III. KẾT LUẬN VÀ KIẾN NGHỊ
### 1. Kết luận
### 2. Kiến nghị

# DANH SÁCH NỘI DUNG CẦN GIÁO VIÊN BỔ SUNG
(Liệt kê các số liệu, minh chứng, hình ảnh cần bổ sung)
"""

# ============================================================
# GIAO DIỆN CHÍNH
# ============================================================
def render_viet_sang_kien(ai_engine):
    st.markdown("### ✍️ Trợ lý Viết Sáng kiến Kinh nghiệm")
    st.info("Tải tài liệu hướng dẫn và minh chứng thực tế. AI sẽ hỗ trợ xây dựng bản thảo dựa trên dữ liệu được cung cấp.")

    with st.expander("📥 Tải tài liệu nền tảng", expanded=True):
        huong_dan = st.file_uploader("Tải công văn/quy định/cấu trúc sáng kiến:", type=["pdf", "docx"], key="sk_huong_dan")
        du_lieu = st.file_uploader("Tải minh chứng/số liệu thực tế:", type=["xlsx", "docx", "pdf"], key="sk_du_lieu")

    st.markdown("#### 📝 Thông tin sáng kiến")
    col1, col2 = st.columns(2)
    with col1:
        nam_hoc = st.text_input("Năm học:", placeholder="Ví dụ: 2025-2026", key="sk_nam_hoc")
        doi_tuong = st.text_input("Đối tượng/Lớp:", placeholder="Ví dụ: Học sinh lớp 7", key="sk_doi_tuong")
    with col2:
        mon_hoc = st.text_input("Môn học:", placeholder="Ví dụ: Khoa học tự nhiên", key="sk_mon_hoc")
        chu_de = st.text_input("Tên đề tài:", placeholder="Nhập tên sáng kiến", key="sk_chu_de")

    if st.button("🚀 PHÁC THẢO DÀN Ý & VIẾT NỘI DUNG", use_container_width=True, type="primary"):
        if ai_engine is None:
            st.error("❌ AI Engine chưa được khởi tạo.")
            return

        if not chu_de.strip():
            st.warning("⚠️ Thầy vui lòng nhập tên đề tài.")
            return

        with st.spinner("📚 Đang đọc tài liệu và xây dựng nội dung..."):
            huong_dan_text = extract_text_from_file(huong_dan) if huong_dan else "Chưa cung cấp tài liệu hướng dẫn."
            du_lieu_text = extract_text_from_file(du_lieu) if du_lieu else "Chưa cung cấp dữ liệu minh chứng."

            MAX_CHARS = 120000
            if len(huong_dan_text) > MAX_CHARS:
                huong_dan_text = huong_dan_text[:MAX_CHARS]
            if len(du_lieu_text) > MAX_CHARS:
                du_lieu_text = du_lieu_text[:MAX_CHARS]

            prompt = build_sang_kien_prompt(nam_hoc, doi_tuong, mon_hoc, chu_de, huong_dan_text, du_lieu_text)

            try:
                # Đã loại bỏ các params (model, temp...) để dùng thẳng hàm chuẩn của app
                result = ai_engine.generate_text(prompt)
                
                # Trích xuất text an toàn từ kết quả AI
                final_text = ""
                if isinstance(result, str):
                    final_text = result.strip()
                elif hasattr(result, "text"):
                    final_text = result.text.strip()
                elif isinstance(result, dict):
                    for key in ["text", "content", "response", "result"]:
                        if key in result:
                            final_text = str(result[key]).strip()
                            break
                            
                st.session_state["sk_viet_result"] = final_text
            except Exception as e:
                st.error(f"❌ Lỗi khi gọi AI Engine: {str(e)}")

    if "sk_viet_result" in st.session_state and st.session_state["sk_viet_result"]:
        st.markdown("---")
        st.markdown("## 📄 BẢN THẢO SÁNG KIẾN")
        st.markdown(st.session_state["sk_viet_result"])

        st.download_button(
            label="📥 Tải bản thảo Markdown",
            data=st.session_state["sk_viet_result"],
            file_name="Ban_thao_sang_kien.md",
            mime="text/markdown",
            use_container_width=True
        )
