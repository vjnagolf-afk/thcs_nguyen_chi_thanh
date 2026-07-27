# -*- coding: utf-8 -*-
"""
============================================================
MODULE: export/word_images.py
Nhiệm vụ: Module chuyên trách xử lý, chuẩn hóa và chèn hình ảnh 
vào tài liệu Word. Hỗ trợ URL, Base64, Path, cấu trúc Metadata từ SGK.
Có Cache URL và cơ chế Tự động tính toán tỷ lệ khung hình an toàn.
============================================================
"""

import os
import re
import json
import base64
import requests
import logging
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

# Bộ nhớ đệm tĩnh lưu trữ các ảnh tải từ HTTP/HTTPS để tránh gọi mạng nhiều lần
_IMAGE_CACHE = {}

def _normalize_image_source(source):
    """
    Chuẩn hóa đa dạng các nguồn ảnh đầu vào thành luồng BytesIO và trích xuất Metadata.
    Hỗ trợ: URL, Local Path, Data URI, Base64, Bytes, BytesIO, Dict, JSON String.
    """
    metadata = {}
    stream = None
    original_source_text = str(source)[:100]

    # 1. Nếu là chuỗi JSON (Được sinh từ AI hoặc module đọc SGK)
    if isinstance(source, str) and source.strip().startswith('{') and source.strip().endswith('}'):
        try:
            source = json.loads(source)
        except Exception:
            pass

    # 2. Nếu là Dictionary chứa Metadata
    if isinstance(source, dict):
        metadata = source
        # Ưu tiên lấy dữ liệu ảnh theo thứ tự: base64 -> url -> path -> data
        actual_source = source.get('base64') or source.get('url') or source.get('path') or source.get('data')
        if not actual_source:
            raise ValueError("Dictionary không chứa trường dữ liệu hình ảnh hợp lệ (base64/url/path).")
        source = actual_source

    # 3. Chuẩn hóa về BytesIO
    if isinstance(source, bytes):
        stream = BytesIO(source)
        
    elif isinstance(source, BytesIO):
        stream = source
        
    elif isinstance(source, str):
        source = source.strip()
        
        # A. Xử lý Data URI (VD: data:image/png;base64,iVBORw0...)
        if source.startswith('data:image'):
            try:
                b64_data = re.sub(r'^data:image/.+;base64,', '', source)
                # Bù padding nếu thiếu
                b64_data += "=" * ((4 - len(b64_data) % 4) % 4)
                stream = BytesIO(base64.b64decode(b64_data))
            except Exception as e:
                raise ValueError(f"Lỗi giải mã Data URI: {e}")
                
        # B. Xử lý Base64 thuần túy (Không có tiền tố)
        elif len(source) > 100 and not source.startswith('http') and not os.path.exists(source):
            try:
                # Bù padding nếu thiếu
                b64_data = source + "=" * ((4 - len(source) % 4) % 4)
                stream = BytesIO(base64.b64decode(b64_data))
            except Exception:
                pass # Chuyển sang thử các nhánh dưới nếu không phải Base64
                
        # C. Xử lý URL HTTP/HTTPS (Có tích hợp Cache)
        if stream is None and re.match(r'^https?://', source):
            if source in _IMAGE_CACHE:
                stream = BytesIO(_IMAGE_CACHE[source])
            else:
                try:
                    resp = requests.get(source, timeout=10)
                    if resp.status_code == 200:
                        _IMAGE_CACHE[source] = resp.content
                        stream = BytesIO(resp.content)
                    else:
                        raise ValueError(f"HTTP Status {resp.status_code} khi tải ảnh.")
                except Exception as e:
                    raise ValueError(f"Lỗi kết nối khi tải URL: {e}")
                    
        # D. Xử lý đường dẫn File Local
        if stream is None and os.path.exists(source):
            try:
                with open(source, 'rb') as f:
                    stream = BytesIO(f.read())
            except Exception as e:
                raise ValueError(f"Lỗi đọc file nội bộ: {e}")

    if stream is None:
        raise ValueError(f"Nguồn ảnh không hợp lệ, không tồn tại hoặc không thể tải: {original_source_text}...")

    return stream, metadata


def insert_image_to_paragraph(paragraph, image_source, width=None, height=None, max_width=Inches(6.0)):
    """
    Chèn hình ảnh vào một đoạn văn (Paragraph) chỉ định. 
    Xử lý thông minh kích thước (chống tràn lề) và chèn chú thích (Caption) từ Metadata.
    """
    try:
        stream, metadata = _normalize_image_source(image_source)
        run = paragraph.add_run()
        
        # Thiết lập thông số kích thước (Smart Sizing)
        kwargs = {}
        if width: kwargs['width'] = width
        if height: kwargs['height'] = height

        # Nếu không ép cứng kích thước, tự động tính toán để không tràn lề giấy
        if not width and not height:
            try:
                from PIL import Image
                img = Image.open(stream)
                w_px, h_px = img.size
                stream.seek(0)
                
                # Giả định phân giải màn hình 96 DPI
                w_in = w_px / 96.0
                limit_w = max_width.inches if max_width else 6.0
                
                # Co nhỏ ảnh nếu bề ngang vượt quá giới hạn lề
                if w_in > limit_w:
                    kwargs['width'] = Inches(limit_w)
                # Nếu ảnh nhỏ hơn max_width, python-docx sẽ tự động giữ nguyên kích thước gốc
            except ImportError:
                # Fallback nếu hệ thống không cài Pillow (PIL)
                kwargs['width'] = max_width or Inches(6.0)
            except Exception:
                kwargs['width'] = max_width or Inches(6.0)

        # Chèn ảnh vào luồng thực thi
        picture = run.add_picture(stream, **kwargs)
        
        # Xử lý chèn Caption (Chú thích ảnh)
        caption = metadata.get('caption') or metadata.get('alt_text') or metadata.get('id')
        if caption and str(caption).lower() != "none" and not str(caption).startswith("IMG_"):
             # Ngắt dòng mềm (Shift+Enter) để giữ caption dính liền với ảnh
            run.add_break()
            cap_run = paragraph.add_run(f"Hình: {caption}")
            cap_run.italic = True
            cap_run.font.size = Pt(11)
            cap_run.font.name = 'Times New Roman'

        return picture

    except Exception as e:
        logger.error(f"Lỗi chèn ảnh: {e}")
        # Chèn cảnh báo text đỏ thẳng vào văn bản nếu lỗi
        err_run = paragraph.add_run(f"\n[Không thể tải hình ảnh: Nguồn ảnh không hợp lệ]\n")
        err_run.italic = True
        err_run.font.color.rgb = RGBColor(255, 0, 0)
        return None

def insert_image_to_docx(doc: Document, image_path_or_url: str):
    """API Tương thích ngược: Chèn một ảnh vào cuối tài liệu Word."""
    if not image_path_or_url:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    insert_image_to_paragraph(p, image_path_or_url)
