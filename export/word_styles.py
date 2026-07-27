# -*- coding: utf-8 -*-
"""
============================================================
MODULE: export/word_styles.py

NHIỆM VỤ
------------------------------------------------------------
Nền tảng định dạng DOCX dùng chung cho toàn bộ hệ thống xuất Word.

MỤC TIÊU
------------------------------------------------------------
1. Thiết lập khổ giấy A4.
2. Thiết lập lề đúng quy định cho văn bản giáo dục:
   - Top    : 1.0 – 1.5 cm
   - Bottom : 1.0 – 1.5 cm
   - Left   : 1.5 – 2.0 cm
   - Right  : 1.0 – 1.5 cm

   Cấu hình mặc định:
   - Top    : 1.5 cm
   - Bottom : 1.5 cm
   - Left   : 2.0 cm
   - Right  : 1.5 cm

3. Không ép căn đều toàn bộ tài liệu một cách máy móc.
4. Căn đều 2 bên cho nội dung văn bản.
5. Không căn đều:
   - Heading
   - Bảng
   - Hình ảnh
   - Công thức hiển thị
   - Danh sách
   - Callout
6. Đồng bộ font bằng cả python-docx và XML.
7. Hỗ trợ xử lý bảng, cell margin, border, shading.
8. Không chứa logic AI, Markdown hay nghiệp vụ.

============================================================
"""

from __future__ import annotations

from typing import Optional

import docx

from docx.enum.section import WD_SECTION
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_LINE_SPACING,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import (
    Cm,
    Inches,
    Pt,
    RGBColor,
)


# ============================================================
# CONSTANTS
# ============================================================

DEFAULT_FONT = "Times New Roman"

DEFAULT_FONT_SIZE = 13

A4_WIDTH_CM = 21.0
A4_HEIGHT_CM = 29.7

# Lề chuẩn được lựa chọn trong phạm vi quy định
TOP_MARGIN_CM = 1.5
BOTTOM_MARGIN_CM = 1.5
LEFT_MARGIN_CM = 2.0
RIGHT_MARGIN_CM = 1.5

# Khoảng cách đoạn
DEFAULT_SPACE_BEFORE_PT = 0
DEFAULT_SPACE_AFTER_PT = 3

# Line spacing
DEFAULT_LINE_SPACING = 1.15


# ============================================================
# XML HELPERS
# ============================================================

class XmlHelpers:
    """
    Các hàm thao tác XML Word dùng chung.

    Mục tiêu:
    - Không tạo thẻ XML trùng lặp.
    - Có thể gọi nhiều lần an toàn.
    - Không làm hỏng cấu trúc DOCX.
    """

    # --------------------------------------------------------
    # FONT
    # --------------------------------------------------------

    @staticmethod
    def set_font_safely(
        run,
        font_name: str = DEFAULT_FONT,
    ) -> None:
        """
        Thiết lập font đầy đủ cho một Run.

        Bao phủ:
        - ascii
        - hAnsi
        - eastAsia
        - cs
        """

        if run is None:
            return

        font_name = font_name or DEFAULT_FONT

        run.font.name = font_name

        r_pr = run._element.get_or_add_rPr()

        r_fonts = r_pr.find(
            qn("w:rFonts")
        )

        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)

        for attr in (
            "ascii",
            "hAnsi",
            "eastAsia",
            "cs",
        ):
            r_fonts.set(
                qn(f"w:{attr}"),
                font_name,
            )

    # --------------------------------------------------------
    # PARAGRAPH ALIGNMENT
    # --------------------------------------------------------

    @staticmethod
    def set_paragraph_alignment(
        paragraph,
        alignment: WD_ALIGN_PARAGRAPH,
    ) -> None:
        """
        Thiết lập căn lề đoạn văn.
        """

        if paragraph is None:
            return

        paragraph.alignment = alignment

    # --------------------------------------------------------
    # SHADING
    # --------------------------------------------------------

    @staticmethod
    def apply_paragraph_shading(
        paragraph,
        color_hex: str = "F5F5F5",
    ) -> None:
        """
        Tô nền toàn bộ đoạn văn.
        """

        if paragraph is None:
            return

        p_pr = paragraph._element.get_or_add_pPr()

        shd = p_pr.find(
            qn("w:shd")
        )

        if shd is None:
            shd = OxmlElement("w:shd")
            p_pr.append(shd)

        shd.set(
            qn("w:val"),
            "clear",
        )

        shd.set(
            qn("w:color"),
            "auto",
        )

        shd.set(
            qn("w:fill"),
            color_hex,
        )

    @staticmethod
    def apply_cell_shading(
        cell,
        color_hex: str = "FFFFFF",
    ) -> None:
        """
        Tô nền cho ô bảng.
        """

        if cell is None:
            return

        tc_pr = cell._tc.get_or_add_tcPr()

        shd = tc_pr.find(
            qn("w:shd")
        )

        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)

        shd.set(
            qn("w:val"),
            "clear",
        )

        shd.set(
            qn("w:color"),
            "auto",
        )

        shd.set(
            qn("w:fill"),
            color_hex,
        )

    # --------------------------------------------------------
    # PARAGRAPH BORDER
    # --------------------------------------------------------

    @staticmethod
    def apply_bottom_border(
        paragraph,
        color_hex: str = "BFBFBF",
        size: int = 6,
        space: int = 1,
    ) -> None:
        """
        Tạo đường kẻ dưới đoạn văn.

        Dùng cho:
        - HR Markdown
        - Phân cách nội dung
        - Một số thành phần hành chính
        """

        if paragraph is None:
            return

        p_pr = paragraph._element.get_or_add_pPr()

        p_bdr = p_pr.find(
            qn("w:pBdr")
        )

        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)

        bottom = p_bdr.find(
            qn("w:bottom")
        )

        if bottom is None:
            bottom = OxmlElement("w:bottom")
            p_bdr.append(bottom)

        bottom.set(
            qn("w:val"),
            "single",
        )

        bottom.set(
            qn("w:sz"),
            str(size),
        )

        bottom.set(
            qn("w:space"),
            str(space),
        )

        bottom.set(
            qn("w:color"),
            color_hex,
        )

    # --------------------------------------------------------
    # TABLE BORDER
    # --------------------------------------------------------

    @staticmethod
    def set_table_borders(
        table,
        color_hex: str = "000000",
        size: int = 4,
        inside: bool = True,
    ) -> None:
        """
        Thiết lập đường viền bảng.
        """

        if table is None:
            return

        tbl = table._tbl

        tbl_pr = tbl.tblPr

        tbl_borders = tbl_pr.find(
            qn("w:tblBorders")
        )

        if tbl_borders is None:
            tbl_borders = OxmlElement("w:tblBorders")
            tbl_pr.append(tbl_borders)

        border_names = [
            "top",
            "left",
            "bottom",
            "right",
        ]

        if inside:
            border_names.extend(
                [
                    "insideH",
                    "insideV",
                ]
            )

        for border_name in border_names:

            border = tbl_borders.find(
                qn(f"w:{border_name}")
            )

            if border is None:
                border = OxmlElement(
                    f"w:{border_name}"
                )
                tbl_borders.append(border)

            border.set(
                qn("w:val"),
                "single",
            )

            border.set(
                qn("w:sz"),
                str(size),
            )

            border.set(
                qn("w:space"),
                "0",
            )

            border.set(
                qn("w:color"),
                color_hex,
            )

    # --------------------------------------------------------
    # TABLE CELL MARGIN
    # --------------------------------------------------------

    @staticmethod
    def set_cell_margins(
        cell,
        top: int = 80,
        start: int = 100,
        bottom: int = 80,
        end: int = 100,
    ) -> None:
        """
        Thiết lập khoảng đệm bên trong ô.

        Đơn vị Word:
        - Twips / DXA
        """

        if cell is None:
            return

        tc_pr = cell._tc.get_or_add_tcPr()

        tc_mar = tc_pr.find(
            qn("w:tcMar")
        )

        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)

        margins = {
            "top": top,
            "start": start,
            "bottom": bottom,
            "end": end,
        }

        for side, value in margins.items():

            node = tc_mar.find(
                qn(f"w:{side}")
            )

            if node is None:
                node = OxmlElement(
                    f"w:{side}"
                )
                tc_mar.append(node)

            node.set(
                qn("w:w"),
                str(value),
            )

            node.set(
                qn("w:type"),
                "dxa",
            )

    # --------------------------------------------------------
    # TABLE WIDTH
    # --------------------------------------------------------

    @staticmethod
    def set_table_width(
        table,
        width_inches: float,
    ) -> None:
        """
        Thiết lập chiều rộng bảng.
        """

        if table is None:
            return

        width_twips = int(
            width_inches * 1440
        )

        tbl_pr = table._tbl.tblPr

        tbl_w = tbl_pr.find(
            qn("w:tblW")
        )

        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)

        tbl_w.set(
            qn("w:w"),
            str(width_twips),
        )

        tbl_w.set(
            qn("w:type"),
            "dxa",
        )

    # --------------------------------------------------------
    # TABLE LAYOUT
    # --------------------------------------------------------

    @staticmethod
    def set_table_fixed_layout(
        table,
    ) -> None:
        """
        Ép bảng dùng layout cố định.

        Giúp tránh:
        - Bảng tự co giãn bất thường.
        - Cột bị vỡ.
        - Nội dung làm thay đổi kích thước cột ngoài ý muốn.
        """

        if table is None:
            return

        tbl_pr = table._tbl.tblPr

        tbl_layout = tbl_pr.find(
            qn("w:tblLayout")
        )

        if tbl_layout is None:
            tbl_layout = OxmlElement(
                "w:tblLayout"
            )
            tbl_pr.append(tbl_layout)

        tbl_layout.set(
            qn("w:type"),
            "fixed",
        )

    # --------------------------------------------------------
    # ROW CONTROL
    # --------------------------------------------------------

    @staticmethod
    def prevent_row_split(
        row,
    ) -> None:
        """
        Không cho phép một hàng bảng bị tách giữa hai trang.
        """

        if row is None:
            return

        tr_pr = row._tr.get_or_add_trPr()

        cant_split = tr_pr.find(
            qn("w:cantSplit")
        )

        if cant_split is None:
            cant_split = OxmlElement(
                "w:cantSplit"
            )
            tr_pr.append(cant_split)

    @staticmethod
    def repeat_table_header(
        row,
    ) -> None:
        """
        Cho phép hàng đầu bảng lặp lại khi sang trang mới.
        """

        if row is None:
            return

        tr_pr = row._tr.get_or_add_trPr()

        tbl_header = tr_pr.find(
            qn("w:tblHeader")
        )

        if tbl_header is None:
            tbl_header = OxmlElement(
                "w:tblHeader"
            )
            tr_pr.append(tbl_header)

    # --------------------------------------------------------
    # KEEP WITH NEXT
    # --------------------------------------------------------

    @staticmethod
    def set_keep_with_next(
        paragraph,
        value: bool = True,
    ) -> None:
        """
        Giữ tiêu đề đi cùng đoạn văn tiếp theo.
        """

        if paragraph is None:
            return

        paragraph.paragraph_format.keep_with_next = value

    # --------------------------------------------------------
    # KEEP TOGETHER
    # --------------------------------------------------------

    @staticmethod
    def set_keep_together(
        paragraph,
        value: bool = True,
    ) -> None:
        """
        Không cho phép đoạn văn bị tách bất hợp lý.
        """

        if paragraph is None:
            return

        paragraph.paragraph_format.keep_together = value


# ============================================================
# BASE STYLE SETUP
# ============================================================

class BaseStyleSetup:
    """
    Thiết lập nền tảng định dạng cho DOCX.

    Đây là lớp trung tâm được gọi trước khi render nội dung.
    """

    # --------------------------------------------------------
    # PAGE SETUP
    # --------------------------------------------------------

    @staticmethod
    def setup_page(
        doc: docx.Document,
    ) -> None:
        """
        Thiết lập khổ A4 và lề trang.

        A4:
        - 21.0 x 29.7 cm

        Lề:
        - Top    : 1.5 cm
        - Bottom : 1.5 cm
        - Left   : 2.0 cm
        - Right  : 1.5 cm
        """

        for section in doc.sections:

            section.page_width = Cm(
                A4_WIDTH_CM
            )

            section.page_height = Cm(
                A4_HEIGHT_CM
            )

            section.top_margin = Cm(
                TOP_MARGIN_CM
            )

            section.bottom_margin = Cm(
                BOTTOM_MARGIN_CM
            )

            section.left_margin = Cm(
                LEFT_MARGIN_CM
            )

            section.right_margin = Cm(
                RIGHT_MARGIN_CM
            )

            # Không chừa thêm gáy.
            section.gutter = Cm(0)

    # --------------------------------------------------------
    # NORMAL STYLE
    # --------------------------------------------------------

    @staticmethod
    def setup_normal_style(
        doc: docx.Document,
    ) -> None:
        """
        Thiết lập Normal Style.
        """

        style = doc.styles["Normal"]

        style.font.name = DEFAULT_FONT
        style.font.size = Pt(
            DEFAULT_FONT_SIZE
        )

        style.font.bold = False
        style.font.italic = False

        style.font.color.rgb = RGBColor(
            0,
            0,
            0,
        )

        XmlHelpers.set_font_safely(
            style._element,
            DEFAULT_FONT,
        )

        paragraph_format = style.paragraph_format

        paragraph_format.alignment = (
            WD_ALIGN_PARAGRAPH.JUSTIFY
        )

        paragraph_format.space_before = Pt(
            DEFAULT_SPACE_BEFORE_PT
        )

        paragraph_format.space_after = Pt(
            DEFAULT_SPACE_AFTER_PT
        )

        paragraph_format.line_spacing = (
            DEFAULT_LINE_SPACING
        )

    # --------------------------------------------------------
    # HEADING STYLES
    # --------------------------------------------------------

    @staticmethod
    def setup_heading_styles(
        doc: docx.Document,
    ) -> None:
        """
        Thiết lập Heading 1, 2, 3.

        Không căn đều heading.
        """

        heading_config = {
            "Heading 1": {
                "size": 16,
                "bold": True,
                "italic": False,
                "before": 10,
                "after": 5,
                "alignment": WD_ALIGN_PARAGRAPH.CENTER,
            },
            "Heading 2": {
                "size": 14,
                "bold": True,
                "italic": False,
                "before": 8,
                "after": 4,
                "alignment": WD_ALIGN_PARAGRAPH.LEFT,
            },
            "Heading 3": {
                "size": 13,
                "bold": True,
                "italic": False,
                "before": 6,
                "after": 3,
                "alignment": WD_ALIGN_PARAGRAPH.LEFT,
            },
        }

        for style_name, config in heading_config.items():

            style = doc.styles[style_name]

            style.font.name = DEFAULT_FONT

            style.font.size = Pt(
                config["size"]
            )

            style.font.bold = config["bold"]

            style.font.italic = config["italic"]

            style.font.color.rgb = RGBColor(
                0,
                0,
                0,
            )

            style.paragraph_format.alignment = (
                config["alignment"]
            )

            style.paragraph_format.space_before = Pt(
                config["before"]
            )

            style.paragraph_format.space_after = Pt(
                config["after"]
            )

            style.paragraph_format.keep_with_next = True

            XmlHelpers.set_font_safely(
                style._element,
                DEFAULT_FONT,
            )

    # --------------------------------------------------------
    # LIST STYLES
    # --------------------------------------------------------

    @staticmethod
    def setup_list_styles(
        doc: docx.Document,
    ) -> None:
        """
        Thiết lập style danh sách.

        Không căn đều để tránh giãn khoảng trắng
        bất thường ở các dòng ngắn.
        """

        for style_name in (
            "List Bullet",
            "List Number",
        ):

            style = doc.styles[style_name]

            style.font.name = DEFAULT_FONT

            style.font.size = Pt(
                DEFAULT_FONT_SIZE
            )

            style.font.bold = False

            style.font.italic = False

            style.paragraph_format.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
            )

            style.paragraph_format.space_before = Pt(
                0
            )

            style.paragraph_format.space_after = Pt(
                2
            )

            style.paragraph_format.line_spacing = (
                DEFAULT_LINE_SPACING
            )

            XmlHelpers.set_font_safely(
                style._element,
                DEFAULT_FONT,
            )

    # --------------------------------------------------------
    # TABLE STYLE
    # --------------------------------------------------------

    @staticmethod
    def setup_table_style(
        doc: docx.Document,
    ) -> None:
        """
        Thiết lập Table Grid nếu có.
        """

        try:

            table_style = doc.styles[
                "Table Grid"
            ]

            table_style.font.name = DEFAULT_FONT

            table_style.font.size = Pt(
                12
            )

            XmlHelpers.set_font_safely(
                table_style._element,
                DEFAULT_FONT,
            )

        except KeyError:
            pass

    # --------------------------------------------------------
    # FULL SETUP
    # --------------------------------------------------------

    @classmethod
    def setup_base_styles(
        cls,
        doc: docx.Document,
    ) -> docx.Document:
        """
        Hàm trung tâm thiết lập toàn bộ style nền.

        Thứ tự:
        1. Page
        2. Normal
        3. Heading
        4. List
        5. Table
        """

        cls.setup_page(doc)

        cls.setup_normal_style(doc)

        cls.setup_heading_styles(doc)

        cls.setup_list_styles(doc)

        cls.setup_table_style(doc)

        return doc


# ============================================================
# BLOCK RENDERER
# ============================================================

class BlockRenderer:
    """
    Render các block cấp cao:
    - Heading
    - List
    - Checkbox
    """

    # --------------------------------------------------------
    # HEADING
    # --------------------------------------------------------

    @classmethod
    def render_heading(
        cls,
        doc,
        node: dict,
        text_renderer,
        math_renderer,
    ):
        """
        Render tiêu đề.
        """

        level = int(
            node.get(
                "level",
                1,
            )
        )

        level = min(
            max(level, 1),
            3,
        )

        paragraph = doc.add_paragraph(
            style=f"Heading {level}"
        )

        if level == 1:

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

        else:

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
            )

        XmlHelpers.set_keep_with_next(
            paragraph,
            True,
        )

        text_renderer.render_inline_tokens(
            paragraph,
            node.get(
                "tokens",
                [],
            ),
            math_renderer,
        )

        for run in paragraph.runs:

            XmlHelpers.set_font_safely(
                run,
                DEFAULT_FONT,
            )

        return paragraph

    # --------------------------------------------------------
    # LIST ITEM
    # --------------------------------------------------------

    @classmethod
    def render_list_item(
        cls,
        doc,
        node: dict,
        text_renderer,
        math_renderer,
    ):
        """
        Render một mục danh sách.

        Hỗ trợ:
        - Bullet
        - Number
        - Nested level
        """

        style_name = (
            "List Number"
            if node.get("style") == "number"
            else "List Bullet"
        )

        paragraph = doc.add_paragraph(
            style=style_name
        )

        level = int(
            node.get(
                "level",
                1,
            )
        )

        level = max(
            level,
            1,
        )

        left_indent = (
            0.25 * level
            + 0.25
        )

        paragraph.paragraph_format.left_indent = (
            Inches(left_indent)
        )

        paragraph.paragraph_format.first_line_indent = (
            Inches(-0.25)
        )

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.LEFT
        )

        text_renderer.render_inline_tokens(
            paragraph,
            node.get(
                "tokens",
                [],
            ),
            math_renderer,
        )

        return paragraph

    # --------------------------------------------------------
    # CHECKBOX
    # --------------------------------------------------------

    @classmethod
    def render_checkbox(
        cls,
        doc,
        node: dict,
        text_renderer,
        math_renderer,
    ):
        """
        Render checkbox.
        """

        paragraph = doc.add_paragraph()

        level = int(
            node.get(
                "level",
                1,
            )
        )

        paragraph.paragraph_format.left_indent = (
            Inches(
                0.25 * level
            )
        )

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.LEFT
        )

        checked = bool(
            node.get(
                "checked",
                False,
            )
        )

        box_char = (
            "☑ "
            if checked
            else "☐ "
        )

        run_box = paragraph.add_run(
            box_char
        )

        XmlHelpers.set_font_safely(
            run_box,
            "Segoe UI Symbol",
        )

        run_box.bold = True

        text_renderer.render_inline_tokens(
            paragraph,
            node.get(
                "tokens",
                [],
            ),
            math_renderer,
        )

        return paragraph


# ============================================================
# CONTAINER RENDERER
# ============================================================

class ContainerRenderer:
    """
    Render các block dạng container:
    - Code
    - Callout
    """

    # --------------------------------------------------------
    # CODE BLOCK
    # --------------------------------------------------------

    @classmethod
    def render_code_block(
        cls,
        doc,
        node: dict,
    ):
        """
        Render code block.
        """

        paragraph = doc.add_paragraph()

        paragraph.paragraph_format.left_indent = (
            Inches(0.25)
        )

        paragraph.paragraph_format.right_indent = (
            Inches(0.25)
        )

        paragraph.paragraph_format.space_before = (
            Pt(4)
        )

        paragraph.paragraph_format.space_after = (
            Pt(4)
        )

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.LEFT
        )

        XmlHelpers.apply_paragraph_shading(
            paragraph,
            "F5F5F5",
        )

        run = paragraph.add_run(
            node.get(
                "text",
                "",
            )
        )

        run.font.size = Pt(
            10.5
        )

        XmlHelpers.set_font_safely(
            run,
            "Courier New",
        )

        return paragraph

    # --------------------------------------------------------
    # CALLOUT
    # --------------------------------------------------------

    @classmethod
    def render_callout(
        cls,
        doc,
        node: dict,
        text_renderer,
        math_renderer,
    ):
        """
        Render khối Callout.

        Không dùng paragraph shading cho block dài.
        Dùng bảng 1 ô để:
        - Giữ nền ổn định.
        - Giữ border.
        - Hạn chế lỗi khi sang trang.
        """

        style = (
            node.get(
                "style",
                "quote",
            )
            or "quote"
        ).lower()

        if style == "warning":

            bg_color = "FFF5F5"
            border_color = "C00000"

        elif style == "tip":

            bg_color = "F0F7FF"
            border_color = "4472C4"

        else:

            bg_color = "F9F9F9"
            border_color = "808080"

        table = doc.add_table(
            rows=1,
            cols=1,
        )

        table.alignment = (
            WD_TABLE_ALIGNMENT.CENTER
        )

        table.autofit = False

        XmlHelpers.set_table_fixed_layout(
            table
        )

        cell = table.cell(
            0,
            0,
        )

        cell.vertical_alignment = (
            WD_CELL_VERTICAL_ALIGNMENT.CENTER
        )

        XmlHelpers.set_cell_margins(
            cell,
            top=100,
            start=140,
            bottom=100,
            end=140,
        )

        XmlHelpers.apply_cell_shading(
            cell,
            bg_color,
        )

        # ----------------------------------------------------
        # BORDER
        # ----------------------------------------------------

        tc_pr = cell._tc.get_or_add_tcPr()

        tc_borders = tc_pr.find(
            qn("w:tcBorders")
        )

        if tc_borders is None:

            tc_borders = OxmlElement(
                "w:tcBorders"
            )

            tc_pr.append(
                tc_borders
            )

        for side in (
            "top",
            "bottom",
            "right",
        ):

            border = tc_borders.find(
                qn(f"w:{side}")
            )

            if border is None:

                border = OxmlElement(
                    f"w:{side}"
                )

                tc_borders.append(
                    border
                )

            border.set(
                qn("w:val"),
                "nil",
            )

        left = tc_borders.find(
            qn("w:left")
        )

        if left is None:

            left = OxmlElement(
                "w:left"
            )

            tc_borders.append(
                left
            )

        left.set(
            qn("w:val"),
            "single",
        )

        left.set(
            qn("w:sz"),
            "20",
        )

        left.set(
            qn("w:space"),
            "0",
        )

        left.set(
            qn("w:color"),
            border_color,
        )

        # ----------------------------------------------------
        # CONTENT
        # ----------------------------------------------------

        first_paragraph = cell.paragraphs[0]

        children = node.get(
            "children",
            [],
        )

        for index, child in enumerate(
            children
        ):

            if index == 0:

                paragraph = first_paragraph

            else:

                paragraph = cell.add_paragraph()

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.JUSTIFY
            )

            paragraph.paragraph_format.space_after = (
                Pt(3)
            )

            if child.get(
                "type"
            ) == "paragraph":

                text_renderer.render_inline_tokens(
                    paragraph,
                    child.get(
                        "tokens",
                        [],
                    ),
                    math_renderer,
                )

            elif child.get(
                "type"
            ) == "heading":

                text_renderer.render_inline_tokens(
                    paragraph,
                    child.get(
                        "tokens",
                        [],
                    ),
                    math_renderer,
                )

                for run in paragraph.runs:

                    run.bold = True

        return table
