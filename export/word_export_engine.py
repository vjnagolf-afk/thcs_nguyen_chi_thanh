# -*- coding: utf-8 -*-

"""
============================================================
WORD EXPORT ENGINE
============================================================

Mục tiêu:

1. Đọc trực tiếp template Word.
2. Bảo toàn cấu trúc template.
3. Thay thế placeholder bằng nội dung rich text.
4. Render Markdown thành DOCX.
5. Hỗ trợ:
   - Văn bản
   - Heading
   - Bold / Italic / Underline
   - Inline code
   - Inline math
   - Công thức Toán
   - Công thức Vật lý
   - Công thức Hóa học
   - Bảng
   - Danh sách
   - Checkbox
   - Hình ảnh
   - Code block
   - Horizontal rule

6. Không phá hỏng template khbd_mau.docx.
7. Không làm mất nội dung TT18 / KHBD.
"""

from __future__ import annotations

import io
import os
import re
import logging
from copy import deepcopy
from pathlib import Path
from typing import Any, Dict, List, Optional

import docx

from docx import Document
from docx.shared import (
    Pt,
    Cm,
    Inches,
    RGBColor
)

from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK
)

from docx.enum.table import (
    WD_TABLE_ALIGNMENT,
    WD_CELL_VERTICAL_ALIGNMENT
)

from docx.enum.style import WD_STYLE_TYPE

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell

try:
    from .markdown_tokenizer import MarkdownTokenizer
except ImportError:
    from export.markdown_tokenizer import MarkdownTokenizer

try:
    from .word_math import ScienceNormalizer
except ImportError:
    from export.word_math import ScienceNormalizer

try:
    from .word_images import ImageRenderer
except ImportError:
    from export.word_images import ImageRenderer

try:
    from .template_loader import TemplateLoader
except ImportError:
    from export.template_loader import TemplateLoader


logger = logging.getLogger("WordExportEngine")


# ============================================================
# CONFIGURATION
# ============================================================

class WordExportConfig:

    FONT_NAME = "Times New Roman"

    BODY_FONT_SIZE = 13

    HEADING_1_SIZE = 16
    HEADING_2_SIZE = 14
    HEADING_3_SIZE = 13

    MATH_FONT_SIZE = 13

    CODE_FONT_NAME = "Courier New"
    CODE_FONT_SIZE = 10.5

    # Lề theo yêu cầu:
    # Top: 1.0 - 1.5 cm
    # Bottom: 1.0 - 1.5 cm
    # Left: 1.5 - 2.0 cm
    # Right: 1.0 - 1.5 cm
    #
    # Chọn giá trị cân bằng:
    TOP_MARGIN = Cm(1.2)
    BOTTOM_MARGIN = Cm(1.2)
    LEFT_MARGIN = Cm(1.8)
    RIGHT_MARGIN = Cm(1.2)

    MAX_TABLE_WIDTH_CM = 17.0

    TABLE_HEADER_COLOR = "D9EAF7"

    TABLE_BORDER_COLOR = "808080"

    CODE_BACKGROUND = "F3F3F3"

    CALLOUT_BACKGROUND = "F8F8F8"


# ============================================================
# XML HELPERS
# ============================================================

class WordXml:

    @staticmethod
    def get_or_create(parent, tag: str):

        element = parent.find(qn(tag))

        if element is None:

            element = OxmlElement(tag)

            parent.append(element)

        return element

    # --------------------------------------------------------

    @staticmethod
    def set_font(
        run,
        font_name: str = WordExportConfig.FONT_NAME
    ):

        run.font.name = font_name

        rPr = run._element.get_or_add_rPr()

        rFonts = rPr.find(qn("w:rFonts"))

        if rFonts is None:

            rFonts = OxmlElement("w:rFonts")

            rPr.append(rFonts)

        for attr in (
            "ascii",
            "hAnsi",
            "eastAsia",
            "cs"
        ):

            rFonts.set(
                qn(f"w:{attr}"),
                font_name
            )

    # --------------------------------------------------------

    @staticmethod
    def set_cell_shading(
        cell,
        fill: str
    ):

        tcPr = cell._tc.get_or_add_tcPr()

        shd = tcPr.find(qn("w:shd"))

        if shd is None:

            shd = OxmlElement("w:shd")

            tcPr.append(shd)

        shd.set(
            qn("w:val"),
            "clear"
        )

        shd.set(
            qn("w:color"),
            "auto"
        )

        shd.set(
            qn("w:fill"),
            fill
        )

    # --------------------------------------------------------

    @staticmethod
    def set_cell_margins(
        cell,
        top: int = 80,
        start: int = 100,
        bottom: int = 80,
        end: int = 100
    ):

        tc = cell._tc

        tcPr = tc.get_or_add_tcPr()

        tcMar = tcPr.first_child_found_in(
            "w:tcMar"
        )

        if tcMar is None:

            tcMar = OxmlElement("w:tcMar")

            tcPr.append(tcMar)

        for margin, value in (
            ("top", top),
            ("start", start),
            ("bottom", bottom),
            ("end", end)
        ):

            node = tcMar.find(
                qn(f"w:{margin}")
            )

            if node is None:

                node = OxmlElement(
                    f"w:{margin}"
                )

                tcMar.append(node)

            node.set(
                qn("w:w"),
                str(value)
            )

            node.set(
                qn("w:type"),
                "dxa"
            )

    # --------------------------------------------------------

    @staticmethod
    def set_cell_border(
        cell,
        color: str = "808080",
        size: str = "4"
    ):

        tc = cell._tc

        tcPr = tc.get_or_add_tcPr()

        tcBorders = tcPr.first_child_found_in(
            "w:tcBorders"
        )

        if tcBorders is None:

            tcBorders = OxmlElement(
                "w:tcBorders"
            )

            tcPr.append(tcBorders)

        for edge in (
            "top",
            "left",
            "bottom",
            "right",
            "insideH",
            "insideV"
        ):

            tag = f"w:{edge}"

            element = tcBorders.find(
                qn(tag)
            )

            if element is None:

                element = OxmlElement(tag)

                tcBorders.append(element)

            element.set(
                qn("w:val"),
                "single"
            )

            element.set(
                qn("w:sz"),
                size
            )

            element.set(
                qn("w:space"),
                "0"
            )

            element.set(
                qn("w:color"),
                color
            )

    # --------------------------------------------------------

    @staticmethod
    def set_table_layout_fixed(
        table
    ):

        tblPr = table._tbl.tblPr

        tblLayout = tblPr.find(
            qn("w:tblLayout")
        )

        if tblLayout is None:

            tblLayout = OxmlElement(
                "w:tblLayout"
            )

            tblPr.append(tblLayout)

        tblLayout.set(
            qn("w:type"),
            "fixed"
        )

    # --------------------------------------------------------

    @staticmethod
    def set_table_width(
        table,
        width_twips: int
    ):

        tblPr = table._tbl.tblPr

        tblW = tblPr.find(
            qn("w:tblW")
        )

        if tblW is None:

            tblW = OxmlElement(
                "w:tblW"
            )

            tblPr.append(tblW)

        tblW.set(
            qn("w:w"),
            str(width_twips)
        )

        tblW.set(
            qn("w:type"),
            "dxa"
        )

    # --------------------------------------------------------

    @staticmethod
    def repeat_table_header(
        row
    ):

        trPr = row._tr.get_or_add_trPr()

        tblHeader = OxmlElement(
            "w:tblHeader"
        )

        tblHeader.set(
            qn("w:val"),
            "true"
        )

        trPr.append(tblHeader)

    # --------------------------------------------------------

    @staticmethod
    def prevent_row_split(
        row
    ):

        trPr = row._tr.get_or_add_trPr()

        cantSplit = OxmlElement(
            "w:cantSplit"
        )

        trPr.append(cantSplit)

    # --------------------------------------------------------

    @staticmethod
    def set_keep_with_next(
        paragraph,
        value: bool = True
    ):

        paragraph.paragraph_format.keep_with_next = value

    # --------------------------------------------------------

    @staticmethod
    def clear_paragraph(
        paragraph
    ):

        p = paragraph._element

        for child in list(p):

            if child.tag != qn("w:pPr"):

                p.remove(child)


# ============================================================
# DOCUMENT STYLE
# ============================================================

class DocumentStyleManager:

    @classmethod
    def setup_document(
        cls,
        doc: Document,
        preserve_template_margins: bool = False
    ):

        cls.setup_page(
            doc,
            preserve_template_margins
        )

        cls.setup_styles(
            doc
        )

    # --------------------------------------------------------

    @classmethod
    def setup_page(
        cls,
        doc: Document,
        preserve_template_margins: bool = False
    ):

        for section in doc.sections:

            # Khi template có định dạng sẵn,
            # vẫn có thể giữ nguyên nếu yêu cầu.
            if preserve_template_margins:

                continue

            section.top_margin = (
                WordExportConfig.TOP_MARGIN
            )

            section.bottom_margin = (
                WordExportConfig.BOTTOM_MARGIN
            )

            section.left_margin = (
                WordExportConfig.LEFT_MARGIN
            )

            section.right_margin = (
                WordExportConfig.RIGHT_MARGIN
            )

    # --------------------------------------------------------

    @classmethod
    def setup_styles(
        cls,
        doc: Document
    ):

        style_names = [
            "Normal",
            "Body Text",
            "Heading 1",
            "Heading 2",
            "Heading 3",
            "List Bullet",
            "List Number"
        ]

        for name in style_names:

            try:

                style = doc.styles[name]

            except KeyError:

                continue

            style.font.name = (
                WordExportConfig.FONT_NAME
            )

            style.font.size = Pt(
                WordExportConfig.BODY_FONT_SIZE
            )

            style.font.color.rgb = (
                RGBColor(0, 0, 0)
            )

            if name == "Normal":

                style.paragraph_format.space_after = Pt(0)

                style.paragraph_format.line_spacing = 1.0

            elif name == "Heading 1":

                style.font.size = Pt(
                    WordExportConfig.HEADING_1_SIZE
                )

                style.font.bold = True

                style.paragraph_format.keep_with_next = True

            elif name == "Heading 2":

                style.font.size = Pt(
                    WordExportConfig.HEADING_2_SIZE
                )

                style.font.bold = True

                style.paragraph_format.keep_with_next = True

            elif name == "Heading 3":

                style.font.size = Pt(
                    WordExportConfig.HEADING_3_SIZE
                )

                style.font.bold = True

                style.paragraph_format.keep_with_next = True


# ============================================================
# INLINE RENDERER
# ============================================================

class InlineRenderer:

    @classmethod
    def render(
        cls,
        paragraph,
        tokens: List[Dict[str, Any]]
    ):

        if not tokens:

            return

        for token in tokens:

            token_type = token.get(
                "type",
                "text"
            )

            content = token.get(
                "content",
                ""
            )

            if content is None:

                content = ""

            content = str(content)

            # ------------------------------------------------
            # INLINE MATH
            # ------------------------------------------------

            if token_type in (
                "inline_math",
                "math"
            ):

                cls.render_math(
                    paragraph,
                    content
                )

                continue

            # ------------------------------------------------
            # TEXT
            # ------------------------------------------------

            run = paragraph.add_run(
                content
            )

            WordXml.set_font(
                run
            )

            if token_type == "bold":

                run.bold = True

            elif token_type == "italic":

                run.italic = True

            elif token_type == "underline":

                run.underline = True

            elif token_type == "strike":

                run.font.strike = True

            elif token_type == "subscript":

                run.font.subscript = True

            elif token_type == "superscript":

                run.font.superscript = True

            elif token_type == "inline_code":

                WordXml.set_font(
                    run,
                    WordExportConfig.CODE_FONT_NAME
                )

                run.font.size = Pt(
                    WordExportConfig.CODE_FONT_SIZE
                )

            elif token_type == "highlight":

                try:

                    run.font.highlight_color = 7

                except Exception:

                    pass

            elif token_type == "link":

                run.underline = True

                run.font.color.rgb = (
                    RGBColor(0, 0, 255)
                )

    # --------------------------------------------------------

    @classmethod
    def render_math(
        cls,
        paragraph,
        latex: str
    ):

        normalized = (
            ScienceNormalizer.normalize(
                latex
            )
        )

        if not normalized:

            return

        run = paragraph.add_run(
            normalized
        )

        WordXml.set_font(
            run,
            WordExportConfig.FONT_NAME
        )

        run.font.size = Pt(
            WordExportConfig.MATH_FONT_SIZE
        )

        run.italic = True


# ============================================================
# TABLE RENDERER
# ============================================================

class TableRenderer:

    @classmethod
    def render(
        cls,
        doc: Document,
        node: Dict[str, Any]
    ):

        headers = node.get(
            "headers",
            []
        )

        rows = node.get(
            "rows",
            []
        )

        cols = node.get(
            "cols",
            0
        )

        if not headers and not rows:

            return None

        if not cols:

            cols = max(
                [
                    len(headers),
                    *[
                        len(row)
                        for row in rows
                    ]
                ]
            )

        if cols <= 0:

            return None

        table = doc.add_table(
            rows=0,
            cols=cols
        )

        table.alignment = (
            WD_TABLE_ALIGNMENT.CENTER
        )

        table.autofit = False

        WordXml.set_table_layout_fixed(
            table
        )

        total_width_twips = int(
            Cm(
                WordExportConfig.MAX_TABLE_WIDTH_CM
            ).twips
        )

        WordXml.set_table_width(
            table,
            total_width_twips
        )

        # ----------------------------------------------------
        # HEADER
        # ----------------------------------------------------

        if headers:

            header_row = table.add_row()

            WordXml.repeat_table_header(
                header_row
            )

            WordXml.prevent_row_split(
                header_row
            )

            for index in range(cols):

                cell = header_row.cells[index]

                WordXml.set_cell_shading(
                    cell,
                    WordExportConfig.TABLE_HEADER_COLOR
                )

                WordXml.set_cell_border(
                    cell,
                    WordExportConfig.TABLE_BORDER_COLOR
                )

                WordXml.set_cell_margins(
                    cell
                )

                cell.vertical_alignment = (
                    WD_CELL_VERTICAL_ALIGNMENT.CENTER
                )

                paragraph = cell.paragraphs[0]

                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                )

                paragraph.paragraph_format.space_after = Pt(0)

                if index < len(headers):

                    InlineRenderer.render(
                        paragraph,
                        headers[index].get(
                            "content",
                            []
                        )
                    )

                for run in paragraph.runs:

                    run.bold = True

        # ----------------------------------------------------
        # BODY
        # ----------------------------------------------------

        for row_index, row_data in enumerate(rows):

            row = table.add_row()

            WordXml.prevent_row_split(
                row
            )

            for col_index in range(cols):

                cell = row.cells[col_index]

                WordXml.set_cell_border(
                    cell,
                    WordExportConfig.TABLE_BORDER_COLOR
                )

                WordXml.set_cell_margins(
                    cell
                )

                cell.vertical_alignment = (
                    WD_CELL_VERTICAL_ALIGNMENT.CENTER
                )

                paragraph = cell.paragraphs[0]

                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                )

                paragraph.paragraph_format.space_after = Pt(0)

                if col_index < len(row_data):

                    InlineRenderer.render(
                        paragraph,
                        row_data[col_index].get(
                            "content",
                            []
                        )
                    )

        # ----------------------------------------------------
        # WIDTH
        # ----------------------------------------------------

        width_each = (
            total_width_twips // cols
        )

        for row in table.rows:

            for cell in row.cells:

                tcPr = cell._tc.get_or_add_tcPr()

                tcW = tcPr.find(
                    qn("w:tcW")
                )

                if tcW is None:

                    tcW = OxmlElement(
                        "w:tcW"
                    )

                    tcPr.append(tcW)

                tcW.set(
                    qn("w:w"),
                    str(width_each)
                )

                tcW.set(
                    qn("w:type"),
                    "dxa"
                )

        return table


# ============================================================
# BLOCK RENDERER
# ============================================================

class BlockRenderer:

    @classmethod
    def render_node(
        cls,
        doc: Document,
        node: Dict[str, Any]
    ):

        node_type = node.get(
            "type",
            "paragraph"
        )

        # ----------------------------------------------------
        # PARAGRAPH
        # ----------------------------------------------------

        if node_type == "paragraph":

            paragraph = doc.add_paragraph()

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.JUSTIFY
            )

            paragraph.paragraph_format.space_after = Pt(0)

            InlineRenderer.render(
                paragraph,
                node.get(
                    "tokens",
                    []
                )
            )

            return paragraph

        # ----------------------------------------------------
        # HEADING
        # ----------------------------------------------------

        if node_type == "heading":

            level = min(
                max(
                    int(
                        node.get(
                            "level",
                            1
                        )
                    ),
                    1
                ),
                3
            )

            paragraph = doc.add_paragraph()

            paragraph.paragraph_format.keep_with_next = True

            paragraph.paragraph_format.space_before = Pt(4)

            paragraph.paragraph_format.space_after = Pt(2)

            if level == 1:

                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                )

            else:

                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                )

            InlineRenderer.render(
                paragraph,
                node.get(
                    "tokens",
                    []
                )
            )

            for run in paragraph.runs:

                run.bold = True

                if level == 1:

                    run.font.size = Pt(
                        WordExportConfig.HEADING_1_SIZE
                    )

                elif level == 2:

                    run.font.size = Pt(
                        WordExportConfig.HEADING_2_SIZE
                    )

                else:

                    run.font.size = Pt(
                        WordExportConfig.HEADING_3_SIZE
                    )

            return paragraph

        # ----------------------------------------------------
        # LIST ITEM
        # ----------------------------------------------------

        if node_type == "list_item":

            paragraph = doc.add_paragraph()

            level = int(
                node.get(
                    "level",
                    1
                )
            )

            indent = Cm(
                min(
                    0.5 * level,
                    2.5
                )
            )

            paragraph.paragraph_format.left_indent = indent

            paragraph.paragraph_format.first_line_indent = Cm(
                -0.35
            )

            paragraph.paragraph_format.alignment = (
                WD_ALIGN_PARAGRAPH.JUSTIFY
            )

            style = node.get(
                "style",
                "bullet"
            )

            prefix = (
                "• "
                if style == "bullet"
                else "1. "
            )

            paragraph.add_run(
                prefix
            )

            InlineRenderer.render(
                paragraph,
                node.get(
                    "tokens",
                    []
                )
            )

            return paragraph

        # ----------------------------------------------------
        # CHECKBOX
        # ----------------------------------------------------

        if node_type == "checkbox":

            paragraph = doc.add_paragraph()

            level = int(
                node.get(
                    "level",
                    1
                )
            )

            paragraph.paragraph_format.left_indent = Cm(
                min(
                    0.5 * level,
                    2.5
                )
            )

            symbol = (
                "☑ "
                if node.get(
                    "checked",
                    False
                )
                else "☐ "
            )

            run = paragraph.add_run(
                symbol
            )

            WordXml.set_font(
                run,
                "Segoe UI Symbol"
            )

            InlineRenderer.render(
                paragraph,
                node.get(
                    "tokens",
                    []
                )
            )

            return paragraph

        # ----------------------------------------------------
        # TABLE
        # ----------------------------------------------------

        if node_type == "table":

            return TableRenderer.render(
                doc,
                node
            )

        # ----------------------------------------------------
        # IMAGE
        # ----------------------------------------------------

        if node_type == "image":

            ImageRenderer.render_image(
                doc,
                node
            )

            return None

        # ----------------------------------------------------
        # CODE
        # ----------------------------------------------------

        if node_type == "code":

            paragraph = doc.add_paragraph()

            paragraph.paragraph_format.left_indent = Cm(
                0.5
            )

            paragraph.paragraph_format.space_before = Pt(3)

            paragraph.paragraph_format.space_after = Pt(3)

            WordXml.set_paragraph_shading(
                paragraph,
                WordExportConfig.CODE_BACKGROUND
            )

            run = paragraph.add_run(
                node.get(
                    "text",
                    ""
                )
            )

            WordXml.set_font(
                run,
                WordExportConfig.CODE_FONT_NAME
            )

            run.font.size = Pt(
                WordExportConfig.CODE_FONT_SIZE
            )

            return paragraph

        # ----------------------------------------------------
        # HORIZONTAL RULE
        # ----------------------------------------------------

        if node_type == "hr":

            paragraph = doc.add_paragraph()

            WordXml.set_paragraph_bottom_border(
                paragraph
            )

            return paragraph

        # ----------------------------------------------------
        # CALLOUT
        # ----------------------------------------------------

        if node_type == "callout":

            table = doc.add_table(
                rows=1,
                cols=1
            )

            table.alignment = (
                WD_TABLE_ALIGNMENT.CENTER
            )

            table.autofit = False

            WordXml.set_table_layout_fixed(
                table
            )

            cell = table.cell(
                0,
                0
            )

            WordXml.set_cell_shading(
                cell,
                WordExportConfig.CALLOUT_BACKGROUND
            )

            WordXml.set_cell_border(
                cell
            )

            WordXml.set_cell_margins(
                cell,
                top=120,
                start=160,
                bottom=120,
                end=160
            )

            children = node.get(
                "children",
                []
            )

            for index, child in enumerate(children):

                if index == 0:

                    paragraph = cell.paragraphs[0]

                else:

                    paragraph = cell.add_paragraph()

                cls.render_node_into_paragraph(
                    paragraph,
                    child
                )

            return table

        return None

    # --------------------------------------------------------

    @classmethod
    def render_node_into_paragraph(
        cls,
        paragraph,
        node
    ):

        node_type = node.get(
            "type",
            "paragraph"
        )

        WordXml.clear_paragraph(
            paragraph
        )

        if node_type == "paragraph":

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.JUSTIFY
            )

            InlineRenderer.render(
                paragraph,
                node.get(
                    "tokens",
                    []
                )
            )

            return paragraph

        if node_type == "heading":

            level = min(
                max(
                    int(
                        node.get(
                            "level",
                            1
                        )
                    ),
                    1
                ),
                3
            )

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if level == 1
                else WD_ALIGN_PARAGRAPH.LEFT
            )

            paragraph.paragraph_format.keep_with_next = True

            InlineRenderer.render(
                paragraph,
                node.get(
                    "tokens",
                    []
                )
            )

            for run in paragraph.runs:

                run.bold = True

                run.font.size = Pt(
                    WordExportConfig.HEADING_1_SIZE
                    if level == 1
                    else (
                        WordExportConfig.HEADING_2_SIZE
                        if level == 2
                        else WordExportConfig.HEADING_3_SIZE
                    )
                )

            return paragraph

        if node_type == "list_item":

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.JUSTIFY
            )

            level = int(
                node.get(
                    "level",
                    1
                )
            )

            paragraph.paragraph_format.left_indent = Cm(
                min(
                    0.5 * level,
                    2.5
                )
            )

            paragraph.paragraph_format.first_line_indent = Cm(
                -0.35
            )

            prefix = (
                "• "
                if node.get(
                    "style"
                ) != "number"
                else "1. "
            )

            paragraph.add_run(
                prefix
            )

            InlineRenderer.render(
                paragraph,
                node.get(
                    "tokens",
                    []
                )
            )

            return paragraph

        return paragraph


# ============================================================
# XML EXTENSIONS
# ============================================================

def _set_paragraph_shading(
    paragraph,
    color_hex: str
):

    pPr = paragraph._element.get_or_add_pPr()

    shd = pPr.find(
        qn("w:shd")
    )

    if shd is None:

        shd = OxmlElement(
            "w:shd"
        )

        pPr.append(shd)

    shd.set(
        qn("w:val"),
        "clear"
    )

    shd.set(
        qn("w:fill"),
        color_hex
    )


def _set_paragraph_bottom_border(
    paragraph,
    color_hex: str = "A0A0A0"
):

    pPr = paragraph._element.get_or_add_pPr()

    pBdr = pPr.find(
        qn("w:pBdr")
    )

    if pBdr is None:

        pBdr = OxmlElement(
            "w:pBdr"
        )

        pPr.append(pBdr)

    bottom = OxmlElement(
        "w:bottom"
    )

    bottom.set(
        qn("w:val"),
        "single"
    )

    bottom.set(
        qn("w:sz"),
        "6"
    )

    bottom.set(
        qn("w:space"),
        "4"
    )

    bottom.set(
        qn("w:color"),
        color_hex
    )

    pBdr.append(
        bottom
    )


WordXml.set_paragraph_shading = staticmethod(
    _set_paragraph_shading
)

WordXml.set_paragraph_bottom_border = staticmethod(
    _set_paragraph_bottom_border
)


# ============================================================
# TEMPLATE NAVIGATION
# ============================================================

class TemplateNavigator:

    PLACEHOLDER_RE = re.compile(
        r"\{\{\s*([A-Za-z0-9_]+)\s*\}\}"
    )

    # --------------------------------------------------------

    @classmethod
    def iter_paragraphs(
        cls,
        doc: Document
    ):

        for paragraph in doc.paragraphs:

            yield paragraph

        for table in doc.tables:

            yield from cls.iter_table_paragraphs(
                table
            )

        for section in doc.sections:

            for paragraph in section.header.paragraphs:

                yield paragraph

            for paragraph in section.footer.paragraphs:

                yield paragraph

    # --------------------------------------------------------

    @classmethod
    def iter_table_paragraphs(
        cls,
        table: Table
    ):

        for row in table.rows:

            for cell in row.cells:

                for paragraph in cell.paragraphs:

                    yield paragraph

                for nested_table in cell.tables:

                    yield from cls.iter_table_paragraphs(
                        nested_table
                    )

    # --------------------------------------------------------

    @classmethod
    def extract_placeholder(
        cls,
        paragraph
    ) -> Optional[str]:

        text = "".join(
            run.text or ""
            for run in paragraph.runs
        )

        if not text:

            return None

        match = cls.PLACEHOLDER_RE.search(
            text
        )

        if not match:

            return None

        return match.group(1)

    # --------------------------------------------------------

    @classmethod
    def find_placeholder_paragraphs(
        cls,
        doc: Document
    ):

        result = []

        for paragraph in cls.iter_paragraphs(
            doc
        ):

            key = cls.extract_placeholder(
                paragraph
            )

            if key:

                result.append(
                    (
                        paragraph,
                        key
                    )
                )

        return result


# ============================================================
# CONTENT NORMALIZER
# ============================================================

class ContentNormalizer:

    @staticmethod
    def value_to_markdown(
        value: Any
    ) -> str:

        if value is None:

            return ""

        if isinstance(
            value,
            list
        ):

            return "\n".join(
                str(item)
                for item in value
            )

        if isinstance(
            value,
            dict
        ):

            return "\n".join(
                f"{key}: {val}"
                for key, val in value.items()
            )

        return str(value).strip()

    # --------------------------------------------------------

    @staticmethod
    def remove_ai_preamble(
        text: str
    ) -> str:

        if not text:

            return ""

        forbidden_prefixes = (
            "Chào bạn",
            "Với vai trò",
            "Tôi là",
            "Lưu ý về"
        )

        lines = []

        for line in text.splitlines():

            stripped = line.strip()

            if any(
                stripped.startswith(
                    prefix
                )
                for prefix in forbidden_prefixes
            ):

                continue

            lines.append(
                line
            )

        return "\n".join(
            lines
        ).strip()


# ============================================================
# TEMPLATE CONTENT INJECTION
# ============================================================

class TemplateContentInjector:

    @classmethod
    def inject(
        cls,
        doc: Document,
        variables: Dict[str, Any]
    ):

        if not variables:

            return doc

        placeholder_paragraphs = (
            TemplateNavigator.find_placeholder_paragraphs(
                doc
            )
        )

        for paragraph, key in placeholder_paragraphs:

            if key not in variables:

                continue

            raw_value = variables.get(
                key
            )

            markdown = (
                ContentNormalizer.value_to_markdown(
                    raw_value
                )
            )

            markdown = (
                ContentNormalizer.remove_ai_preamble(
                    markdown
                )
            )

            if not markdown:

                WordXml.clear_paragraph(
                    paragraph
                )

                continue

            nodes = MarkdownTokenizer.parse(
                markdown
            )

            cls.replace_paragraph_with_nodes(
                paragraph,
                nodes
            )

        return doc

    # --------------------------------------------------------

    @classmethod
    def replace_paragraph_with_nodes(
        cls,
        paragraph,
        nodes: List[Dict[str, Any]]
    ):

        if not nodes:

            WordXml.clear_paragraph(
                paragraph
            )

            return

        # Render node đầu tiên vào paragraph
        first_node = nodes[0]

        BlockRenderer.render_node_into_paragraph(
            paragraph,
            first_node
        )

        # Các node tiếp theo được chèn ngay sau
        current_element = paragraph._element

        for node in nodes[1:]:

            new_paragraph = cls.insert_paragraph_after(
                paragraph
            )

            if node.get(
                "type"
            ) == "table":

                # Table không thể render trực tiếp
                # vào paragraph.
                #
                # Tạo một paragraph rỗng trước,
                # sau đó chèn table sau paragraph.
                table = TableRenderer.render(
                    paragraph.part.document,
                    node
                )

                if table is not None:

                    paragraph._element.addnext(
                        table._element
                    )

                paragraph = new_paragraph

            else:

                BlockRenderer.render_node_into_paragraph(
                    new_paragraph,
                    node
                )

                paragraph = new_paragraph

    # --------------------------------------------------------

    @staticmethod
    def insert_paragraph_after(
        paragraph
    ):

        new_p = OxmlElement(
            "w:p"
        )

        paragraph._p.addnext(
            new_p
        )

        new_paragraph = Paragraph(
            new_p,
            paragraph._parent
        )

        return new_paragraph


# ============================================================
# MAIN ENGINE
# ============================================================

class WordExportEngine:

    # --------------------------------------------------------

    @classmethod
    def _resolve_template_path(
        cls,
        template_path: Optional[str] = None
    ) -> Optional[str]:

        if template_path:

            return template_path

        candidates = [

            Path(
                "templates"
            ) / "khbd_mau.docx",

            Path(
                "template"
            ) / "khbd_mau.docx",

            Path(
                "khbd_mau.docx"
            )
        ]

        for candidate in candidates:

            if candidate.exists():

                return str(
                    candidate
                )

        return None

    # --------------------------------------------------------

    @classmethod
    def load_template(
        cls,
        template_path: Optional[str] = None
    ) -> Document:

        resolved = cls._resolve_template_path(
            template_path
        )

        if resolved:

            logger.info(
                "Đang sử dụng template: %s",
                resolved
            )

            return TemplateLoader.load(
                resolved
            )

        logger.warning(
            "Không tìm thấy template. "
            "Khởi tạo DOCX mới."
        )

        return Document()

    # --------------------------------------------------------

    @classmethod
    def convert_markdown_to_docx_bytes(
        cls,
        markdown_text: str,
        template_path: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None,
        preserve_template_margins: bool = False
    ) -> bytes:

        metadata = metadata or {}

        # ----------------------------------------------------
        # TRƯỜNG HỢP 1:
        # Có template và metadata variables
        # ----------------------------------------------------

        resolved_template = (
            cls._resolve_template_path(
                template_path
            )
        )

        if resolved_template and metadata:

            doc = cls.load_template(
                resolved_template
            )

            DocumentStyleManager.setup_document(
                doc,
                preserve_template_margins
            )

            variables = dict(
                metadata
            )

            # Nếu toàn bộ nội dung AI được lưu
            # trong ai_generated_content,
            # dùng làm fallback cho các placeholder
            #
            # Tuy nhiên KHÔNG đổ toàn bộ nội dung
            # vào template vì sẽ làm mất cấu trúc.
            #
            TemplateContentInjector.inject(
                doc,
                variables
            )

            bio = io.BytesIO()

            doc.save(
                bio
            )

            return bio.getvalue()

        # ----------------------------------------------------
        # TRƯỜNG HỢP 2:
        # Không có metadata:
        # Render Markdown tự do
        # ----------------------------------------------------

        doc = Document()

        DocumentStyleManager.setup_document(
            doc,
            preserve_template_margins=False
        )

        nodes = MarkdownTokenizer.parse(
            markdown_text or ""
        )

        for node in nodes:

            BlockRenderer.render_node(
                doc,
                node
            )

        bio = io.BytesIO()

        doc.save(
            bio
        )

        return bio.getvalue()

    # --------------------------------------------------------

    @classmethod
    def export_to_word(
        cls,
        data_cache: Dict[str, Any]
    ) -> bytes:

        metadata = dict(
            data_cache or {}
        )

        markdown_content = (
            metadata.get(
                "ai_generated_content",
                ""
            )
        )

        template_path = (
            metadata.get(
                "template_path"
            )
        )

        return cls.convert_markdown_to_docx_bytes(
            markdown_text=markdown_content,
            template_path=template_path,
            metadata=metadata
        )


# ============================================================
# PUBLIC API
# ============================================================

def export_word(
    markdown_text: str,
    template_path: Optional[str] = None,
    metadata: Optional[Dict[str, Any]] = None
) -> bytes:

    return WordExportEngine.convert_markdown_to_docx_bytes(
        markdown_text=markdown_text,
        template_path=template_path,
        metadata=metadata
    )
