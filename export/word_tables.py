# -*- coding: utf-8 -*-
"""
============================================================
MODULE: export/word_tables.py
Nhiệm vụ: Phân tích bảng Markdown, khởi tạo bảng Word Native,
xử lý nội dung từng ô (Văn bản, Toán OMML, Hình ảnh, Markdown).
TUYỆT ĐỐI KHÔNG dùng cell.text để bảo toàn định dạng và AST.
============================================================
"""

import re
import logging
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# Tích hợp an toàn với các API nội bộ
try:
    from .word_math import insert_math_to_paragraph
except ImportError:
    insert_math_to_paragraph = None

try:
    from .word_images import insert_image_to_paragraph
except ImportError:
    insert_image_to_paragraph = None


def _set_cell_shading(cell, fill_hex: str = "F2F2F2"):
    """Đổ màu nền cho ô (thường dùng cho tiêu đề bảng)."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def _draw_table_borders(table):
    """Vẽ viền bảng đen chuẩn bằng XML để tương thích mọi Template."""
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4') 
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000') 
            tblBorders.append(border)
        tblPr[0].append(tblBorders)


def _parse_and_fill_cell(cell, cell_text: str, is_header: bool = False, metadata: dict = None):
    """
    Trình phân tích nội dung ô (Inline Cell Parser).
    Nhận diện Văn bản, Toán học OMML, Hình ảnh và định dạng Markdown.
    """
    if not cell.paragraphs:
        p = cell.add_paragraph()
    else:
        p = cell.paragraphs[0]
        
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(4)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # Chuyển đổi ngắt dòng HTML/Markdown thành ngắt dòng thực
    cell_text = cell_text.replace('<br>', '\n').replace('<br/>', '\n')

    # Regex siêu nhạy: Bắt Toán, Hình ảnh, Bold, Italic
    pattern = re.compile(
        r'(\$\$(.*?)\$\$)|'                       # 1, 2: Block math (Ép thành inline trong bảng)
        r'(\$([^$]+?)\$)|'                        # 3, 4: Inline math
        r'(\[IMAGE\s*(?:-\s*ID:\s*)?([^\]]+)\])|' # 5, 6: Image tag (Sinh từ Kỷ luật thép)
        r'(!\[.*?\]\((.*?)\))|'                   # 7, 8: Markdown image
        r'(\*\*([^*]+?)\*\*)|'                    # 9, 10: Bold
        r'(\*([^*]+?)\*)'                         # 11, 12: Italic
    )
    
    last_idx = 0
    for match in pattern.finditer(cell_text):
        # 1. Ghi văn bản thường đứng trước
        text_before = cell_text[last_idx:match.start()]
        if text_before:
            r = p.add_run(text_before.replace('\\|', '|'))
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            if is_header: r.bold = True
            
        # 2. Xử lý phần tử đặc biệt
        if match.group(1) or match.group(3): 
            # CÔNG THỨC TOÁN HỌC (Gọi OMML)
            math_content = match.group(2) or match.group(4)
            if insert_math_to_paragraph:
                # Trong bảng, luôn ép thành is_block=False để không bị ngắt dòng vô duyên
                insert_math_to_paragraph(p, math_content, is_block=False)
            else:
                r = p.add_run(math_content)
                r.font.name = 'Cambria Math'
                r.italic = True
                
        elif match.group(5) or match.group(7): 
            # HÌNH ẢNH TRONG Ô BẢNG
            img_id_or_url = match.group(6) or match.group(8)
            img_src = None
            
            # Tìm dữ liệu Base64 từ Metadata nếu có
            if metadata and "pages" in metadata:
                for page in metadata["pages"]:
                    for img in page.get("images", []):
                        if img.get("id") == img_id_or_url:
                            img_src = {"base64": img.get("base64"), "caption": img_id_or_url}
                            break
                    if img_src: break
            
            if insert_image_to_paragraph:
                # Giới hạn kích thước ảnh trong bảng (Inches 1.8) để chống vỡ cột
                insert_image_to_paragraph(p, img_src if img_src else img_id_or_url, max_width=Inches(1.8))
            else:
                r = p.add_run(f"[Hình ảnh: {img_id_or_url}]")
                r.font.name = 'Times New Roman'
                r.italic = True
                
        elif match.group(9): 
            # IN ĐẬM
            r = p.add_run(match.group(10).replace('\\|', '|'))
            r.font.name, r.font.size, r.bold = 'Times New Roman', Pt(12), True
            
        elif match.group(11): 
            # IN NGHIÊNG
            r = p.add_run(match.group(12).replace('\\|', '|'))
            r.font.name, r.font.size, r.italic = 'Times New Roman', Pt(12), True
            if is_header: r.bold = True
            
        last_idx = match.end()
        
    # 3. Ghi phần văn bản còn sót lại
    if last_idx < len(cell_text):
        text_after = cell_text[last_idx:]
        r = p.add_run(text_after.replace('\\|', '|'))
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        if is_header: r.bold = True


def process_and_draw_markdown_table(doc: Document, table_lines: list, metadata: dict = None):
    """
    API Công khai: Chuyển đổi danh sách dòng Markdown table thành bảng Word Native.
    Bảo toàn cấu trúc, viền bảng, tự động co giãn và định dạng chi tiết từng ô.
    """
    if not table_lines or len(table_lines) < 2:
        return
        
    try:
        # 1. Phân tích an toàn hàng/cột (Không vỡ bảng khi gặp công thức chứa dấu | như $|x|$)
        parsed_rows = []
        for line in table_lines:
            line = line.strip()
            if not line: continue
            
            # Xóa ống | bao ngoài viền
            if line.startswith('|'): line = line[1:]
            if line.endswith('|'): line = line[:-1]
            
            # Cắt cột an toàn bằng regex Negative Lookbehind (bỏ qua \|)
            cells = re.split(r'(?<!\\)\|', line)
            cells = [c.strip() for c in cells]
            parsed_rows.append(cells)
            
        # 2. Lọc bỏ dòng phân cách Markdown (vd: |---|---|)
        filtered_rows = []
        for row in parsed_rows:
            if row and all(re.match(r'^[\-\:]+$', c.replace(' ', '')) for c in row if c):
                continue
            filtered_rows.append(row)
            
        if not filtered_rows:
            return
            
        # 3. Chuẩn hóa ma trận bảng (Bù đắp ô trống nếu cột lệch)
        num_rows = len(filtered_rows)
        num_cols = max(len(r) for r in filtered_rows)
        
        for r in filtered_rows:
            while len(r) < num_cols:
                r.append("")
                
        # 4. Khởi tạo Bảng Word Native
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        
        _draw_table_borders(table)

        # 5. Phân tích và render nội dung từng ô thông minh
        for r_idx, row_data in enumerate(filtered_rows):
            row_cells = table.rows[r_idx].cells
            is_header = (r_idx == 0)
            
            for c_idx, cell_value in enumerate(row_data):
                cell = row_cells[c_idx]
                if is_header:
                    _set_cell_shading(cell, fill_hex='F2F2F2') # Xám nhạt cho Header
                
                # Gọi bộ xử lý nội dung đa thành phần
                _parse_and_fill_cell(cell, cell_value, is_header=is_header, metadata=metadata)

        # Cách 1 khoảng paragraph dưới bảng để không bị dính văn bản
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        
    except Exception as e:
        logger.error(f"Lỗi vẽ bảng Markdown: {e}")
        # Fallback an toàn: Trả văn bản thô lại vào Word nếu có lỗi ngoại lệ
        for line in table_lines:
            doc.add_paragraph(line)
