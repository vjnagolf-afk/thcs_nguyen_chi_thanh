# -*- coding: utf-8 -*-
"""
============================================================
MODULE: export/word_math.py

BỘ BIÊN DỊCH CÔNG THỨC KHOA HỌC SANG OMML NATIVE CỦA MICROSOFT WORD

Mục tiêu:
- Toán học: phân số, căn, căn bậc n, số mũ, chỉ số,
  chỉ số + số mũ, ký hiệu Hy Lạp, toán tử, ngoặc.
- Vật lý: công thức và đơn vị vật lý.
- Hóa học: công thức hóa học, điện tích ion, đồng vị,
  chỉ số nguyên tử, phân tử nước kết tinh.
- Chèn công thức trực tiếp vào Word dưới dạng OMML Native.
- Không để lọt các chuỗi $...$, \( ... \), \[ ... \] ra văn bản.
- Không dùng Unicode superscript/subscript làm phương án chính.
- Tương thích với python-docx.

============================================================
"""

from __future__ import annotations

import html
import logging
import re
from typing import List, Optional, Tuple

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn


logger = logging.getLogger("WordMath")


MATH_NS = (
    "http://schemas.openxmlformats.org/"
    "officeDocument/2006/math"
)


# ============================================================
# 1. BẢNG KÝ HIỆU LATEX → UNICODE
# ============================================================

SYMBOLS = {
    # Greek
    "alpha": "α",
    "beta": "β",
    "gamma": "γ",
    "delta": "δ",
    "epsilon": "ε",
    "varepsilon": "ε",
    "zeta": "ζ",
    "eta": "η",
    "theta": "θ",
    "vartheta": "ϑ",
    "iota": "ι",
    "kappa": "κ",
    "lambda": "λ",
    "mu": "μ",
    "nu": "ν",
    "xi": "ξ",
    "omicron": "ο",
    "pi": "π",
    "varpi": "ϖ",
    "rho": "ρ",
    "varrho": "ϱ",
    "sigma": "σ",
    "varsigma": "ς",
    "tau": "τ",
    "upsilon": "υ",
    "phi": "φ",
    "varphi": "ϕ",
    "chi": "χ",
    "psi": "ψ",
    "omega": "ω",

    # Uppercase Greek
    "Gamma": "Γ",
    "Delta": "Δ",
    "Theta": "Θ",
    "Lambda": "Λ",
    "Xi": "Ξ",
    "Pi": "Π",
    "Sigma": "Σ",
    "Upsilon": "Υ",
    "Phi": "Φ",
    "Psi": "Ψ",
    "Omega": "Ω",

    # Relations
    "le": "≤",
    "leq": "≤",
    "ge": "≥",
    "geq": "≥",
    "neq": "≠",
    "ne": "≠",
    "approx": "≈",
    "equiv": "≡",
    "sim": "∼",
    "simeq": "≃",
    "cong": "≅",
    "propto": "∝",

    # Operators
    "pm": "±",
    "mp": "∓",
    "times": "×",
    "cdot": "·",
    "div": "÷",
    "ast": "∗",
    "star": "⋆",
    "circ": "∘",
    "bullet": "•",

    # Sets / logic
    "in": "∈",
    "notin": "∉",
    "subset": "⊂",
    "subseteq": "⊆",
    "supset": "⊃",
    "supseteq": "⊇",
    "cup": "∪",
    "cap": "∩",
    "emptyset": "∅",
    "infty": "∞",
    "forall": "∀",
    "exists": "∃",
    "nexists": "∄",

    # Geometry
    "perp": "⊥",
    "parallel": "∥",
    "angle": "∠",
    "triangle": "△",
    "square": "□",

    # Arrows
    "rightarrow": "→",
    "to": "→",
    "leftarrow": "←",
    "leftrightarrow": "↔",
    "Rightarrow": "⇒",
    "Leftarrow": "⇐",
    "Leftrightarrow": "⇔",
    "uparrow": "↑",
    "downarrow": "↓",

    # Calculus / algebra
    "sum": "∑",
    "prod": "∏",
    "int": "∫",
    "oint": "∮",
    "partial": "∂",
    "nabla": "∇",

    # Text spacing
    "quad": " ",
    "qquad": "  ",
    "enspace": " ",
    "thinspace": " ",
    ",": " ",
    ";": " ",
    ":": " ",
    "!": "",

    # Named functions
    "sin": "sin",
    "cos": "cos",
    "tan": "tan",
    "cot": "cot",
    "sec": "sec",
    "csc": "csc",
    "arcsin": "arcsin",
    "arccos": "arccos",
    "arctan": "arctan",
    "log": "log",
    "ln": "ln",
    "lim": "lim",
    "max": "max",
    "min": "min",
}


# ============================================================
# 2. HÀM TIỆN ÍCH
# ============================================================

def escape_xml(text: str) -> str:
    """
    Escape XML an toàn cho OMML.
    """

    if text is None:
        return ""

    return html.escape(
        str(text),
        quote=True
    )


def _empty_math_run() -> str:
    return (
        "<m:r>"
        "<m:t xml:space=\"preserve\"></m:t>"
        "</m:r>"
    )


def _normalize_math_input(value: str) -> str:
    """
    Loại bỏ các lớp delimiter bên ngoài nhưng không phá
    các delimiter nằm bên trong biểu thức.
    """

    if not value:
        return ""

    s = str(value).strip()

    # $$ ... $$
    if len(s) >= 4 and s.startswith("$$") and s.endswith("$$"):
        return s[2:-2].strip()

    # $ ... $
    if len(s) >= 2 and s.startswith("$") and s.endswith("$"):
        return s[1:-1].strip()

    # \[ ... \]
    if len(s) >= 4 and s.startswith(r"\[") and s.endswith(r"\]"):
        return s[2:-2].strip()

    # \( ... \)
    if len(s) >= 4 and s.startswith(r"\(") and s.endswith(r"\)"):
        return s[2:-2].strip()

    return s


# ============================================================
# 3. LATEX PARSER
# ============================================================

class LatexParser:
    """
    Parser đệ quy đơn giản nhưng an toàn cho LaTeX giáo dục.

    AST node:

        ("text", "x")
        ("group", [...])
        ("frac", numerator, denominator)
        ("sqrt", expression)
        ("root", degree, expression)
        ("sup", base, exponent)
        ("sub", base, subscript)
        ("subsup", base, subscript, exponent)
        ("normal_text", [...])
        ("accent", accent, expression)
    """

    def __init__(self, source: str):
        self.source = source or ""
        self.pos = 0
        self.length = len(self.source)

    # --------------------------------------------------------
    # CORE
    # --------------------------------------------------------

    def peek(self) -> str:
        if self.pos >= self.length:
            return ""

        return self.source[self.pos]

    def get(self) -> str:
        char = self.peek()

        if char:
            self.pos += 1

        return char

    def parse(self) -> list:
        nodes = []

        while self.pos < self.length:

            char = self.peek()

            # ------------------------------------------------
            # COMMAND
            # ------------------------------------------------

            if char == "\\":
                command_node = self.parse_command()

                if command_node is not None:
                    nodes.append(command_node)

                continue

            # ------------------------------------------------
            # GROUP
            # ------------------------------------------------

            if char == "{":
                nodes.append(
                    (
                        "group",
                        self.parse_group_content()
                    )
                )
                continue

            # ------------------------------------------------
            # KẾT THÚC GROUP
            # ------------------------------------------------

            if char == "}":
                break

            # ------------------------------------------------
            # SUPERSCRIPT
            # ------------------------------------------------

            if char == "^":
                self.get()

                exponent = self.parse_argument()

                if nodes:
                    base = nodes.pop()
                    nodes.append(
                        (
                            "sup",
                            base,
                            exponent
                        )
                    )

                continue

            # ------------------------------------------------
            # SUBSCRIPT
            # ------------------------------------------------

            if char == "_":
                self.get()

                subscript = self.parse_argument()

                if nodes:

                    base = nodes.pop()

                    if base[0] == "sup":

                        nodes.append(
                            (
                                "subsup",
                                base[1],
                                subscript,
                                base[2]
                            )
                        )

                    else:

                        nodes.append(
                            (
                                "sub",
                                base,
                                subscript
                            )
                        )

                continue

            # ------------------------------------------------
            # NORMAL CHARACTER
            # ------------------------------------------------

            nodes.append(
                (
                    "text",
                    self.get()
                )
            )

        return self.combine_text_nodes(nodes)

    # --------------------------------------------------------
    # COMMAND
    # --------------------------------------------------------

    def parse_command(self) -> Optional[tuple]:

        self.get()

        if self.pos >= self.length:
            return ("text", "\\")

        match = re.match(
            r"[A-Za-z]+|.",
            self.source[self.pos:]
        )

        if not match:
            return ("text", "\\")

        command = match.group(0)
        self.pos += len(command)

        # ----------------------------------------------------
        # FRACTION
        # ----------------------------------------------------

        if command == "frac":

            numerator = self.parse_argument()
            denominator = self.parse_argument()

            return (
                "frac",
                numerator,
                denominator
            )

        # ----------------------------------------------------
        # SQUARE ROOT / NTH ROOT
        # ----------------------------------------------------

        if command == "sqrt":

            degree = None

            if self.peek() == "[":
                self.get()

                degree = self.parse_until("]")

            expression = self.parse_argument()

            if degree is not None:

                return (
                    "root",
                    degree,
                    expression
                )

            return (
                "sqrt",
                expression
            )

        # ----------------------------------------------------
        # TEXT
        # ----------------------------------------------------

        if command in (
            "text",
            "mbox"
        ):

            return (
                "normal_text",
                self.parse_argument()
            )

        # ----------------------------------------------------
        # FONT / GROUPING COMMANDS
        # ----------------------------------------------------

        if command in (
            "mathrm",
            "mathbf",
            "mathit",
            "mathsf",
            "mathtt",
            "operatorname",
            "boldsymbol",
            "bm"
        ):

            return (
                "group",
                self.parse_argument()
            )

        # ----------------------------------------------------
        # ACCENTS
        # ----------------------------------------------------

        if command in (
            "widehat",
            "hat",
            "bar",
            "vec",
            "overline",
            "underline"
        ):

            return (
                "accent",
                command,
                self.parse_argument()
            )

        # ----------------------------------------------------
        # DELIMITER CONTROL
        # ----------------------------------------------------

        if command in (
            "left",
            "right"
        ):

            delimiter = self.parse_delimiter()

            return (
                "text",
                delimiter
            )

        # ----------------------------------------------------
        # SYMBOL
        # ----------------------------------------------------

        if command in SYMBOLS:

            return (
                "text",
                SYMBOLS[command]
            )

        # ----------------------------------------------------
        # UNKNOWN COMMAND
        # ----------------------------------------------------

        return (
            "text",
            "\\" + command
        )

    # --------------------------------------------------------
    # ARGUMENT
    # --------------------------------------------------------

    def parse_argument(self) -> list:

        while self.peek() in (
            " ",
            "\t",
            "\n",
            "\r"
        ):

            self.get()

        if not self.peek():
            return []

        if self.peek() == "{":

            return self.parse_group_content()

        if self.peek() == "\\":

            node = self.parse_command()

            return [node] if node else []

        return [
            (
                "text",
                self.get()
            )
        ]

    # --------------------------------------------------------
    # GROUP
    # --------------------------------------------------------

    def parse_group_content(self) -> list:

        if self.peek() != "{":
            return []

        self.get()

        start = self.pos
        depth = 1

        while self.pos < self.length:

            char = self.get()

            if char == "{":
                depth += 1

            elif char == "}":

                depth -= 1

                if depth == 0:

                    content = self.source[
                        start:self.pos - 1
                    ]

                    return LatexParser(
                        content
                    ).parse()

        # Không có dấu đóng:
        # vẫn giữ nội dung thay vì làm mất công thức.
        content = self.source[start:self.pos]

        return LatexParser(
            content
        ).parse()

    # --------------------------------------------------------
    # UNTIL
    # --------------------------------------------------------

    def parse_until(self, end_char: str) -> list:

        start = self.pos

        while (
            self.pos < self.length
            and self.peek() != end_char
        ):

            self.get()

        content = self.source[
            start:self.pos
        ]

        if self.peek() == end_char:
            self.get()

        return LatexParser(
            content
        ).parse()

    # --------------------------------------------------------
    # DELIMITER
    # --------------------------------------------------------

    def parse_delimiter(self) -> str:

        if not self.peek():
            return ""

        if self.peek() == "\\":

            self.get()

            match = re.match(
                r"[A-Za-z]+|.",
                self.source[self.pos:]
            )

            if match:

                delimiter = match.group(0)
                self.pos += len(delimiter)

                return {
                    "lbrace": "{",
                    "rbrace": "}",
                    "langle": "⟨",
                    "rangle": "⟩",
                    "vert": "|",
                    "Vert": "‖",
                }.get(
                    delimiter,
                    SYMBOLS.get(
                        delimiter,
                        delimiter
                    )
                )

        return self.get()

    # --------------------------------------------------------
    # COMBINE TEXT
    # --------------------------------------------------------

    @staticmethod
    def combine_text_nodes(nodes: list) -> list:

        result = []
        buffer = []

        for node in nodes:

            if node[0] == "text":

                buffer.append(
                    node[1]
                )

            else:

                if buffer:

                    result.append(
                        (
                            "text",
                            "".join(buffer)
                        )
                    )

                    buffer = []

                result.append(node)

        if buffer:

            result.append(
                (
                    "text",
                    "".join(buffer)
                )
            )

        return result


# ============================================================
# 4. OMML RENDERER
# ============================================================

def render_omml(nodes: list) -> str:

    if not nodes:
        return ""

    xml_parts = []

    for node in nodes:

        node_type = node[0]

        # ----------------------------------------------------
        # TEXT
        # ----------------------------------------------------

        if node_type == "text":

            xml_parts.append(
                "<m:r>"
                "<m:t xml:space=\"preserve\">"
                f"{escape_xml(node[1])}"
                "</m:t>"
                "</m:r>"
            )

        # ----------------------------------------------------
        # NORMAL TEXT
        # ----------------------------------------------------

        elif node_type == "normal_text":

            content = render_omml(
                node[1]
            )

            content = content.replace(
                "<m:r>",
                (
                    "<m:r>"
                    "<m:rPr>"
                    "<m:nor/>"
                    "</m:rPr>"
                )
            )

            xml_parts.append(content)

        # ----------------------------------------------------
        # GROUP
        # ----------------------------------------------------

        elif node_type == "group":

            xml_parts.append(
                render_omml(
                    node[1]
                )
            )

        # ----------------------------------------------------
        # FRACTION
        # ----------------------------------------------------

        elif node_type == "frac":

            numerator = (
                render_omml(node[1])
                or _empty_math_run()
            )

            denominator = (
                render_omml(node[2])
                or _empty_math_run()
            )

            xml_parts.append(
                "<m:f>"
                f"<m:num>{numerator}</m:num>"
                f"<m:den>{denominator}</m:den>"
                "</m:f>"
            )

        # ----------------------------------------------------
        # SQUARE ROOT
        # ----------------------------------------------------

        elif node_type == "sqrt":

            expression = (
                render_omml(node[1])
                or _empty_math_run()
            )

            xml_parts.append(
                "<m:rad>"
                "<m:deg/>"
                f"<m:e>{expression}</m:e>"
                "</m:rad>"
            )

        # ----------------------------------------------------
        # NTH ROOT
        # ----------------------------------------------------

        elif node_type == "root":

            degree = (
                render_omml(node[1])
                or _empty_math_run()
            )

            expression = (
                render_omml(node[2])
                or _empty_math_run()
            )

            xml_parts.append(
                "<m:rad>"
                f"<m:deg>{degree}</m:deg>"
                f"<m:e>{expression}</m:e>"
                "</m:rad>"
            )

        # ----------------------------------------------------
        # SUPERSCRIPT
        # ----------------------------------------------------

        elif node_type == "sup":

            base = render_omml(
                [node[1]]
            ) or _empty_math_run()

            exponent = render_omml(
                node[2]
            ) or _empty_math_run()

            xml_parts.append(
                "<m:sSup>"
                f"<m:e>{base}</m:e>"
                f"<m:sup>{exponent}</m:sup>"
                "</m:sSup>"
            )

        # ----------------------------------------------------
        # SUBSCRIPT
        # ----------------------------------------------------

        elif node_type == "sub":

            base = render_omml(
                [node[1]]
            ) or _empty_math_run()

            subscript = render_omml(
                node[2]
            ) or _empty_math_run()

            xml_parts.append(
                "<m:sSub>"
                f"<m:e>{base}</m:e>"
                f"<m:sub>{subscript}</m:sub>"
                "</m:sSub>"
            )

        # ----------------------------------------------------
        # SUB + SUP
        # ----------------------------------------------------

        elif node_type == "subsup":

            base = render_omml(
                [node[1]]
            ) or _empty_math_run()

            subscript = render_omml(
                node[2]
            ) or _empty_math_run()

            exponent = render_omml(
                node[3]
            ) or _empty_math_run()

            xml_parts.append(
                "<m:sSubSup>"
                f"<m:e>{base}</m:e>"
                f"<m:sub>{subscript}</m:sub>"
                f"<m:sup>{exponent}</m:sup>"
                "</m:sSubSup>"
            )

        # ----------------------------------------------------
        # ACCENT
        # ----------------------------------------------------

        elif node_type == "accent":

            accent_name = node[1]

            expression = render_omml(
                node[2]
            ) or _empty_math_run()

            accent_map = {
                "widehat": "^",
                "hat": "^",
                "bar": "¯",
                "overline": "¯",
                "underline": "_",
                "vec": "→",
            }

            accent_char = accent_map.get(
                accent_name,
                "^"
            )

            xml_parts.append(
                "<m:acc>"
                "<m:accPr>"
                f"<m:chr m:val=\"{escape_xml(accent_char)}\"/>"
                "</m:accPr>"
                f"<m:e>{expression}</m:e>"
                "</m:acc>"
            )

    return "".join(xml_parts)


# ============================================================
# 5. LATEX → OMML
# ============================================================

def latex_to_omml_xml(
    latex_str: str
) -> str:

    s = _normalize_math_input(
        latex_str
    )

    if not s:

        return (
            f'<m:oMath xmlns:m="{MATH_NS}">'
            f"{_empty_math_run()}"
            "</m:oMath>"
        )

    try:

        parser = LatexParser(
            s
        )

        nodes = parser.parse()

        body = render_omml(
            nodes
        )

        if not body:
            body = _empty_math_run()

        return (
            f'<m:oMath xmlns:m="{MATH_NS}">'
            f"{body}"
            "</m:oMath>"
        )

    except Exception as exc:

        logger.exception(
            "Lỗi biên dịch OMML: %s | %s",
            exc,
            latex_str
        )

        safe_text = escape_xml(
            s
        )

        return (
            f'<m:oMath xmlns:m="{MATH_NS}">'
            "<m:r>"
            f'<m:t xml:space="preserve">'
            f"{safe_text}"
            "</m:t>"
            "</m:r>"
            "</m:oMath>"
        )


# ============================================================
# 6. CHÈN OMML VÀO WORD
# ============================================================

def insert_math_to_paragraph(
    paragraph,
    latex_content: str,
    is_block: bool = False
):
    """
    Chèn công thức OMML native vào Paragraph.

    Không dùng add_run cho công thức nếu có thể biên dịch OMML.
    """

    if not latex_content:
        return

    content = str(
        latex_content
    ).strip()

    if not content:
        return

    try:

        if is_block:

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

        xml_string = latex_to_omml_xml(
            content
        )

        omml_element = parse_xml(
            xml_string
        )

        paragraph._p.append(
            omml_element
        )

    except Exception as exc:

        logger.exception(
            "Không thể chèn OMML: %s",
            exc
        )

        run = paragraph.add_run(
            f" {content} "
        )

        run.font.name = (
            "Cambria Math"
        )

        run.italic = True


# ============================================================
# 7. API TƯƠNG THÍCH NGƯỢC
# ============================================================

class MathRenderer:
    """
    API tương thích với các module cũ.
    """

    @staticmethod
    def render_inline_math(
        paragraph,
        latex_str: str
    ):

        insert_math_to_paragraph(
            paragraph,
            latex_str,
            is_block=False
        )

    @staticmethod
    def render_display_math(
        doc,
        latex_str: str
    ):

        paragraph = doc.add_paragraph()

        insert_math_to_paragraph(
            paragraph,
            latex_str,
            is_block=True
        )

        return paragraph


class ScienceNormalizer:
    """
    Tương thích ngược với code cũ.

    Lưu ý:
    Không còn dùng normalize() để biến công thức thành Unicode
    trước khi chèn vào Word. Công thức chính thức phải đi qua
    OMML.
    """

    @classmethod
    def normalize(
        cls,
        text: str
    ) -> str:

        if not text:
            return ""

        return _normalize_math_input(
            text
        )
