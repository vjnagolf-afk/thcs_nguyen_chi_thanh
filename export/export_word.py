# -*- coding: utf-8 -*-
"""
============================================================
MODULE: export/export_word.py
Nhiệm vụ: Bộ điều phối trung tâm kết xuất Markdown / AI Generated Content 
thành file Word (.docx) chuẩn 5512.
(Bản hoàn hảo: Tự gọi Metadata từ Session State, tự biến Backtick thành Toán)
============================================================
"""

import io
import re
import json
import logging
import streamlit as st
from typing import List, Dict, Any, Optional

import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

try:
    from .markdown_tokenizer import MarkdownTokenizer
except ImportError:
    try:
        from export.markdown_tokenizer import MarkdownTokenizer
    except ImportError:
        MarkdownTokenizer = None

try:
    from .word_math import insert_math_to_paragraph
except ImportError:
    try:
        from export.word_math import insert_math_to_paragraph
    except ImportError:
        insert_math_to_paragraph = None

try:
    from .word_tables import process_and_draw_markdown_table
except ImportError:
    try:
        from export.word_tables import process_and_draw_markdown_table
    except ImportError:
        process_and_draw_markdown_table = None

try:
    from .word_images import insert_image_to_paragraph, insert_image_to_docx
except ImportError:
    try:
        from export.word_images import insert_image_to_paragraph, insert_image_to_docx
    except ImportError:
        insert_image_to_paragraph = None
        insert_image_to_docx = None

try:
    from .template_loader import get_word_document
except ImportError:
    def get_word_document(): 
        return Document()


def _sanitize_and_fix_math(text: str) -> str:
    if not text: return ""

    # AUTO-FIX: Biến mọi dấu nháy ngược (`) do AI sinh nhầm thành dấu $ để module OMML xử lý Toán học
    text = re.sub(r'`([^`\n]+)`', r'$\1$', text)

    text = re.sub(r'\|(sqrt|frac|approx|times|sin|cos|tan|cot|lim|log|ln|alpha|beta|gamma|Delta|pi)\b', r'\\\1', text)

    protected = []
    def protect_math(match):
        protected.append(match.group(0))
        return f"@@MATH_PROTECTED_{len(protected) - 1}@@"

    text = re.sub(r"\$\$[\s\S]*?\$\$", protect_math, text)
    text = re.sub(r"\$(?!\$)[^$\n]+?\$(?!\$)", protect_math, text)

    formula_regex = r"(?<![\w$])((?:[a-zA-Z][a-zA-Z0-9_]*|\d+)\^\{?[^{}\n]+\}?\s*(?:=|<|>|\\le|\\ge|\\approx|\\neq)\s*[^.,;:\n]+?(?=[.,;:]?(?:\s|$))|[A-Za-z](?:_\{[^{}\n]+\}|_\d+|\d+)\s*=\s*(?:\\frac\s*\{[^{}\n]+\}\s*\{[^{}\n]+\})|\\sqrt\s*(?:\[[^\]]*\])?\s*\{[^{}\n]+\}\s*(?:=|\\approx)\s*[^.,;:\n]+?(?=[.,;:]?(?:\s|$)))(?![\w$])"
    formula_pattern = re.compile(formula_regex)

    def wrap_formula(match):
        formula = match.group(1).strip()
        formula = re.sub(r"\b([A-Za-z])(\d{1,3})\b", r"\1_{\2}", formula)
        return f"${formula}$"

    text = formula_pattern.sub(wrap_formula, text)
    text = re.sub(r"(?<![$\w])([a-zA-Z]\^\{?[^{}\n\s]+\}?)(?![$])", r"$\1$", text)
    text = re.sub(r"(?<![$\\])(\\frac\s*\{[^{}\n]*\}\s*\{[^{}\n]*\})(?![$])", r"$\1$", text)
    text = re.sub(r"(?<![$\\])(\\sqrt(?:\[[^\]]*\])?\s*\{[^{}\n]*\})(?![$])", r"$\1$", text)

    for idx, original in enumerate(protected):
        text = text.replace(f"@@MATH_PROTECTED_{idx}@@", original)

    return text


def _parse_inline_with_math_and_images(text: str) -> List[Dict[str, Any]]:
    tokens = []
    pattern_str = r'(\$\$(.*?)\$\$)|(\$([^$]+?)\$)|(\[IMAGE\s*(?:-\s*ID:\s*)?:?\s*([^\]]+)\])|(\*\*([^*]+?)\*\*)|(\*([^*]+?)\*)'
    pattern = re.compile(pattern_str)
    last_idx = 0
    
    for match in pattern.finditer(text):
        if match.start() > last_idx:
            tokens.append({"type": "text", "content": text[last_idx:match.start()]})
        
        if match.group(1):
            tokens.append({"type": "block_math", "content": match.group(2)})
        elif match.group(3):
            tokens.append({"type": "inline_math", "content": match.group(4)})
        elif match.group(5):
            tokens.append({"type": "image", "content": match.group(6).strip()})
        elif match.group(7):
            tokens.append({"type": "bold", "content": match.group(8)})
        elif match.group(9):
            tokens.append({"type": "italic", "content": match.group(10)})
        last_idx = match.end()
        
    if last_idx < len(text):
        tokens.append({"type": "text", "content": text[last_idx:]})
    return tokens


def _fallback_parse_markdown(markdown_text: str) -> List[Dict[str, Any]]:
    ast_nodes = []
    lines = (markdown_text or "").splitlines()
    table_buffer = []
    in_code = False
    code_lang = ""
    code_buffer = []

    def flush_table():
        if table_buffer:
            ast_nodes.append({"type": "table_raw_lines", "lines": list(table_buffer)})
            table_buffer.clear()

    for line in lines:
        s_line = line.strip()
        if s_line.startswith("```"):
            if in_code:
                in_code = False
                ast_nodes.append({"type": "code", "language": code_lang, "text": "\n".join(code_buffer)})
                code_buffer.clear()
            else:
                flush_table()
                in_code = True
                code_lang = s_line.lstrip("`").strip()
            continue
        
        if in_code:
            code_buffer.append(line)
            continue
            
        if s_line.startswith("|"):
            table_buffer.append(s_line)
            continue
        else:
            flush_table()
            
        if not s_line:
            continue
        
        if s_line.startswith("$$") and s_line.endswith("$$") and len(s_line) > 4:
            ast_nodes.append({"type": "block_math", "content": s_line[2:-2].strip()})
            continue
            
        if s_line.startswith("#"):
            match = re.match(r'^(#{1,6})\s+(.*)', s_line)
            if match:
                ast_nodes.append({
                    "type": "heading", 
                    "level": len(match.group(1)), 
                    "text": match.group(2), 
                    "tokens": _parse_inline_with_math_and_images(match.group(2))
                })
                continue
                
        if s_line.startswith("- ") or s_line.startswith("* "):
            ast_nodes.append({
                "type": "list_item", 
                "style": "bullet", 
                "level": 1, 
                "tokens": _parse_inline_with_math_and_images(s_line[2:].strip())
            })
            continue
            
        num_match = re.match(r'^\d+\.\s+(.*)', s_line)
        if num_match:
            ast_nodes.append({
                "type": "list_item", 
                "style": "number", 
                "level": 1, 
                "tokens": _parse_inline_with_math_and_images(num_match.group(1))
            })
            continue

        img_match = re.match(r'^\[IMAGE\s*(?:-\s*ID:\s*)?:?\s*([^\]]+)\]$', s_line)
        if img_match:
            ast_nodes.append({"type": "image", "content": img_match.group(1).strip()})
            continue
            
        tbl_match = re.match(r'^\[TABLE\s*(?:-\s*ID:\s*)?:?\s*([^\]]+)\]$', s_line)
        if tbl_match:
            ast_nodes.append({"type": "table_ref", "id": tbl_match.group(1).strip()})
            continue
            
        img_match_md = re.match(r'^!\[(.*?)\]\((.*?)\)', s_line)
        if img_match_md:
            ast_nodes.append({"type": "image", "content": img_match_md.group(2)})
            continue

        ast_nodes.append({"type": "paragraph", "text": s_line, "tokens": _parse_inline_with_math_and_images(s_line)})

    flush_table()
    return ast_nodes


class WordExportEngine:

    @staticmethod
    def _set_font(run, font_name="Times New Roman"):
        try:
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)
            rFonts.set(qn('w:cs'), font_name)
        except Exception:
            pass

    @classmethod
    def _render_inline_tokens(cls, p, tokens: List[Dict[str, Any]], export_errors: List[Dict[str, Any]], data_cache: dict = None):
        if not tokens:
            return

        for idx, token in enumerate(tokens):
            try:
                ttype = token.get("type", "text")
                content = token.get("content") or token.get("text") or ""

                if ttype in ["text", "bold", "italic", "underline", "strike", "highlight", "inline_code", "code"]:
                    if ttype in ["inline_code", "code"]:
                        run = p.add_run(content)
                        cls._set_font(run, "Courier New")
                        run.font.size = Pt(11)
                    else:
                        sub_tokens = _parse_inline_with_math_and_images(_sanitize_and_fix_math(content))
                        if len(sub_tokens) > 1 or (sub_tokens and sub_tokens[0].get("type") != "text"):
                            cls._render_inline_tokens(p, sub_tokens, export_errors, data_cache)
                        else:
                            run = p.add_run(content)
                            cls._set_font(run, "Times New Roman")
                            run.font.size = Pt(13)
                            
                            if ttype == "bold": 
                                run.bold = True
                            elif ttype == "italic": 
                                run.italic = True
                            elif ttype == "underline": 
                                run.underline = True
                            elif ttype == "strike": 
                                run.font.strike = True
                            elif ttype == "highlight": 
                                run.font.color.rgb = RGBColor(199, 37, 78)
                    continue

                if ttype in ["inline_math", "math_inline", "math", "math_block", "block_math"]:
                    if insert_math_to_paragraph:
                        is_block = (ttype in ["math_block", "block_math"])
                        insert_math_to_paragraph(p, content, is_block=is_block)
                    else:
                        run = p.add_run(f" {content} ")
                        cls._set_font(run, "Cambria Math")
                        run.italic = True

                elif ttype == "image":
                    img_id = content
                    img_src = None
                    if data_cache and "pages" in data_cache:
                        for page in data_cache["pages"]:
                            for img in page.get("images", []):
                                if img.get("id") == img_id:
                                    img_src = {"base64": img.get("base64"), "caption": img_id}
                                    break
                            if img_src: break
                            
                    if insert_image_to_paragraph: 
                        if img_src:
                            insert_image_to_paragraph(p, img_src)
                        else:
                            insert_image_to_paragraph(p, img_id)
                    else: 
                        run_img = p.add_run(f"[Hình ảnh: {img_id}]")
                        run_img.italic = True

                else:
                    run = p.add_run(str(content))
                    cls._set_font(run, "Times New Roman")
                    run.font.size = Pt(13)

            except Exception as e:
                export_errors.append({"type": "inline_token", "index": idx, "token": token, "error": str(e)})
                try: 
                    fallback_text = str(token.get("content") or token.get("text") or "")
                    p.add_run(fallback_text)
                except Exception: 
                    pass

    @classmethod
    def _render_khbd_header(cls, doc: Document, metadata: dict):
        try:
            table = doc.add_table(rows=1, cols=2)
            tblPr = table._element.xpath('w:tblPr')
            if tblPr:
                tblBorders = OxmlElement('w:tblBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'none')
                    tblBorders.append(border)
                tblPr[0].append(tblBorders)

            c0 = table.cell(0, 0)
            p0 = c0.paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r0 = p0.add_run("TRƯỜNG: ....................................\nTỔ: ........................................")
            r0.bold = True
            cls._set_font(r0)

            c1 = table.cell(0, 1)
            p1 = c1.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            mon = str(metadata.get('mon', '.......')).upper()
            lop = str(metadata.get('lop', '.......'))
            r1 = p1.add_run(f"HỌ VÀ TÊN GIÁO VIÊN: ..........................\nMÔN: {mon}\nLỚP: {lop}")
            r1.bold = True
            cls._set_font(r1)
            doc.add_paragraph()

            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title = str(metadata.get('title', '.........................')).upper()
            rt = p_title.add_run(f"TÊN BÀI DẠY: {title}")
            rt.bold = True
            rt.font.size = Pt(16)
            cls._set_font(rt)

            p_time = doc.add_paragraph()
            p_time.alignment = WD_ALIGN_PARAGRAPH.CENTER
            so_tiet = str(metadata.get('so_tiet', '...'))
            r_time = p_time.add_run(f"Môn học/Hoạt động giáo dục: {mon}; Thời gian thực hiện: {so_tiet} tiết")
            r_time.font.italic = True
            cls._set_font(r_time)
            doc.add_paragraph()
        except Exception:
            pass

    @classmethod
    def convert_markdown_to_docx_bytes(cls, markdown_text: str, metadata: dict = None) -> bytes:
        export_errors = []
        doc = get_word_document()

        markdown_text = _sanitize_and_fix_math(markdown_text or "")

        ast_nodes = []
        if MarkdownTokenizer and hasattr(MarkdownTokenizer, 'parse'):
            try:
                ast_nodes = MarkdownTokenizer.parse(markdown_text)
            except Exception as tok_err:
                export_errors.append({"type": "tokenizer_error", "error": str(tok_err)})
                ast_nodes = _fallback_parse_markdown(markdown_text)
        else:
            ast_nodes = _fallback_parse_markdown(markdown_text)

        for node_idx, node in enumerate(ast_nodes):
            try:
                ntype = node.get("type", "paragraph")

                if ntype == "paragraph":
                    raw_text = str(node.get("text", "")).strip()

                    img_match = re.match(r'^\[IMAGE\s*(?:-\s*ID:\s*)?:?\s*([^\]]+)\]$', raw_text)
                    if img_match:
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        img_id = img_match.group(1).strip()
                        img_src = None
                        if metadata and "pages" in metadata:
                            for page in metadata["pages"]:
                                for img in page.get("images", []):
                                    if img.get("id") == img_id:
                                        img_src = {"base64": img.get("base64"), "caption": img_id}
                                        break
                                if img_src: break
                        if insert_image_to_paragraph: insert_image_to_paragraph(p, img_src if img_src else img_id)
                        else: p.add_run(f"[Hình ảnh: {img_id}]").italic = True
                        continue
                        
                    tbl_match = re.match(r'^\[TABLE\s*(?:-\s*ID:\s*)?:?\s*([^\]]+)\]$', raw_text)
                    if tbl_match:
                        tbl_id = tbl_match.group(1).strip()
                        tbl_data = None
                        if metadata and "pages" in metadata:
                            for page in metadata["pages"]:
                                for tab in page.get("tables", []):
                                    if tab.get("id") == tbl_id:
                                        tbl_data = tab
                                        break
                                if tbl_data: break
                        
                        if tbl_data:
                            lines = []
                            lines.append("| " + " | ".join(tbl_data["headers"]) + " |")
                            lines.append("|" + "|".join(["---"] * len(tbl_data["headers"])) + "|")
                            for row in tbl_data["rows"]:
                                lines.append("| " + " | ".join(row) + " |")
                            if process_and_draw_markdown_table:
                                process_and_draw_markdown_table(doc, lines, metadata=metadata)
                        else:
                            p = doc.add_paragraph()
                            p.add_run(f"[Hệ thống: Không tìm thấy dữ liệu bảng {tbl_id}]").italic = True
                        continue

                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(4)
                    tokens = node.get("tokens", [])
                    if tokens:
                        cls._render_inline_tokens(p, tokens, export_errors, metadata)
                    elif raw_text:
                        inline_tokens = _parse_inline_with_math_and_images(_sanitize_and_fix_math(raw_text))
                        cls._render_inline_tokens(p, inline_tokens, export_errors, metadata)

                elif ntype == "table_ref":
                    tbl_id = node.get("id")
                    tbl_data = None
                    if metadata and "pages" in metadata:
                        for page in metadata["pages"]:
                            for tab in page.get("tables", []):
                                if tab.get("id") == tbl_id:
                                    tbl_data = tab
                                    break
                            if tbl_data: break
                    if tbl_data:
                        lines = []
                        lines.append("| " + " | ".join(tbl_data["headers"]) + " |")
                        lines.append("|" + "|".join(["---"] * len(tbl_data["headers"])) + "|")
                        for row in tbl_data["rows"]:
                            lines.append("| " + " | ".join(row) + " |")
                        if process_and_draw_markdown_table:
                            process_and_draw_markdown_table(doc, lines, metadata=metadata)
                    else:
                        p = doc.add_paragraph()
                        p.add_run(f"[Hệ thống: Không tìm thấy dữ liệu bảng {tbl_id}]").italic = True

                elif ntype == "heading":
                    level = min(max(node.get("level", 1), 1), 6)
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(10)
                    p.paragraph_format.space_after = Pt(4)
                    p.paragraph_format.keep_with_next = True
                    
                    if level == 1:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    tokens = node.get("tokens", [])
                    if tokens:
                        cls._render_inline_tokens(p, tokens, export_errors, metadata)
                    else:
                        inline_tokens = _parse_inline_with_math_and_images(_sanitize_and_fix_math(str(node.get("text") or "")))
                        cls._render_inline_tokens(p, inline_tokens, export_errors, metadata)

                    for r in p.runs:
                        r.bold = True
                        cls._set_font(r, "Times New Roman")
                        if level == 1: r.font.size = Pt(16)
                        elif level == 2: r.font.size = Pt(14)
                        else: r.font.size = Pt(13)

                elif ntype in ["block_math", "math_block"]:
                    math_content = node.get("content") or node.get("text") or ""
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    if insert_math_to_paragraph:
                        insert_math_to_paragraph(p, math_content, is_block=True)
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r = p.add_run(math_content)
                        cls._set_font(r, "Cambria Math")
                        r.italic = True

                elif ntype in ["list_item", "bullet_list", "numbered_list"]:
                    style_name = 'List Number' if (node.get("style") == "number" or ntype == "numbered_list") else 'List Bullet'
                    level = node.get("level", 1)
                    p = doc.add_paragraph(style=style_name)
                    p.paragraph_format.left_indent = Inches(0.25 * level + 0.25)
                    p.paragraph_format.first_line_indent = Inches(-0.25)
                    p.paragraph_format.space_after = Pt(3)

                    tokens = node.get("tokens", [])
                    if tokens:
                        cls._render_inline_tokens(p, tokens, export_errors, metadata)
                    else:
                        inline_tokens = _parse_inline_with_math_and_images(_sanitize_and_fix_math(str(node.get("text") or "")))
                        cls._render_inline_tokens(p, inline_tokens, export_errors, metadata)

                elif ntype == "image":
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    img_id = node.get("content") or node.get("url") or node.get("path")
                    img_src = None
                    if metadata and "pages" in metadata:
                        for page in metadata["pages"]:
                            for img in page.get("images", []):
                                if img.get("id") == img_id:
                                    img_src = {"base64": img.get("base64"), "caption": img_id}
                                    break
                            if img_src: break
                    if insert_image_to_paragraph: insert_image_to_paragraph(p, img_src if img_src else img_id)
                    else: p.add_run(f"[Hình ảnh: {img_id}]").italic = True

                elif ntype in ["table", "table_raw_lines"]:
                    lines = node.get("lines", [])
                    if process_and_draw_markdown_table and lines:
                        process_and_draw_markdown_table(doc, lines, metadata=metadata)
                    else:
                        for line in lines:
                            p = doc.add_paragraph()
                            p.add_run(str(line))

                elif ntype in ["code", "code_block"]:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.4)
                    r = p.add_run(node.get("text", ""))
                    cls._set_font(r, "Courier New")
                    r.font.size = Pt(10.5)

                elif ntype in ["hr", "horizontal_rule"]:
                    p = doc.add_paragraph()
                    pPr = p._element.get_or_add_pPr()
                    pb = OxmlElement("w:pBdr")
                    bottom = OxmlElement("w:bottom")
                    bottom.set(qn("w:val"), "single")
                    bottom.set(qn("w:sz"), "8")
                    bottom.set(qn("w:color"), "CCCCCC")
                    pb.append(bottom)
                    pPr.append(pb)

                elif ntype == "page_break":
                    doc.add_page_break()

            except Exception as node_err:
                export_errors.append({"type": "ast_node", "node_index": node_idx, "node": node, "error": str(node_err)})

        output_stream = io.BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)
        return output_stream.getvalue()

    @classmethod
    def export_to_word(cls, data_cache: Dict[str, Any]) -> bytes:
        # KIỂM TRA SESSION STATE ĐỂ TỰ ĐỘNG LẤY METADATA
        metadata = {}
        if isinstance(data_cache, dict):
            metadata = data_cache.copy()
            md_text = metadata.get("ai_generated_content", "")
        else:
            md_text = str(data_cache)
            
        if "current_source_metadata" in st.session_state:
            metadata["pages"] = st.session_state["current_source_metadata"].get("pages", [])
            
        return cls.convert_markdown_to_docx_bytes(md_text, metadata=metadata)


def export_word(markdown_text_or_cache) -> bytes:
    try:
        # BỌC LỚP BẢO VỆ CUỐI CÙNG ĐỂ GỌI SESSION STATE TỪ MỌI NƠI
        metadata = {}
        if isinstance(markdown_text_or_cache, dict):
            metadata = markdown_text_or_cache.copy()
            md_text = metadata.get("ai_generated_content", "")
        else:
            md_text = str(markdown_text_or_cache)
            
        if "current_source_metadata" in st.session_state:
            metadata["pages"] = st.session_state["current_source_metadata"].get("pages", [])
            
        return WordExportEngine.convert_markdown_to_docx_bytes(md_text, metadata=metadata)
        
    except Exception as fatal_err:
        fallback_doc = Document()
        fallback_doc.add_paragraph("KẾ HOẠCH BÀI DẠY (BẢN PHỤC HỒI)")
        fallback_doc.add_paragraph(f"Lỗi: {fatal_err}")
        bio = io.BytesIO()
        fallback_doc.save(bio)
        bio.seek(0)
        return bio.getvalue()
