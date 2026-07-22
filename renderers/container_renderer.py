# -*- coding: utf-8 -*-

import docx

from docx.shared import (
    Pt,
    Inches
)

from docx.oxml import OxmlElement

from docx.oxml.ns import qn

from styles.xml_helpers import XmlHelpers


class ContainerRenderer:

    @classmethod
    def render_code_block(
        cls,
        doc,
        node
    ):

        paragraph = doc.add_paragraph()

        paragraph.paragraph_format.left_indent = Inches(
            0.4
        )

        paragraph.paragraph_format.space_before = Pt(
            4
        )

        paragraph.paragraph_format.space_after = Pt(
            4
        )

        XmlHelpers.apply_paragraph_shading(
            paragraph,
            "F5F5F5"
        )

        run = paragraph.add_run(
            node.get(
                "text",
                ""
            )
        )

        run.font.size = Pt(
            10.5
        )

        XmlHelpers.set_font_safely(
            run,
            "Courier New"
        )

    @classmethod
    def render_callout(
        cls,
        doc,
        node,
        text_renderer,
        math_renderer
    ):

        style = node.get(
            "style",
            "quote"
        )

        if style == "warning":

            bg_color = "FFF5F5"
            border_color = "FF3B30"

        elif style == "tip":

            bg_color = "F0F7FF"
            border_color = "007AFF"

        else:

            bg_color = "F9F9F9"
            border_color = "8E8E93"

        table = doc.add_table(
            rows=1,
            cols=1
        )

        table.alignment = (
            docx.enum.table
            .WD_TABLE_ALIGNMENT.CENTER
        )

        table.autofit = False

        table.columns[0].width = Inches(
            6.3
        )

        cell = table.cell(
            0,
            0
        )

        tcPr = cell._element.get_or_add_tcPr()

        # Background
        shd = OxmlElement(
            "w:shd"
        )

        shd.set(
            qn("w:val"),
            "clear"
        )

        shd.set(
            qn("w:fill"),
            bg_color
        )

        tcPr.append(shd)

        # Borders
        borders = OxmlElement(
            "w:tcBorders"
        )

        left = OxmlElement(
            "w:left"
        )

        left.set(
            qn("w:val"),
            "single"
        )

        left.set(
            qn("w:sz"),
            "24"
        )

        left.set(
            qn("w:color"),
            border_color
        )

        borders.append(
            left
        )

        for side in (
            "top",
            "bottom",
            "right"
        ):

            border = OxmlElement(
                f"w:{side}"
            )

            border.set(
                qn("w:val"),
                "none"
            )

            borders.append(
                border
            )

        tcPr.append(
            borders
        )

        # Nội dung
        first = True

        for child in node.get(
            "children",
            []
        ):

            if first:

                paragraph = cell.paragraphs[0]

                first = False

            else:

                paragraph = cell.add_paragraph()

            paragraph.paragraph_format.space_after = Pt(
                4
            )

            if child.get(
                "type"
            ) == "paragraph":

                text_renderer.render_inline_tokens(
                    paragraph,
                    child.get(
                        "tokens",
                        []
                    ),
                    math_renderer
                )
