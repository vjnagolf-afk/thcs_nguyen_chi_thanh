# -*- coding: utf-8 -*-

from docx.shared import Inches

from docx.enum.text import (
    WD_ALIGN_PARAGRAPH
)

from styles.xml_helpers import XmlHelpers


class BlockRenderer:

    @classmethod
    def render_heading(
        cls,
        doc,
        node,
        text_renderer,
        math_renderer
    ):

        level = min(
            max(
                node.get(
                    "level",
                    1
                ),
                1
            ),
            3
        )

        paragraph = doc.add_paragraph(
            style=f"Heading {level}"
        )

        if level == 1:

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

        text_renderer.render_inline_tokens(
            paragraph,
            node.get(
                "tokens",
                []
            ),
            math_renderer
        )

        for run in paragraph.runs:

            XmlHelpers.set_font_safely(
                run,
                "Times New Roman"
            )

    @classmethod
    def render_list_item(
        cls,
        doc,
        node,
        text_renderer,
        math_renderer
    ):

        style_name = (
            "List Number"
            if node.get("style")
            == "number"
            else "List Bullet"
        )

        paragraph = doc.add_paragraph(
            style=style_name
        )

        level = node.get(
            "level",
            1
        )

        paragraph.paragraph_format.left_indent = Inches(
            0.25 * level + 0.25
        )

        paragraph.paragraph_format.first_line_indent = Inches(
            -0.25
        )

        text_renderer.render_inline_tokens(
            paragraph,
            node.get(
                "tokens",
                []
            ),
            math_renderer
        )

    @classmethod
    def render_checkbox(
        cls,
        doc,
        node,
        text_renderer,
        math_renderer
    ):

        paragraph = doc.add_paragraph()

        paragraph.paragraph_format.left_indent = Inches(
            0.25 * node.get(
                "level",
                1
            )
        )

        box = (
            "☑ "
            if node.get(
                "checked",
                False
            )
            else "☐ "
        )

        run = paragraph.add_run(
            box
        )

        XmlHelpers.set_font_safely(
            run,
            "MS Gothic"
        )

        run.bold = True

        text_renderer.render_inline_tokens(
            paragraph,
            node.get(
                "tokens",
                []
            ),
            math_renderer
        )
