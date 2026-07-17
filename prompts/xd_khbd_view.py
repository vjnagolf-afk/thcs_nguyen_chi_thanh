import streamlit as st
import sys
import os
from pathlib import Path
from string import Template
from pypdf import PdfReader

# Nhận diện module gốc
root_path = Path(__file__).resolve().parents[1]
if str(root_path) not in sys.path:
    sys.path.insert(0, str(root_path))

from models.khbd import KHBD

def load_prompt(filename):
    file_path = root_path / "prompts" / filename
    if not file_path.exists():
        return ""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def extract_text_from_file(uploaded_file):
    if not uploaded_file: return ""
    try:
        if uploaded_file.name.endswith('.pdf'):
            return "\n".join([p.extract_text() for p in PdfReader(uploaded_file).pages if p.extract_text()])
        elif uploaded_file.name.endswith('.docx'):
            import docx
            doc = docx.Document(uploaded_file)
            return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        elif uploaded_file.name.endswith('.txt'):
            return uploaded_file.read().decode("utf-8")
    except Exception as e:
        st.warning(f"Không thể đọc tài liệu: {e}")
    return ""

def read_khbd_template():
    """Tự động đọc file khbd_mau.docx từ thư mục templates để ép AI làm theo"""
    template_path = root_path / "templates" / "khbd_mau.docx"
    if not template_path.exists():
        return "Viết theo chuẩn Công văn 5512 của Bộ GD&ĐT (I. Mục tiêu, II. Thiết bị, III. Tiến trình...)"
    try:
        import docx
        doc = docx.Document(template_path)
        return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    except:
        return "Viết theo chuẩn Công văn 5512 của Bộ GD&ĐT."

def render_xd_khbd(ai_engine):
    st.markdown("### 📝 Xây dựng Kế hoạch bài dạy (AI Hỗ trợ - Chuẩn 5512)")

    ds_mon = [
        "Ngữ văn", "Toán", "Ngoại ngữ", "Giáo dục công dân", "Lịch sử và Địa lý", 
        "Khoa học tự nhiên", "Vật lí", "Hoá học", "Sinh học", "Công nghệ", 
        "Tin học", "Giáo dục thể chất", "Nghệ thuật", "Giáo dục địa phương", 
        "Hoạt động trải nghiệm, hướng nghiệp"
    ]

    col1, col2, col3 = st.columns([1, 1, 1])
    mon_hoc = col1.selectbox("Môn học", ds_mon)
    lop = col2.selectbox("Lớp", ["Lớp 6", "Lớp 7", "Lớp 8", "Lớp 9", "Lớp 10", "Lớp 11", "Lớp 12"])
    hinh_thuc = col3.selectbox("Chọn hình thức", ["KHBD chi tiết (Chuẩn 5512)"])
    
    so_tiet = st.number_input("Số tiết", min_value=1, max_value=20, value=2)
    ten_bai_hoc = st.text_input("Tên chủ đề / Tên bài học")
    
    bam_sat = st.checkbox("Bám sát nội dung file tải lên", value=True)
    yeu_cau = st.text_area("Yêu cầu bổ sung (Ví dụ: Lồng ghép Năng lực số, tích hợp AI...)")
    file_tai_len = st.file_uploader("Tài liệu tham khảo (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
    
    c1, c2 = st.columns(2)
    if c1.button("🚀 KHỞI TẠO TIẾN TRÌNH", type="primary"):
        if not ten_bai_hoc.strip():
            st.error("⚠️ Vui lòng nhập 'Tên chủ đề / Tên bài học'!")
        else:
            with st.spinner("⏳ AI đang học mẫu 5512 và thiết kế giáo án..."):
                file_context = extract_text_from_file(file_tai_len) if (file_tai_len and bam_sat) else "Dùng kiến thức chuẩn."
                mau_cau_truc = read_khbd_template() # Tự động lấy cấu trúc mẫu
                
                # Load và format Prompt
                task_template = load_prompt("task_config_khbd.txt")
                prompt = Template(task_template).substitute(
                    mon_hoc=mon_hoc, lop=lop, ten_bai_hoc=ten_bai_hoc, so_tiet=so_tiet,
                    yeu_cau=yeu_cau if yeu_cau else "Không có",
                    file_context=file_context[:10000], mau_cau_truc=mau_cau_truc[:5000]
                )
                
                try:
                    content = ai_engine.generate_text(prompt)
                    # Gói vào Data Model
                    current_khbd = KHBD(ten_bai_hoc, mon_hoc, lop, so_tiet, yeu_cau, content)
                    st.session_state['khbd_data'] = current_khbd.to_dict()
                    st.rerun()
                except Exception as e:
                    st.error(f"Lỗi hệ thống AI: {e}")

    if c2.button("🗑️ XÓA DỮ LIỆU"):
        st.session_state.pop('khbd_data', None)
        st.rerun()

    # HIỂN THỊ VÀ XUẤT WORD
    if 'khbd_data' in st.session_state:
        khbd_obj = KHBD.from_dict(st.session_state['khbd_data'])
        
        st.markdown("---")
        st.markdown(khbd_obj.ai_generated_content)
        
        try:
            from export.export_word import WordExportEngine

            # Tạo dictionary đúng form WordExportEngine yêu cầu (kế thừa logic cũ của thầy)
            data_export = st.session_state['khbd_data'].copy()
            data_export["is_khbd"] = True
            
            word_bytes = WordExportEngine.export_to_word(data_export)
            
            st.download_button(
                label="📥 TẢI FILE KẾ HOẠCH BÀI DẠY (WORD)", 
                data=word_bytes, 
                file_name=f"KHBD_{khbd_obj.ten_bai_hoc}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary"
            )
        except Exception as e:
            st.error(f"❌ Lỗi xuất file Word: {e}")
