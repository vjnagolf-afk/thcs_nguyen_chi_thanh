# -*- coding: utf-8 -*-
r"""
============================================================
MODULE: views/xd_de_kt_data.py
Logic xử lý Xây dựng Đề kiểm tra (Đọc dữ liệu, Bóc tách Metadata, Gọi AI)
Hỗ trợ 2 chế độ: Chuẩn CV 7991 và Tùy chọn (4 mức độ, tùy chỉnh điểm).
Có cơ chế Auto-Retry chống lỗi 503 Overloaded từ Google.
============================================================
"""

import os
import re
import json
import time
import logging
import base64
from io import BytesIO
import streamlit as st
from docx import Document

logger = logging.getLogger(__name__)

def init_session_state_de_kt():
    defaults = {
        "dkt_result": None,
        "dkt_processing": False,
        "current_dkt_metadata": {}
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

def reset_dkt_result():
    st.session_state["dkt_result"] = None
    st.session_state["current_dkt_metadata"] = {}

def safe_text(value):
    if value is None: return ""
    text = str(value).replace("\x00", "").replace("\ufeff", "").replace("\u200b", "")
    return re.sub(r"[ ]{2,}", " ", text.replace("\r", "").replace("\t", " ")).strip()

def parse_pdf_structured_dkt(uploaded_file):
    if uploaded_file is None: return {"source_name": "unknown", "pages": []}
    content = uploaded_file.getvalue() if hasattr(uploaded_file, "getvalue") else uploaded_file.read()
    file_name = getattr(uploaded_file, 'name', 'document.pdf')
    pages_data = []
    try:
        import fitz
        doc = fitz.open(stream=content, filetype="pdf")
        for i in range(len(doc)):
            page = doc[i]
            page_images = []
            for img_idx, img in enumerate(page.get_images(full=True)):
                try:
                    base_image = doc.extract_image(img[0])
                    page_images.append({"id": f"IMG_DKT_P{i+1}_{img_idx+1}", "page": i + 1, "ext": base_image["ext"], "base64": base64.b64encode(base_image["image"]).decode("utf-8")})
                except: pass
            page_tables = []
            try:
                tabs = page.find_tables()
                if tabs and tabs.tables:
                    for t_idx, tab in enumerate(tabs.tables):
                        extracted_df = tab.extract()
                        if extracted_df and len(extracted_df) > 0:
                            page_tables.append({"id": f"TAB_DKT_P{i+1}_{t_idx+1}", "page": i + 1, "headers": [str(c) if c is not None else "" for c in extracted_df[0]], "rows": [[str(c) if c is not None else "" for c in r] for r in extracted_df[1:]]})
            except: pass
            pages_data.append({"page_number": i + 1, "text": safe_text(page.get_text("text")), "images": page_images, "tables": page_tables})
    except Exception as e:
        logger.error(f"Lỗi đọc PDF trong module Đề KT: {e}")
    return {"source_name": file_name, "pages": pages_data}

def parse_docx_structured_dkt(uploaded_file):
    if uploaded_file is None: return {"source_name": "unknown", "pages": []}
    file_name = getattr(uploaded_file, 'name', 'document.docx')
    paragraphs_text = []
    tables_data = []
    try:
        source = uploaded_file.getvalue() if hasattr(uploaded_file, "getvalue") else uploaded_file.read()
        doc = Document(BytesIO(source))
        for p in doc.paragraphs:
            if p.text.strip(): paragraphs_text.append(p.text.strip())
        for t_idx, table in enumerate(doc.tables):
            rows_raw = [[cell.text.strip().replace('\n', ' ') for cell in row.cells] for row in table.rows]
            if rows_raw: tables_data.append({"id": f"TAB_DKT_DOCX_{t_idx+1}", "page": 1, "headers": rows_raw[0], "rows": rows_raw[1:] if len(rows_raw) > 1 else []})
    except Exception as e:
        logger.error(f"Lỗi đọc DOCX trong module Đề KT: {e}")
    return {"source_name": file_name, "pages": [{"page_number": 1, "text": "\n".join(paragraphs_text), "images": [], "tables": tables_data}]}

def read_multiple_files_dkt(files):
    combined = {"source_name": "multi_files", "pages": []}
    offset = 0
    for f in files or []:
        if f.name.lower().endswith('.pdf'): parsed = parse_pdf_structured_dkt(f) 
        else: parsed = parse_docx_structured_dkt(f)
        for p in parsed["pages"]:
            p["page_number"] += offset
            combined["pages"].append(p)
        offset += len(parsed["pages"])
        
    st.session_state["current_dkt_metadata"] = combined
    
    source_lines = []
    for page in combined["pages"]:
        p_num = page["page_number"]
        source_lines.append(f"=== TRANG {p_num} ===")
        if page["text"]: source_lines.append(f"[TEXT - Trang {p_num}]\n{page['text']}")
        for tab in page.get("tables", []):
            source_lines.append(f"BẢNG DỮ LIỆU [TABLE: {tab['id']}]\nHeaders: {' | '.join(tab['headers'])}\nRows:\n" + "\n".join([' | '.join(r) for r in tab['rows']]))
        for img in page.get("images", []): 
            source_lines.append(f"HÌNH MINH HỌA [IMAGE: {img['id']}]")
            
    return "\n\n".join(source_lines)

# ============================================================
# PROMPT: CHUẨN CV 7991 (3 Mức độ)
# ============================================================
def get_prompt_cv7991():
    return r"""
Bạn là chuyên gia xây dựng Đề kiểm tra và Ma trận. BẠN PHẢI TUÂN THỦ TUYỆT ĐỐI CÔNG VĂN 7991/BGDĐT-GDTrH.

1. BẮT BUỘC 3 MỨC ĐỘ ĐÁNH GIÁ: Biết (Nhận biết), Hiểu (Thông hiểu), Vận dụng. CẤM tách thêm mức độ Vận dụng cao.
2. BẮT BUỘC 4 DẠNG CÂU HỎI TRONG MA TRẬN VÀ ĐỀ: Nhiều lựa chọn, Đúng - Sai, Trả lời ngắn, Tự luận.

# 1. MA TRẬN ĐỀ KIỂM TRA ĐỊNH KÌ (Dùng Bảng Markdown)
| TT | Chủ đề/Chương | Nội dung/đơn vị kiến thức | Nhiều lựa chọn (Biết) | Nhiều lựa chọn (Hiểu) | Nhiều lựa chọn (Vận dụng) | Đúng - Sai (Biết) | Đúng - Sai (Hiểu) | Đúng - Sai (Vận dụng) | Trả lời ngắn (Biết) | Trả lời ngắn (Hiểu) | Trả lời ngắn (Vận dụng) | Tự luận (Biết) | Tự luận (Hiểu) | Tự luận (Vận dụng) | Tổng |

# 2. BẢN ĐẶC TẢ ĐỀ KIỂM TRA ĐỊNH KÌ (Dùng Bảng Markdown)
| TT | Chủ đề/Chương | Nội dung/đơn vị kiến thức | Yêu cầu cần đạt | Nhiều lựa chọn | Đúng - Sai | Trả lời ngắn | Tự luận |

# 3. ĐỀ KIỂM TRA CHI TIẾT
Tuân thủ đúng 4 phần: Trắc nghiệm nhiều lựa chọn, Trắc nghiệm đúng sai, Trắc nghiệm trả lời ngắn, Tự luận.

# 4. HƯỚNG DẪN CHẤM VÀ BẢNG ĐÁP ÁN

[KỶ LUẬT THÉP VỀ TOÁN HỌC VÀ HÌNH ẢNH]
- CẤM dùng dấu backtick (`). MỌI biểu thức Toán/Lý/Hóa BẮT BUỘC bọc trong cặp dấu $...$ (VD: $x^2 = 49$, $\frac{a}{b}$).
- Bắt buộc nhúng thẻ `[IMAGE: ID]` hoặc `[TABLE: ID]` được trích xuất từ dữ liệu trung gian vào đề bài nếu cần dùng ảnh/bảng.
"""

# ============================================================
# PROMPT: TÙY CHỌN (4 Mức độ, phân bổ điểm tự do)
# ============================================================
def get_prompt_tuy_chon(config):
    tn = config["trac_nghiem"]
    tl = config["tu_luan"]
    tl_details = ", ".join([f"Câu {i+1} ({diem}đ)" for i, diem in enumerate(tl["chi_tiet_diem"])])

    return f"""
Bạn là chuyên gia xây dựng Đề kiểm tra. Hãy lập Ma trận, Đặc tả, Đề thi và Đáp án TUYỆT ĐỐI TUÂN THỦ CẤU TRÚC ĐIỂM SAU ĐÂY:

1. MỨC ĐỘ NHẬN THỨC (4 mức độ): 
   Nhận biết ({config['muc_do']['nb']}%), Thông hiểu ({config['muc_do']['th']}%), Vận dụng ({config['muc_do']['vd']}%), Vận dụng cao ({config['muc_do']['vdc']}%).

2. CẤU TRÚC TRẮC NGHIỆM (Tổng {tn['tong_diem']} điểm):
   - Nhiều lựa chọn: {tn['nlc_cau']} câu (Mỗi câu {tn['nlc_diem']} điểm).
   - Đúng/Sai: {tn['ds_cau']} câu (Mỗi câu {tn['ds_diem']} điểm).
   - Điền khuyết: {tn['dk_cau']} câu (Mỗi câu {tn['dk_diem']} điểm).
   - Trả lời ngắn: {tn['tln_cau']} câu (Mỗi câu {tn['tln_diem']} điểm).

3. CẤU TRÚC TỰ LUẬN (Tổng {tl['tong_diem']} điểm - Bao gồm {tl['so_cau']} câu):
   - Chi tiết: {tl_details}.

# 1. MA TRẬN ĐỀ KIỂM TRA (Dùng Bảng Markdown)
| TT | Chủ đề | Nội dung kiến thức | Nhận biết (TN/TL) | Thông hiểu (TN/TL) | Vận dụng (TN/TL) | Vận dụng cao (TN/TL) | Tổng |
Lưu ý: Phân bổ số câu hỏi vào các ô sao cho tổng điểm và tỷ lệ % khớp với cấu hình ở trên.

# 2. BẢN ĐẶC TẢ ĐỀ KIỂM TRA (Dùng Bảng Markdown)
| TT | Chủ đề | Nội dung kiến thức | Mức độ | Yêu cầu cần đạt | Số câu TN | Số câu TL |

# 3. ĐỀ KIỂM TRA CHI TIẾT
- Phần I: Trắc nghiệm (Bao gồm các loại câu TN như cấu hình).
- Phần II: Tự luận (Bao gồm đúng {tl['so_cau']} câu với số điểm đã cho).

# 4. HƯỚNG DẪN CHẤM VÀ BẢNG ĐÁP ÁN

[KỶ LUẬT THÉP VỀ TOÁN HỌC VÀ HÌNH ẢNH]
- CẤM dùng dấu backtick (`). MỌI biểu thức Toán/Lý/Hóa BẮT BUỘC bọc trong cặp dấu $...$ (VD: $x^2 = 49$).
- Bắt buộc nhúng thẻ `[IMAGE: ID]` hoặc `[TABLE: ID]` được trích xuất từ dữ liệu trung gian vào đề bài nếu cần dùng ảnh/bảng.
"""

def generate_dkt_ai(model_name: str, config: dict, source_text: str):
    # Phân luồng Prompt dựa trên chế độ
    if config.get("mode") == "tuy_chon":
        prompt_instruction = get_prompt_tuy_chon(config)
        info_header = f"--- THÔNG TIN CẤU HÌNH ---\nMôn: {config.get('mon_hoc')}\nLớp: {config.get('lop')}\nTên bài/Đề số: {config.get('ten_bai')}\nThời gian: {config.get('thoi_gian')} phút\nYêu cầu khác: {config.get('yeu_cau_dac_biet')}"
    else:
        prompt_instruction = get_prompt_cv7991()
        info_header = f"--- THÔNG TIN CẤU HÌNH ---\nMôn: {config.get('mon_hoc')}\nLớp: {config.get('lop')}\nThời gian: {config.get('thoi_gian')} phút\nTỉ lệ: {config.get('ty_le')}\nYêu cầu khác: {config.get('yeu_cau_dac_biet')}"

    full_prompt = f"{info_header}\n\n--- HƯỚNG DẪN BIÊN SOẠN ---\n{prompt_instruction}\n\n--- NGUỒN TÀI LIỆU TRUNG GIAN ---\n{source_text[:40000]}"

    api_key = st.session_state.get("user_api_key", "").strip() or os.environ.get("GEMINI_API_KEY", "").strip()
    try: api_key = api_key or st.secrets.get("GEMINI_API_KEY", "").strip()
    except: pass

    text_out = ""
    max_retries = 3
    
    for attempt in range(max_retries):
        try:
            if api_key.startswith("sk-") or "proj-" in api_key:
                import openai
                oai_client = openai.OpenAI(api_key=api_key)
                response = oai_client.chat.completions.create(
                    model="gpt-4o" if "Pro" in model_name else "gpt-4o-mini",
                    messages=[{"role": "system", "content": "Bạn là chuyên gia xây dựng Đề kiểm tra."}, {"role": "user", "content": full_prompt}],
                    max_tokens=8192
                )
                text_out = response.choices[0].message.content.strip()
            else:
                from google import genai
                client_genai = genai.Client(api_key=api_key or "AIzaSy_dummy")
                api_model = "gemini-2.5-pro" if "Pro" in model_name else "gemini-2.5-flash"
                response = client_genai.models.generate_content(model=api_model, contents=full_prompt)
                text_out = response.text.strip()
            break
        except Exception as e:
            error_msg = str(e)
            logger.error(f"AI Generation Error (Lần thử {attempt + 1}): {error_msg}")
            if ("503" in error_msg or "429" in error_msg or "UNAVAILABLE" in error_msg) and attempt < max_retries - 1:
                time.sleep((attempt + 1) * 3)
                continue
            raise RuntimeError(f"Lỗi kết nối AI: {error_msg}. Vui lòng chờ vài phút rồi tạo lại.")

    if not text_out: raise RuntimeError("Không thể nhận phản hồi từ AI.")
    text_out = re.sub(r'^```markdown\s*', '', text_out, flags=re.MULTILINE)
    text_out = re.sub(r'^```\s*$', '', text_out, flags=re.MULTILINE)
    return text_out
