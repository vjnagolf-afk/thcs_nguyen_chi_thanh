# -*- coding: utf-8 -*-
r"""
============================================================
MODULE: views/xd_ma_tran_tu_de_data.py
Logic: Phân tích ngược Đề kiểm tra có sẵn để lập Ma trận & Đặc tả.
Bỏ cơ chế JSON cũ. Dùng Markdown Table + Bộ xuất Word Native.
Tích hợp Khóa chốt API trực tiếp và Auto-Retry.
============================================================
"""

import os
import re
import time
import logging
import base64
from io import BytesIO
import streamlit as st
from docx import Document

logger = logging.getLogger(__name__)

def init_session_state_mt():
    defaults = {
        "mt_result": None,
        "mt_processing": False,
        "current_mt_metadata": {}
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

def reset_mt_result():
    st.session_state["mt_result"] = None
    st.session_state["current_mt_metadata"] = {}

def safe_text(value):
    if value is None: return ""
    text = str(value).replace("\x00", "").replace("\ufeff", "").replace("\u200b", "")
    return re.sub(r"[ ]{2,}", " ", text.replace("\r", "").replace("\t", " ")).strip()

# ============================================================
# BỘ ĐỌC ĐỀ THI CŨ TƯƠNG ĐƯƠNG MODULE ĐỀ KIỂM TRA
# ============================================================
def parse_pdf_mt(uploaded_file):
    content = uploaded_file.getvalue() if hasattr(uploaded_file, "getvalue") else uploaded_file.read()
    file_name = getattr(uploaded_file, 'name', 'document.pdf')
    pages_data = []
    try:
        import fitz
        doc = fitz.open(stream=content, filetype="pdf")
        for i in range(len(doc)):
            page = doc[i]
            page_images = []
            for img_idx, img in enumerate(page.get_images(full=True)):
                try:
                    base_image = doc.extract_image(img[0])
                    page_images.append({
                        "id": f"IMG_MT_P{i+1}_{img_idx+1}",
                        "page": i + 1, "ext": base_image["ext"],
                        "base64": base64.b64encode(base_image["image"]).decode("utf-8")
                    })
                except: pass
            page_tables = []
            try:
                tabs = page.find_tables()
                if tabs and tabs.tables:
                    for t_idx, tab in enumerate(tabs.tables):
                        extracted_df = tab.extract()
                        if extracted_df and len(extracted_df) > 0:
                            page_tables.append({
                                "id": f"TAB_MT_P{i+1}_{t_idx+1}", "page": i + 1,
                                "headers": [str(c) if c is not None else "" for c in extracted_df[0]],
                                "rows": [[str(c) if c is not None else "" for c in r] for r in extracted_df[1:]]
                            })
            except: pass
            pages_data.append({"page_number": i + 1, "text": safe_text(page.get_text("text")), "images": page_images, "tables": page_tables})
    except Exception as e:
        logger.error(f"Lỗi đọc PDF: {e}")
    return {"source_name": file_name, "pages": pages_data}

def parse_docx_mt(uploaded_file):
    file_name = getattr(uploaded_file, 'name', 'document.docx')
    paragraphs_text = []
    tables_data = []
    try:
        source = uploaded_file.getvalue() if hasattr(uploaded_file, "getvalue") else uploaded_file.read()
        doc = Document(BytesIO(source))
        for p in doc.paragraphs:
            if p.text.strip(): paragraphs_text.append(p.text.strip())
        for t_idx, table in enumerate(doc.tables):
            rows_raw = [[cell.text.strip().replace('\n', ' ') for cell in row.cells] for row in table.rows]
            if rows_raw: tables_data.append({"id": f"TAB_MT_DOCX_{t_idx+1}", "page": 1, "headers": rows_raw[0], "rows": rows_raw[1:] if len(rows_raw) > 1 else []})
    except Exception as e:
        logger.error(f"Lỗi đọc DOCX: {e}")
    return {"source_name": file_name, "pages": [{"page_number": 1, "text": "\n".join(paragraphs_text), "images": [], "tables": tables_data}]}

def extract_exam_source(uploaded_file):
    if not uploaded_file: return ""
    file_name = uploaded_file.name.lower()
    combined = {"source_name": file_name, "pages": []}
    
    if file_name.endswith('.pdf'):
        parsed = parse_pdf_mt(uploaded_file)
    elif file_name.endswith('.docx'):
        parsed = parse_docx_mt(uploaded_file)
    else:
        # Nếu là txt
        try:
            text = uploaded_file.getvalue().decode('utf-8')
            parsed = {"source_name": file_name, "pages": [{"page_number": 1, "text": text, "images": [], "tables": []}]}
        except:
            return ""

    combined["pages"] = parsed["pages"]
    st.session_state["current_mt_metadata"] = combined
    
    source_lines = []
    for page in combined["pages"]:
        if page["text"]: source_lines.append(page['text'])
        for tab in page.get("tables", []):
            source_lines.append(f"[BẢNG DỮ LIỆU: {tab['id']}]")
        for img in page.get("images", []): 
            source_lines.append(f"[HÌNH ẢNH: {img['id']}]")
            
    return "\n\n".join(source_lines)

# ============================================================
# GỌI AI PHÂN TÍCH NGƯỢC (REVERSE ENGINEERING)
# ============================================================
def generate_matrix_ai(model_name: str, mon_hoc: str, lop: str, exam_text: str):
    system_prompt = r"""
Bạn là một Chuyên gia Khảo thí. Bạn được cung cấp nội dung của một Đề kiểm tra có sẵn.
NHIỆM VỤ CỦA BẠN: Phân tích ngược đề thi này để lập ra MA TRẬN và BẢN ĐẶC TẢ chuẩn xác theo Công văn 7991/BGDĐT-GDTrH.

1. BƯỚC 1 - QUÉT ĐỀ: Tự động gom nhóm các câu hỏi theo Chủ đề / Bài học. Nhận diện dạng câu hỏi (Nhiều lựa chọn, Đúng-Sai, Trả lời ngắn, Tự luận). Đánh giá mức độ của từng câu (Biết, Hiểu, Vận dụng).
2. BƯỚC 2 - LẬP MA TRẬN (Dùng Markdown Table):
| TT | Chủ đề/Chương | Nội dung/đơn vị kiến thức | Nhiều lựa chọn (Biết) | Nhiều lựa chọn (Hiểu) | Nhiều lựa chọn (Vận dụng) | Đúng - Sai (Biết) | Đúng - Sai (Hiểu) | Đúng - Sai (Vận dụng) | Trả lời ngắn (Biết) | Trả lời ngắn (Hiểu) | Trả lời ngắn (Vận dụng) | Tự luận (Biết) | Tự luận (Hiểu) | Tự luận (Vận dụng) | Tổng |
3. BƯỚC 3 - LẬP BẢN ĐẶC TẢ (Dùng Markdown Table):
| TT | Chủ đề/Chương | Nội dung/đơn vị kiến thức | Yêu cầu cần đạt | Số câu Nhiều lựa chọn | Số câu Đúng - Sai | Số câu Trả lời ngắn | Số câu Tự luận |

[KỶ LUẬT ĐỊNH DẠNG]
- XUẤT TRỰC TIẾP DƯỚI DẠNG VĂN BẢN MARKDOWN (Dùng Heading # cho tiêu đề). 
- KHÔNG TRẢ VỀ JSON.
- Trong các bảng, số lượng câu ghi bằng số (vd: 2). Ở cột Yêu cầu cần đạt, tóm tắt nội dung câu hỏi yêu cầu.
- Mọi công thức hóa học, toán học trong bảng phải bọc bằng dấu $...$ (VD: $H_2O$, $x^2$). TUYỆT ĐỐI KHÔNG dùng dấu backtick (`).
"""
    
    full_prompt = f"Môn học: {mon_hoc}\nLớp: {lop}\n\nNỘI DUNG ĐỀ KIỂM TRA:\n{exam_text[:40000]}"

    # Lấy API Key trực tiếp từ Session State, loại bỏ sự phụ thuộc vào ai_engine cũ
    api_key = st.session_state.get("user_api_key", "").strip() or os.environ.get("GEMINI_API_KEY", "").strip()
    try: api_key = api_key or st.secrets.get("GEMINI_API_KEY", "").strip()
    except: pass

    text_out = ""
    for attempt in range(3):
        try:
            if api_key.startswith("sk-") or "proj-" in api_key:
                import openai
                oai_client = openai.OpenAI(api_key=api_key)
                response = oai_client.chat.completions.create(
                    model="gpt-4o" if "Pro" in model_name else "gpt-4o-mini",
                    messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": full_prompt}],
                    max_tokens=8192
                )
                text_out = response.choices[0].message.content.strip()
            else:
                from google import genai
                client_genai = genai.Client(api_key=api_key or "AIzaSy_dummy")
                api_model = "gemini-2.5-pro" if "Pro" in model_name else "gemini-2.5-flash"
                response = client_genai.models.generate_content(
                    model=api_model, contents=system_prompt + "\n\n" + full_prompt
                )
                text_out = response.text.strip()
            break
        except Exception as e:
            error_msg = str(e)
            logger.error(f"Lỗi AI phân tích đề (Lần {attempt+1}): {error_msg}")
            if ("503" in error_msg or "429" in error_msg or "UNAVAILABLE" in error_msg) and attempt < 2:
                time.sleep(3)
                continue
            raise RuntimeError(f"Lỗi phân tích: {error_msg}. Vui lòng thử lại.")

    if not text_out: raise RuntimeError("AI không phản hồi.")
    text_out = re.sub(r'^```markdown\s*', '', text_out, flags=re.MULTILINE)
    text_out = re.sub(r'^```\s*$', '', text_out, flags=re.MULTILINE)
    return text_out
