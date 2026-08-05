# -*- coding: utf-8 -*-
r"""
============================================================
DATA & LOGIC: XÂY DỰNG KẾ HOẠCH BÀI DẠY
FILE: views/xd_khbd_data.py
Nâng cấp: Đa mô hình, Định tuyến chéo (Cross-Fallback) chống đứt gãy.
============================================================
"""

import streamlit as st
import os
import re
import json
import logging
import base64
from docx import Document
from io import BytesIO

logger = logging.getLogger(__name__)

NLS_GV_VAN_BAN_MAC_DINH = "Thông tư 18/2026/TT-BGDĐT"

MODE_LABELS = {
    "chinh_sua": "Chỉnh sửa và nâng cấp giáo án gốc",
    "tu_dong": "Tự động soạn từ SGK",
}

MIN_SOURCE_CHARS = 100

def init_session_state():
    defaults = {
        "khbd_mode": "tu_dong",
        "khbd_result": None,
        "khbd_nls_list": [],
        "khbd_hoat_dong_list": [],
        "khbd_processing": False,
        "khbd_nls_noi_dung": "",
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

def reset_ket_qua():
    st.session_state["khbd_result"] = None

def reset_toan_bo_khbd():
    st.session_state["khbd_result"] = None
    st.session_state["khbd_nls_list"] = []
    st.session_state["khbd_hoat_dong_list"] = []
    st.session_state["khbd_nls_noi_dung"] = ""
    st.session_state["khbd_mode"] = "tu_dong"
    st.session_state["khbd_processing"] = False

def set_mode(mode: str):
    if mode not in MODE_LABELS:
        raise ValueError(f"Chế độ soạn không hợp lệ: {mode}")
    st.session_state.khbd_mode = mode

# --- KHUNG NĂNG LỰC SỐ ---
KHUNG_NLS_GV = {
    "1. TỔ CHỨC DẠY HỌC, GIÁO DỤC TRONG MÔI TRƯỜNG SỐ": {
        "1.1. Dạy học và giáo dục trong môi trường số": {
            "Cơ bản": "Sử dụng được các chức năng và công cụ cơ bản của nền tảng quản lí học tập (LMS). Sử dụng được các công cụ hỗ trợ dạy học trực tuyến để triển khai các phiên học trực tuyến, cung cấp tài liệu học tập và tổ chức các hoạt động dạy học đơn giản.",
            "Thành thạo": "Xây dựng được kế hoạch bài dạy theo tiếp cận công nghệ. Thiết kế và triển khai được các hoạt động dạy học theo mô hình kết hợp (blended learning) hiệu quả. Lựa chọn và áp dụng được các công cụ và tài nguyên số phù hợp với mục tiêu dạy học. Hướng dẫn và triển khai được phương pháp dạy học theo dự án trực tuyến, lớp học đảo ngược trên nền tảng quản lí học tập LMS.",
            "Nâng cao": "Sáng tạo và đổi mới các mô hình dạy học ứng dụng công nghệ số. Hướng dẫn được đồng nghiệp trong việc thiết kế các trải nghiệm học tập số hóa tiên tiến."
        },
        "1.2. Hướng dẫn, hỗ trợ học tập": {
            "Cơ bản": "Sử dụng được các kênh giao tiếp số (email, diễn đàn) để trả lời câu hỏi và hỗ trợ người học khi cần thiết.",
            "Thành thạo": "Sử dụng được các kênh giao tiếp số (diễn đàn, LMS, nhóm chat...) để tương tác, giải đáp thắc mắc, cung cấp tài liệu hỗ trợ người học. Sử dụng được dữ liệu học tập số để xác định người học cần hỗ trợ và lựa chọn biện pháp can thiệp phù hợp. Thiết lập được các hoạt động, môi trường số cho phép tương tác, hỗ trợ cá nhân hóa và kịp thời cho người học.",
            "Nâng cao": "Hướng dẫn được đồng nghiệp xây dựng văn hóa hỗ trợ học tập tích cực trên nền tảng số. Sáng tạo, thử nghiệm, phát triển được các công cụ/phương pháp hỗ trợ dạy học thông minh trên nền tảng số."
        },
        "1.3. Cá nhân hóa người học": {
            "Cơ bản": "Xác định được các nhu cầu, sự khác biệt của người học trong sử dụng công nghệ số. Lựa chọn, điều chỉnh được các công cụ và nội dung số phù hợp với nhu cầu của người học.",
            "Thành thạo": "Sử dụng được công nghệ số để thiết kế lộ trình học tập linh hoạt, cho phép người học tự chủ về tốc độ, nội dung và cách thức học tập dựa trên nền tảng số. Thiết kế được các nhiệm vụ học tập phân hóa theo trình độ, năng lực và sở thích của người học trong môi trường trực tuyến. Tích hợp đa dạng, linh hoạt các công cụ số hỗ trợ hoạt động học tập thích ứng, điều chỉnh theo kết quả học tập của người học.",
            "Nâng cao": "Thiết kế được môi trường học tập số cá nhân hóa, cung cấp công cụ cập nhật cho người học tự định hướng, tự điều chỉnh quá trình học tập. Đánh giá được hiệu quả của các chiến lược dạy học cá nhân hóa bằng công nghệ."
        },
        "1.4. Học tập cộng tác": {
            "Cơ bản": "Sử dụng được công cụ số cơ bản để tổ chức cho người học làm việc nhóm đơn giản. Thiết kế được các nhiệm vụ học tập, chia sẻ tài liệu, ý tưởng trên nền tảng số.",
            "Thành thạo": "Thiết kế được các nhiệm vụ học tập yêu cầu người học sử dụng đa dạng công cụ số để cùng xây dựng nội dung, kiến thức. Hướng dẫn được người học kĩ năng giao tiếp, hợp tác hiệu quả trong môi trường số. Sử dụng được công nghệ để quản lý và đánh giá hiệu quả làm việc nhóm, đánh giá quá trình và sản phẩm cộng tác của nhóm.",
            "Nâng cao": "Xây dựng được các dự án cộng tác phức tạp, kết nối người học theo mô hình học tập cộng tác trên nền tảng số. Xây dựng và quản lí được các cộng đồng học tập trực tuyến. Hướng dẫn được đồng nghiệp triển khai học tập cộng tác dựa trên công nghệ số trong nhà trường."
        }
    },
    "2. KIỂM TRA, ĐÁNH GIÁ": {
        "2.1. Phương thức đánh giá": {
            "Cơ bản": "Sử dụng được hình thức kiểm tra, đánh giá truyền thống, có thể nhập điểm vào hệ thống số. Áp dụng được một số công cụ tạo bài kiểm tra online đơn giản trong hoạt động dạy học.",
            "Thành thạo": "Sử dụng được các công cụ số phổ biến để tạo bài kiểm tra, khảo sát nhằm đánh giá quá trình và đánh giá tổng kết. Kết hợp được một số hình thức đánh giá số đơn giản vào quá trình dạy học. Thiết kế và áp dụng được các hình thức đa dạng, công cụ đánh giá số phù hợp với mục tiêu và nội dung học tập.",
            "Nâng cao": "Sáng tạo triển khai các phương pháp, mô hình đánh giá số tiên tiến, đáp ứng yêu cầu đánh giá năng lực phức hợp. Hướng dẫn được đồng nghiệp xây dựng và áp dụng các chiến lược đánh giá số hiệu quả, công bằng trong nhà trường."
        },
        "2.2. Phân tích kết quả học tập": {
            "Cơ bản": "Sử dụng được các chức năng cơ bản của LMS/công cụ đánh giá để xem báo cáo về hoạt động, kết quả của người học.",
            "Thành thạo": "Phân tích được dữ liệu từ các hệ thống đánh giá số để nhận diện quá trình tiến bộ và thành tích học tập của người học. Sử dụng được dữ liệu, công cụ trực quan hóa dữ liệu để xây dựng báo cáo đánh giá sự tiến bộ của người học. Xây dựng được các bảng điều khiển tự động (dashboard) dữ liệu học tập trực quan.",
            "Nâng cao": "Áp dụng được các kĩ thuật phân tích dữ liệu học tập nâng cao để dự đoán xu hướng, phát hiện sớm các vấn đề và đề xuất can thiệp. Hướng dẫn được đồng nghiệp về cách khai thác và diễn giải dữ liệu để cải tiến dạy học."
        },
        "2.3. Phản hồi và đánh giá cải tiến": {
            "Cơ bản": "Sử dụng được các chức năng cung cấp phản hồi trên hệ thống LMS. Cung cấp phản hồi kịp thời cho người học bằng văn bản hoặc điểm số thông qua các nền tảng số.",
            "Thành thạo": "Sử dụng được đa dạng công cụ số (ghi âm, video ngắn, bình luận trực tiếp trên tài liệu) để đưa ra phản hồi chi tiết, kịp thời. Thiết kế được các qui trình phản hồi và đánh giá cải tiến có sự tham gia của người học (tự đánh giá, đánh giá chéo) bằng công nghệ.",
            "Nâng cao": "Sử dụng được dữ liệu phân tích để điều chỉnh kế hoạch bài dạy, phương pháp và cung cấp nhiệm vụ học tập hỗ trợ cá nhân người học. Hướng dẫn được đồng nghiệp sử dụng phản hồi bằng công cụ số và dữ liệu học tập để cải tiến liên tục chương trình và hoạt động giáo dục."
        }
    },
    "3. TRAO QUYỀN CHO NGƯỜI HỌC": {
        "3.1. Tiếp cận và hòa nhập": {
            "Cơ bản": "Sử dụng được các công cụ số cơ bản để hỗ trợ người học gặp khó khăn trong học tập. Lựa chọn và sử dụng được các công cụ, tài nguyên số có tính đến sự đa dạng của người học (đa dạng ngôn ngữ, phong cách học, người học khuyết tật). Đảm bảo mọi người học có cơ hội sử dụng thiết bị, hạ tầng số của nhà trường khi cần thiết.",
            "Thành thạo": "Khai thác, lựa chọn và điều chỉnh được tài nguyên, đa dạng hóa công cụ số để đáp ứng nhu cầu đặc biệt của người học. Thiết kế được nội dung, tài nguyên số đảm bảo tính tiếp cận và hòa nhập trong môi trường số.",
            "Nâng cao": "Hướng dẫn được đồng nghiệp về chiến lược và công nghệ số hỗ trợ giáo dục hòa nhập."
        },
        "3.2. Giải quyết vấn đề": {
            "Cơ bản": "Thiết kế được các nhiệm vụ học tập yêu cầu người học sử dụng Internet để tìm kiếm thông tin để trả lời câu hỏi hoặc giải quyết vấn đề học tập đơn giản.",
            "Thành thạo": "Thiết kế được các nhiệm vụ dự án học tập yêu cầu người học sử dụng công nghệ số để xác định vấn đề, thu thập, phân tích thông tin và đề xuất giải pháp. Tổ chức được các hoạt động học tập dựa trên vấn đề (problem-based) hoặc dự án (project-based) phức tạp, trong đó công nghệ số là công cụ thiết yếu để nghiên cứu, hợp tác và tạo ra sản phẩm.",
            "Nâng cao": "Hướng dẫn được đồng nghiệp xây dựng hệ sinh thái học tập số, kết nối người học với các chuyên gia và vấn đề thực tiễn bên ngoài nhà trường để giải quyết các vấn đề thực tế của cộng đồng."
        },
        "3.3. Khuyến khích sự tham gia tích cực của người học": {
            "Cơ bản": "Sáng tạo và điều phối tương tác số đơn giản để thu hút sự chú ý và khuyến khích người học tham gia vào hoạt động học tập. Tích hợp được công nghệ số trong dạy học nhằm trực quan hóa và tăng hiệu quả trình bày nội dung dạy học.",
            "Thành thạo": "Tích hợp được các yếu tố trò chơi hóa, tương tác và các công cụ sáng tạo nội dung để thúc đẩy người học chủ động, tích cực tham gia vào bài học. Thiết kế được hoạt động khuyến khích người học tự tạo ra nội dung số, chia sẻ kiến thức thông qua các nền tảng số, giải quyết vấn đề bằng mô phỏng, thí nghiệm ảo, thực tế ảo, thực tế ảo tăng cường.",
            "Nâng cao": "Sử dụng được các công cụ để thiết kế môi trường học tập số năng động, lấy người học làm trung tâm. Hướng dẫn được đồng nghiệp sáng tạo triển khai các hoạt động học tập tích cực bằng công nghệ số."
        }
    },
    "4. KĨ NĂNG CÔNG NGHỆ SỐ": {
        "4.1. Kĩ năng thông tin và dữ liệu": {
            "Cơ bản": "Sử dụng được công cụ tìm kiếm để tìm thông tin, tài liệu phục vụ bài giảng. Lưu trữ và sắp xếp một cách khoa học các dữ liệu trên máy tính hoặc đám mây.",
            "Thành thạo": "Đánh giá được độ tin cậy của nguồn tin trên Internet, mạng xã hội. Sử dụng được các kĩ thuật tìm kiếm nâng cao. Hướng dẫn được người học các kĩ năng tư duy phản biện khi tìm kiếm, xử lí, tiếp nhận thông tin số từ các nguồn khác nhau. Tổ chức được các nhiệm vụ học tập nâng cao cho phép người học chủ động tìm kiếm và xử lí thông tin trong môi trường số.",
            "Nâng cao": "Sử dụng được công cụ để thu thập và trực quan hóa dữ liệu đơn giản, phân tích và đánh giá độ tin cậy của thông tin trong quá trình dạy học. Hướng dẫn được đồng nghiệp tích hợp phát triển năng lực thông tin vào chương trình dạy học."
        },
        "4.2. Sáng tạo nội dung số": {
            "Cơ bản": "Sử dụng được các công cụ số phổ biến để tạo nội dung dạy học theo định dạng số khác nhau. Tích hợp được các định dạng số trong nội dung thực hiện nhiệm vụ của người học.",
            "Thành thạo": "Tích hợp được công nghệ số trong hoạt động sáng tạo nội dung số, xây dựng kho học liệu số. Hướng dẫn được người học sử dụng các công cụ cơ bản để tạo nội dung số, thực hiện quyền tác giả, giấy phép, cách trích dẫn, sử dụng và chia sẻ tài nguyên số hợp pháp. Sử dụng được nền tảng, công cụ số đa dạng để tạo và chia sẻ nội dung hợp lệ.",
            "Nâng cao": "Sử dụng thành thạo các công cụ chuyên dụng để tạo ra các học liệu số có tính tương tác cao. Hướng dẫn được đồng nghiệp phát triển nền tảng học tập tích hợp AI, thực tế ảo, thực tế ảo tăng cường trong sáng tạo nội dung số vào các môn học."
        },
        "4.3. An toàn": {
            "Cơ bản": "Có hiểu biết về bảo vệ sức khỏe thể chất, tinh thần, đảm bảo an sinh số trong hoạt động dạy học. Bố trí, sắp xếp được không gian, thời gian sử dụng thiết bị, công cụ số hợp lí cho người học. Nhận diện và xử lí được các tình huống bắt nạt trực tuyến.",
            "Thành thạo": "Tích hợp được kiến thức, kĩ năng nhận diện và phòng tránh các rủi ro phổ biến trên mạng trong quá trình dạy học. Áp dụng được các biện pháp bảo vệ dữ liệu cá nhân và của người học. Thực hiện được các biện pháp cơ bản đảm bảo an toàn thiết bị, tài khoản trong lớp học và hướng dẫn người học cách bảo vệ dữ liệu cá nhân, định danh số, quản lí dấu vết số.",
            "Nâng cao": "Thực hiện được các biện pháp đảm bảo an toàn sức khỏe, áp dụng các phương pháp dạy học giảm căng thăng trong môi trường số cho người học. Cập nhật và phổ biến các xu hướng, mối đe dọa mới và cách phòng chống cho cộng đồng giáo viên, phụ huynh. Hướng dẫn được đồng nghiệp xây dựng môi trường học tập số an toàn, lành mạnh trong lớp học và nhà trường."
        }
    },
    "5. PHÁT TRIỂN CHUYÊN MÔN": {
        "5.1. Giao tiếp trong tổ chức": {
            "Cơ bản": "Sử dụng được email, nhóm chat của trường/tổ để trao đổi thông tin công việc và giao tiếp với phụ huynh.",
            "Thành thạo": "Sử dụng hiệu quả các kênh giao tiếp số chính thức của trường để tương tác với các bên liên quan, phù hợp với từng đối tượng và mục đích giáo dục. Sử dụng được công cụ số cơ bản để giao tiếp, chia sẻ thông tin, dữ liệu và tham gia hoạt động chuyên môn với đồng nghiệp.",
            "Nâng cao": "Xây dựng và quản lí được chiến lược truyền thông số, các kênh truyền thông số chính thức của trường để chia sẻ và kết nối cộng đồng. Hướng dẫn được đồng nghiệp đổi mới cách thức giao tiếp trong tổ chức bằng công nghệ số tăng cường tính minh bạch, sự tham gia của các bên liên quan."
        },
        "5.2. Hợp tác phát triển chuyên môn": {
            "Cơ bản": "Chủ động tham gia các cộng đồng học tập trực tuyến. Tự đánh giá được khó khăn, thách thức và thuận lợi ứng dụng công nghệ số trong công việc.",
            "Thành thạo": "Xây dựng được kế hoạch cải tiến, đổi mới ứng dụng công nghệ số trong hoạt động chuyên môn. Chủ động tìm kiếm và tham gia các khóa học cơ bản để cập nhật kiến thức, kĩ năng số. Tự đánh giá, cải tiến ứng dụng công nghệ số trong dạy học trong môi trường số. Tham gia chia sẻ, học tập và cập nhật kĩ năng ứng dụng công nghệ số với đồng nghiệp.",
            "Nâng cao": "Cập nhật được các xu hướng công nghệ và phương pháp sư phạm số mới, áp dụng kiến thức, kĩ năng số vào thực tiễn dạy học. Hướng dẫn được đồng nghiệp xây dựng các yêu cầu về dạy học trong môi trường số, phát triển các công cụ, phương pháp hỗ trợ tự phản ánh về năng lực số."
        },
        "5.3. Phát triển, sử dụng, chia sẻ và quản lí học liệu số": {
            "Cơ bản": "Sử dụng được các công cụ tìm kiếm phổ biến để tìm kiếm tài nguyên, kho học liệu số, thư viện trực tuyến, tài nguyên giáo dục mở (OER). Lựa chọn được tài nguyên phù hợp với mục tiêu bài học.",
            "Thành thạo": "Lựa chọn và sử dụng được tài nguyên, học liệu số phù hợp với đối tượng đa dạng của người học. Tạo được tài nguyên số phục vụ cho môn học dựa từ các nguồn có sẵn. Tổ chức lưu trữ, quản lí và chia sẻ được kho học liệu số cá nhân một cách khoa học, an toàn. Đánh giá được chất lượng, độ tin cậy, tính pháp lí và sư phạm của tài nguyên, học liệu số.",
            "Nâng cao": "Xây dựng và quản trị được các hệ thống quản lí, chia sẻ tài nguyên số cho tổ/trường. Hướng dẫn được đồng nghiệp xây dựng và quản trị kho học liệu số mở của nhà trường."
        }
    },
    "6. ỨNG DỤNG TRÍ TUỆ NHÂN TẠO (AI)": {
        "6.1. Tư duy lấy con người làm trung tâm": {
            "Cơ bản": "Nhận diện được cách vận hành của AI và các công nghệ có tích hợp AI.",
            "Thành thạo": "Thiết kế được các hoạt động dạy học tích hợp AI một cách sáng tạo và có trách nhiệm.",
            "Nâng cao": "Triển khai đổi mới phương pháp dạy học mới có tích hợp sâu AI đáp ứng cá nhân hóa và dạy học thích ứng. Hướng dẫn, lựa chọn và đề xuất sử dụng các công cụ AI phù hợp cho đồng nghiệp."
        },
        "6.2. Đạo đức AI": {
            "Cơ bản": "Nhận diện được các khả năng tích hợp sử dụng AI trong hỗ trợ hoạt động dạy học. Sử dụng được các công cụ AI đơn giản (chủ yếu là AI tạo sinh) để hỗ trợ dạy học và kiểm tra đánh giá. Nhận diện được khả năng thu thập dữ liệu và thông tin cá nhân khi sử dụng công cụ AI, những tiềm ẩn rủi ro khi sử dụng AI không đúng cách. Thể hiện sự cẩn trọng và có trách nhiệm đối với quyền riêng tư của người học, có trách nhiệm khi sử dụng công cụ AI trong dạy học, kiểm tra đánh giá. Thiết kế các hoạt động giáo dục tích hợp AI, cân bằng giữa tương tác công nghệ và tương tác xã hội, phát triển tư duy phản biện.",
            "Thành thạo": "Khai thác hiệu quả các công cụ AI chuyên biệt cho giáo dục để tạo học liệu số tương tác đa dạng, cá nhân hóa một phần nội dung/bài tập, hỗ trợ chấm điểm tự động. Hướng dẫn được người học sử dụng AI có trách nhiệm, nhận biết ưu/nhược điểm và các rủi ro liên quan khi sử dụng AI. Thực hiện được các biện pháp cần thiết phòng ngừa rủi ro và về các vấn đề đạo đức cơ bản khi sử dụng AI. Lựa chọn, đánh giá được các ứng dụng AI dựa trên tiêu chí về đạo đức, chính sách bảo mật, sự công bằng trong tiếp cận và tác động khác trong dạy học, kiểm tra đánh giá. Thiết kế và tích hợp hoạt động hướng dẫn người học sử dụng AI an toàn và có đạo đức trong hoạt động học tập.",
            "Nâng cao": "Đánh giá được ưu/nhược điểm và các vấn đề đạo đức của công cụ AI trong giáo dục, cập nhật, hướng dẫn và chia sẻ với đồng nghiệp các vấn đề về đạo đức sử dụng AI. Tham gia xây dựng chính sách, hướng dẫn về sử dụng AI có đạo đức trong nhà trường."
        },
        "6.3. Sư phạm AI": {
            "Cơ bản": "Nhận diện được khả năng tích hợp AI theo hướng cá nhân hóa và lấy người học làm trung tâm. Có hiểu biết về lợi ích sư phạm của công cụ AI để hỗ trợ dạy học.",
            "Thành thạo": "Ứng dụng được công cụ AI linh hoạt trong các bước dạy học đảm bảo nguyên tắc lấy người học làm trung tâm, giáo dục hòa nhập. Lựa chọn và ứng dụng được các hệ thống, công cụ AI phù hợp, giảm thiểu rủi ro trong thiết kế dạy học, kiểm tra đánh giá. Tổ chức và quản lý được hoạt động tương tác 3 chiều giữa giáo viên, người học với các công cụ AI trong dạy học.",
            "Nâng cao": "Xây dựng được nguyên tắc sư phạm sử dụng AI trong hoạt động dạy học. Hướng dẫn được đồng nghiệp thiết kế và sử dụng AI theo tiếp cận đồng sáng tạo, lấy con người làm trung tâm trong các hoạt động sư phạm."
        },
        "6.4. AI cho phát triển chuyên môn": {
            "Cơ bản": "Nhận diện được sự cân bằng giữa vai trò của người giáo viên và nhiệm vụ phát triển năng lực số, năng lực AI trong dạy học. Sử dụng được công cụ AI phù hợp để lập kế hoạch, theo dõi và phân tích quá trình phát triển chuyên môn của bản thân.",
            "Thành thạo": "Sử dụng được các công cụ AI đơn giản (chủ yếu là AI tạo sinh) để hỗ trợ học tập suốt đời và phát triển chuyên môn nghiệp vụ bản thân. Đề xuất được các hướng sử dụng hiệu quả các nền tảng AI để tìm kiếm tài nguyên, tham gia cộng đồng thực hành hỗ trợ phát triển bản thân. Đánh giá được các rủi ro đạo đức từ các nền tảng AI và triển khai các biện pháp phòng ngừa giảm thiểu tác động tiêu cực.",
            "Nâng cao": "Hướng dẫn đổi mới sáng tạo cho đồng nghiệp dựa trên các nền tảng AI phù hợp và tiếp cận sư phạm số. Xây dựng hoặc sử dụng được các bộ công cụ AI tạo sinh hỗ trợ phát triển chuyên môn của đồng nghiệp."
        }
    }
}

KHUNG_NLS_HS = {
    "I. Khai thác dữ liệu và thông tin": {
        "1.1. Duyệt, tìm kiếm và lọc dữ liệu, thông tin và nội dung số": "Xác định được nhu cầu thông tin, tìm kiếm được dữ liệu, thông tin và nội dung trong môi trường số, truy cập chúng và khai thác được kết quả tìm kiếm. Tạo và cập nhật được chiến lược tìm kiếm.",
        "1.2. Đánh giá dữ liệu, thông tin và nội dung số": "Phân tích, so sánh và đánh giá được độ tin cậy và tính xác thực của nguồn dữ liệu, thông tin và nội dung số. Phân tích, giải thích và đánh giá được dữ liệu, thông tin và nội dung số.",
        "1.3. Quản lý dữ liệu, thông tin và nội dung số": "Tổ chức, lưu trữ và truy xuất được dữ liệu, thông tin và nội dung trong môi trường số. Tổ chức và sắp xếp được chúng trong một môi trường có cấu trúc."
    },
    "II. Giao tiếp và hợp tác trong môi trường số": {
        "2.1. Tương tác thông qua công nghệ số": "Tương tác thông qua các công nghệ số khác nhau và nhận biết được phương tiện giao tiếp số nào phù hợp cho một bối cảnh cụ thể.",
        "2.2. Chia sẻ thông tin và nội dung thông qua công nghệ số": "Chia sẻ dữ liệu, thông tin và nội dung số với người khác thông qua các công nghệ số phù hợp. Đóng vai trò là người trung gian, hiểu biết về thực hành trích dẫn và ghi chú nguồn.",
        "2.3. Sử dụng công nghệ số để thực hiện trách nhiệm công dân": "Tham gia đóng góp cho xã hội thông qua việc sử dụng các dịch vụ số công và tư. Tìm kiếm được cơ hội, để trao quyền và thu hút công dân thông qua các công nghệ số phù hợp.",
        "2.4. Hợp tác thông qua công nghệ số": "Sử dụng được các công cụ và công nghệ số cho các quá trình hợp tác cũng như để cùng xây dựng và đồng sáng tạo dữ liệu, tài nguyên và kiến thức.",
        "2.5. Thực hiện quy tắc ứng xử trên mạng": "Nhận thức được các chuẩn mực hành vi và kiến thức khi sử dụng công nghệ số và tương tác trong môi trường số. Điều chỉnh được các chiến lược giao tiếp phù hợp với đối tượng cụ thể và nhận thức được sự đa dạng về văn hóa và thế hệ trong môi trường số.",
        "2.6. Quản lý danh tính số": "Tạo và quản lý được một hoặc nhiều danh tính số để bảo vệ danh tiếng của bản thân, làm việc với dữ liệu mà một người tạo ra bằng nhiều công cụ, môi trường và dịch vụ số."
    },
    "III. Sáng tạo nội dung số": {
        "3.1. Phát triển nội dung số": "Tạo và chỉnh sửa được nội dung số ở các định dạng khác nhau, nhằm thể hiện bản thân thông qua các phương tiện số.",
        "3.2. Tích hợp và tạo lập lại nội dung số": "Sửa đổi, tinh chỉnh và tích hợp được thông tin và nội dung mới vào khối kiến thức và tài nguyên hiện có để tạo ra nội dung và kiến thức mới, độc đáo và phù hợp.",
        "3.3. Thực thi bản quyền và giấy phép": "Hiểu được cách áp dụng bản quyền và giấy phép cho thông tin và nội dung số.",
        "3.4. Lập trình": "Lập kế hoạch và phát triển được một chuỗi các câu lệnh dễ hiểu cho một hệ thống máy tính để giải quyết một vấn đề nhất định hoặc thực hiện một nhiệm vụ cụ thể."
    },
    "IV. An toàn": {
        "4.1. Bảo vệ thiết bị": "Bảo vệ được các thiết bị và nội dung số cũng như hiểu rõ các rủi ro và mối đe dọa trong môi trường kỹ thuật số. Biết được các biện pháp an toàn và bảo mật cũng như có sự quan tâm đúng mức đến độ tin cậy và quyền riêng tư.",
        "4.2. Bảo vệ dữ liệu cá nhân và quyền riêng tư": "Bảo vệ được dữ liệu cá nhân và quyền riêng tư trong môi trường số. Hiểu được cách sử dụng và chia sẻ thông tin định danh cá nhân một cách an toàn, có khả năng bảo vệ bản thân và người khác. Hiểu được cách các dịch vụ số sử dụng “Chính sách Quyền riêng tư để thông báo phương thức sử dụng dữ liệu cá nhân.",
        "4.3. Bảo vệ sức khỏe và an sinh số": "Tránh được rủi ro và đe dọa đến sức khỏe thể chất và tinh thần khi sử dụng công nghệ số. Bảo vệ được bản thân và người khác khỏi nguy cơ trong môi trường số (ví dụ: bắt nạt trên mạng). Nhận biết được những công nghệ số giúp tăng cường thịnh vượng xã hội và sự hòa hợp trong xã hội.",
        "4.4. Bảo vệ môi trường": "Nhận thức được tác động của công nghệ số và việc sử dụng công nghệ số đối với môi trường."
    },
    "V. Giải quyết vấn đề": {
        "5.1. Giải quyết các vấn đề kỹ thuật": "Xác định được các vấn đề kỹ thuật khi vận hành thiết bị, sử dụng môi trường số và giải quyết chúng (từ xử lý sự cố đến giải quyết các vấn đề phức tạp hơn).",
        "5.2. Xác định nhu cầu và giải pháp công nghệ": "Đánh giá được nhu cầu và xác định, đánh giá, lựa chọn, sử dụng các công cụ số cùng với các giải pháp công nghệ khả thi để giải quyết chúng. Điều chỉnh và tùy chỉnh được môi trường số theo nhu cầu cá nhân.",
        "5.3. Sử dụng sáng tạo công nghệ số": "Sử dụng được các công cụ và công nghệ số để tạo ra kiến thức, đổi mới quy trình và sản phẩm. Gắn kết cá nhân và tập thể vào quá trình xử lý nhận thức để hiểu và giải quyết các vấn đề mang tính khái niệm và các tình huống có vấn đề trong môi trường số.",
        "5.4. Xác định các vấn đề cần cải thiện về năng lực số": "Hiểu được năng lực số của chính mình cần được cải thiện hoặc cập nhật ở đâu. Có thể hỗ trợ người khác phát triển năng lực số của họ. Tìm kiếm được cơ hội phát triển bản thân và cập nhật sự phát triển số."
    },
    "VI. Ứng dụng trí tuệ nhân tạo": {
        "6.1. Hiểu biết về AI (trong đó có Gen AI)": "Hiểu được cách AI ảnh hưởng đến cuộc sống hằng ngày và vai trò của AI trong các lĩnh vực khác nhau. Nắm vững được nguyên tắc hoạt động của AI, khả năng và hạn chế của AI.",
        "6.2. Sử dụng AI có đạo đức và trách nhiệm": "Sử dụng hiệu quả các hệ thống AI và hiểu rõ ứng dụng thực tế của chúng. Sử dụng được AI để tạo nội dung, khám phá kiến thức và giải quyết các vấn đề trong công việc, học tập và cuộc sống hàng ngày.",
        "6.3. Đánh giá các công cụ AI": "Đánh giá và lọc được thông tin từ các nguồn được tạo ra hoặc xử lý bằng AI, để hiểu rõ hơn về tính đáng tin cậy và cách sử dụng thông tin đó. Đánh giá được AI trên các khía cạnh minh bạch, an toàn, đạo đức và tác động."
    }
}

def get_nls_framework(loai_khung): 
    return KHUNG_NLS_GV if loai_khung == "Giáo viên (Thông tư 18)" else KHUNG_NLS_HS

def get_nls_domains(loai_khung): 
    return list(get_nls_framework(loai_khung).keys())

def get_nls_components(loai_khung, linh_vuc): 
    return list(get_nls_framework(loai_khung).get(linh_vuc, {}).keys())

def get_nls_levels(loai_khung, linh_vuc, thanh_phan):
    try:
        data_thanh_phan = get_nls_framework(loai_khung).get(linh_vuc, {}).get(thanh_phan, {})
        if isinstance(data_thanh_phan, str):
            return ["Chuẩn TT 02/2025"]
        if isinstance(data_thanh_phan, dict):
            return list(data_thanh_phan.keys())
        return []
    except Exception:
        return []

def get_nls_content(loai_khung, linh_vuc, thanh_phan, muc_do):
    try:
        data_thanh_phan = get_nls_framework(loai_khung).get(linh_vuc, {}).get(thanh_phan, {})
        if isinstance(data_thanh_phan, str):
            return data_thanh_phan
        if isinstance(data_thanh_phan, dict):
            return data_thanh_phan.get(muc_do, "")
        return ""
    except Exception:
        return ""

def add_nls():
    linh_vuc = safe_text(st.session_state.get("khbd_nls_linh_vuc", ""))
    thanh_phan = safe_text(st.session_state.get("khbd_nls_thanh_phan", ""))
    muc_do = safe_text(st.session_state.get("khbd_nls_muc_do", ""))
    noi_dung = safe_text(st.session_state.get("khbd_nls_noi_dung", ""))
    
    if not noi_dung: 
        return
        
    van_ban = NLS_GV_VAN_BAN_MAC_DINH if st.session_state.get("khbd_loai_khung_nls") == "Giáo viên (Thông tư 18)" else "Khung DigComp"
    item = {"van_ban": van_ban, "linh_vuc": linh_vuc, "thanh_phan": thanh_phan, "muc_do": muc_do, "noi_dung": noi_dung}
    if item not in st.session_state.khbd_nls_list: 
        st.session_state.khbd_nls_list.append(item)

def format_nls():
    items = st.session_state.get("khbd_nls_list", [])
    if not items:
        return "Không có yêu cầu đặc thù về Năng lực số."
    return "\n".join([f"- Năng lực {item['linh_vuc']} > {item['thanh_phan']} ({item['muc_do']}): {item['noi_dung']}" for item in items])

def safe_text(value):
    if value is None: 
        return ""
    text = str(value).replace("\x00", "").replace("\ufeff", "").replace("\u200b", "")
    text = text.replace("\r", "").replace("\t", " ")
    return re.sub(r"[ ]{2,}", " ", text).strip()

def diagnose_source_quality(text, source_name="Tài liệu nguồn"):
    text = safe_text(text)
    if len(text) < MIN_SOURCE_CHARS:
        return {"status": "insufficient", "message": f"{source_name} quá ngắn hoặc không đọc được chữ."}
    return {"status": "valid", "message": "Dữ liệu hợp lệ."}

def parse_pdf_structured(uploaded_file, range_str=""):
    if uploaded_file is None:
        return {"source_name": "unknown", "pages": []}
    content = uploaded_file.getvalue() if hasattr(uploaded_file, "getvalue") else uploaded_file.read()
    file_name = getattr(uploaded_file, 'name', 'document.pdf')
    pages_data = []
    try:
        import fitz
        doc = fitz.open(stream=content, filetype="pdf")
        target_pages = set()
        if range_str.strip():
            for part in range_str.split(','):
                if '-' in part:
                    try:
                        start, end = map(int, part.split('-'))
                        for p in range(start, end + 1): 
                            target_pages.add(p - 1)
                    except: 
                        pass
                else:
                    try: 
                        target_pages.add(int(part.strip()) - 1)
                    except: 
                        pass
        for i in range(len(doc)):
            if target_pages and i not in target_pages: 
                continue
            page = doc[i]
            page_text = safe_text(page.get_text("text"))
            
            page_images = []
            for img_idx, img in enumerate(page.get_images(full=True)):
                try:
                    base_image = doc.extract_image(img[0])
                    page_images.append({
                        "id": f"IMG_P{i+1}_{img_idx+1}",
                        "page": i + 1,
                        "ext": base_image["ext"],
                        "base64": base64.b64encode(base_image["image"]).decode("utf-8")
                    })
                except: 
                    pass
                
            page_tables = []
            try:
                tabs = page.find_tables()
                if tabs and tabs.tables:
                    for t_idx, tab in enumerate(tabs.tables):
                        extracted_df = tab.extract()
                        if extracted_df and len(extracted_df) > 0:
                            page_tables.append({
                                "id": f"TAB_P{i+1}_{t_idx+1}",
                                "page": i + 1,
                                "headers": [str(c) if c is not None else "" for c in extracted_df[0]],
                                "rows": [[str(c) if c is not None else "" for c in r] for r in extracted_df[1:]]
                            })
            except: 
                pass
            
            pages_data.append({"page_number": i + 1, "text": page_text, "images": page_images, "tables": page_tables, "figures": [], "charts": []})
    except Exception as e:
        logger.error(f"Lỗi đọc cấu trúc PDF: {e}")
    return {"source_name": file_name, "pages": pages_data}

def parse_docx_structured(uploaded_file):
    if uploaded_file is None: 
        return {"source_name": "unknown", "pages": []}
    file_name = getattr(uploaded_file, 'name', 'document.docx')
    paragraphs_text = []
    tables_data = []
    try:
        source = uploaded_file.getvalue() if hasattr(uploaded_file, "getvalue") else uploaded_file.read()
        doc = Document(BytesIO(source))
        for p in doc.paragraphs:
            if p.text.strip(): 
                paragraphs_text.append(p.text.strip())
        for t_idx, table in enumerate(doc.tables):
            rows_raw = [[cell.text.strip().replace('\n', ' ') for cell in row.cells] for row in table.rows]
            if rows_raw:
                tables_data.append({"id": f"TAB_DOCX_{t_idx+1}", "page": 1, "headers": rows_raw[0], "rows": rows_raw[1:] if len(rows_raw) > 1 else []})
    except Exception as e:
        logger.error(f"Lỗi đọc cấu trúc DOCX: {e}")
    return {"source_name": file_name, "pages": [{"page_number": 1, "text": "\n".join(paragraphs_text), "images": [], "tables": tables_data, "figures": [], "charts": []}]}

def build_intermediate_knowledge_source(structured_data):
    if not structured_data or not structured_data.get("pages"): 
        return "Không có dữ liệu nguồn."
    source_lines = []
    for page in structured_data["pages"]:
        p_num = page["page_number"]
        source_lines.append(f"=== TRANG {p_num} ===")
        if page["text"]: 
            source_lines.append(f"[TEXT - Trang {p_num}]\n{page['text']}")
        for tab in page.get("tables", []):
            source_lines.append(f"BẢNG DỮ LIỆU [TABLE: {tab['id']}]\nHeaders: {' | '.join(tab['headers'])}\nRows:\n" + "\n".join([' | '.join(r) for r in tab['rows']]))
        for img in page.get("images", []): 
            source_lines.append(f"HÌNH MINH HỌA [IMAGE: {img['id']}]")
    return "\n\n".join(source_lines)

def read_multiple_files(files, range_str="", is_pdf_target=False):
    combined = {"source_name": "multi_files", "pages": []}
    offset = 0
    for f in files or []:
        if f.name.lower().endswith('.pdf'):
            parsed = parse_pdf_structured(f, range_str) 
        else:
            parsed = parse_docx_structured(f)
        for p in parsed["pages"]:
            p["page_number"] += offset
            combined["pages"].append(p)
        offset += len(parsed["pages"])
        
    st.session_state["current_source_metadata"] = combined
    return build_intermediate_knowledge_source(combined)

def generate_ai(ai_engine, prompt, model_name="Gemini 1.5 Flash (Tốc độ)"):
    """
    Định tuyến đa mô hình (Cross-Routing).
    - Ưu tiên OpenAI nếu tên mô hình có "GPT".
    - Ưu tiên Gemini nếu tên mô hình có "Gemini".
    - Tự động nhảy sang nền tảng đối lập nếu gặp lỗi (Quota/429).
    """
    system_instruction = r"""
[KỶ LUẬT THÉP XÂY DỰNG KHBD CHUẨN CÔNG VĂN 5512]:

1. BẮT BUỘC TUÂN THỦ CẤU TRÚC 4 HOẠT ĐỘNG:
Bạn phải giữ nguyên cấu trúc chuẩn sau đây:
- Hoạt động 1: Mở đầu
- Hoạt động 2: Hình thành kiến thức mới
- Hoạt động 3: Luyện tập
- Hoạt động 4: Vận dụng
*Lưu ý: Tại mỗi hoạt động bắt buộc phải có đủ 4 mục: a) Mục tiêu, b) Nội dung, c) Sản phẩm, d) Tổ chức thực hiện.

2. QUY ĐỊNH VỀ CÔNG THỨC TOÁN HỌC (NẾU CÓ):
- CẤM TUYỆT ĐỐI dùng dấu backtick (`) cho công thức Toán, Lý, Hóa. 
- BẮT BUỘC dùng dấu $...$ cho TẤT CẢ công thức Toán, Lý, Hóa.

3. QUY TRÌNH "TỔ CHỨC THỰC HIỆN" (Giao - Thực - Báo - Kết):
Tại phần d) Tổ chức thực hiện của mỗi hoạt động, trình bày rõ 4 bước: Giao nhiệm vụ, Thực hiện, Báo cáo thảo luận, Kết luận nhận định.
"""
    full_prompt = system_instruction + "\n\n" + prompt

    is_openai_preferred = "GPT" in model_name
    is_pro_tier = "Pro" in model_name or "Cao cấp" in model_name or "GPT-4o (" in model_name

    openai_model = "gpt-4o" if is_pro_tier else "gpt-4o-mini"
    
    def call_openai():
        api_key = None
        for key, val in st.session_state.items():
            if isinstance(val, str) and val.startswith("sk-"):
                api_key = val
                break
        if not api_key:
            for k in ["user_api_key", "api_key", "openai_api_key", "sk_key"]:
                if st.session_state.get(k) and str(st.session_state.get(k)).startswith("sk-"):
                    api_key = st.session_state.get(k)
                    break
        if not api_key and "OPENAI_API_KEY" in st.secrets:
            api_key = st.secrets["OPENAI_API_KEY"]

        if api_key:
            import openai
            client = openai.OpenAI(api_key=str(api_key).strip())
            response = client.chat.completions.create(
                model=openai_model,
                messages=[
                    {"role": "system", "content": system_instruction},
                    {"role": "user", "content": prompt}
                ]
            )
            return response.choices[0].message.content.strip()
        raise RuntimeError("Không tìm thấy khóa API OpenAI (sk-).")

    def call_gemini():
        try:
            from utils.ai_engine_2 import AIEngine2
            gemini_model = "gemini-2.5-pro" if is_pro_tier else "gemini-2.5-flash"
            engine_v2 = AIEngine2(default_model=gemini_model)
            res = engine_v2.generate_text(full_prompt)
            if res and not res.startswith("❌") and not res.startswith("⚠️") and "429" not in res and "RESOURCE_EXHAUSTED" not in res:
                return res
        except ImportError:
            if ai_engine and hasattr(ai_engine, "generate_text"):
                res = ai_engine.generate_text(full_prompt)
                if res and "429" not in res and "RESOURCE_EXHAUSTED" not in res and not res.startswith("❌"):
                    if isinstance(res, dict): return res.get("text", str(res))
                    elif hasattr(res, "text"): return res.text
                    return res
        raise RuntimeError("Gemini bị quá tải hoặc lỗi 429.")

    text_out = None
    if is_openai_preferred:
        try:
            text_out = call_openai()
        except Exception:
            text_out = call_gemini()
    else:
        try:
            text_out = call_gemini()
        except Exception:
            text_out = call_openai()

    return process_output_format(text_out)

def process_output_format(text_out):
    if not text_out: 
        raise RuntimeError("Không thể nhận phản hồi từ AI.")
        
    if "# TÊN BÀI HỌC:" in text_out: 
        text_out = text_out[text_out.find("# TÊN BÀI HỌC:"):]
        
    text_out = text_out.replace("**", "")
    text_out = re.sub(r'([^\n])\s*([a-d]\)\s+)', r'\1\n\n\2', text_out)
    text_out = re.sub(r'([^\n])\s*(\*(?:Chuyển giao nhiệm vụ học tập|Thực hiện nhiệm vụ học tập|Báo cáo kết quả và thảo luận|Kết luận)[^:\n]*:)', r'\1\n\n\2', text_out)
    
    return text_out

def validate_khbd_result(text):
    if not text or len(text) < 500: 
        return False, "Nội dung giáo án quá ngắn hoặc trống."
    return True, "Hợp lệ"

def build_prompt(thong_tin, noi_dung_chinh, noi_dung_ga, nls_str, tich_hop_ai, tich_hop_hoa_nhap, nhu_cau_hoa_nhap, mode, so_tiet):
    source = safe_text(noi_dung_chinh)[:35000]
    ga_block = f"--- GIÁO ÁN CŨ ĐỂ CHỈNH SỬA ---\n{safe_text(noi_dung_ga)[:10000]}\n" if mode == "chinh_sua" else ""
    return f"--- THÔNG TIN CHUNG ---\n{thong_tin}\n\nNHIỆM VỤ: SOẠN KẾ HOẠCH BÀI DẠY {so_tiet} TIẾT BÁM SÁT NGUỒN SAU:\n\n--- NGUỒN KIẾN THỨC TRUNG GIAN (CHỨA TEXT, ẢNH VÀ BẢNG) ---\n{source}\n\n{ga_block}"
