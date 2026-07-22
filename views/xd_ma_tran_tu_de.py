# -*- coding: utf-8 -*-

import streamlit as st
import json
import re
import pandas as pd


# ============================================================
# IMPORT EXPORT WORD - PIPELINE CHUẨN
# ============================================================

try:

    from export.export_word import export_word

except ImportError as e:

    export_word = None

    st.error(
        "❌ Không import được hệ thống xuất Word.\n\n"
        f"Chi tiết: {e}"
    )


# ============================================================
# SERVICE 1: ĐỌC VÀ TRÍCH XUẤT VĂN BẢN
# ============================================================

class ExamTextExtractor:

    @staticmethod
    def extract(uploaded_file):

        if not uploaded_file:

            return ""

        file_name = uploaded_file.name.lower()

        file_bytes = uploaded_file.getvalue()

        try:

            # ====================================================
            # PDF
            # ====================================================

            if file_name.endswith(".pdf"):

                from pypdf import PdfReader

                reader = PdfReader(
                    uploaded_file
                )

                pages = []

                for page in reader.pages:

                    text = page.extract_text()

                    if text:

                        pages.append(
                            text.strip()
                        )

                return "\n".join(pages)


            # ====================================================
            # DOCX
            # ====================================================

            if file_name.endswith(".docx"):

                from docx import Document

                doc = Document(
                    uploaded_file
                )

                result = []

                # -------------------------------
                # ĐOẠN VĂN
                # -------------------------------

                for paragraph in doc.paragraphs:

                    text = paragraph.text.strip()

                    if text:

                        result.append(text)

                # -------------------------------
                # BẢNG
                # -------------------------------

                for table_index, table in enumerate(
                    doc.tables,
                    start=1
                ):

                    result.append(
                        f"\n[BẢNG {table_index}]"
                    )

                    for row in table.rows:

                        cells = []

                        for cell in row.cells:

                            cell_text = (
                                cell.text
                                .replace(
                                    "\n",
                                    " "
                                )
                                .strip()
                            )

                            cells.append(
                                cell_text
                            )

                        row_text = " | ".join(
                            cells
                        )

                        if row_text.strip():

                            result.append(
                                row_text
                            )

                return "\n".join(result)


            # ====================================================
            # TXT
            # ====================================================

            if file_name.endswith(".txt"):

                for encoding in (
                    "utf-8",
                    "utf-8-sig",
                    "cp1258"
                ):

                    try:

                        return file_bytes.decode(
                            encoding
                        ).strip()

                    except UnicodeDecodeError:

                        continue

                raise ValueError(
                    "Không thể giải mã tệp TXT."
                )


            raise ValueError(
                f"Định dạng không được hỗ trợ: "
                f"{file_name}"
            )


        except Exception as e:

            raise RuntimeError(
                f"Lỗi đọc tệp {file_name}: {e}"
            )


    # ========================================================
    # CHUẨN HÓA
    # ========================================================

    @staticmethod
    def normalize(text):

        if not text:

            return ""

        text = re.sub(
            r"\r\n?",
            "\n",
            text
        )

        text = re.sub(
            r"[ \t]+",
            " ",
            text
        )

        text = re.sub(
            r"\n{3,}",
            "\n\n",
            text
        )

        return text.strip()[:120000]


# ============================================================
# SERVICE 2: XỬ LÝ JSON
# ============================================================

class MatrixCalculator:

    # ========================================================
    # PARSE JSON AI
    # ========================================================

    @staticmethod
    def parse_ai_json(result_text):

        if not result_text:

            raise ValueError(
                "AI không trả về dữ liệu."
            )

        text = result_text.strip()

        # ------------------------------------
        # Loại markdown code fence
        # ------------------------------------

        text = re.sub(
            r"^```json\s*",
            "",
            text,
            flags=re.IGNORECASE
        )

        text = re.sub(
            r"^```\s*",
            "",
            text
        )

        text = re.sub(
            r"\s*```$",
            "",
            text
        )

        text = text.strip()

        # ------------------------------------
        # Tách JSON nếu AI có văn bản thừa
        # ------------------------------------

        if not text.startswith("{"):

            start = text.find("{")

            end = text.rfind("}")

            if start == -1 or end == -1:

                raise ValueError(
                    "Không tìm thấy JSON hợp lệ."
                )

            text = text[
                start:end + 1
            ]

        try:

            data = json.loads(text)

        except json.JSONDecodeError as e:

            raise ValueError(
                f"JSON không hợp lệ: {e}"
            )

        if not isinstance(data, dict):

            raise ValueError(
                "JSON phải là một object."
            )

        return data


    # ========================================================
    # CHUYỂN SỐ
    # ========================================================

    @staticmethod
    def to_number(value):

        if value is None:

            return 0

        if isinstance(
            value,
            bool
        ):

            return int(value)

        if isinstance(
            value,
            (int, float)
        ):

            return value

        if isinstance(
            value,
            str
        ):

            value = value.strip()

            if not value:

                return 0

            value = value.replace(
                ",",
                "."
            )

            try:

                number = float(
                    value
                )

                if number.is_integer():

                    return int(number)

                return number

            except ValueError:

                return 0

        return 0


    # ========================================================
    # CHUẨN HÓA MA TRẬN
    # ========================================================

    @staticmethod
    def normalize_matrix(
        rows
    ):

        result = []

        if not isinstance(
            rows,
            list
        ):

            return result

        for item in rows:

            if not isinstance(
                item,
                dict
            ):

                continue

            nb = MatrixCalculator.to_number(
                item.get(
                    "nb",
                    0
                )
            )

            th = MatrixCalculator.to_number(
                item.get(
                    "th",
                    0
                )
            )

            vd = MatrixCalculator.to_number(
                item.get(
                    "vd",
                    0
                )
            )

            vdc = MatrixCalculator.to_number(
                item.get(
                    "vdc",
                    0
                )
            )

            tong_so_cau = (
                nb
                + th
                + vd
                + vdc
            )

            result.append({

                "chu_de": str(
                    item.get(
                        "chu_de",
                        ""
                    )
                ).strip(),

                "noi_dung": str(
                    item.get(
                        "noi_dung",
                        ""
                    )
                ).strip(),

                "nb": nb,

                "th": th,

                "vd": vd,

                "vdc": vdc,

                "tong_so_cau": tong_so_cau,

                "tong_diem":
                    MatrixCalculator.to_number(
                        item.get(
                            "tong_diem",
                            0
                        )
                    )

            })

        return result


    # ========================================================
    # CHUẨN HÓA ĐẶC TẢ
    # ========================================================

    @staticmethod
    def normalize_specification(
        rows
    ):

        result = []

        if not isinstance(
            rows,
            list
        ):

            return result

        for index, item in enumerate(
            rows,
            start=1
        ):

            if not isinstance(
                item,
                dict
            ):

                continue

            result.append({

                "stt": item.get(
                    "stt",
                    index
                ),

                "chu_de": str(
                    item.get(
                        "chu_de",
                        ""
                    )
                ).strip(),

                "bai_hoc": str(
                    item.get(
                        "bai_hoc",
                        item.get(
                            "noi_dung",
                            ""
                        )
                    )
                ).strip(),

                "yccd": str(
                    item.get(
                        "yccd",
                        ""
                    )
                ).strip(),

                "tn_nb":
                    MatrixCalculator.to_number(
                        item.get(
                            "tn_nb",
                            0
                        )
                    ),

                "tn_hieu":
                    MatrixCalculator.to_number(
                        item.get(
                            "tn_hieu",
                            0
                        )
                    ),

                "tn_vd":
                    MatrixCalculator.to_number(
                        item.get(
                            "tn_vd",
                            0
                        )
                    ),

                "ds_nb":
                    MatrixCalculator.to_number(
                        item.get(
                            "ds_nb",
                            0
                        )
                    ),

                "ds_hieu":
                    MatrixCalculator.to_number(
                        item.get(
                            "ds_hieu",
                            0
                        )
                    ),

                "ds_vd":
                    MatrixCalculator.to_number(
                        item.get(
                            "ds_vd",
                            0
                        )
                    ),

                "tl_biet":
                    MatrixCalculator.to_number(
                        item.get(
                            "tl_biet",
                            0
                        )
                    ),

                "tl_hieu":
                    MatrixCalculator.to_number(
                        item.get(
                            "tl_hieu",
                            0
                        )
                    ),

                "tl_vd":
                    MatrixCalculator.to_number(
                        item.get(
                            "tl_vd",
                            0
                        )
                    ),

                "tong_diem":
                    MatrixCalculator.to_number(
                        item.get(
                            "tong_diem",
                            0
                        )
                    )

            })

        return result


    # ========================================================
    # CONTEXT CHUẨN
    # ========================================================

    @staticmethod
    def prepare_context(
        parsed_data,
        mon_hoc,
        lop
    ):

        if not isinstance(
            parsed_data,
            dict
        ):

            raise ValueError(
                "Dữ liệu AI không đúng cấu trúc."
            )

        ma_tran = MatrixCalculator.normalize_matrix(
            parsed_data.get(
                "ma_tran",
                []
            )
        )

        dac_ta = MatrixCalculator.normalize_specification(
            parsed_data.get(
                "dac_ta",
                []
            )
        )

        return {

            # ------------------------------------
            # Thông tin chung
            # ------------------------------------

            "MON_HOC": mon_hoc,

            "LOP": lop,

            # ------------------------------------
            # Dữ liệu chính
            # ------------------------------------

            "MA_TRAN": ma_tran,

            "DAC_TA": dac_ta,

            # ------------------------------------
            # Alias tương thích
            # ------------------------------------

            "ma_tran_data": ma_tran,

            "dac_ta_data": dac_ta

        }


# ============================================================
# VIEW CHÍNH
# ============================================================

def render_xd_ma_tran_tu_de(
    ai_engine
):

    st.markdown(
        "### 🧩 Sinh Ma trận & Đặc tả Đề kiểm tra"
    )

    # ========================================================
    # THÔNG TIN
    # ========================================================

    col1, col2 = st.columns(2)

    with col1:

        mon_hoc = st.selectbox(

            "Môn học",

            [

                "Khoa học Tự nhiên",

                "Toán học",

                "Ngữ văn",

                "Ngoại ngữ",

                "Khác"

            ],

            key="mt_mon"

        )

    with col2:

        lop = st.selectbox(

            "Lớp",

            [

                "Lớp 6",

                "Lớp 7",

                "Lớp 8",

                "Lớp 9",

                "Lớp 10",

                "Lớp 11",

                "Lớp 12"

            ],

            index=2,

            key="mt_lop"

        )


    # ========================================================
    # UPLOAD
    # ========================================================

    file_de = st.file_uploader(

        "Tải lên đề kiểm tra",

        type=[

            "pdf",

            "docx",

            "txt"

        ],

        key="mt_file"

    )


    # ========================================================
    # PHÂN TÍCH
    # ========================================================

    if st.button(

        "PHÂN TÍCH ĐỀ & LẬP MA TRẬN",

        type="primary",

        use_container_width=True

    ):

        if not file_de:

            st.warning(
                "Vui lòng tải lên đề kiểm tra."
            )

            return

        try:

            with st.spinner(
                "AI đang phân tích đề kiểm tra..."
            ):

                # ------------------------------------
                # Đọc đề
                # ------------------------------------

                raw_text = (
                    ExamTextExtractor.extract(
                        file_de
                    )
                )

                exam_text = (
                    ExamTextExtractor.normalize(
                        raw_text
                    )
                )

                if not exam_text:

                    raise ValueError(
                        "Không đọc được nội dung đề."
                    )


                # ------------------------------------
                # Schema
                # ------------------------------------

                json_schema = """

{
  "ma_tran": [
    {
      "chu_de": "Tên chủ đề",
      "noi_dung": "Tên bài học",
      "nb": 0,
      "th": 0,
      "vd": 0,
      "vdc": 0,
      "tong_diem": 0
    }
  ],
  "dac_ta": [
    {
      "stt": 1,
      "chu_de": "Tên chủ đề",
      "bai_hoc": "Tên bài học",
      "yccd": "Yêu cầu cần đạt",
      "tn_nb": 0,
      "tn_hieu": 0,
      "tn_vd": 0,
      "ds_nb": 0,
      "ds_hieu": 0,
      "ds_vd": 0,
      "tl_biet": 0,
      "tl_hieu": 0,
      "tl_vd": 0,
      "tong_diem": 0
    }
  ]
}

"""


                # ------------------------------------
                # Prompt
                # ------------------------------------

                prompt = f"""

BẠN LÀ CHUYÊN GIA KHẢO THÍ
THEO CHƯƠNG TRÌNH GDPT 2018.

Hãy phân tích đề kiểm tra
môn {mon_hoc},
{lop}.

MỤC TIÊU:

1. Xác định toàn bộ nội dung kiến thức
   thực sự xuất hiện trong đề.

2. Phân tích đầy đủ từng câu hỏi.

3. Phân loại đúng mức độ:

- nb: Nhận biết
- th: Thông hiểu
- vd: Vận dụng
- vdc: Vận dụng cao

4. Không được bỏ sót câu hỏi.

5. Không được thêm nội dung
   không xuất hiện trong đề.

6. Không suy đoán chương,
   bài học hoặc kiến thức
   nếu đề không đủ căn cứ.

7. Chỉ trả về JSON hợp lệ.

MA TRẬN:

- Mỗi dòng tương ứng với một
  nội dung kiến thức thực tế.

- tong_so_cau =
  nb + th + vd + vdc.

- tong_diem là tổng điểm
  của các câu trong dòng đó.

ĐẶC TẢ:

- yccd phải cụ thể.
- Phải bám sát câu hỏi.
- Phải thể hiện đúng mức độ nhận thức.
- Không viết chung chung.

ĐỀ KIỂM TRA:

------------------------------

{exam_text}

------------------------------

JSON BẮT BUỘC:

{json_schema}

CHỈ TRẢ VỀ JSON.
"""


                # ------------------------------------
                # Gọi AI
                # ------------------------------------

                result = ai_engine.generate_text(
                    prompt
                )


                # ------------------------------------
                # Parse
                # ------------------------------------

                parsed_data = (
                    MatrixCalculator.parse_ai_json(
                        result
                    )
                )


                # ------------------------------------
                # Chuẩn hóa
                # ------------------------------------

                context = (
                    MatrixCalculator.prepare_context(
                        parsed_data,
                        mon_hoc,
                        lop
                    )
                )


                # ------------------------------------
                # LƯU
                # ------------------------------------

                st.session_state[
                    "processed_matrix_data"
                ] = context


                st.session_state[
                    "mt_mon_hoc_file"
                ] = mon_hoc


                st.session_state[
                    "mt_lop_file"
                ] = lop


                # ------------------------------------
                # XUẤT WORD
                # ------------------------------------

                if export_word is None:

                    raise ImportError(
                        "export_word chưa được import."
                    )


                word_bytes = export_word(

                    data=context,

                    template_name=(
                        "ma_tran_dac_ta_mau.docx"
                    )

                )


                st.session_state[
                    "download_word_bytes"
                ] = word_bytes


                st.success(
                    "🎉 Đã phân tích và xuất Word thành công."
                )


        except json.JSONDecodeError:

            st.error(
                "❌ AI trả về JSON không hợp lệ."
            )


        except Exception as e:

            st.error(
                f"❌ Lỗi xử lý: {e}"
            )


    # ========================================================
    # HIỂN THỊ KẾT QUẢ
    # ========================================================

    if (
        "processed_matrix_data"
        in st.session_state
    ):

        data = st.session_state[
            "processed_matrix_data"
        ]

        st.divider()

        st.markdown(
            "#### 👁️ Xem trước dữ liệu"
        )


        # ====================================================
        # MA TRẬN
        # ====================================================

        if data.get(
            "MA_TRAN"
        ):

            st.markdown(
                "**1. MA TRẬN**"
            )

            st.dataframe(

                pd.DataFrame(
                    data[
                        "MA_TRAN"
                    ]
                ),

                use_container_width=True

            )


        # ====================================================
        # ĐẶC TẢ
        # ====================================================

        if data.get(
            "DAC_TA"
        ):

            st.markdown(
                "**2. ĐẶC TẢ**"
            )

            st.dataframe(

                pd.DataFrame(
                    data[
                        "DAC_TA"
                    ]
                ),

                use_container_width=True

            )


        # ====================================================
        # DOWNLOAD
        # ====================================================

        safe_mon = st.session_state.get(

            "mt_mon_hoc_file",

            "Mon"

        )

        safe_lop = st.session_state.get(

            "mt_lop_file",

            "Lop"

        )


        st.download_button(

            label=(
                "📥 TẢI XUỐNG FILE WORD"
            ),

            data=st.session_state[
                "download_word_bytes"
            ],

            file_name=(

                f"Ma_tran_Dac_ta_"
                f"{safe_mon}_"
                f"{safe_lop}.docx"

            ),

            mime=(
                "application/vnd.openxmlformats-"
                "officedocument.wordprocessingml.document"
            ),

            use_container_width=True,

            type="primary"

        )
