# -*- coding: utf-8 -*-
"""
app.py - Bộ công cụ đọc tệp và xử lý dữ liệu đề cương.
"""
import streamlit as st
import sys
from pathlib import Path
from io import BytesIO
import re

def extract_text_from_file(uploaded_file):
    """Bóc tách văn bản thô từ file PDF, Word hoặc TXT tải lên, chống trùng bảng."""
    if uploaded_file is None:
        return ""
    try:
        file_name = uploaded_file.name.lower()
        file_bytes = uploaded_file.getvalue()
        if not file_bytes:
            return ""
            
        # 1. Xử lý tệp định dạng PDF
        if file_name.endswith(".pdf"):
            from pypdf import PdfReader
            reader = PdfReader(BytesIO(file_bytes))
            pages = []
            for page_number, page in enumerate(reader.pages, start=1):
                text = page.extract_text()
                if text and text.strip():
                    pages.append(f"\n--- TRANG {page_number} ---\n{text.strip()}")
            return "\n\n".join(pages).strip()
            
        # 2. Xử lý tệp định dạng DOCX (Chống lặp nội dung ô bảng)
        elif file_name.endswith(".docx"):
            from docx import Document
            document = Document(BytesIO(file_bytes))
            contents = []
            seen_texts = set()
            
            # Đọc Paragraphs (văn bản nằm ngoài bảng)
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if text and text not in seen_texts:
                    contents.append(text)
                    seen_texts.add(text)
                    
            # Đọc Tables (dữ liệu nằm gọn trong ô bảng)
            for table in document.tables:
                for row in table.rows:
                    row_data = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    row_text = " | ".join(filter(None, row_data))
                    if row_text.strip() and row_text not in seen_texts:
                        contents.append(row_text)
                        seen_texts.add(row_text)
            return "\n".join(contents).strip()
            
        # 3. Xử lý tệp định dạng văn bản thuần TXT
        elif file_name.endswith(".txt"):
            for encoding in ["utf-8", "utf-8-sig", "cp1258", "latin-1"]:
                try:
                    return file_bytes.decode(encoding).strip()
                except Exception:
                    continue
            return ""
    except Exception as e:
        st.error(f"❌ Lỗi đọc tài liệu {uploaded_file.name}: {e}")
        return ""
    return ""

def normalize_outline(text):
    """Chuẩn hóa dữ liệu chữ và giới hạn an toàn 6,000 từ để chống tràn Token."""
    if not text:
        return ""
    clean_text = re.sub(r"\s+", " ", text).strip()
    words = clean_text.split(" ")
    safe_text = " ".join(words[:6000])
    
    if len(words) > 6000:
        st.warning(f"⚠️ Dữ liệu tải lên rất dài ({len(words):,} từ). AI đã tự động lọc 6,000 từ để bảo vệ hệ thống.")
    return safe_text

def render_xd_de_kt(ai_engine):
    """Hiển thị giao diện soạn thảo cấu hình đề kiểm tra chuẩn 5512."""
    st.markdown("### Soạn thảo Ma trận, Đặc tả & Đề KT (Chuẩn 5512)")
    
    # --- KHỐI THÔNG TIN CHUNG ---
    c1, c2, c3, c4, c5, c6 = st.columns([1, 0.8, 1.2, 1, 2, 0.8])
    mon_hoc = c1.selectbox(
        "Môn", 
        [
            "Toán học", "Ngữ văn", "Ngoại ngữ", "Khoa học Tự nhiên", 
            "Lịch sử và Địa lý", "Lịch sử", "Địa lý", "Vật lý", "Hóa học", 
            "Sinh học", "Giáo dục công dân", "Giáo dục kinh tế và pháp luật", 
            "Tin học", "Công nghệ", "Giáo dục thể chất", "Khác"
        ], 
        key="de_kt_mon_hoc"
    )
    lop = c2.selectbox("Lớp", ["Lớp 6", "Lớp 7", "Lớp 8", "Lớp 9", "Lớp 10", "Lớp 11", "Lớp 12"], index=2, key="de_kt_lop")
    hinh_thuc = c3.selectbox("Hình thức", ["Trắc nghiệm & Tự luận", "100% Trắc nghiệm", "100% Tự luận"], key="de_kt_hinh_thuc")
    thoi_gian = c4.selectbox("Thời gian", ["15 phút", "45 phút", "60 phút", "90 phút", "120 phút"], index=3, key="de_kt_thoi_gian")
    ten_de = c5.text_input("Tên bài kiểm tra", key="de_kt_ten_de")
    with c6:
        st.markdown("<div style='margin-top: 25px;'></div>", unsafe_allow_html=True)
        bam_sat = st.checkbox("Bám sát đề cương", value=True, key="de_kt_bam_sat")
        
    # --- KHỐI TẢI NHIỀU TÀI LIỆU CÙNG LÚC ---
    files_de = st.file_uploader("📥 Tải đề cương / tài liệu (Có thể chọn nhiều file cùng lúc)", type=["pdf", "docx", "txt"], accept_multiple_files=True, key="de_kt_file_de_cuong")
    
    # --- KHỐI THIẾT LẬP TỶ LỆ NHẬN THỨC ---
    with st.expander("📊 Cấu hình Tỷ lệ & Số câu", expanded=True):
        r1, r2, r3, r4 = st.columns(4)
        nb = r1.number_input("Nhận biết (%)", min_value=0, max_value=100, value=40, step=5, key="de_kt_nb")
        th = r2.number_input("Thông hiểu (%)", min_value=0, max_value=100, value=30, step=5, key="de_kt_th")
        vd = r3.number_input("Vận dụng (%)", min_value=0, max_value=100, value=20, step=5, key="de_kt_vd")
        vdc = r4.number_input("Vận dụng cao (%)", min_value=0, max_value=100, value=10, step=5, key="de_kt_vdc")
        
        tong_ty_le = nb + th + vd + vdc
        if tong_ty_le != 100:
            st.warning(f"⚠️ Tổng tỷ lệ mức độ hiện tại là {tong_ty_le}%. Phải bằng 100%.")
            
        # --- CẤU TRÚC ĐỀ TRẮC NGHIỆM ---
        st.markdown("#### Cấu trúc các dạng câu hỏi")
        cols = st.columns(8)
        n_nlc = cols[0].number_input("NLC", min_value=0, value=10, key="de_kt_n_nlc")
        d_nlc = cols[1].number_input("Đ.NLC", min_value=0.0, value=0.25, step=0.25, key="de_kt_d_nlc")
        n_ds = cols[2].number_input("Đ/S", min_value=0, value=2, key="de_kt_n_ds")
        d_ds = cols[3].number_input("Đ.Đ/S", min_value=0.0, value=0.25, step=0.25, key="de_kt_d_ds")
        n_dk = cols[4].number_input("Điền K", min_value=0, value=2, key="de_kt_n_dk")
        d_dk = cols[5].number_input("Đ.DK", min_value=0.0, value=0.25, step=0.25, key="de_kt_d_dk")
        n_ngan = cols[6].number_input("TL Ngắn", min_value=0, value=2, key="de_kt_n_ngan")
        d_ngan = cols[7].number_input("Đ.TLN", min_value=0.0, value=0.50, step=0.25, key="de_kt_d_ngan")
        
        total_diem_tn = (n_nlc * d_nlc) + (n_ds * d_ds) + (n_dk * d_dk) + (n_ngan * d_ngan)
        
        # --- CẤU TRÚC ĐỀ TỰ LUẬN ---
        st.markdown("#### PHẦN TỰ LUẬN")
        num_tl = st.number_input("Số câu Tự luận", min_value=0, max_value=10, value=3, key="de_kt_num_tl")
        
        tl_points = []
        if num_tl > 0:
            rows_needed = (num_tl + 3) // 4
            for r in range(rows_needed):
                tl_cols = st.columns(4)
                for c in range(4):
                    idx = r * 4 + c
                    if idx < num_tl:
                        p = tl_cols[c].number_input(f"Câu {idx + 1} (đ)", min_value=0.0, value=2.0, step=0.25, key=f"de_kt_tl_p_{idx}")
                        tl_points.append(p)
                        
        total_diem_tl = sum(tl_points)
        total_diem = total_diem_tn + total_diem_tl
        
        st.markdown("---")
        res_cols = st.columns(3)
        res_cols[0].metric("Tổng điểm Trắc nghiệm", f"{total_diem_tn:.2f}")
        res_cols[1].metric("Tổng điểm Tự luận", f"{total_diem_tl:.2f}")
        res_cols[2].metric("TỔNG ĐIỂM ĐỀ", f"{total_diem:.2f} / 10")

    # --- KHỐI XỬ LÝ KÍCH HOẠT VÀ RENDER DỮ LIỆU ĐẦU RA ---
    if st.button("🚀 TẠO MA TRẬN & ĐỀ KIỂM TRA", type="primary", use_container_width=True, key="de_kt_btn_generate"):
        if tong_ty_le != 100:
            st.error("❌ Tổng tỷ lệ Nhận biết + Thông hiểu + Vận dụng + Vận dụng cao phải bằng 100%.")
            st.stop()
        if bam_sat and not files_de:
            st.error("❌ Thầy/Cô đã tích chọn 'Bám sát đề cương' nhưng chưa tải lên tệp căn cứ nào.")
            st.stop()
        if abs(total_diem - 10.0) > 0.01:
            st.error(f"❌ Tổng điểm thiết lập hiện tại là {total_diem:.2f}/10. Vui lòng điều chỉnh lại cấu hình số câu/điểm.")
            st.stop()
            
        with st.spinner("⏳ Hệ thống đang bóc tách và xử lý dữ liệu từ (các) tài liệu tải lên..."):
            if bam_sat and files_de:
                raw_outline = ""
                for f in files_de:
                    raw_outline += f"\n--- TÀI LIỆU: {f.name} ---\n"
                    raw_outline += extract_text_from_file(f)
                    
                outline_text = normalize_outline(raw_outline)
                if not outline_text:
                    st.error("❌ Không trích xuất được nội dung chữ từ các file đã tải lên.")
                    st.stop()
            else:
                outline_text = "Không cung cấp đề cương. AI tự động bám sát chương trình GDPT 2018 theo Môn học và Lớp."
        
        # --- ĐỊNH VỊ SỐ THỨ TỰ CÂU BẰNG PYTHON ---
        tong_cau_tn = n_nlc + n_ds + n_dk + n_ngan
        idx_nlc_start, idx_nlc_end = 1, n_nlc
        idx_ds_start, idx_ds_end = idx_nlc_end + 1, idx_nlc_end + n_ds
        idx_dk_start, idx_dk_end = idx_ds_end + 1, idx_ds_end + n_dk
        idx_ngan_start, idx_ngan_end = idx_dk_end + 1, idx_dk_end + n_ngan
        idx_tl_start = tong_cau_tn + 1
        
        chi_tiet_tu_luan = ""
        for i, p in enumerate(tl_points):
            chi_tiet_tu_luan += f"├── Câu {idx_tl_start + i} = {p} điểm\n"
            
        # --- XÂY DỰNG RÀNG BUỘC PHẲNG ---
        # --- XÂY DỰNG RÀNG BUỘC PHẲNG ---
        base_prompt = (
            "BẠN LÀ CHUYÊN GIA BIÊN SOẠN ĐỀ KIỂM TRA THEO CHUẨN GDPT 2018.\n"
            "NHIỆM VỤ: Soạn thảo Ma trận, Đặc tả, Đề kiểm tra và Đáp án. Bạn bắt buộc phải trả về ĐỊNH DẠNG VĂN BẢN MARKDOWN THUẦN TÚY (Tuyệt đối không bọc trong JSON hay Code Block).\n\n"
            
            # ... (Các phần trên giữ nguyên) ...

            "============================================================\n"
            "QUY TẮC NGHIÊM NGẶT (VI PHẠM LÀ LỖI HỆ THỐNG)\n"
            "============================================================\n"
            "1. SỐ THỨ TỰ CÂU HỎI: Phải liên tục từ Câu 1 đến Câu {idx_tl_end}. Tuyệt đối KHÔNG đánh số lại từ Câu 1 khi chuyển sang phần Tự luận.\n"
            
            # --- ĐÂY LÀ ĐOẠN ĐÃ ĐƯỢC NÂNG CẤP ĐỂ SỬA LỖI WORD ---
            "2. CÔNG THỨC TOÁN HỌC VÀ KÝ HIỆU (QUAN TRỌNG):\n"
            "   - MỌI biểu thức, số liệu, lũy thừa (VD: a^2), đơn vị (VD: cm^2) BẮT BUỘC phải bọc trong cặp dấu $. (Ví dụ đúng: $S = a^2 = 16\\text{{ cm}}^2$).\n"
            "   - Áp dụng quy tắc bọc dấu $ này cho CẢ PHẦN ĐỀ BÀI LẪN PHẦN ĐÁP ÁN.\n"
            "   - TUYỆT ĐỐI KHÔNG để lộ mã thô như a^2 hay cm^2 ra ngoài văn bản thường.\n"
            "   - Đối với ký hiệu SONG SONG, TUYỆT ĐỐI KHÔNG dùng mã \\parallel. Bắt buộc phải dùng ký hiệu // thông thường (Ví dụ viết: $AB // CD$).\n"
            # ---------------------------------------------------

            "3. XỬ LÝ HÌNH HỌC VÀ HÌNH VẼ (ĐẶC BIỆT LƯU Ý):\n"
            # ... (Các phần dưới giữ nguyên) ...
            "============================================================\n"
            "QUY TẮC NGHIÊM NGẶT (VI PHẠM LÀ LỖI HỆ THỐNG)\n"
            "============================================================\n"
            "1. SỐ THỨ TỰ CÂU HỎI: Phải liên tục từ Câu 1 đến Câu {idx_tl_end}. Tuyệt đối KHÔNG đánh số lại từ Câu 1 khi chuyển sang phần Tự luận.\n"
            "2. CÔNG THỨC TOÁN HỌC (LaTeX): BẮT BUỘC bọc biểu thức trong cặp dấu $. (Ví dụ: $y = x^2 + 2$, $\\Delta$).\n"
            "3. XỬ LÝ HÌNH HỌC VÀ HÌNH VẼ (ĐẶC BIỆT LƯU Ý):\n"
            "   - HỆ THỐNG KHÔNG THỂ XUẤT ẢNH. Tuyệt đối KHÔNG dùng các cụm từ như 'Cho hình vẽ bên', 'Theo hình vẽ' trong Đề thi.\n"
            "   - Nếu trong ĐỀ CƯƠNG có câu hỏi dạng 'Cho hình vẽ', BẠN PHẢI TỰ ĐỘNG CHUYỂN ĐỔI nó thành một bài toán mô tả hoàn toàn bằng chữ (Ví dụ: Cho tam giác ABC vuông tại A, có đường cao AH...).\n"
            "   - NẾU BẮT BUỘC phải cần hình minh họa để giải bài, hãy chèn DÒNG MÃ SAU dưới câu hỏi để giáo viên tự chèn hình: `[GIÁO VIÊN CHÈN HÌNH VẼ TẠI ĐÂY: <Mô tả chi tiết bằng chữ để giáo viên dễ vẽ>]`.\n\n"
            "TRÌNH BÀY ĐẦU RA BẰNG VĂN BẢN THEO ĐÚNG CÁC TIÊU ĐỀ SAU:\n"
            "# PHẦN I. PHẠM VI KIẾN THỨC SỬ DỤNG\n"
            "# PHẦN II. MA TRẬN ĐỀ KIỂM TRA\n"
            "# PHẦN III. BẢN ĐẶC TẢ\n"
            "# PHẦN IV. ĐỀ KIỂM TRA\n"
            "# PHẦN V. ĐÁP ÁN VÀ HƯỚNG DẪN CHẤM\n\n"
            "============================================================\n"
            "TẬP HỢP TÀI LIỆU VÀ ĐỀ CƯƠNG KIẾN THỨC ĐƯỢC PHÉP DÙNG\n"
            "============================================================\n"
            "{outline_text}"
        )
        
        # Đã bổ sung đẩy đủ các biến idx_ngan_start và idx_ngan_end
        strict_prompt = base_prompt.format(
            mon_hoc=mon_hoc, lop=lop, ten_de=ten_de, thoi_gian=thoi_gian, nb=nb, th=th, vd=vd, vdc=vdc,
            tong_cau_tn=tong_cau_tn, total_diem_tn=total_diem_tn, n_nlc=n_nlc, d_nlc=d_nlc,
            idx_nlc_start=idx_nlc_start, idx_nlc_end=idx_nlc_end, n_ds=n_ds, d_ds=d_ds,
            idx_ds_start=idx_ds_start, idx_ds_end=idx_ds_end, n_dk=n_dk, d_dk=d_dk,
            idx_dk_start=idx_dk_start, idx_dk_end=idx_dk_end, 
            n_ngan=n_ngan, d_ngan=d_ngan, idx_ngan_start=idx_ngan_start, idx_ngan_end=idx_ngan_end,
            num_tl=num_tl, total_diem_tl=total_diem_tl, chi_tiet_tu_luan=chi_tiet_tu_luan.strip(),
            idx_tl_end=idx_tl_start + num_tl - 1, outline_text=outline_text
        )
        
        with st.spinner("🤖 AI đang phân tích dữ liệu đa luồng, soạn thảo ma trận và định dạng đề thi..."):
            try:
                # Gọi đến hàm xử lý văn bản thuần túy của Engine
                result = ai_engine.generate_text(strict_prompt)
                if not result or not result.strip():
                    st.error("❌ AI trả về kết quả rỗng.")
                    st.stop()
                    
                st.session_state["de_kt_content"] = result
                st.session_state["de_kt_config"] = {
                    "mon_hoc": mon_hoc, "lop": lop, "ten_de": ten_de,
                    "hinh_thuc": hinh_thuc, "thoi_gian": thoi_gian,
                    "tong_diem": total_diem, "bam_sat": bam_sat
                }
                st.success("✅ Đã xử lý xong toàn bộ tài liệu! Đề kiểm tra đã được tối ưu hóa.")
                st.rerun()
            except Exception as e:
                st.error(f"❌ Lỗi sinh đề: {e}")

    # --- KHỐI HIỂN THỊ KẾT QUẢ VÀ KẾT XUẤT TÀI LIỆU WORD ---
    if "de_kt_content" in st.session_state:
        st.divider()
        st.markdown("## KẾT QUẢ ĐỀ KIỂM TRA")
        if st.button("🗑️ XÓA ĐỀ HIỆN TẠI", key="de_kt_delete"):
            st.session_state.pop("de_kt_content", None)
            st.session_state.pop("de_kt_config", None)
            st.rerun()
            
        st.markdown(st.session_state["de_kt_content"])
        
        try:
            root_path = str(Path(__file__).resolve().parents[2])
            if root_path not in sys.path:
                sys.path.insert(0, root_path)
                
            from export.export_word import WordExportEngine
            config = st.session_state.get("de_kt_config", {})
            word_bytes = WordExportEngine.export_to_word({
                "ai_generated_content": st.session_state["de_kt_content"],
                "is_de_kt": True,
                "title": config.get("ten_de", "Đề kiểm tra")
            })
            st.download_button("📥 TẢI XUỐNG FILE WORD (.DOCX)", data=word_bytes, file_name="De_Thi_5512.docx", use_container_width=True, key="de_kt_download_word")
        except Exception as e:
            st.warning(f"⚠️ Tính năng xuất bản Word gặp sự cố hoặc chưa cấu hình thư viện export_word: {e}")
