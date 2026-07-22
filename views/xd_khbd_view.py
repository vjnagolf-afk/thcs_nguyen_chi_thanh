# -*- coding: utf-8 -*-

import streamlit as st
import PyPDF2
import docx
import pandas as pd
import io
import os
import re
import json
import base64
import traceback
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# =========================================================
# 0. IMPORT TÙY CHỌN CHO CÔNG THỨC WORD
# =========================================================

try:
    import latex2mathml.converter
    LATEX2MATHML_AVAILABLE = True
except ImportError:
    LATEX2MATHML_AVAILABLE = False


try:
    from lxml import etree
    LXML_AVAILABLE = True
except ImportError:
    LXML_AVAILABLE = False


# =========================================================
# 1. SESSION STATE
# =========================================================

def init_session_state():
    """Khởi tạo toàn bộ trạng thái của module."""

    defaults = {
        "hoat_dong_list": [],
        "soan_mode": "chinh_sua",
        "nls_list": [],
        "ket_qua_giao_an": None,
        "khbd_last_error": None,
        "khbd_processing": False,
    }

    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


def reset_ket_qua():
    """Xóa kết quả cũ trước khi xử lý lại."""

    st.session_state["ket_qua_giao_an"] = None
    st.session_state["khbd_last_error"] = None


def set_mode(mode: str):
    st.session_state.soan_mode = mode


def add_hoat_dong():
    new_hd = st.session_state.get("new_hoat_dong", "").strip()

    if new_hd and new_hd not in st.session_state.hoat_dong_list:
        st.session_state.hoat_dong_list.append(new_hd)

    st.session_state["new_hoat_dong"] = ""


def add_nls_item():
    tp = st.session_state.get("nls_tp", "")
    md = st.session_state.get("nls_md", "")
    nd = st.session_state.get("nls_nd_input", "").strip()

    if nd:
        st.session_state.nls_list.append({
            "thanh_phan": tp,
            "muc_do": md,
            "noi_dung": nd
        })

    st.session_state["nls_nd_input"] = ""


# =========================================================
# 2. HÀM TIỆN ÍCH XỬ LÝ FILE
# =========================================================

def _get_file_extension(uploaded_file) -> str:
    """
    Lấy phần mở rộng file một cách an toàn.
    """

    if not uploaded_file:
        return ""

    name = getattr(uploaded_file, "name", "")

    if "." not in name:
        return ""

    return name.rsplit(".", 1)[-1].lower().strip()


def normalize_uploaded_files(uploaded_files) -> List[Any]:
    """
    Chuẩn hóa:
    - None -> []
    - UploadedFile -> [UploadedFile]
    - list UploadedFile -> list
    """

    if uploaded_files is None:
        return []

    if isinstance(uploaded_files, (list, tuple)):
        return list(uploaded_files)

    return [uploaded_files]


def _safe_seek(uploaded_file):
    """
    Đưa con trỏ file về đầu.
    """

    try:
        uploaded_file.seek(0)
    except Exception:
        pass


def _safe_read_bytes(uploaded_file) -> bytes:
    """
    Đọc bytes file một cách an toàn.
    """

    _safe_seek(uploaded_file)

    try:
        data = uploaded_file.read()

        _safe_seek(uploaded_file)

        return data or b""

    except Exception:
        return b""


# =========================================================
# 3. ĐỌC PDF
# =========================================================

def doc_pdf(uploaded_file) -> str:
    """
    Đọc nội dung PDF.

    Ưu tiên:
    - PdfReader
    - Từng trang
    - Giữ ngắt trang để AI hiểu cấu trúc tài liệu.
    """

    try:

        _safe_seek(uploaded_file)

        reader = PyPDF2.PdfReader(uploaded_file)

        pages = []

        for page_index, page in enumerate(reader.pages, start=1):

            try:
                text = page.extract_text() or ""

                text = text.replace("\x00", " ")

                if text.strip():

                    pages.append(
                        f"\n--- TRANG PDF {page_index} ---\n{text.strip()}"
                    )

            except Exception as page_error:

                pages.append(
                    f"\n[⚠️ Không đọc được trang PDF {page_index}: "
                    f"{page_error}]"
                )

        return "\n".join(pages)

    except Exception as e:

        return (
            f"\n[⚠️ LỖI ĐỌC PDF: {str(e)}]\n"
        )


# =========================================================
# 4. ĐỌC DOCX
# =========================================================

def _extract_docx_paragraphs(doc: Document) -> List[str]:
    """
    Đọc toàn bộ đoạn văn trong DOCX.
    """

    result = []

    for index, paragraph in enumerate(doc.paragraphs, start=1):

        text = paragraph.text.strip()

        if text:

            result.append(
                f"[ĐOẠN {index}] {text}"
            )

    return result


def _extract_docx_tables(doc: Document) -> List[str]:
    """
    Đọc toàn bộ bảng DOCX.

    Bảo toàn:
    - Thứ tự bảng
    - Hàng
    - Cột
    - Nội dung từng ô
    """

    result = []

    for table_index, table in enumerate(doc.tables, start=1):

        result.append(
            f"\n--- BẢNG DOCX {table_index} ---"
        )

        for row_index, row in enumerate(table.rows, start=1):

            cells = []

            for cell in row.cells:

                cell_text = cell.text

                cell_text = cell_text.replace(
                    "\n",
                    " "
                ).strip()

                cells.append(cell_text)

            result.append(
                f"Hàng {row_index}: "
                + " | ".join(cells)
            )

    return result


def doc_docx(uploaded_file) -> str:
    """
    Đọc DOCX đầy đủ đoạn văn và bảng.
    """

    try:

        _safe_seek(uploaded_file)

        doc = Document(uploaded_file)

        paragraphs = _extract_docx_paragraphs(doc)

        tables = _extract_docx_tables(doc)

        return "\n".join(
            paragraphs + tables
        )

    except Exception as e:

        return (
            f"\n[⚠️ LỖI ĐỌC DOCX: {str(e)}]\n"
        )


# =========================================================
# 5. ĐỌC EXCEL
# =========================================================

def doc_excel(uploaded_file) -> str:
    """
    Đọc toàn bộ các sheet Excel.

    Bảo toàn:
    - Tên sheet
    - Tên cột
    - Dữ liệu bảng
    """

    try:

        _safe_seek(uploaded_file)

        all_sheets = pd.read_excel(
            uploaded_file,
            sheet_name=None
        )

        result = []

        for sheet_name, df in all_sheets.items():

            result.append(
                f"\n--- SHEET: {sheet_name} ---"
            )

            if df.empty:

                result.append(
                    "[Sheet không có dữ liệu]"
                )

                continue

            df = df.fillna("")

            result.append(
                df.to_string(
                    index=False,
                    max_rows=10000,
                    max_cols=100
                )
            )

        return "\n".join(result)

    except Exception as e:

        return (
            f"\n[⚠️ LỖI ĐỌC EXCEL: {str(e)}]\n"
        )


# =========================================================
# 6. ĐỌC FILE BẤT KỲ
# =========================================================

def doc_noi_dung_file(
    uploaded_file,
    ai_engine=None
) -> str:
    """
    Đọc file PDF, DOCX, XLSX, XLS.

    Với ảnh:
    - Không cố trích xuất OCR giả.
    - Trả metadata để Vision AI xử lý.
    """

    if not uploaded_file:

        return ""

    ext = _get_file_extension(uploaded_file)

    try:

        if ext == "pdf":

            return doc_pdf(uploaded_file)

        elif ext == "docx":

            return doc_docx(uploaded_file)

        elif ext in ["xlsx", "xls"]:

            return doc_excel(uploaded_file)

        elif ext in ["jpg", "jpeg", "png", "webp"]:

            return (
                f"\n--- HÌNH ẢNH: "
                f"{getattr(uploaded_file, 'name', 'unknown')} ---\n"
                "[Đây là file ảnh. Cần sử dụng Vision AI để phân tích nội dung hình ảnh.]\n"
            )

        else:

            return (
                f"\n[⚠️ Định dạng file chưa được hỗ trợ: {ext}]\n"
            )

    except Exception as e:

        return (
            f"\n[⚠️ LỖI ĐỌC FILE "
            f"{getattr(uploaded_file, 'name', 'unknown')}: {str(e)}]\n"
        )


# =========================================================
# 7. ĐỌC NỘI DUNG NHIỀU FILE
# =========================================================

def doc_nhieu_file(
    uploaded_files,
    ai_engine=None
) -> str:
    """
    Đọc nhiều file và ghép nội dung có đánh dấu rõ ràng.
    """

    files = normalize_uploaded_files(uploaded_files)

    if not files:

        return ""

    result = []

    for index, file in enumerate(files, start=1):

        file_name = getattr(
            file,
            "name",
            f"file_{index}"
        )

        result.append(
            f"\n\n========== "
            f"FILE {index}: {file_name} "
            f"==========\n"
        )

        result.append(
            doc_noi_dung_file(
                file,
                ai_engine=ai_engine
            )
        )

    return "\n".join(result)


# =========================================================
# 8. ĐỌC FILE MẪU KHBD
# =========================================================

def doc_file_mau_local(
    path: str = "templates/KHBD_Mau.docx"
) -> str:
    """
    Đọc file mẫu KHBD trên máy chủ.

    Dùng cho:
    - Phân tích cấu trúc
    - Phân tích tiêu đề
    - Phân tích bảng
    """

    if not os.path.exists(path):

        return ""

    try:

        with open(path, "rb") as f:

            doc = Document(f)

        result = []

        result.extend(
            _extract_docx_paragraphs(doc)
        )

        result.extend(
            _extract_docx_tables(doc)
        )

        return "\n".join(result)

    except Exception:

        return ""


# =========================================================
# 9. PHÂN TÍCH CẤU TRÚC FILE MẪU
# =========================================================

def phan_tich_cau_truc_docx(
    uploaded_file
) -> Dict[str, Any]:
    """
    Phân tích cấu trúc file mẫu.

    Trả về:
    - Paragraph
    - Heading
    - Table
    - Số dòng
    - Số cột
    """

    result = {
        "paragraphs": [],
        "tables": [],
        "headings": [],
    }

    if not uploaded_file:

        return result

    try:

        _safe_seek(uploaded_file)

        doc = Document(uploaded_file)

        for paragraph in doc.paragraphs:

            text = paragraph.text.strip()

            if not text:

                continue

            result["paragraphs"].append(text)

            if paragraph.style:

                style_name = paragraph.style.name or ""

                if (
                    "Heading" in style_name
                    or text.isupper()
                ):

                    result["headings"].append(
                        text
                    )

        for table in doc.tables:

            table_data = []

            for row in table.rows:

                row_data = []

                for cell in row.cells:

                    row_data.append(
                        cell.text.strip()
                    )

                table_data.append(
                    row_data
                )

            result["tables"].append(
                table_data
            )

        return result

    except Exception:

        return result


# =========================================================
# 10. AI ENGINE COMPATIBILITY LAYER
# =========================================================

def _extract_ai_text(result) -> str:
    """
    Chuẩn hóa kết quả AI Engine.

    Hỗ trợ:
    - str
    - dict
    - object có content/text
    """

    if result is None:

        return ""

    if isinstance(result, str):

        return result.strip()

    if isinstance(result, dict):

        for key in [
            "text",
            "content",
            "response",
            "output",
            "answer"
        ]:

            value = result.get(key)

            if value:

                return str(value).strip()

    for attr in [
        "text",
        "content",
        "response",
        "output"
    ]:

        try:

            value = getattr(
                result,
                attr,
                None
            )

            if value:

                return str(value).strip()

        except Exception:

            pass

    return str(result).strip()


def _is_ai_error(text: str) -> bool:
    """
    Kiểm tra kết quả AI có phải thông báo lỗi hay không.
    """

    if not text:

        return True

    error_markers = [
        "❌",
        "API Error",
        "API error",
        "quota",
        "insufficient_quota",
        "invalid_api_key",
        "401",
        "429",
        "500 Internal",
        "503",
        "timeout",
        "timed out",
    ]

    lower_text = text.lower()

    return any(
        marker.lower() in lower_text
        for marker in error_markers
    )


def ai_generate_text(
    ai_engine,
    prompt: str
) -> str:
    """
    Gọi AI Engine hiện tại.

    Ưu tiên:
    ai_engine.generate_text(prompt)

    Không tự tạo API Key mới.
    Không phá vỡ kiến trúc AI Engine hiện tại.
    """

    if ai_engine is None:

        raise RuntimeError(
            "AI Engine chưa được truyền vào render_xd_khbd()."
        )

    if not hasattr(
        ai_engine,
        "generate_text"
    ):

        raise AttributeError(
            "AI Engine không có hàm generate_text()."
        )

    result = ai_engine.generate_text(
        prompt
    )

    text = _extract_ai_text(result)

    if _is_ai_error(text):

        return ""

    return text


def ai_generate_vision(
    ai_engine,
    prompt: str,
    image_files: List[Any]
) -> str:
    """
    Gọi Vision AI nếu AI Engine hỗ trợ.

    Hỗ trợ một số kiểu interface:

    generate_vision(
        prompt,
        images
    )

    hoặc

    generate_vision(
        prompt=prompt,
        images=images
    )
    """

    if not image_files:

        return ""

    if ai_engine is None:

        return ""

    if not hasattr(
        ai_engine,
        "generate_vision"
    ):

        return ""

    images = []

    for image_file in image_files:

        try:

            image_bytes = _safe_read_bytes(
                image_file
            )

            if image_bytes:

                images.append({
                    "name": getattr(
                        image_file,
                        "name",
                        "image"
                    ),
                    "mime_type": getattr(
                        image_file,
                        "type",
                        "image/png"
                    ),
                    "data": image_bytes,
                })

        except Exception:

            continue

    if not images:

        return ""

    try:

        result = ai_engine.generate_vision(
            prompt=prompt,
            images=images
        )

        return _extract_ai_text(
            result
        )

    except TypeError:

        try:

            result = ai_engine.generate_vision(
                prompt,
                images
            )

            return _extract_ai_text(
                result
            )

        except Exception:

            return ""

    except Exception:

        return ""


# =========================================================
# 11. KIỂM SOÁT ĐỘ DÀI TÀI LIỆU
# =========================================================

def cat_noi_dung_an_toan(
    text: str,
    max_chars: int = 120000
) -> str:

    if not text:

        return ""

    if len(text) <= max_chars:

        return text

    return (
        text[:max_chars]
        + "\n\n"
        + "[⚠️ TÀI LIỆU ĐÃ ĐƯỢC CẮT BỚT "
        "DO VƯỢT GIỚI HẠN DỮ LIỆU]"
    )


# =========================================================
# 12. XỬ LÝ LATEX
# =========================================================

def clean_latex_formula(
    formula: str
) -> str:

    formula = formula.strip()

    formula = formula.replace(
        "\\left",
        ""
    )

    formula = formula.replace(
        "\\right",
        ""
    )

    return formula


def latex_to_omml(
    latex: str
):
    """
    Chuyển LaTeX → MathML → OMML.

    Word sử dụng OMML cho Equation.
    """

    if not LATEX2MATHML_AVAILABLE:

        return None

    if not LXML_AVAILABLE:

        return None

    try:

        latex = clean_latex_formula(
            latex
        )

        mathml = (
            latex2mathml.converter.convert(
                latex
            )
        )

        mathml_tree = etree.fromstring(
            mathml.encode("utf-8")
        )

        ns = {
            "m": "http://schemas.openxmlformats.org/officeDocument/2006/math"
        }

        # Một số phiên bản converter trả MathML.
        # Word không nhận trực tiếp MathML.
        # Hàm này chuẩn bị cấu trúc để mở rộng chuyển đổi.
        #
        # Nếu không chuyển đổi được hoàn toàn,
        # trả None để fallback về text.

        return None

    except Exception:

        return None


def add_formula_as_word_equation(
    paragraph,
    formula: str
):
    """
    Thêm công thức vào paragraph.

    Hiện tại:
    - Nếu có bộ chuyển đổi OMML hoàn chỉnh: dùng Equation.
    - Nếu không: thêm công thức dạng text dễ đọc.

    Hàm được thiết kế độc lập để có thể nâng cấp
    sang MathML → OMML mà không thay đổi pipeline.
    """

    formula = formula.strip()

    omml = latex_to_omml(
        formula
    )

    if omml is not None:

        paragraph._p.append(
            omml
        )

        return

    run = paragraph.add_run(
        formula
    )

    run.font.name = (
        "Cambria Math"
    )

    run.font.size = Pt(12)


# =========================================================
# 13. CHUYỂN MARKDOWN TABLE
# =========================================================

def is_markdown_table_line(
    line: str
) -> bool:

    line = line.strip()

    return (
        line.startswith("|")
        and line.endswith("|")
        and "|" in line[1:-1]
    )


def is_markdown_separator(
    line: str
) -> bool:

    line = line.strip()

    if "|" not in line:

        return False

    cells = line.strip(
        "|"
    ).split("|")

    for cell in cells:

        cell = cell.strip()

        if cell and not re.fullmatch(
            r":?-{3,}:?",
            cell
        ):

            return False

    return True


def parse_markdown_table(
    lines: List[str]
) -> List[List[str]]:

    rows = []

    for line in lines:

        if not is_markdown_table_line(
            line
        ):

            continue

        if is_markdown_separator(
            line
        ):

            continue

        cells = [
            cell.strip()
            for cell in line.strip(
                "|"
            ).split("|")
        ]

        cells = [
            cell.replace(
                "**",
                ""
            ).replace(
                "__",
                ""
            )
            for cell in cells
        ]

        rows.append(
            cells
        )

    return rows


# =========================================================
# 14. ĐỊNH DẠNG BẢNG WORD
# =========================================================

def set_cell_text(
    cell,
    text: str,
    bold: bool = False
):

    cell.text = ""

    paragraph = cell.paragraphs[0]

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.LEFT
    )

    run = paragraph.add_run(
        text
    )

    run.bold = bold

    run.font.name = (
        "Times New Roman"
    )

    run.font.size = Pt(12)

    cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.CENTER
    )


def format_word_table(
    table,
    header_rows: int = 1
):

    table.style = "Table Grid"

    table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    for row_index, row in enumerate(
        table.rows
    ):

        for cell in row.cells:

            cell.vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )

            for paragraph in cell.paragraphs:

                paragraph.paragraph_format.space_after = Pt(0)

                for run in paragraph.runs:

                    run.font.name = (
                        "Times New Roman"
                    )

                    run.font.size = Pt(11)

                    if row_index < header_rows:

                        run.bold = True


# =========================================================
# 15. XỬ LÝ CÔNG THỨC TRONG VĂN BẢN
# =========================================================

def split_text_and_formulas(
    text: str
) -> List[Tuple[str, str]]:
    """
    Trả về danh sách:
    ("text", "..."),
    ("formula", "...")
    """

    pattern = re.compile(
        r"(\$\$.*?\$\$|\$.*?\$|\\\[.*?\\\]|\\\(.*?\\\))",
        re.DOTALL
    )

    result = []

    last_end = 0

    for match in pattern.finditer(
        text
    ):

        if match.start() > last_end:

            result.append(
                (
                    "text",
                    text[
                        last_end:
                        match.start()
                    ]
                )
            )

        formula = match.group(0)

        formula = formula.strip(
            "$"
        )

        formula = formula.replace(
            "\\[",
            ""
        ).replace(
            "\\]",
            ""
        ).replace(
            "\\(",
            ""
        ).replace(
            "\\)",
            ""
        )

        result.append(
            (
                "formula",
                formula.strip()
            )
        )

        last_end = match.end()

    if last_end < len(text):

        result.append(
            (
                "text",
                text[last_end:]
            )
        )

    if not result:

        result.append(
            (
                "text",
                text
            )
        )

    return result


# =========================================================
# 16. XUẤT MARKDOWN → WORD
# =========================================================

def tao_file_word_hoan_hao(
    van_ban: str,
    template_file=None
):

    """
    Xuất giáo án Word.

    Nếu có template_file:
    - Đọc template làm nền.
    - Giữ lại toàn bộ bảng biểu có sẵn.
    - Giữ lại header/footer.
    - Giữ lại page setup.
    - Thêm nội dung AI sau nội dung mẫu.

    Nếu không có template:
    - Tạo Word mới theo chuẩn A4.
    """

    if template_file:

        try:

            _safe_seek(
                template_file
            )

            doc_word = Document(
                template_file
            )

        except Exception:

            doc_word = Document()

    else:

        doc_word = Document()

    # -----------------------------------------------------
    # PAGE SETUP
    # -----------------------------------------------------

    section = doc_word.sections[0]

    section.page_height = Cm(29.7)

    section.page_width = Cm(21.0)

    section.left_margin = Cm(2.0)

    section.right_margin = Cm(2.0)

    section.top_margin = Cm(2.0)

    section.bottom_margin = Cm(2.0)

    # -----------------------------------------------------
    # NORMAL STYLE
    # -----------------------------------------------------

    try:

        style = doc_word.styles["Normal"]

        style.font.name = (
            "Times New Roman"
        )

        style.font.size = Pt(13)

        style.paragraph_format.line_spacing = 1.15

        style.paragraph_format.alignment = (
            WD_ALIGN_PARAGRAPH.JUSTIFY
        )

    except Exception:

        pass

    # -----------------------------------------------------
    # NGĂN CÁCH PHẦN TEMPLATE VÀ NỘI DUNG AI
    # -----------------------------------------------------

    if template_file:

        doc_word.add_page_break()

    lines = van_ban.splitlines()

    index = 0

    while index < len(lines):

        raw_line = lines[index]

        line = raw_line.strip()

        # -------------------------------------------------
        # BỎ DÒNG TRỐNG
        # -------------------------------------------------

        if not line:

            index += 1

            continue

        # -------------------------------------------------
        # MARKDOWN TABLE
        # -------------------------------------------------

        if is_markdown_table_line(
            line
        ):

            table_lines = []

            while (
                index < len(lines)
                and is_markdown_table_line(
                    lines[index].strip()
                )
            ):

                table_lines.append(
                    lines[index]
                )

                index += 1

            table_data = parse_markdown_table(
                table_lines
            )

            if table_data:

                max_cols = max(
                    len(row)
                    for row in table_data
                )

                table = doc_word.add_table(
                    rows=len(table_data),
                    cols=max_cols
                )

                for row_index, row in enumerate(
                    table_data
                ):

                    for col_index in range(
                        max_cols
                    ):

                        value = ""

                        if col_index < len(
                            row
                        ):

                            value = row[
                                col_index
                            ]

                        set_cell_text(
                            table.cell(
                                row_index,
                                col_index
                            ),
                            value,
                            bold=(
                                row_index == 0
                            )
                        )

                format_word_table(
                    table
                )

            continue

        # -------------------------------------------------
        # HEADING
        # -------------------------------------------------

        if line.startswith(
            "### "
        ):

            text = line[4:].strip()

            paragraph = doc_word.add_paragraph()

            run = paragraph.add_run(
                text
            )

            run.bold = True

            run.font.name = (
                "Times New Roman"
            )

            run.font.size = Pt(13)

            index += 1

            continue

        if line.startswith(
            "## "
        ):

            text = line[3:].strip()

            paragraph = doc_word.add_paragraph()

            run = paragraph.add_run(
                text
            )

            run.bold = True

            run.font.name = (
                "Times New Roman"
            )

            run.font.size = Pt(14)

            index += 1

            continue

        if line.startswith(
            "# "
        ):

            text = line[2:].strip()

            paragraph = doc_word.add_heading(
                text,
                level=1
            )

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

            index += 1

            continue

        # -------------------------------------------------
        # BULLET
        # -------------------------------------------------

        if line.startswith(
            "- "
        ) or line.startswith(
            "* "
        ):

            text = line[2:].strip()

            paragraph = doc_word.add_paragraph(
                style="List Bullet"
            )

            parts = split_text_and_formulas(
                text
            )

            for part_type, part_value in parts:

                if part_type == "text":

                    run = paragraph.add_run(
                        part_value
                    )

                    run.font.name = (
                        "Times New Roman"
                    )

                    run.font.size = Pt(13)

                else:

                    add_formula_as_word_equation(
                        paragraph,
                        part_value
                    )

            index += 1

            continue

        # -------------------------------------------------
        # NUMBERED LIST
        # -------------------------------------------------

        if re.match(
            r"^\d+[\.\)]\s+",
            line
        ):

            text = re.sub(
                r"^\d+[\.\)]\s+",
                "",
                line
            )

            paragraph = doc_word.add_paragraph(
                style="List Number"
            )

            parts = split_text_and_formulas(
                text
            )

            for part_type, part_value in parts:

                if part_type == "text":

                    run = paragraph.add_run(
                        part_value
                    )

                    run.font.name = (
                        "Times New Roman"
                    )

                    run.font.size = Pt(13)

                else:

                    add_formula_as_word_equation(
                        paragraph,
                        part_value
                    )

            index += 1

            continue

        # -------------------------------------------------
        # ĐOẠN VĂN THƯỜNG
        # -------------------------------------------------

        clean_line = line

        clean_line = clean_line.replace(
            "**",
            ""
        )

        clean_line = clean_line.replace(
            "__",
            ""
        )

        clean_line = clean_line.replace(
            "`",
            ""
        )

        paragraph = doc_word.add_paragraph()

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.JUSTIFY
        )

        parts = split_text_and_formulas(
            clean_line
        )

        for part_type, part_value in parts:

            if part_type == "text":

                run = paragraph.add_run(
                    part_value
                )

                run.font.name = (
                    "Times New Roman"
                )

                run.font.size = Pt(13)

            else:

                add_formula_as_word_equation(
                    paragraph,
                    part_value
                )

        index += 1

    bio = io.BytesIO()

    doc_word.save(
        bio
    )

    bio.seek(0)

    return bio


# =========================================================
# 17. TẠO PHẠM VI KIẾN THỨC ĐƯỢC PHÉP
# =========================================================

def tao_pham_vi_kien_thuc(
    noi_dung_chinh: str,
    noi_dung_ppct: str = "",
    noi_dung_ai_file: str = ""
) -> str:

    return f"""
================ PHẠM VI KIẾN THỨC ĐƯỢC PHÉP ================

NGUỒN 1 - TÀI LIỆU CHÍNH:
{noi_dung_chinh}

NGUỒN 2 - PPCT:
{noi_dung_ppct or "Không cung cấp."}

NGUỒN 3 - TÀI LIỆU BỔ SUNG:
{noi_dung_ai_file or "Không cung cấp."}

==============================================================

QUY TẮC PHẠM VI:

1. Chỉ sử dụng kiến thức xuất hiện trực tiếp hoặc có thể suy ra hợp lý
   từ tài liệu chính.

2. Không đưa kiến thức của bài học khác vào.

3. Không tự ý đổi tên bài.

4. Không tự ý thay đổi khái niệm khoa học.

5. Không tự thêm công thức, định luật, thuật ngữ hoặc số liệu
   nếu không có căn cứ từ tài liệu.

6. Nếu tài liệu không đủ thông tin:
   - Không được bịa.
   - Ghi rõ: "Chưa có đủ dữ liệu trong tài liệu đầu vào."

7. PPCT chỉ dùng để:
   - Đối chiếu tên bài.
   - Đối chiếu thời lượng.
   - Đối chiếu yêu cầu cần đạt.

8. Tài liệu bảng tích hợp AI chỉ dùng để xác định:
   - Vị trí tích hợp AI.
   - Hoạt động AI.
   - Sản phẩm AI.

9. Nếu có mâu thuẫn giữa tài liệu:
   - Ưu tiên nội dung SGK/giáo án gốc.
   - Không tự ý hòa trộn hai nội dung trái ngược.
   - Nêu rõ điểm cần giáo viên kiểm tra.

==============================================================
"""


# =========================================================
# 18. PROMPT CHUYÊN SÂU THEO CV 5512
# =========================================================

def tao_prompt_khbd(
    *,
    thong_tin_bai_day: str,
    pham_vi_kien_thuc: str,
    noi_dung_mau_khbd: str,
    lenh_nhiem_vu: str,
    tich_hop_nls_str: str,
    tich_hop_ai_str: str,
    tich_hop_kt_str: str,
    hoat_dong_str: str,
    mode: str
) -> str:

    cau_truc_mau = ""

    if noi_dung_mau_khbd:

        cau_truc_mau = f"""

================ MẪU GIÁO ÁN CỦA TRƯỜNG ================

{noi_dung_mau_khbd}

==========================================================

YÊU CẦU VỀ MẪU:

- Phải giữ nguyên tên các đề mục chính.
- Phải giữ nguyên thứ tự các phần.
- Phải giữ nguyên logic của các bảng.
- Không tự ý đổi mẫu sang một bố cục khác.
- Nếu mẫu có bảng Hoạt động dạy học:
  phải sử dụng đúng cấu trúc bảng đó.
- Nếu mẫu có các mục:
  + Mục tiêu
  + Thiết bị dạy học và học liệu
  + Tiến trình dạy học
  + Hoạt động
  + Sản phẩm
  + Tổ chức thực hiện
  thì phải giữ đầy đủ.

==========================================================
"""

    return f"""
BẠN LÀ CHUYÊN GIA XÂY DỰNG KẾ HOẠCH BÀI DẠY
THEO CHƯƠNG TRÌNH GDPT 2018 VÀ CÔNG VĂN 5512.

NHIỆM VỤ:
{lenh_nhiem_vu}

==========================================================
I. THÔNG TIN BÀI DẠY
==========================================================

{thong_tin_bai_day}

==========================================================
II. PHẠM VI KIẾN THỨC ĐƯỢC PHÉP SỬ DỤNG
==========================================================

{pham_vi_kien_thuc}

==========================================================
III. YÊU CẦU TÍCH HỢP
==========================================================

NĂNG LỰC SỐ:
{tich_hop_nls_str}

NĂNG LỰC AI:
{tich_hop_ai_str}

DẠY HỌC HÒA NHẬP:
{tich_hop_kt_str}

HOẠT ĐỘNG GIÁO VIÊN YÊU CẦU:
{hoat_dong_str}

{cau_truc_mau}

==========================================================
IV. NGUYÊN TẮC CHUYÊN MÔN BẮT BUỘC
==========================================================

1. Bám sát tuyệt đối tài liệu đầu vào.

2. Không bịa kiến thức ngoài phạm vi.

3. Không tự ý thay đổi tên bài.

4. Không tự ý thêm nội dung của bài khác.

5. Nếu nội dung tài liệu không rõ:
   phải đánh dấu [CẦN GIÁO VIÊN KIỂM TRA],
   không được tự bịa.

6. Thời lượng các hoạt động phải khớp tổng thời lượng.

7. Mỗi hoạt động phải thể hiện rõ:

   - Mục tiêu.
   - Nội dung.
   - Sản phẩm.
   - Tổ chức thực hiện.

8. Phần tổ chức thực hiện phải có:

   - Giáo viên nói gì?
   - Giáo viên giao nhiệm vụ gì?
   - Học sinh làm gì?
   - Học sinh dự kiến trả lời gì?
   - Giáo viên nhận xét, chốt kiến thức như thế nào?

9. Không được viết chung chung:

   CẤM:
   "GV hướng dẫn HS thực hiện nhiệm vụ."

   PHẢI VIẾT:
   "GV yêu cầu HS quan sát ...,
   đọc thông tin ...,
   thảo luận trong ... phút
   và trả lời câu hỏi ..."

10. Phải mô tả sản phẩm học tập cụ thể.

11. Không sử dụng LaTeX trong phần văn bản chính
    nếu không thực sự cần thiết.

12. Khi cần biểu diễn công thức,
    có thể dùng định dạng:

    $F = ma$

    hoặc

    $$F = ma$$

    để hệ thống xử lý khi xuất Word.

13. Không viết lời chào.

14. Không viết lời dạo đầu.

15. Không viết nhận xét ngoài giáo án.

16. Chỉ trả về nội dung giáo án hoàn chỉnh.

==========================================================
V. KIỂM TRA TRƯỚC KHI TRẢ KẾT QUẢ
==========================================================

Trước khi xuất kết quả, tự kiểm tra:

[ ] Đúng tên bài.
[ ] Đúng khối lớp.
[ ] Đúng môn học.
[ ] Đúng thời lượng.
[ ] Không có kiến thức ngoài phạm vi.
[ ] Đã sử dụng hoạt động giáo viên yêu cầu.
[ ] Đã tích hợp NLS nếu được yêu cầu.
[ ] Đã tích hợp AI nếu được yêu cầu.
[ ] Đã điều chỉnh dạy học hòa nhập nếu được yêu cầu.
[ ] Đúng cấu trúc mẫu.
[ ] Đủ mục tiêu.
[ ] Đủ thiết bị/học liệu.
[ ] Đủ tiến trình.
[ ] Có sản phẩm học tập.
[ ] Có hoạt động GV.
[ ] Có hoạt động HS.

BÂY GIỜ HÃY TẠO GIÁO ÁN.
"""


# =========================================================
# 19. KIỂM TRA KẾT QUẢ AI
# =========================================================

def kiem_tra_ket_qua_ai(
    text: str,
    ten_bai: str,
    so_tiet: str
) -> Dict[str, Any]:

    result = {
        "hop_le": True,
        "canh_bao": [],
    }

    if not text:

        result["hop_le"] = False

        result["canh_bao"].append(
            "AI không trả về nội dung."
        )

        return result

    if len(text) < 500:

        result["canh_bao"].append(
            "Kết quả AI khá ngắn."
        )

    if ten_bai and ten_bai != "Không cung cấp. Căn cứ hoàn toàn vào tài liệu":

        if ten_bai.lower() not in text.lower():

            result["canh_bao"].append(
                "Chưa tìm thấy rõ tên bài trong kết quả."
            )

    required_keywords = [
        "Mục tiêu",
        "Hoạt động",
        "Sản phẩm",
    ]

    for keyword in required_keywords:

        if keyword.lower() not in text.lower():

            result["canh_bao"].append(
                f"Thiếu hoặc chưa rõ mục: {keyword}"
            )

    return result


# =========================================================
# 20. GIAO DIỆN CHÍNH
# =========================================================

def render_xd_khbd(
    ai_engine=None
):

    init_session_state()

    # -----------------------------------------------------
    # CSS
    # -----------------------------------------------------

    st.markdown(
        """
        <style>

        .stButton button[kind="primary"] {
            background-color: #9333ea;
            color: white;
            border: none;
            border-radius: 8px;
            font-weight: bold;
        }

        .stButton button[kind="primary"]:hover {
            background-color: #7e22ce;
        }

        .stButton button[kind="secondary"] {
            color: #6b7280;
            border: 1px solid #e5e7eb;
            border-radius: 8px;
            font-weight: 600;
            background-color: #f9fafb;
        }

        .upload-card {
            text-align: center;
            padding: 10px;
        }

        .upload-icon {
            font-size: 2.5rem;
            color: #9333ea;
            margin-bottom: 10px;
        }

        .upload-title {
            font-weight: bold;
            font-size: 1.1rem;
            color: #1f2937;
        }

        .upload-desc {
            font-size: 0.85rem;
            color: #6b7280;
        }

        </style>
        """,
        unsafe_allow_html=True
    )

    # -----------------------------------------------------
    # DANH MỤC NLS
    # -----------------------------------------------------

    THANH_PHAN_NLS = [

        "1.1. Duyệt, tìm kiếm, lọc dữ liệu",
        "1.2. Đánh giá dữ liệu",
        "1.3. Quản lý dữ liệu",

        "2.1. Tương tác công nghệ số",
        "2.2. Chia sẻ thông tin",
        "2.3. Thực hiện trách nhiệm công dân",
        "2.4. Hợp tác công nghệ số",
        "2.5. Quy tắc ứng xử",
        "2.6. Quản lý danh tính số",

        "3.1. Phát triển nội dung",
        "3.2. Tích hợp nội dung",
        "3.3. Bản quyền, giấy phép",
        "3.4. Lập trình",

        "4.1. Bảo vệ thiết bị",
        "4.2. Bảo vệ dữ liệu cá nhân",
        "4.3. Bảo vệ sức khỏe",
        "4.4. Bảo vệ môi trường",

        "5.1. Giải quyết vấn đề kỹ thuật",
        "5.2. Giải pháp công nghệ",
        "5.3. Sáng tạo công nghệ",
        "5.4. Xác định vấn đề NLS",

        "6.1. Hiểu biết AI",
        "6.2. Sử dụng AI",
        "6.3. Đánh giá AI",

    ]

    MUC_DO_NLS = [

        "-- Tự nhập --",

        "CB1a",
        "CB1b",
        "CB1c",

        "CB2a",
        "CB2b",
        "CB2c",
        "CB2d",

        "TC1a",
        "TC1b",
        "TC1c",
        "TC1d",

        "TC2a",
        "TC2b",
        "TC2c",
        "TC2d",

        "NC1a",
        "NC1b",
        "NC1c",
        "NC1d",

    ]

    tu_dien_nls = {

        "CB1a":
        "Xác định được nhu cầu thông tin, tìm kiếm dữ liệu, thông tin và nội dung thông qua tìm kiếm đơn giản trong môi trường số.",

        "CB1b":
        "Biết cách duyệt qua các trang web hoặc tài liệu số cơ bản.",

        "NC1b":
        "Áp dụng được kỹ thuật tìm kiếm để lấy được dữ liệu, thông tin và nội dung trong môi trường số.",

    }

    # -----------------------------------------------------
    # THÔNG TIN BÀI DẠY
    # -----------------------------------------------------

    st.markdown(
        "### 🎛️ Thông tin bài dạy"
    )

    c_khoi, c_mon = st.columns(2)

    with c_khoi:

        st.selectbox(
            "KHỐI LỚP",
            [
                "Lớp 6",
                "Lớp 7",
                "Lớp 8",
                "Lớp 9",
                "Lớp 10",
                "Lớp 11",
                "Lớp 12",
            ],
            key="khbd_khoi_lop"
        )

    with c_mon:

        st.selectbox(
            "MÔN HỌC",
            [
                "Toán",
                "Ngữ văn",
                "Tiếng Anh",
                "Khoa học tự nhiên",
                "Lịch sử và Địa lí",
                "Vật lí",
                "Hóa học",
                "Sinh học",
                "Lịch sử",
                "Địa lí",
                "Giáo dục công dân",
                "Tin học",
                "Công nghệ",
                "Khác",
            ],
            key="khbd_mon_hoc"
        )

    # -----------------------------------------------------
    # CHẾ ĐỘ TÍCH HỢP
    # -----------------------------------------------------

    st.markdown(
        "#### ✨ Chế độ tích hợp"
    )

    c_th1, c_th2, c_th3 = st.columns(3)

    with c_th1:

        with st.container(
            border=True
        ):

            tich_hop_nls = st.checkbox(
                "Tích hợp Năng lực số",
                key="chk_nls"
            )

    with c_th2:

        with st.container(
            border=True
        ):

            tich_hop_ai = st.checkbox(
                "Tích hợp Năng lực AI",
                key="chk_ai"
            )

    with c_th3:

        with st.container(
            border=True
        ):

            tich_hop_kt = st.checkbox(
                "Dạy học khuyết tật",
                key="chk_kt"
            )

    # -----------------------------------------------------
    # CHỌN CHẾ ĐỘ
    # -----------------------------------------------------

    c_btn1, c_btn2 = st.columns(2)

    with c_btn1:

        st.button(
            "📄 CHỈNH SỬA GIÁO ÁN GỐC",
            type=(
                "primary"
                if st.session_state.soan_mode == "chinh_sua"
                else "secondary"
            ),
            use_container_width=True,
            on_click=set_mode,
            args=("chinh_sua",)
        )

    with c_btn2:

        st.button(
            "⚡ TỰ ĐỘNG SOẠN TỪ SGK",
            type=(
                "primary"
                if st.session_state.soan_mode == "tu_dong"
                else "secondary"
            ),
            use_container_width=True,
            on_click=set_mode,
            args=("tu_dong",)
        )

    st.divider()

    # -----------------------------------------------------
    # CHỈNH SỬA GIÁO ÁN
    # -----------------------------------------------------

    if st.session_state.soan_mode == "chinh_sua":

        with st.container(
            border=True
        ):

            st.markdown(
                "### 📤 Tài liệu đầu vào"
            )

            st.caption(
                "Khuyến nghị: tải lên giáo án 1 bài hoặc 1 tiết."
            )

            c_up1, c_up2, c_up3 = st.columns(3)

            with c_up1:

                with st.container(
                    border=True
                ):

                    st.markdown(
                        """
                        <div class="upload-card">
                            <div class="upload-icon">📄</div>
                            <div class="upload-title">Giáo án gốc</div>
                            <div class="upload-desc">
                                Word, PDF, JPG, PNG
                            </div>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )

                    st.file_uploader(
                        "Upload GA",
                        type=[
                            "docx",
                            "pdf",
                            "jpg",
                            "jpeg",
                            "png",
                        ],
                        accept_multiple_files=True,
                        label_visibility="collapsed",
                        key="file_ga"
                    )

            with c_up2:

                with st.container(
                    border=True
                ):

                    st.markdown(
                        """
                        <div class="upload-card">
                            <div class="upload-icon">📊</div>
                            <div class="upload-title">PPCT</div>
                            <div class="upload-desc">
                                PDF, Word, Excel
                            </div>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )

                    st.file_uploader(
                        "Upload PPCT",
                        type=[
                            "pdf",
                            "docx",
                            "xlsx",
                            "xls",
                        ],
                        label_visibility="collapsed",
                        key="file_ppct"
                    )

            with c_up3:

                with st.container(
                    border=True
                ):

                    st.markdown(
                        """
                        <div class="upload-card">
                            <div class="upload-icon">📋</div>
                            <div class="upload-title">
                                Bảng tích hợp AI
                            </div>
                            <div class="upload-desc">
                                PDF, Word, Excel
                            </div>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )

                    st.file_uploader(
                        "Upload AI",
                        type=[
                            "pdf",
                            "docx",
                            "xlsx",
                            "xls",
                        ],
                        label_visibility="collapsed",
                        key="file_ai"
                    )

    # -----------------------------------------------------
    # SOẠN MỚI TỪ SGK
    # -----------------------------------------------------

    else:

        st.markdown(
            "### 📄 Thông tin giáo án soạn mới"
        )

        c_cap, c_mau = st.columns(2)

        with c_cap:

            st.selectbox(
                "Cấp học",
                [
                    "THCS",
                    "Tiểu học",
                    "THPT",
                ],
                key="khbd_cap_hoc"
            )

        with c_mau:

            st.selectbox(
                "Mẫu giáo án",
                [
                    "Công văn 5512 (Chuẩn Bộ)",
                    "Mẫu rút gọn",
                    "Mẫu tư duy",
                ],
                key="khbd_mau_giao_an"
            )

        c_ten, c_tg = st.columns(2)

        with c_ten:

            st.text_input(
                "Tên bài dạy",
                placeholder="VD: Phân thức đại số",
                key="khbd_ten_bai"
            )

        with c_tg:

            st.text_input(
                "Thời lượng",
                placeholder="VD: 2 tiết",
                key="khbd_so_tiet"
            )

        st.markdown(
            "**Hình ảnh / PDF SGK cơ sở**"
        )

        with st.container(
            border=True
        ):

            st.file_uploader(
                "Tải lên Sách Giáo Khoa",
                type=[
                    "pdf",
                    "jpg",
                    "jpeg",
                    "png",
                ],
                accept_multiple_files=True,
                key="file_sgk"
            )

        st.markdown(
            "**📄 File Mẫu Giáo Án của trường**"
        )

        with st.container(
            border=True
        ):

            st.file_uploader(
                "Tải lên File Word mẫu",
                type=[
                    "docx"
                ],
                key="file_template_custom"
            )

            st.caption(
                "Nếu không tải lên, hệ thống sử dụng "
                "templates/KHBD_Mau.docx."
            )

        st.markdown(
            "**Kế hoạch Hoạt động (Tùy chọn)**"
        )

        c_input, c_add = st.columns(
            [4, 1]
        )

        with c_input:

            st.text_input(
                "Nhập hoạt động",
                placeholder=(
                    "VD: Tìm hiểu cấu trúc..."
                ),
                key="new_hoat_dong",
                label_visibility="collapsed",
                on_change=add_hoat_dong
            )

        with c_add:

            st.button(
                "Thêm",
                on_click=add_hoat_dong,
                type="primary",
                use_container_width=True
            )

        if st.session_state.hoat_dong_list:

            for i, hd in enumerate(
                st.session_state.hoat_dong_list
            ):

                c_tag1, c_tag2 = st.columns(
                    [11, 1]
                )

                with c_tag1:

                    st.info(
                        f"📍 {hd}"
                    )

                with c_tag2:

                    if st.button(
                        "Xóa",
                        key=f"del_hd_{i}"
                    ):

                        st.session_state.hoat_dong_list.pop(
                            i
                        )

                        st.rerun()

        if tich_hop_nls or tich_hop_ai:

            st.markdown(
                "### 📤 Tài liệu tích hợp bổ sung"
            )

            c_tl1, c_tl2 = st.columns(2)

            if tich_hop_nls:

                with c_tl1:

                    st.file_uploader(
                        "📄 PPCT",
                        type=[
                            "pdf",
                            "docx",
                            "xlsx",
                            "xls",
                        ],
                        key="file_ppct_tu_dong"
                    )

            if tich_hop_ai:

                with c_tl2:

                    st.file_uploader(
                        "📋 Bảng tích hợp AI",
                        type=[
                            "pdf",
                            "docx",
                            "xlsx",
                            "xls",
                        ],
                        key="file_ai_tu_dong"
                    )

    # -----------------------------------------------------
    # DẠY HỌC HÒA NHẬP
    # -----------------------------------------------------

    dang_khuyet_tat_chon = []

    if tich_hop_kt:

        with st.container(
            border=True
        ):

            st.markdown(
                "#### 🎯 Dạng khuyết tật hòa nhập"
            )

            dang_khuyet_tat_chon = st.pills(
                "Chọn dạng khuyết tật",
                [
                    "Vận động",
                    "Nghe",
                    "Nói",
                    "Nhìn",
                    "Thần kinh",
                    "Tâm thần",
                    "Trí tuệ",
                    "Tự kỷ",
                    "Khác",
                    "Chung",
                ],
                selection_mode="multi",
                default=["Chung"],
                key="pill_kt"
            )

    # -----------------------------------------------------
    # NĂNG LỰC SỐ
    # -----------------------------------------------------

    if tich_hop_nls:

        with st.container(
            border=True
        ):

            if st.checkbox(
                "🎯 Yêu cầu Năng lực số cụ thể",
                value=True,
                key="chk_nls_ct"
            ):

                c_tp, c_md, c_nd = st.columns(
                    [1.5, 1, 2.5]
                )

                with c_tp:

                    st.selectbox(
                        "1. THÀNH PHẦN",
                        THANH_PHAN_NLS,
                        key="nls_tp"
                    )

                with c_md:

                    chon_md = st.selectbox(
                        "2. MỨC ĐỘ",
                        MUC_DO_NLS,
                        key="nls_md"
                    )

                with c_nd:

                    noi_dung_mac_dinh = ""

                    if (
                        chon_md != "-- Tự nhập --"
                        and chon_md in tu_dien_nls
                    ):

                        noi_dung_mac_dinh = (
                            tu_dien_nls[
                                chon_md
                            ]
                        )

                    st.text_area(
                        "3. NỘI DUNG YÊU CẦU",
                        value=noi_dung_mac_dinh,
                        placeholder="Mô tả...",
                        key="nls_nd_input",
                        height=90
                    )

                c_space, c_btn_add = st.columns(
                    [3, 1]
                )

                with c_btn_add:

                    st.button(
                        "➕ Thêm vào danh sách",
                        type="primary",
                        on_click=add_nls_item,
                        use_container_width=True
                    )

                if st.session_state.nls_list:

                    for i, item in enumerate(
                        st.session_state.nls_list
                    ):

                        with st.container(
                            border=True
                        ):

                            c_info, c_del = st.columns(
                                [11, 1]
                            )

                            with c_info:

                                st.write(
                                    f"**{item['thanh_phan']}** "
                                    f"(`{item['muc_do']}`) "
                                    f"👉 {item['noi_dung']}"
                                )

                            with c_del:

                                if st.button(
                                    "Xóa",
                                    key=f"del_nls_{i}"
                                ):

                                    st.session_state.nls_list.pop(
                                        i
                                    )

                                    st.rerun()

    # -----------------------------------------------------
    # NGÔN NGỮ
    # -----------------------------------------------------

    st.write("")

    with st.container(
        border=True
    ):

        is_english = st.checkbox(
            "Giáo án viết bằng ngôn ngữ Tiếng Anh",
            key="khbd_ngon_ngu"
        )

    # =====================================================
    # NÚT XỬ LÝ AI
    # =====================================================

    st.write("")

    if st.button(
        "⚡ KÍCH HOẠT XỬ LÝ AI",
        type="primary",
        use_container_width=True
    ):

        reset_ket_qua()

        # -------------------------------------------------
        # KIỂM TRA AI ENGINE
        # -------------------------------------------------

        if not ai_engine:

            st.error(
                "❌ Chưa truyền AI Engine vào module."
            )

            st.stop()

        if not hasattr(
            ai_engine,
            "generate_text"
        ):

            st.error(
                "❌ AI Engine không có hàm generate_text()."
            )

            st.stop()

        # -------------------------------------------------
        # KIỂM TRA FILE
        # -------------------------------------------------

        if (
            st.session_state.soan_mode
            == "chinh_sua"
            and not st.session_state.get(
                "file_ga"
            )
        ):

            st.error(
                "⚠️ Vui lòng tải lên ít nhất một giáo án gốc."
            )

            st.stop()

        if (
            st.session_state.soan_mode
            == "tu_dong"
            and not st.session_state.get(
                "file_sgk"
            )
        ):

            st.error(
                "⚠️ Vui lòng tải lên SGK."
            )

            st.stop()

        try:

            with st.spinner(
                "🧠 AI đang phân tích tài liệu "
                "và xây dựng KHBD..."
            ):

                # =========================================
                # ĐỌC TÀI LIỆU
                # =========================================

                noi_dung_chinh = ""

                noi_dung_ppct = ""

                noi_dung_ai_file = ""

                template_file_for_export = None

                image_files = []

                if (
                    st.session_state.soan_mode
                    == "chinh_sua"
                ):

                    ga_files = normalize_uploaded_files(
                        st.session_state.get(
                            "file_ga"
                        )
                    )

                    for file in ga_files:

                        ext = _get_file_extension(
                            file
                        )

                        if ext in [
                            "jpg",
                            "jpeg",
                            "png",
                            "webp",
                        ]:

                            image_files.append(
                                file
                            )

                        noi_dung_chinh += (
                            "\n\n--- GIÁO ÁN GỐC: "
                            f"{file.name} ---\n"
                        )

                        noi_dung_chinh += (
                            doc_noi_dung_file(
                                file,
                                ai_engine
                            )
                        )

                    noi_dung_ppct = (
                        doc_noi_dung_file(
                            st.session_state.get(
                                "file_ppct"
                            ),
                            ai_engine
                        )
                    )

                    noi_dung_ai_file = (
                        doc_noi_dung_file(
                            st.session_state.get(
                                "file_ai"
                            ),
                            ai_engine
                        )
                    )

                else:

                    sgk_files = normalize_uploaded_files(
                        st.session_state.get(
                            "file_sgk"
                        )
                    )

                    for file in sgk_files:

                        ext = _get_file_extension(
                            file
                        )

                        if ext in [
                            "jpg",
                            "jpeg",
                            "png",
                            "webp",
                        ]:

                            image_files.append(
                                file
                            )

                        noi_dung_chinh += (
                            "\n\n--- SGK: "
                            f"{file.name} ---\n"
                        )

                        noi_dung_chinh += (
                            doc_noi_dung_file(
                                file,
                                ai_engine
                            )
                        )

                    noi_dung_ppct = (
                        doc_noi_dung_file(
                            st.session_state.get(
                                "file_ppct_tu_dong"
                            ),
                            ai_engine
                        )
                    )

                    noi_dung_ai_file = (
                        doc_noi_dung_file(
                            st.session_state.get(
                                "file_ai_tu_dong"
                            ),
                            ai_engine
                        )
                    )

                    template_file_for_export = (
                        st.session_state.get(
                            "file_template_custom"
                        )
                    )

                    if not template_file_for_export:

                        local_template = (
                            "templates/KHBD_Mau.docx"
                        )

                        if os.path.exists(
                            local_template
                        ):

                            template_file_for_export = (
                                local_template
                            )

                # =========================================
                # VISION AI
                # =========================================

                vision_text = ""

                if image_files:

                    vision_prompt = """
Hãy phân tích chính xác nội dung các hình ảnh tài liệu
được cung cấp.

Yêu cầu:

1. Đọc toàn bộ chữ trong ảnh.
2. Nhận diện tiêu đề, mục, bảng, sơ đồ, công thức.
3. Không tự suy diễn nội dung không nhìn thấy.
4. Giữ nguyên thuật ngữ chuyên môn.
5. Nếu có công thức Toán, Vật lí, Hóa học:
   hãy mô tả rõ công thức.
6. Nếu có hình ảnh SGK:
   hãy trích xuất nội dung phục vụ xây dựng kế hoạch bài dạy.

Chỉ trả về nội dung đã quan sát được.
"""

                    vision_text = ai_generate_vision(
                        ai_engine,
                        vision_prompt,
                        image_files
                    )

                    if vision_text:

                        noi_dung_chinh += (
                            "\n\n"
                            "========== NỘI DUNG PHÂN TÍCH TỪ ẢNH ==========\n"
                            + vision_text
                        )

                # =========================================
                # ĐỌC FILE MẪU
                # =========================================

                noi_dung_mau_khbd = ""

                if (
                    st.session_state.soan_mode
                    == "tu_dong"
                ):

                    custom_template = (
                        st.session_state.get(
                            "file_template_custom"
                        )
                    )

                    if custom_template:

                        noi_dung_mau_khbd = (
                            doc_noi_dung_file(
                                custom_template
                            )
                        )

                    else:

                        noi_dung_mau_khbd = (
                            doc_file_mau_local()
                        )

                # =========================================
                # GIỚI HẠN DỮ LIỆU
                # =========================================

                noi_dung_chinh = (
                    cat_noi_dung_an_toan(
                        noi_dung_chinh,
                        120000
                    )
                )

                noi_dung_ppct = (
                    cat_noi_dung_an_toan(
                        noi_dung_ppct,
                        50000
                    )
                )

                noi_dung_ai_file = (
                    cat_noi_dung_an_toan(
                        noi_dung_ai_file,
                        50000
                    )
                )

                noi_dung_mau_khbd = (
                    cat_noi_dung_an_toan(
                        noi_dung_mau_khbd,
                        60000
                    )
                )

                # =========================================
                # THÔNG TIN BÀI DẠY
                # =========================================

                cap_hoc = st.session_state.get(
                    "khbd_cap_hoc",
                    "THCS"
                )

                khoi_lop = st.session_state.get(
                    "khbd_khoi_lop",
                    "Không xác định"
                )

                mon_hoc = st.session_state.get(
                    "khbd_mon_hoc",
                    "Không xác định"
                )

                ten_bai = st.session_state.get(
                    "khbd_ten_bai",
                    "Không cung cấp. Căn cứ hoàn toàn vào tài liệu"
                )

                so_tiet = st.session_state.get(
                    "khbd_so_tiet",
                    "1 tiết"
                )

                mau_giao_an = st.session_state.get(
                    "khbd_mau_giao_an",
                    "Công văn 5512"
                )

                ngon_ngu = (
                    "Tiếng Anh"
                    if is_english
                    else "Tiếng Việt"
                )

                thong_tin_bai_day = f"""
Cấp học: {cap_hoc}
Khối lớp: {khoi_lop}
Môn học: {mon_hoc}
Tên bài dạy: {ten_bai}
Thời lượng: {so_tiet}
Mẫu giáo án: {mau_giao_an}
Ngôn ngữ: {ngon_ngu}
"""

                # =========================================
                # NHIỆM VỤ
                # =========================================

                if (
                    st.session_state.soan_mode
                    == "chinh_sua"
                ):

                    lenh_nhiem_vu = """
CHỈNH SỬA VÀ NÂNG CẤP GIÁO ÁN GỐC.

Yêu cầu:

1. Giữ nguyên tên bài.
2. Giữ nguyên kiến thức đúng.
3. Giữ lại các hoạt động tốt.
4. Sửa lỗi chuyên môn nếu phát hiện.
5. Sửa lỗi logic.
6. Bổ sung mục tiêu theo yêu cầu cần đạt.
7. Bổ sung hoạt động GV thật cụ thể.
8. Bổ sung hoạt động HS thật cụ thể.
9. Bổ sung sản phẩm học tập.
10. Bảo đảm tổng thời lượng đúng yêu cầu.
"""

                else:

                    lenh_nhiem_vu = """
SOẠN MỚI KẾ HOẠCH BÀI DẠY TỪ SGK.

Yêu cầu:

1. Bám sát trực tiếp nội dung SGK.
2. Không thêm kiến thức ngoài phạm vi.
3. Xây dựng giáo án hoàn chỉnh theo mẫu.
4. Mô tả chi tiết kịch bản lên lớp.
5. Giáo viên phải có lời dẫn và câu hỏi cụ thể.
6. Học sinh phải có thao tác và câu trả lời dự kiến.
7. Mỗi hoạt động phải có sản phẩm rõ ràng.
"""

                # =========================================
                # TÍCH HỢP
                # =========================================

                if tich_hop_nls:

                    tich_hop_nls_str = (
                        "CÓ.\n"
                        + json.dumps(
                            st.session_state.nls_list,
                            ensure_ascii=False,
                            indent=2
                        )
                    )

                else:

                    tich_hop_nls_str = (
                        "KHÔNG TÍCH HỢP."
                    )

                if tich_hop_ai:

                    tich_hop_ai_str = """
CÓ.

Phải chỉ rõ:

- Học sinh dùng AI ở hoạt động nào.
- Học sinh nhập yêu cầu gì.
- AI hỗ trợ nhiệm vụ gì.
- Sản phẩm AI tạo ra là gì.
- Học sinh kiểm chứng và đánh giá kết quả AI ra sao.
"""

                else:

                    tich_hop_ai_str = (
                        "KHÔNG TÍCH HỢP."
                    )

                if (
                    tich_hop_kt
                    and dang_khuyet_tat_chon
                ):

                    tich_hop_kt_str = (
                        ", ".join(
                            dang_khuyet_tat_chon
                        )
                    )

                else:

                    tich_hop_kt_str = (
                        "KHÔNG TÍCH HỢP."
                    )

                hoat_dong_str = (
                    json.dumps(
                        st.session_state.hoat_dong_list,
                        ensure_ascii=False
                    )
                    if st.session_state.hoat_dong_list
                    else "Không có."
                )

                # =========================================
                # PHẠM VI KIẾN THỨC
                # =========================================

                pham_vi_kien_thuc = (
                    tao_pham_vi_kien_thuc(
                        noi_dung_chinh,
                        noi_dung_ppct,
                        noi_dung_ai_file
                    )
                )

                # =========================================
                # PROMPT
                # =========================================

                prompt = tao_prompt_khbd(
                    thong_tin_bai_day=(
                        thong_tin_bai_day
                    ),
                    pham_vi_kien_thuc=(
                        pham_vi_kien_thuc
                    ),
                    noi_dung_mau_khbd=(
                        noi_dung_mau_khbd
                    ),
                    lenh_nhiem_vu=(
                        lenh_nhiem_vu
                    ),
                    tich_hop_nls_str=(
                        tich_hop_nls_str
                    ),
                    tich_hop_ai_str=(
                        tich_hop_ai_str
                    ),
                    tich_hop_kt_str=(
                        tich_hop_kt_str
                    ),
                    hoat_dong_str=(
                        hoat_dong_str
                    ),
                    mode=(
                        st.session_state.soan_mode
                    )
                )

                # =========================================
                # GỌI AI ENGINE
                # =========================================

                ket_qua_ai = ai_generate_text(
                    ai_engine,
                    prompt
                )

                if not ket_qua_ai:

                    st.error(
                        "❌ AI Engine không trả về nội dung "
                        "hợp lệ hoặc API đang gặp lỗi."
                    )

                    st.stop()

                # =========================================
                # KIỂM TRA KẾT QUẢ
                # =========================================

                kiem_tra = (
                    kiem_tra_ket_qua_ai(
                        ket_qua_ai,
                        ten_bai,
                        so_tiet
                    )
                )

                for warning in (
                    kiem_tra["canh_bao"]
                ):

                    st.warning(
                        "⚠️ " + warning
                    )

                # =========================================
                # LƯU KẾT QUẢ
                # =========================================

                st.session_state[
                    "ket_qua_giao_an"
                ] = ket_qua_ai

                st.success(
                    "🎉 Đã tạo Kế hoạch bài dạy thành công!"
                )

        except Exception as e:

            st.session_state[
                "khbd_last_error"
            ] = traceback.format_exc()

            st.error(
                f"❌ Lỗi hệ thống: {str(e)}"
            )

            with st.expander(
                "🔍 Chi tiết lỗi kỹ thuật"
            ):

                st.code(
                    st.session_state[
                        "khbd_last_error"
                    ]
                )

    # =====================================================
    # HIỂN THỊ KẾT QUẢ
    # =====================================================

    if st.session_state.get(
        "ket_qua_giao_an"
    ):

        st.markdown(
            "### 📝 Kết quả Kế hoạch bài dạy"
        )

        with st.container(
            border=True
        ):

            st.markdown(
                st.session_state[
                    "ket_qua_giao_an"
                ]
            )

        # -------------------------------------------------
        # CHỌN FILE MẪU KHI XUẤT
        # -------------------------------------------------

        template_export = None

        if (
            st.session_state.soan_mode
            == "tu_dong"
        ):

            template_export = (
                st.session_state.get(
                    "file_template_custom"
                )
            )

            if not template_export:

                local_template = (
                    "templates/KHBD_Mau.docx"
                )

                if os.path.exists(
                    local_template
                ):

                    template_export = (
                        local_template
                    )

        try:

            word_file = (
                tao_file_word_hoan_hao(
                    st.session_state[
                        "ket_qua_giao_an"
                    ],
                    template_file=(
                        template_export
                    )
                )
            )

            st.download_button(
                "📥 Tải xuống Giáo án Word",
                data=word_file,
                file_name=(
                    "Giao_An_Thong_Minh.docx"
                ),
                mime=(
                    "application/vnd.openxmlformats-officedocument"
                    ".wordprocessingml.document"
                ),
                use_container_width=True
            )

        except Exception as e:

            st.error(
                f"❌ Lỗi xuất file Word: {str(e)}"
            )
